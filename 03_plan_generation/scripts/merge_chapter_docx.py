from __future__ import annotations

import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docxcompose.composer import Composer


ROOT = Path(__file__).resolve().parents[2]
MERGE_DIR = ROOT / "合并"
CHAPTER_DIR = ROOT / "02_source_assets" / "chapter_drafts"
ORDERED_DIR = MERGE_DIR / "ordered_sources"

ORDERED_SOURCES = [
    ("00_01.docx", MERGE_DIR / "0.1.docx"),
    ("02_03.docx", CHAPTER_DIR / "chapter_02_03.docx"),
    ("04_05.docx", CHAPTER_DIR / "chapter_04_05.docx"),
    ("06.docx", CHAPTER_DIR / "chapter_06.docx"),
    ("07_08.docx", CHAPTER_DIR / "chapter_07_08.docx"),
    ("09_10.docx", CHAPTER_DIR / "chapter_09_10.docx"),
    ("11_12.docx", CHAPTER_DIR / "chapter_11_12.docx"),
    ("13_14.docx", CHAPTER_DIR / "chapter_13_14.docx"),
]

OUTPUT = MERGE_DIR / "SkyGuard_商业计划书_0-14合并版.docx"


def require_sources() -> None:
    missing = [
        str(src)
        for target_name, src in ORDERED_SOURCES
        if not src.exists() and not (ORDERED_DIR / target_name).exists()
    ]
    if missing:
        raise FileNotFoundError("Missing source documents:\n" + "\n".join(missing))


def copy_ordered_sources() -> list[Path]:
    ORDERED_DIR.mkdir(parents=True, exist_ok=True)
    copied: list[Path] = []
    for target_name, src in ORDERED_SOURCES:
        dst = ORDERED_DIR / target_name
        actual_src = src if src.exists() else dst
        if actual_src.resolve() != dst.resolve():
            shutil.copy2(actual_src, dst)
        copied.append(dst)
    return copied


def add_field_char(paragraph, field_char_type: str) -> None:
    field_char = OxmlElement("w:fldChar")
    field_char.set(qn("w:fldCharType"), field_char_type)
    paragraph.add_run()._r.append(field_char)


def add_instruction_text(paragraph, text: str) -> None:
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = text
    paragraph.add_run()._r.append(instr)


def style_page_number_run(result) -> None:
    result.font.name = "宋体"
    result.font.size = Pt(10.5)
    result.font.bold = True
    result.font.color.rgb = RGBColor(0, 0, 0)
    result._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    shade = OxmlElement("w:shd")
    shade.set(qn("w:val"), "clear")
    shade.set(qn("w:color"), "auto")
    shade.set(qn("w:fill"), "FFFFFF")
    result._r.get_or_add_rPr().append(shade)


def add_nested_field(paragraph, code: str, result_text: str = "1") -> None:
    add_field_char(paragraph, "begin")
    add_instruction_text(paragraph, f" {code} ")
    add_field_char(paragraph, "separate")
    style_page_number_run(paragraph.add_run(result_text))
    add_field_char(paragraph, "end")


def add_page_field(paragraph) -> None:
    """Show the page number except on the final page."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_field_char(paragraph, "begin")
    add_instruction_text(paragraph, " IF ")
    add_nested_field(paragraph, "PAGE")
    add_instruction_text(paragraph, " = ")
    add_nested_field(paragraph, "NUMPAGES")
    add_instruction_text(paragraph, ' "" "')
    add_nested_field(paragraph, "PAGE")
    add_instruction_text(paragraph, '" ')
    add_field_char(paragraph, "separate")
    style_page_number_run(paragraph.add_run("1"))
    add_field_char(paragraph, "end")


def clear_story(story) -> None:
    element = story._element
    for child in list(element):
        element.remove(child)


def reset_footer(story) -> None:
    clear_story(story)
    paragraph = story.add_paragraph()
    add_page_field(paragraph)


def normalize_footers(doc: Document) -> None:
    settings = doc.settings.element
    for even_odd in settings.xpath("./w:evenAndOddHeaders"):
        settings.remove(even_odd)

    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    for section in doc.sections:
        section.different_first_page_header_footer = True
        section.footer.is_linked_to_previous = False
        section.first_page_footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False
        reset_footer(section.footer)
        clear_story(section.first_page_footer)
        reset_footer(section.even_page_footer)


def merge_documents(sources: list[Path], output: Path) -> None:
    master = Document(str(sources[0]))
    composer = Composer(master)

    for src in sources[1:]:
        master.add_page_break()
        composer.append(Document(str(src)))

    composer.save(str(output))

    merged = Document(str(output))
    normalize_footers(merged)
    merged.save(str(output))


def assert_no_old_footer_text(path: Path) -> None:
    old_tokens = ["资料口径", "第 ", "页"]
    has_page_field = False
    with ZipFile(path) as zf:
        footer_names = [
            name
            for name in zf.namelist()
            if name.startswith("word/footer") and name.endswith(".xml")
        ]
        if not footer_names:
            raise AssertionError("No footer XML parts found in merged document.")
        for name in footer_names:
            text = zf.read(name).decode("utf-8", errors="ignore")
            has_page_field = has_page_field or "PAGE" in text
            for token in old_tokens:
                if token in text:
                    raise AssertionError(f"{name} still contains old footer token: {token!r}")
    if not has_page_field:
        raise AssertionError("No footer XML part contains a PAGE field.")


def main() -> None:
    require_sources()
    sources = copy_ordered_sources()
    merge_documents(sources, OUTPUT)
    assert_no_old_footer_text(OUTPUT)
    print(f"ordered_sources={MERGE_DIR / 'ordered_sources'}")
    print(f"output={OUTPUT}")
    print("order=" + " -> ".join(path.name for path in sources))


if __name__ == "__main__":
    main()
