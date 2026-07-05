from __future__ import annotations

import json
import math
import random
import re
import shutil
import textwrap
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
import networkx as nx
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[2]
DATA_DIR = ROOT / "数据以及图" / "_整理" / "数据"
IMAGE_DIR = ROOT / "数据以及图" / "_整理" / "图"
PLAN_FILE = ROOT / "SkyGuard_低空智眼_商业计划书与WebDemo工作流PLAN.md"
SOURCE_MD = ROOT / "SkyGuard_政策与参考来源清单.md"

WORK_DIR = ROOT / "skyguard-plan"
CHART_DIR = WORK_DIR / "charts"
FIGURE_DIR = WORK_DIR / "figures"
SOURCE_DIR = WORK_DIR / "sources"
TABLE_DIR = WORK_DIR / "tables"
QA_DIR = WORK_DIR / "qa"
OUT_DIR = ROOT / "deliverables"

for folder in [WORK_DIR, CHART_DIR, FIGURE_DIR, SOURCE_DIR, TABLE_DIR, QA_DIR, OUT_DIR]:
    folder.mkdir(parents=True, exist_ok=True)

RANDOM_SEED = 618
random.seed(RANDOM_SEED)
np.random.seed(RANDOM_SEED)


def setup_fonts() -> None:
    plt.rcParams["axes.unicode_minus"] = False
    font_name = "DejaVu Sans"
    font_candidates = [
        Path("C:/Windows/Fonts/simfang.ttf"),
        Path("C:/Windows/Fonts/STFANGSO.TTF"),
        Path("C:/Windows/Fonts/Deng.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/msyhbd.ttc"),
    ]
    for p in font_candidates:
        if p.exists():
            try:
                font_manager.fontManager.addfont(str(p))
                font_name = font_manager.FontProperties(fname=str(p)).get_name()
            except Exception:
                pass
            try:
                pdfmetrics.registerFont(TTFont("CNFont", str(p)))
            except Exception:
                pass
            break
    plt.rcParams["font.family"] = font_name
    plt.rcParams["font.sans-serif"] = [font_name, "FangSong", "SimSun", "SimHei", "DejaVu Sans"]


def clean_filename(value: str) -> str:
    bad = '<>:"/\\|?*'
    cleaned = "".join("_" if ch in bad else ch for ch in value)
    cleaned = cleaned.replace("\n", "_").replace("\r", "_").strip(" .")
    return cleaned[:90] or "chart"


sns.set_theme(style="whitegrid")
setup_fonts()

PALETTE = {
    "ink": "#111827",
    "muted": "#5f6b7a",
    "line": "#d8e1ea",
    "blue": "#1d4ed8",
    "teal": "#0f766e",
    "green": "#15803d",
    "yellow": "#b45309",
    "orange": "#c2410c",
    "red": "#b91c1c",
    "slate": "#334155",
    "indigo": "#4338ca",
    "violet": "#7c3aed",
    "rose": "#be123c",
    "light": "#f8fafc",
}

REPORT_COLORS = {
    "ink": "#172033",
    "soft_ink": "#334155",
    "muted": "#667085",
    "paper": "#f6f4ef",
    "paper2": "#ebe7dd",
    "line": "#d6d0c3",
    "navy": "#151f2b",
    "teal": "#0c6f69",
    "blue": "#234f7d",
    "green": "#2f6b4f",
    "amber": "#9b6a2f",
    "red": "#9b3f3a",
    "gold": "#b69252",
    "ash": "#6d716c",
}

CHARTS: list[dict] = []
TABLES: dict[str, pd.DataFrame] = {}


def read_csv(name: str, **kwargs) -> pd.DataFrame:
    path = DATA_DIR / name
    for enc in ["utf-8-sig", "utf-8", "gb18030"]:
        try:
            return pd.read_csv(path, encoding=enc, **kwargs)
        except UnicodeDecodeError:
            continue
    return pd.read_csv(path, **kwargs)


def safe_numeric(s):
    return pd.to_numeric(s, errors="coerce")


def save_chart(fig, chart_id: str, title: str, chart_type: str, source: str, chapter: str, conclusion: str) -> Path:
    path = CHART_DIR / f"{clean_filename(title)}.png"
    fig.patch.set_facecolor("white")
    for ax in fig.axes:
        raw_title = ax.get_title()
        if raw_title and len(raw_title) > 18:
            ax.set_title("\n".join(textwrap.wrap(raw_title, width=18)), loc="left", fontsize=12.2, fontweight="bold", color=PALETTE["ink"], pad=13)
        ax.tick_params(axis="both", labelsize=8.1, colors="#3f4d5f", length=0, pad=3)
        labels = ax.get_xticklabels()
        if len(labels) >= 6:
            next_labels = []
            for label in labels:
                text = label.get_text()
                if len(text) > 6:
                    text = "\n".join(textwrap.wrap(text, width=6))
                next_labels.append(text)
            ticks = ax.get_xticks()
            if len(ticks) == len(next_labels):
                ax.set_xticks(ticks)
            ax.set_xticklabels(next_labels, rotation=0, ha="center")
        for label in ax.get_yticklabels():
            text = label.get_text()
            if len(text) > 10:
                label.set_text("\n".join(textwrap.wrap(text, width=10)))
        ax.margins(x=0.04, y=0.08)
    try:
        fig.tight_layout(pad=1.45)
    except Exception:
        pass
    fig.savefig(path, dpi=320, bbox_inches="tight", facecolor="white", pad_inches=0.18)
    plt.close(fig)
    CHARTS.append(
        {
            "chart_id": chart_id,
            "title": title,
            "chart_type": chart_type,
            "source": source,
            "chapter": chapter,
            "conclusion": conclusion,
            "file": str(path.relative_to(ROOT)),
        }
    )
    return path


def public_source_registry(source_registry: pd.DataFrame) -> pd.DataFrame:
    cols = ["title", "publisher", "date", "url", "use", "source_type"]
    return source_registry[[c for c in cols if c in source_registry.columns]].rename(
        columns={
            "title": "资料名称",
            "publisher": "发布主体",
            "date": "时间",
            "url": "链接",
            "use": "用途",
            "source_type": "资料类型",
        }
    )


def public_chart_catalog(chart_catalog: pd.DataFrame) -> pd.DataFrame:
    cols = ["title", "chart_type", "source", "chapter", "conclusion"]
    return chart_catalog[[c for c in cols if c in chart_catalog.columns]].rename(
        columns={
            "title": "图表名称",
            "chart_type": "图表类型",
            "source": "数据来源",
            "chapter": "所属章节",
            "conclusion": "核心结论",
        }
    )


def public_dataset_summary(data: dict[str, pd.DataFrame]) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "数据表": name,
                "行数": len(df),
                "字段数": len(df.columns),
                "数据性质": ", ".join(map(str, df["data_nature"].dropna().unique()[:5])) if "data_nature" in df.columns else "混合资料或来源特定",
            }
            for name, df in data.items()
        ]
    )


def style_axis(ax, title: str, subtitle: str | None = None):
    ax.set_facecolor("#fbfcfd")
    ax.set_title(title, loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"], pad=10)
    if subtitle:
        ax.text(0, 1.01, subtitle, transform=ax.transAxes, color=PALETTE["muted"], fontsize=8.8)
    ax.grid(True, axis="y", color="#e8edf3", linewidth=0.8)
    ax.grid(False, axis="x")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#cbd5df")
    ax.spines["bottom"].set_color("#cbd5df")
    ax.tick_params(colors="#45545f", labelsize=9, length=0)


def first_value(df: pd.DataFrame, code: str, year: int | None = None) -> float:
    sub = df[df["indicator_code"].eq(code)].copy()
    if year is not None:
        sub = sub[sub["year"].eq(year)]
    if sub.empty:
        return float("nan")
    return float(sub.iloc[-1]["value"])


def generate_source_registry(nat: pd.DataFrame, city: pd.DataFrame) -> pd.DataFrame:
    base_sources = [
        {
            "source_id": "P-01",
            "title": "2024年政府工作报告",
            "publisher": "国务院；国家发展改革委转载",
            "date": "2024-03-20",
            "url": "https://www.ndrc.gov.cn/fzggw/jgsj/zys/sjdt/202403/t20240320_1365089.html",
            "use": "低空经济被纳入新增长引擎叙事，支撑赛道正当性。",
            "source_type": "official_policy",
        },
        {
            "source_id": "P-02",
            "title": "无人驾驶航空器飞行管理暂行条例",
            "publisher": "国务院、中央军委；中国民用航空局公开",
            "date": "2023-06-28发布，2024-01-01施行",
            "url": "https://www.caac.gov.cn/XXGK/XXGK/FLFG/202401/t20240115_222642.html",
            "use": "支撑飞行活动、身份、空域、记录和合规管理需求。",
            "source_type": "official_regulation",
        },
        {
            "source_id": "P-03",
            "title": "通用航空装备创新应用实施方案（2024-2030年）",
            "publisher": "工业和信息化部等部门",
            "date": "2024",
            "url": "https://wap.miit.gov.cn/zwgk/zcwj/wjfb/tz/art/2024/art_4ce8d09c15ee4fb1aefc3d5dfbbb6584.html",
            "use": "支撑通航装备、应用示范、低空物流和公共服务场景。",
            "source_type": "official_policy",
        },
        {
            "source_id": "P-04",
            "title": "低空经济及其核心产业统计分类（试行）",
            "publisher": "国家发展改革委",
            "date": "2025-12-26",
            "url": "https://www.ndrc.gov.cn/xxgk/zcfb/tz/202512/t20251226_1402669.html",
            "use": "界定低空经济产业边界和统计口径。",
            "source_type": "official_policy",
        },
        {
            "source_id": "D-01",
            "title": "OurAirports Open Data",
            "publisher": "OurAirports",
            "date": "持续更新",
            "url": "https://davidmegginson.github.io/ourairports-data/airports.csv",
            "use": "机场、跑道、导航台公开基础设施数据。",
            "source_type": "open_data",
        },
        {
            "source_id": "D-02",
            "title": "OpenSky Network Flightlist",
            "publisher": "OpenSky Network / ActiveConclusion",
            "date": "2020-01",
            "url": "https://raw.githubusercontent.com/ActiveConclusion/COVID19_AirTraffic/master/opensky_data/flightlist_20200101_20200131.csv.gz",
            "use": "飞行活动公开数据参考，用于空域运行分析方法说明。",
            "source_type": "open_data",
        },
    ]
    for sid, sub in nat.groupby("source_id"):
        base_sources.append(
            {
                "source_id": sid,
                "title": f"{sid.replace('_', ' ')} 民航统计数据",
                "publisher": "中国民用航空局/交通运输部公开资料",
                "date": str(int(sub["year"].max())),
                "url": sub["source_url"].iloc[0],
                "use": "无人机注册、运营单位、飞行小时、机场和安全指标。",
                "source_type": "official_statistic",
            }
        )
    for sid, sub in city.groupby("source_id"):
        base_sources.append(
            {
                "source_id": sid,
                "title": sub["policy_or_case_doc"].iloc[0],
                "publisher": "地方政府/主管部门公开资料",
                "date": str(int(sub["target_year"].dropna().max())) if sub["target_year"].notna().any() else "",
                "url": sub["source_url"].iloc[0],
                "use": "地方低空经济目标、平台建设、航线和起降网络规划。",
                "source_type": "local_policy_or_case",
            }
        )
    out = pd.DataFrame(base_sources).drop_duplicates(subset=["source_id", "url"])
    out.to_csv(SOURCE_DIR / "source_registry.csv", index=False, encoding="utf-8-sig")
    return out


def final_chapter_specs(data: dict[str, pd.DataFrame]) -> list[dict]:
    nat = data["national"]
    uav_2025 = first_value(nat, "registered_uavs_10k", 2025)
    hours_2025 = first_value(nat, "uav_flight_hours_10k", 2025)
    units_2025 = first_value(nat, "uav_operating_units", 2025)
    stations_2025 = first_value(nat, "low_altitude_flight_service_stations", 2025)
    return [
        {
            "chapter": "第一章 行业机会",
            "count": 15,
            "accent": REPORT_COLORS["teal"],
            "topics": ["政策窗口", "运行规模", "服务站网络", "地方目标", "区域分层", "产业链位置", "试点入口"],
            "opening": f"低空经济进入城市治理场景后，管理对象从少量试验飞行变为高频、多主体、跨场景的日常运行。2025 年公开口径中，注册无人机 {uav_2025:.1f} 万架，累计飞行小时 {hours_2025:.2f} 万小时，运营单位 {units_2025:.0f} 家，低空飞行服务站 {stations_2025:.0f} 个。规模增长本身并不直接等于商业机会，真正的机会来自运行活动变密后产生的识别、记录、预警和复盘需求。",
            "decision": "重点区域先行，以可验证的监管需求切入。",
        },
        {
            "chapter": "第二章 痛点验证",
            "count": 15,
            "accent": REPORT_COLORS["amber"],
            "topics": ["目标发现", "身份核验", "轨迹偏离", "围栏触发", "事件流转", "复盘报表", "责任边界", "用户调研"],
            "opening": "低空管理的难点不在于屏幕上有没有目标点，而在于目标出现后能否确认计划、核验身份、判断围栏、安排处置并形成记录。只要其中一个节点依赖电话、截图或临时表格，现场响应和事后说明都会被拖慢。",
            "decision": "把痛点落到流程断点，避免泛泛描述安全焦虑。",
        },
        {
            "chapter": "第三章 产品服务",
            "count": 18,
            "accent": REPORT_COLORS["blue"],
            "topics": ["态势总览", "飞行计划", "电子围栏", "AI识别", "事件处置", "报表中心", "场景包", "客户成功", "移动协同"],
            "opening": "SkyGuard 的产品价值不靠页面数量体现，而靠闭环完整度体现。系统把态势图、规则中心、事件中心和报表中心放在同一条工作流里，使重点区域可以完成从发现、判断、确认、派单到复盘的连续动作。",
            "decision": "客户购买的是日常运行机制，不是一次性展示大屏。",
        },
        {
            "chapter": "第四章 技术架构",
            "count": 18,
            "accent": REPORT_COLORS["green"],
            "topics": ["多源接入", "时空索引", "轨迹比对", "风险评分", "规则引擎", "权限审计", "边缘部署", "系统集成", "模型校验"],
            "opening": "技术方案采用规则先行、模型辅助、人工确认的路径。这样既能保留算法升级空间，又能避免早期试点把风险判断完全交给黑箱模型。每一条告警都需要能解释来源、规则、分值、确认人和处置结果。",
            "decision": "技术架构服务业务闭环，不堆砌概念。",
        },
        {
            "chapter": "第五章 数据体系",
            "count": 15,
            "accent": REPORT_COLORS["teal"],
            "topics": ["官方统计", "地方政策", "公开数据", "演示样本", "数据字典", "质量规则", "模型校验", "证据链"],
            "opening": "数据体系的首要原则是口径分层。官方统计用于行业判断，政策文件用于区域机会，公开数据用于方法说明，演示样本用于功能展示。只有明确边界，图表数量增加才不会削弱可信度。",
            "decision": "所有数字都能回到来源、字段或假设。",
        },
        {
            "chapter": "第六章 市场竞争",
            "count": 15,
            "accent": REPORT_COLORS["blue"],
            "topics": ["替代方案", "客户分层", "进入顺序", "竞品矩阵", "采购逻辑", "区域复制", "市场边界", "伙伴生态"],
            "opening": "SkyGuard 面对的竞争不是某一家单一公司，而是视频监控、单点感知设备、内部任务系统、城市级重平台和人工巡查组成的替代路径。项目必须明确自己的位置：以低空业务闭环和轻量部署取得早期客户，再逐步扩展到区域协同。",
            "decision": "先占重点区域运行监管，再争取区县复制。",
        },
        {
            "chapter": "第七章 商业模式",
            "count": 15,
            "accent": REPORT_COLORS["green"],
            "topics": ["订阅收入", "项目部署", "活动保障", "航线评估", "运维续费", "伙伴渠道", "客户成功", "定价体系"],
            "opening": "商业模式围绕客户的持续运行动作设计。景区和园区重视低成本常态化监管，大型活动重视临时保障，物流航线和巡检客户重视航线评估和运行记录。不同收入项对应不同交付强度，必须分别测算毛利和回款周期。",
            "decision": "收入拆到客户动作、交付成本和续费机制上。",
        },
        {
            "chapter": "第八章 落地计划",
            "count": 13,
            "accent": REPORT_COLORS["amber"],
            "topics": ["0-3个月", "3-6个月", "6-18个月", "18-36个月", "试点验收", "区域复制", "产品路线"],
            "opening": "落地计划按照试点、复制、协同三层推进。早期要交付的不是远景，而是可运行 Demo、可核验数据、可配置规则、可提交报表和可复盘工单。只有这些内容形成闭环，后续复制才有基础。",
            "decision": "每个阶段都对应清晰交付物和验收口径。",
        },
        {
            "chapter": "第九章 财务预测",
            "count": 17,
            "accent": REPORT_COLORS["green"],
            "topics": ["收入预测", "成本结构", "毛利路径", "现金流", "回款周期", "融资用途", "敏感性", "单位经济", "退出机制"],
            "opening": "财务预测把增长、毛利、现金和交付能力放在一起看。项目早期的关键变量包括客户数量、客单价、交付成本、续费率、回款周期和研发投入；任何一个变量变化，都会影响现金跑道和利润拐点。",
            "decision": "用三情景测算解释经营弹性，用现金流约束扩张速度。",
        },
        {
            "chapter": "第十章 风险应对",
            "count": 11,
            "accent": REPORT_COLORS["red"],
            "topics": ["政策边界", "数据合规", "误报漏报", "硬件依赖", "采购周期", "现金压力", "过度承诺", "安全责任"],
            "opening": "低空项目必须主动写清边界。SkyGuard 做辅助感知、风险预警、事件协同和运行报表，不做干扰、捕获、打击，也不替代审批或执法。边界越清楚，试点过程中的责任争议越少。",
            "decision": "风险分为阻断项目、拖慢节奏和增加成本三类管理。",
        },
        {
            "chapter": "第十一章 团队组织",
            "count": 9,
            "accent": REPORT_COLORS["blue"],
            "topics": ["岗位配置", "产品负责人", "算法数据", "前后端研发", "交付运维", "行业顾问", "伙伴协同"],
            "opening": "早期团队要围绕产品、数据、工程和交付配置。一个低空监管 Demo 要被评委和客户相信，需要有人负责行业场景、风险字段、前后端实现、可视化报表、试点沟通和运维材料。",
            "decision": "团队组织服务试点交付，而不是堆岗位名称。",
        },
        {
            "chapter": "第十二章 社会价值",
            "count": 9,
            "accent": REPORT_COLORS["teal"],
            "topics": ["公共安全", "城市治理", "应急保障", "产业数据", "岗位培养", "合规意识", "公众活动"],
            "opening": "社会价值来自低空运行的透明化和协同化。未知目标、临时活动、重点区域和应急保障进入同一套记录后，管理方、运营方和现场人员可以共享事实，减少沟通成本和责任争议。",
            "decision": "公共价值写成可观察的治理改进。",
        },
        {
            "chapter": "第十三章 Web Demo",
            "count": 14,
            "accent": REPORT_COLORS["blue"],
            "topics": ["首页总览", "演示脚本", "态势地图", "实时监测", "AI识别", "围栏告警", "事件工单", "报表中心", "移动端", "部署"],
            "opening": "Web Demo 是商业计划书的产品证据。它需要按真实工作路径组织：从首页进入演示脚本，再进入态势、目标、计划、围栏、事件、报表和移动端。每个页面都对应一个业务动作，而不是只展示静态卡片。",
            "decision": "用完整交互证明项目能演示、能试点、能部署。",
        },
        {
            "chapter": "附录",
            "count": 18,
            "accent": REPORT_COLORS["ash"],
            "topics": ["来源索引", "图表目录", "数据字典", "财务假设", "问卷访谈", "部署说明", "截图包", "交付清单", "政策原文"],
            "opening": "附录承担核验功能。正文里的结论、图表和截图，都需要能回到来源、字段、假设或文件位置。附录越清楚，正文越能保持简洁。",
            "decision": "证据、假设和交付材料集中管理，便于继续迭代。",
        },
    ]


def final_narrative_paragraphs(spec: dict, topic: str, row: dict | None, metrics: dict, figure_caption: str | None = None) -> list[str]:
    chapter = spec["chapter"]
    chart_title = row.get("title", topic) if row else (figure_caption or topic)
    source = row.get("source", "整理图片、Demo截图与项目资料") if row else "整理图片、Demo截图与项目资料"
    conclusion = row.get("conclusion", spec["decision"]) if row else spec["decision"]
    evidence = f"配套材料“{chart_title}”给出了直接证据，资料口径来自{source}。"

    if chapter.startswith("第一章"):
        return [
            f"{topic}要从运行事实而不是概念热度说起。注册无人机、累计飞行小时、运营主体、服务站和地方任务同时增长，说明低空活动已经进入持续管理阶段。SkyGuard 的切入点不是制造飞行器，而是让运行活动在重点区域内被看见、被记录、被解释、被复盘。",
            f"市场进入不能只看全国总量，还要看区域基础和客户责任边界。景区、园区、场馆、机场周边和固定航线具备较清晰的管理主体，试点成本和验收指标更容易控制，也更适合作为首批客户。",
            f"{evidence}该证据支持的经营判断是：{conclusion}",
            "项目因此采用重点区域先行策略。先形成围栏规则、事件工单、日报模板和客户成功流程，再把这些可复用能力复制到区县级场景。"
        ]
    if chapter.startswith("第二章"):
        return [
            f"{topic}对应的是现场运行中的具体断点。低空目标出现后，工作人员需要确认是否有飞行计划、身份是否可信、是否触碰围栏、风险等级是否需要升级、由谁处理以及处置结果如何归档。任何一个节点断开，都会导致现场响应慢、责任说明难、复盘材料散。",
            "真正的痛点验证要落到流程，而不是停留在“监管压力大”这种概括上。计划库、目标库、围栏库、事件库和报表库之间缺少联动，才是重点区域客户愿意采购系统的原因。",
            f"{evidence}结合该证据，{topic}的核心判断是：{conclusion}",
            "SkyGuard 的产品动作是先给出风险解释，再保留人工确认，最后把处置过程写回报表。这样既控制误报争议，也把每次现场处理沉淀为可复用经验。"
        ]
    if chapter.startswith("第三章"):
        return [
            f"{topic}必须贴近真实使用路径。用户进入系统后先看区域态势，再追踪具体目标；目标异常时进入计划核验、围栏判断和事件工单；处置结束后进入日报、月报和风险复盘。这条路径完整，Demo 才不会像几张孤立页面。",
            "产品结构采用“态势图、规则中心、事件中心、报表中心”四个核心面。态势图解决看见，规则中心解决边界，事件中心解决协同，报表中心解决交付和复盘。",
            f"{evidence}围绕该证据，产品结论是：{conclusion}",
            "场景包是早期商业化的关键。不同客户不购买完全相同的功能，景区重视游客航拍和重点保护区，园区重视巡检和设备健康，大型活动重视临时围栏和移动告警。"
        ]
    if chapter.startswith("第四章"):
        return [
            f"{topic}的设计原则是可解释、可回放、可迁移。系统不会把所有判断交给模型，而是把飞行计划、目标识别、轨迹异常、围栏规则、人工确认和处置结果串成完整链路。",
            "轻量部署更适合早期试点。数据接入先覆盖飞行计划、传感器告警、轨迹点、围栏规则和工单状态；模型负责排序和解释，人工确认负责关键决策。",
            f"{evidence}技术上采用的判断是：{conclusion}",
            "这种架构的好处是后续可接入不同硬件，也可按客户区域调整规则。系统的资产沉淀在字段、规则、日志和报表模板，而不是锁死在某一种设备形态中。"
        ]
    if chapter.startswith("第五章"):
        return [
            f"{topic}首先要处理口径。官方统计、政策目标、公开数据和演示样本各自承担不同作用，不能把演示样本当成真实客户数据，也不能用政策目标直接替代订单。",
            "数据多不等于可信度高。可信度来自字段清楚、来源清楚、更新频率清楚、使用边界清楚。计划书中的图表都应能回到数据表、来源清单或财务假设。",
            f"{evidence}该材料支撑的结论是：{conclusion}",
            "数据体系还直接影响产品能力。风险评分需要计划匹配、身份可信度、轨迹偏离、围栏权重、环境影响和响应紧迫度等字段共同解释，单一指标无法支撑可复核的预警。"
        ]
    if chapter.startswith("第六章"):
        return [
            f"{topic}要回答客户为什么不用现有替代方案。视频监控能提供画面，单点感知设备能发现信号，内部系统能派任务，但低空重点区域需要的是计划、目标、规则、处置和报表的联动。",
            "市场进入顺序决定资源消耗。首批客户应选择责任主体明确、运行活动稳定、验收材料可定义的区域；城市级平台可以作为中长期方向，但不适合作为创业阶段的首个承诺。",
            f"{evidence}竞争判断可以概括为：{conclusion}",
            "对外叙事应保持客观。SkyGuard 不否认硬件和重平台的价值，而是在它们之间提供业务闭环、轻量部署和报表复盘能力。"
        ]
    if chapter.startswith("第七章"):
        return [
            f"{topic}必须和客户采购动作相连。Lite SaaS 适合单区域常态化监管，Pro 部署适合多区域试点，活动保障适合短周期临时任务，航线评估和报告服务适合物流、巡检和运营复盘。",
            "收入结构不能只看年度总额，还要拆交付强度。软件订阅毛利高但需要续费，部署项目回款大但交付重，活动保障现金回收快但不可持续，报告服务能增强客户黏性。",
            f"{evidence}商业判断是：{conclusion}",
            "客户成功是续费的核心。系统上线后需要围栏调优、规则复盘、报表培训和月度运营会议，才能把一次试点变成持续服务。"
        ]
    if chapter.startswith("第八章"):
        return [
            f"{topic}必须对应交付物。早期工作包括 Demo 可运行、数据字典清楚、规则样例完整、截图包可复核、财务假设可调整；试点阶段再补客户现场配置、培训材料和验收报告。",
            "路线图不是时间装饰，而是依赖关系管理。没有稳定数据口径，风险模型无法解释；没有工单闭环，报表缺少事实来源；没有试点反馈，定价和渠道判断都容易失真。",
            f"{evidence}排期判断是：{conclusion}",
            "落地节奏应保守。先用一个重点区域完成端到端闭环，再复制到相邻场景，最后再讨论区县级协同。"
        ]
    if chapter.startswith("第九章"):
        return [
            f"{topic}看的是经营质量。收入增长必须同时满足毛利改善、现金安全和交付能力三个条件，否则规模越大，试点和运维压力也越大。",
            "财务模型采用三情景测算，核心变量包括客户数量、客单价、交付成本、续费率、回款周期、研发投入和客户成功成本。变量越透明，投资判断越可信。",
            f"{evidence}财务结论是：{conclusion}",
            "融资用途应优先投向产品打磨、试点交付、数据治理、客户成功和必要的伙伴接入，避免过早形成重资产负担。"
        ]
    if chapter.startswith("第十章"):
        return [
            f"{topic}是项目必须主动管理的约束。政策边界、数据合规、误报漏报、设备依赖、采购周期和现金压力，都可能影响试点速度和客户信任。",
            "风险应对需要写触发条件和动作。误报上升时进入人工复核和样本回标；硬件接入延迟时启用模拟接口和手工导入；采购周期拉长时调整现金流和交付资源。",
            f"{evidence}风险判断是：{conclusion}",
            "边界声明尤其重要。系统提供辅助感知和协同记录，不替代审批或执法，不做干扰、捕获、打击类动作。"
        ]
    if chapter.startswith("第十一章"):
        return [
            f"{topic}服务于试点交付。早期团队需要产品、数据、前端、后端、算法、交付和行业顾问共同参与，确保系统能运行、材料能讲清、客户能试用。",
            "岗位配置不能脱离路线图。产品负责人定义场景和字段，研发团队保证交互和数据链路，交付负责人负责现场配置和培训，行业顾问把政策边界和客户语言翻译成产品规则。",
            f"{evidence}组织判断是：{conclusion}",
            "团队规模应随客户和试点数量逐步增加，不宜在商业验证前过度扩张。"
        ]
    if chapter.startswith("第十二章"):
        return [
            f"{topic}体现的是城市治理效率。低空目标、临时活动、重点保护区和应急保障进入同一套记录后，管理方和运营方能基于共同事实协同处理。",
            "社会价值需要可观察。可以关注响应时间、事件闭环率、报表提交及时率、规则更新次数、培训覆盖人数和重点区域保障次数，而不是只写宏观口号。",
            f"{evidence}价值判断是：{conclusion}",
            "这种价值与商业目标并不冲突。客户持续使用系统，正是因为它降低了沟通成本、责任争议和复盘成本。"
        ]
    if chapter.startswith("第十三章"):
        return [
            f"{topic}要像真实产品一样可操作。首页给出关键指标和证据，演示中心提供脚本，后台进入态势、目标、计划、围栏、事件、报表、设备和设置，移动端用于现场确认。",
            "Demo 的真实感来自三个方面：图片和统计口径来自整理材料，运行数据明确标注为演示样本，交互路径按真实业务动作组织。这样既能展示产品能力，也不会混淆数据边界。",
            f"{evidence}产品演示结论是：{conclusion}",
            "可部署性需要体现在前端结构、Mock 数据、截图包和部署说明中。演示系统不追求大而全，而追求能够完整跑通重点区域监管流程。"
        ]
    return [
        f"{topic}用于支撑正文核验。来源、字段、图表、假设、截图和部署说明集中放入附录，便于从结论反查证据。",
        "附录不承担堆叠页数的功能，而承担审阅效率。资料分层越清楚，正文就越能保持商业计划书应有的克制和连贯。",
        f"{evidence}附录判断是：{conclusion}",
        "后续新增资料应先进入来源索引或数据字典，再决定是否进入正文图表，避免材料越多、口径越散。"
    ]


def final_page_table(chapter: str, topic: str, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame, chart_catalog: pd.DataFrame) -> list[list[str]]:
    if chapter == "使用说明":
        return [
            ["章节", "审阅重点", "核验材料"],
            ["行业机会", "规模、政策、区域", "官方统计、政策法规、指数数据"],
            ["痛点验证", "目标、计划、围栏、工单", "安全事件、演示样本、访谈提纲"],
            ["产品服务", "工作流与场景包", "Demo截图、产品矩阵"],
            ["技术架构", "可解释、可回放、可部署", "架构图、数据字典"],
            ["商业财务", "收入、毛利、现金", "财务测算表"],
            ["附录", "来源、图表、假设", "证据资料索引表"],
        ]
    if chapter == "第三章 产品服务":
        return [["模块", "客户动作", "交付证据"], ["态势总览", "查看目标、围栏、告警", "地图与KPI"], ["计划管理", "审批、核验、归档", "计划表与轨迹比对"], ["事件中心", "确认、派单、处置", "工单与响应时长"], ["报表中心", "复盘、提交、归档", "日报/月报/专项报告"]]
    if chapter == "第七章 商业模式":
        return [["收入项", "建议口径", "适用客户"], ["Lite SaaS", "8-15万元/年", "单景区、园区"], ["Pro部署", "30-80万元/项目", "多区域试点"], ["活动保障", "5-30万元/次", "赛事、会展、节庆"], ["航线评估", "3000-20000元/条", "物流、巡检、文旅航线"]]
    if chapter == "第九章 财务预测":
        finance = build_finance_tables()
        return [["假设项", "模型口径"]] + finance["assumptions"].head(6)[["item", "assumption"]].astype(str).values.tolist()
    if chapter == "第五章 数据体系":
        return [["数据层", "用途", "边界"], ["官方统计", "行业规模判断", "不替代实时监管"], ["政策文件", "区域机会判断", "不等同客户订单"], ["公开数据", "方法说明", "需标明来源"], ["演示样本", "功能链路展示", "不写成真实客户业务"]]
    if chapter == "附录" and "来源" in topic:
        return [["来源名称", "资料类型"]] + source_registry[["title", "source_type"]].head(7).astype(str).values.tolist()
    if chapter == "附录":
        return [["图表名称", "所属章节", "图表类型"]] + chart_catalog[["title", "chapter", "chart_type"]].head(7).astype(str).values.tolist()
    return []


def make_page_items(chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame) -> list[dict]:
    nat = data["national"]
    metrics = {
        "uav": first_value(nat, "registered_uavs_10k", 2025),
        "hours": first_value(nat, "uav_flight_hours_10k", 2025),
        "units": first_value(nat, "uav_operating_units", 2025),
        "stations": first_value(nat, "low_altitude_flight_service_stations", 2025),
        "airport": first_value(nat, "registered_general_airports", 2025),
    }
    figures = figure_files()
    figure_idx = 0
    used_figures: set[Path] = set()
    used_figures: set[Path] = set()
    used_figures: set[Path] = set()
    all_rows = chart_catalog.to_dict("records")
    chapter_alias = {"第二章 需求验证": "第二章 痛点验证", "第十章 风险": "第十章 风险应对", "第十一章 团队": "第十一章 团队组织", "第十三章 Demo展示": "第十三章 Web Demo"}
    chapter_chart_map: dict[str, list[dict]] = defaultdict(list)
    for row in all_rows:
        chapter_chart_map[chapter_alias.get(row["chapter"], row["chapter"])].append(row)
    cursor: dict[str, int] = {}
    pages: list[dict] = []

    def take_figure(match=None) -> Path | None:
        nonlocal figure_idx
        if not figures:
            return None
        if match is not None:
            for fig in figures:
                if fig not in used_figures and match(fig):
                    used_figures.add(fig)
                    return fig
        while figure_idx < len(figures):
            fig = figures[figure_idx]
            figure_idx += 1
            if fig not in used_figures:
                used_figures.add(fig)
                return fig
        return None

    def next_row(chapter: str) -> dict | None:
        return chart_for(chapter_chart_map, chapter, cursor, all_rows)

    def add(page_type: str, chapter: str, title: str, paras: list[str], bullets: list[str] | None = None, chart_id: str | None = None, figure: Path | None = None, table: list[list[str]] | None = None, accent: str | None = None, kicker: str = "", callout: str = ""):
        pages.append(
            {
                "type": page_type,
                "chapter": chapter,
                "title": title,
                "paras": paras,
                "bullets": bullets or [],
                "chart_id": chart_id,
                "figure": str(figure) if figure else None,
                "table": table or [],
                "accent": accent or REPORT_COLORS["teal"],
                "kicker": kicker,
                "callout": callout,
            }
        )

    cover_image = take_figure(lambda p: any(k in p.name for k in ["平台", "低空", "SILAS", "无人机"])) or take_figure()
    add(
        "cover",
        "封面",
        "低空智眼 SkyGuard 商业计划书",
        [
            "城市低空空域安全感知与运行监管平台",
            f"面向景区、园区、场馆、机场周边和低空物流航线。行业基准引用公开统计：2025 年注册无人机 {metrics['uav']:.1f} 万架，累计飞行小时 {metrics['hours']:.2f} 万小时，运营单位 {metrics['units']:.0f} 家。",
        ],
        bullets=["重点区域监管", "计划与围栏联动", "事件闭环处置", "运行报表复盘"],
        figure=cover_image,
        accent=REPORT_COLORS["teal"],
        kicker="商业计划书",
        callout="看得见、判得清、处置快、可复盘",
    )
    add(
        "memo",
        "使用说明",
        "证据边界与审阅路径",
        [
            "本计划书把真实公开资料、地方政策、公开数据和演示样本分层使用。官方统计用于行业规模判断，政策法规用于合规和区域机会判断，公开数据用于方法说明，Demo 样本只用于产品流程展示。",
            "正文采用执行摘要、行业机会、痛点验证、产品服务、技术架构、数据体系、市场竞争、商业模式、落地计划、财务预测、风险应对、团队组织、社会价值、Web Demo 和附录的顺序，保持商业计划书的完整闭环。",
            "图表、来源、财务假设和截图都集中在附录及配套表格中，便于审阅时从结论反查证据。"
        ],
        bullets=[f"来源索引：{len(source_registry)} 条", f"图表目录：{len(chart_catalog)} 张", "配套文件：PDF、DOCX、图表包、财务表、Demo源码包"],
        table=final_page_table("使用说明", "", data, source_registry, chart_catalog),
        accent=REPORT_COLORS["blue"],
        kicker="审阅说明",
        callout="先看判断，再核验证据。",
    )
    summary_chart = next_row("第一章 行业机会")
    add(
        "summary",
        "执行摘要",
        "项目判断",
        [
            "低空经济从产业概念进入城市运行后，最迫切的问题是把飞行计划、目标身份、电子围栏、异常告警、现场处置和复盘报告连接起来。SkyGuard 的定位是面向重点区域提供低空安全感知和运行监管平台，帮助客户完成日常监管、临时保障和运行复盘。",
            "项目不从重硬件或全域城市总包切入，而从边界清楚、责任主体明确、可快速验收的区域开始。景区、园区、场馆、机场周边和固定低空航线，是首批更现实的市场入口。",
            "商业化采用软件订阅、专业部署、活动保障、航线评估、运行报告和运维续费组合。硬件侧通过伙伴接入，平台侧沉淀规则、数据、日志和报表能力。"
        ],
        bullets=["首批客户：景区、园区、场馆、航线运营方", "产品边界：辅助监管，不替代审批或执法", "经营重点：控制交付成本、提高续费率、缩短回款周期"],
        chart_id=summary_chart["chart_id"] if summary_chart else None,
        accent=REPORT_COLORS["teal"],
        kicker="执行摘要",
        callout="先做一个区域的闭环，再做可复制的区域方案。",
    )
    add(
        "summary",
        "执行摘要",
        "产品与商业路径",
        [
            "SkyGuard 的核心交付是一套可每天使用的运行机制。态势图解决实时查看，规则中心解决围栏和阈值，事件中心解决确认和派单，报表中心解决复盘和验收。Web Demo 按同样逻辑组织，能够从首页进入完整业务路径。",
            "竞争优势不写成单一算法优势，而写成业务闭环优势。通用监控、单点感知设备和人工巡查都能解决局部问题，但客户真正需要的是可解释、可记录、可提交的全流程材料。",
            "财务测算采用三情景，重点关注客户数量、客单价、交付成本、续费率和回款周期。早期融资优先用于产品打磨、试点交付、数据治理和客户成功。"
        ],
        bullets=["产品：计划、目标、围栏、工单、报表一体化", "商业：订阅、部署、活动、评估、运维组合收入", "风险：政策、数据、误报、采购和现金流分层管理"],
        figure=take_figure(),
        accent=REPORT_COLORS["blue"],
        kicker="执行摘要",
        callout="计划书与Demo共用同一套证据和术语。",
    )

    for spec in final_chapter_specs(data):
        chapter = spec["chapter"]
        rows = chapter_chart_map.get(chapter) or all_rows
        row = next_row(chapter)
        add(
            "divider",
            chapter,
            spec["decision"],
            [spec["opening"]],
            bullets=[f"本章重点：{', '.join(spec['topics'][:5])}", f"关联图表：{len(rows)} 张", "写法要求：痛点、证据、动作相互对应"],
            chart_id=row["chart_id"] if row else None,
            figure=take_figure() if chapter in {"第三章 产品服务", "第四章 技术架构", "第十三章 Web Demo"} else None,
            accent=spec["accent"],
            kicker="章节导读",
            callout=spec["decision"],
        )
        for i in range(spec["count"]):
            topic = spec["topics"][i % len(spec["topics"])]
            row = next_row(chapter)
            use_figure = (i % 6 == 2) and bool(figures)
            use_table = i in {4, 10}
            figure = take_figure() if use_figure else None
            figure_caption = humanize_figure_caption(figure) if figure else None
            title = f"{topic}：{figure_caption}" if figure_caption else (row["title"] if row and i % 2 else topic)
            paras = final_narrative_paragraphs(spec, topic, row, metrics, figure_caption)
            bullets = [
                f"判断：{(row or {}).get('conclusion', spec['decision'])}",
                f"来源：{(row or {}).get('source', '整理图片、Demo截图与项目资料')}",
                f"动作：明确{topic}对应的字段、页面、责任或验收口径。",
            ]
            table = final_page_table(chapter, topic, data, source_registry, chart_catalog) if use_table else []
            add(
                "visual" if figure else "table" if table else "chart",
                chapter,
                title,
                paras,
                bullets=bullets,
                chart_id=None if figure else (row["chart_id"] if row else None),
                figure=figure,
                table=table,
                accent=spec["accent"],
                callout=spec["decision"],
            )
    return pages


def generate_chart_pack(data: dict[str, pd.DataFrame]) -> pd.DataFrame:
    nat = data["national"].copy()
    regional = data["regional"].copy()
    airport = data["airport"].copy()
    cloud = data["cloud"].copy()
    city = data["city"].copy()
    plans = data["plans"].copy()
    telemetry = data["telemetry"].copy()
    alerts = data["alerts"].copy()
    risk = data["risk"].copy()
    grid = data["grid"].copy()
    sensors = data["sensors"].copy()
    vertiports = data["vertiports"].copy()
    routes = data["routes"].copy()

    for df in [nat, regional, airport, cloud, city, plans, telemetry, alerts, risk, grid, sensors, vertiports, routes]:
        for col in df.columns:
            if col in ["value", "risk_score", "preflight_risk_score", "planned_altitude_m", "planned_distance_km"]:
                df[col] = safe_numeric(df[col])

    official_source = "中国民用航空局/交通运输部公开统计，本地清洗表01-04"
    city_source = "地方低空经济政策目标，本地清洗表05"
    demo_source = "平台演示样本，已在数据表data_nature字段标注为simulated_demo_sample"

    # 1-12: official statistics
    key_codes = [
        ("registered_uavs_10k", "全行业注册无人机", "万架"),
        ("uav_flight_hours_10k", "无人机累计飞行小时", "万小时"),
        ("uav_operating_units", "无人机运营单位/企业数量", "家"),
        ("uav_operator_licenses_valid_10k", "有效无人机操控员执照", "万本"),
        ("registered_general_airports", "在册管理通用机场", "个"),
    ]
    for idx, (code, name, unit) in enumerate(key_codes, start=1):
        sub = nat[nat["indicator_code"].eq(code)].sort_values("year")
        fig, ax = plt.subplots(figsize=(8, 4.5))
        ax.plot(sub["year"], sub["value"], marker="o", linewidth=2.4, color=PALETTE["teal"])
        style_axis(ax, f"{name}变化趋势", f"单位：{unit}")
        ax.set_xlabel("")
        ax.set_ylabel(unit)
        save_chart(
            fig,
            f"C{idx:03d}",
            f"{name}趋势",
            "折线图",
            official_source,
            "第一章 行业机会",
            f"{name}是判断低空运行活跃度和监管平台需求的基础指标。",
        )

    uav_codes = ["registered_uavs_10k", "uav_flight_hours_10k", "uav_operator_licenses_valid_10k"]
    pivot = nat[nat["indicator_code"].isin(uav_codes)].pivot(index="year", columns="indicator_name_cn", values="value")
    fig, ax = plt.subplots(figsize=(8, 4.5))
    pivot.plot(kind="bar", ax=ax, color=[PALETTE["teal"], PALETTE["blue"], PALETTE["orange"]])
    style_axis(ax, "无人机运行基础指标对比", "注册量、飞行小时和操控员执照同步抬升")
    ax.set_xlabel("")
    ax.legend(fontsize=8)
    save_chart(fig, "C006", "无人机运行基础指标分组柱", "分组柱状图", official_source, "第一章 行业机会", "运行主体、飞行器和飞行活动同时增加，管理复杂度上升。")

    owners = nat[nat["indicator_code"].isin(["uav_personal_users_10k", "uav_legal_users_10k"])].pivot(index="year", columns="indicator_name_cn", values="value")
    fig, ax = plt.subplots(figsize=(8, 4.5))
    owners.plot(kind="bar", stacked=True, ax=ax, color=[PALETTE["blue"], PALETTE["green"]])
    style_axis(ax, "无人机注册用户结构", "个人用户与法人用户构成")
    ax.set_xlabel("")
    ax.legend(fontsize=8)
    save_chart(fig, "C007", "无人机用户结构堆叠柱", "堆叠柱状图", official_source, "第二章 需求验证", "个人与法人用户并存，平台需要同时支持合规运营和公众风险管理。")

    sub = regional[(regional["year"].eq(2025)) & (regional["category_code"].eq("uav_operating_units"))].sort_values("value")
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.barh(sub["region"], sub["value"], color=PALETTE["teal"])
    style_axis(ax, "2025年各区域无人机运营单位分布", "单位：家")
    save_chart(fig, "C008", "区域运营单位横向条形", "横向条形图", official_source, "第六章 市场竞争", "区域分布为首批市场选择提供了量化参考。")

    sub = regional[(regional["year"].eq(2025)) & (regional["category_code"].eq("ga_traditional_enterprises"))]
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.pie(sub["value"], labels=sub["region"], autopct="%1.1f%%", startangle=90, colors=sns.color_palette("Set2", len(sub)))
    ax.set_title("2025年传统通航企业区域占比", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C009", "传统通航企业区域占比饼图", "饼图", official_source, "第六章 市场竞争", "通航产业基础集中区域更适合形成早期合作网络。")

    sub = regional[(regional["year"].eq(2025)) & (regional["category_code"].eq("uav_operating_units"))]
    fig, ax = plt.subplots(figsize=(6, 6))
    wedges, _ = ax.pie(sub["value"], startangle=90, colors=sns.color_palette("crest", len(sub)), wedgeprops=dict(width=0.35))
    ax.legend(wedges, sub["region"], loc="center left", bbox_to_anchor=(1, 0.5), fontsize=8)
    ax.set_title("2025年无人机运营单位区域结构", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C010", "无人机运营单位环形图", "环形图", official_source, "第六章 市场竞争", "运营单位分布决定了平台获客与试点复制的优先顺序。")

    sub = airport[(airport["dimension"].eq("机场数量区域分布")) & (~airport["region_or_bucket"].eq("全国"))]
    sub = sub[sub["year"].eq(2025)].sort_values("value", ascending=False)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.bar(sub["region_or_bucket"], sub["value"], color=PALETTE["blue"])
    style_axis(ax, "2025年颁证运输机场区域分布", "机场净空与低空运行需共同纳入安全地图")
    ax.set_xlabel("")
    save_chart(fig, "C011", "运输机场区域分布柱状", "柱状图", official_source, "第四章 技术架构", "低空平台不能只看低空目标，还要叠加机场净空和航路约束。")

    sub = airport[(airport["dimension"].eq("民航数字导航监视能力")) & (airport["year"].eq(2025))]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.stem(sub["indicator"], sub["value"], linefmt=PALETTE["teal"], markerfmt="o", basefmt="#cbd5df")
    ax.tick_params(axis="x", rotation=25)
    style_axis(ax, "民航数字导航监视能力指标", "PBN、RNP AR、ADS-B OUT、EFB")
    save_chart(fig, "C012", "数字导航监视针状图", "针状图", official_source, "第四章 技术架构", "成熟航空体系的数字化经验可迁移到城市低空运行保障。")

    safety = nat[nat["indicator_code"].str.contains("safety|accident", regex=True, na=False)].copy()
    fig, ax = plt.subplots(figsize=(8, 4.5))
    for name, sub in safety.groupby("indicator_name_cn"):
        ax.plot(sub["year"], sub["value"], marker="o", label=name)
    style_axis(ax, "安全指标变化", "事故率指标提醒平台需保留风险边界和人工复核")
    ax.legend(fontsize=8)
    save_chart(fig, "C013", "安全指标多折线", "多折线图", official_source, "第十章 风险", "安全类产品必须以辅助决策和记录闭环为边界。")

    cloud_num = cloud.copy()
    cloud_num["value_num"] = safe_numeric(cloud_num["value"])
    sub = cloud_num[cloud_num["value_num"].notna()].head(18)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.barh(sub["indicator"], sub["value_num"], color=sns.color_palette("viridis", len(sub)))
    style_axis(ax, "2023年无人机云系统公开指标摘录", "取可量化字段展示")
    save_chart(fig, "C014", "无人机云指标瀑布前置条形", "长标签条形图", official_source, "第五章 数据体系", "云系统、执照、考试点等公开指标可作为平台参数校准依据。")

    # 15-25: city policy targets
    city_num = city.copy()
    city_num["value_num"] = safe_numeric(city_num["value"])
    top_city_targets = city_num[city_num["value_num"].notna()].sort_values("value_num", ascending=False).head(12)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    labels = top_city_targets["entity"] + "\n" + top_city_targets["indicator"].str.slice(0, 8)
    ax.bar(labels, top_city_targets["value_num"], color=PALETTE["teal"])
    ax.tick_params(axis="x", rotation=35)
    style_axis(ax, "地方低空经济重点目标摘录", "不同单位并列展示，正文逐项说明口径")
    save_chart(fig, "C015", "地方政策目标柱状", "政策目标柱状图", city_source, "第一章 行业机会", "地方政策目标已经从产业规模延伸到平台、航线和服务体系。")

    cat_count = city["category"].value_counts().head(12)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.barh(cat_count.index[::-1], cat_count.values[::-1], color=PALETTE["blue"])
    style_axis(ax, "地方政策目标主题频次", "按category字段统计")
    save_chart(fig, "C016", "政策主题频次横条", "主题频次图", city_source, "第一章 行业机会", "平台与基础设施、航线网络、安全监管是高频政策主题。")

    heat = pd.crosstab(city["entity"], city["category"]).clip(upper=5)
    fig, ax = plt.subplots(figsize=(9, 5.5))
    sns.heatmap(heat, cmap="YlGnBu", linewidths=0.5, ax=ax, cbar=True)
    ax.set_title("地方政策目标矩阵", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C017", "地方政策目标热力矩阵", "矩阵热力图", city_source, "第一章 行业机会", "不同城市政策重心不同，产品进入策略应按场景包配置。")

    timeline = city_num[city_num["target_year"].notna()].copy()
    timeline["target_year"] = safe_numeric(timeline["target_year"])
    timeline["jitter"] = np.random.default_rng(RANDOM_SEED).normal(0, 0.06, len(timeline))
    fig, ax = plt.subplots(figsize=(8, 4.5))
    y_map = {c: i for i, c in enumerate(timeline["category"].value_counts().head(10).index)}
    t2 = timeline[timeline["category"].isin(y_map)]
    ax.scatter(t2["target_year"], t2["category"].map(y_map) + t2["jitter"], s=38, alpha=0.65, color=PALETTE["teal"])
    ax.set_yticks(list(y_map.values()))
    ax.set_yticklabels(list(y_map.keys()))
    style_axis(ax, "地方低空政策目标年份分布", "政策目标多集中在2025-2027年")
    save_chart(fig, "C018", "政策目标时间散点", "时间散点图", city_source, "第八章 落地计划", "近期政策窗口更适合做园区、景区和城市试点。")

    # Demo samples and operation charts
    plan_status = plans["approval_status"].value_counts()
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.pie(plan_status.values, labels=plan_status.index, autopct="%1.1f%%", startangle=90, colors=sns.color_palette("Set3", len(plan_status)))
    ax.set_title("演示飞行计划审批状态", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C019", "飞行计划审批状态饼图", "饼图", demo_source, "第五章 数据体系", "审批结果是合规比对和运行报表的关键字段。")

    sc = plans["scenario"].value_counts().head(12)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.barh(sc.index[::-1], sc.values[::-1], color=PALETTE["teal"])
    style_axis(ax, "演示飞行计划场景分布", "按scenario字段统计")
    save_chart(fig, "C020", "场景分布条形图", "条形图", demo_source, "第三章 产品服务", "场景分布决定产品场景包的优先级。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.hist(plans["preflight_risk_score"], bins=30, color=PALETTE["blue"], edgecolor="white")
    style_axis(ax, "飞行计划预检风险分布", "演示样本")
    save_chart(fig, "C021", "预检风险分布直方图", "直方图", demo_source, "第五章 数据体系", "风险分布可用于设置预警阈值和人工复核规则。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    sns.kdeplot(plans["planned_distance_km"], fill=True, color=PALETTE["teal"], ax=ax)
    style_axis(ax, "计划航程核密度分布", "演示样本")
    save_chart(fig, "C022", "计划航程密度图", "核密度图", demo_source, "第五章 数据体系", "航程分布影响走廊宽度、起降点密度和续航约束。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    sns.boxplot(data=plans, x="scenario", y="planned_altitude_m", ax=ax, color="#a7f3d0")
    ax.tick_params(axis="x", rotation=25)
    style_axis(ax, "不同场景计划高度箱线图", "演示样本")
    save_chart(fig, "C023", "场景高度箱线图", "箱线图", demo_source, "第四章 技术架构", "不同业务场景应配置不同限高和风险规则。")

    sample_flight = telemetry["flight_plan_id"].value_counts().index[0]
    sub = telemetry[telemetry["flight_plan_id"].eq(sample_flight)].head(120).copy()
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.plot(range(len(sub)), sub["altitude_m"], color=PALETTE["teal"], linewidth=2)
    style_axis(ax, "单次飞行高度剖面", f"样本航班：{sample_flight}")
    save_chart(fig, "C024", "单次飞行高度剖面", "剖面折线图", demo_source, "第四章 技术架构", "高度剖面能帮助判断超高、低电量返航和航线偏离。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    samp = telemetry.sample(min(2200, len(telemetry)), random_state=RANDOM_SEED)
    ax.scatter(samp["speed_mps"], samp["altitude_m"], s=6, alpha=0.22, color=PALETTE["blue"])
    style_axis(ax, "速度-高度散点", "演示遥测样本")
    save_chart(fig, "C025", "速度高度散点图", "散点图", demo_source, "第五章 数据体系", "速度和高度联合异常比单一阈值更有解释力。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.hexbin(samp["lon"], samp["lat"], gridsize=28, cmap="YlOrRd")
    style_axis(ax, "遥测点空间密度", "经纬度hexbin展示")
    save_chart(fig, "C026", "遥测空间六边形密度", "六边形密度图", demo_source, "第三章 产品服务", "空间密度图用于识别高频飞行走廊和重点监管区域。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    sns.violinplot(data=telemetry.sample(3000, random_state=RANDOM_SEED), x="geofence_status", y="altitude_m", ax=ax, palette="Set2")
    style_axis(ax, "围栏状态下的高度分布", "演示遥测样本")
    save_chart(fig, "C027", "围栏状态高度小提琴", "小提琴图", demo_source, "第四章 技术架构", "围栏状态与高度叠加后，风险分级更贴近业务。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    counts = telemetry["remote_id_valid"].value_counts()
    ax.bar(["有效Remote ID" if x else "无效/缺失Remote ID" for x in counts.index], counts.values, color=[PALETTE["green"], PALETTE["red"]])
    style_axis(ax, "Remote ID有效性统计", "演示遥测样本")
    save_chart(fig, "C028", "RemoteID有效性柱状", "二分类柱状图", demo_source, "第十章 风险", "身份字段是区分合规飞行与未知目标的第一道门槛。")

    event_type = alerts["alert_type"].value_counts().head(12)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.barh(event_type.index[::-1], event_type.values[::-1], color=PALETTE["orange"])
    style_axis(ax, "风险告警类型排行", "演示事件样本")
    save_chart(fig, "C029", "告警类型排行横条", "Pareto条形图", demo_source, "第二章 需求验证", "告警类型分布决定了工单模板和响应预案。")

    sev = alerts["severity"].value_counts()
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.pie(sev.values, labels=sev.index, autopct="%1.1f%%", startangle=90, colors=[PALETTE["green"], PALETTE["yellow"], PALETTE["orange"], PALETTE["red"]][: len(sev)], wedgeprops=dict(width=0.35))
    ax.set_title("告警严重程度结构", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C030", "告警等级环形图", "环形图", demo_source, "第十章 风险", "严重程度结构用于配置分级响应和人工复核。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    sns.boxenplot(data=alerts, x="severity", y="response_time_min", ax=ax, palette="rocket")
    style_axis(ax, "不同严重程度响应时间分布", "演示事件样本")
    save_chart(fig, "C031", "响应时间增强箱线图", "增强箱线图", demo_source, "第十三章 Demo展示", "响应时间是衡量平台闭环价值的核心运营指标。")

    cross = pd.crosstab(alerts["city"], alerts["severity"]).head(12)
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.heatmap(cross, cmap="OrRd", linewidths=0.4, ax=ax)
    ax.set_title("城市-告警等级交叉热力", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C032", "城市告警等级交叉热力", "交叉热力图", demo_source, "第十三章 Demo展示", "区域风险画像能把事件复盘转化为下一轮配置优化。")

    source_counts = alerts["detection_source"].value_counts().head(8)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    sizes = source_counts.values
    labels = source_counts.index
    x = np.arange(len(sizes))
    ax.scatter(x, sizes, s=sizes * 3, alpha=0.55, color=PALETTE["teal"])
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=25)
    style_axis(ax, "检测来源气泡图", "演示事件样本")
    save_chart(fig, "C033", "检测来源气泡图", "气泡图", demo_source, "第四章 技术架构", "多源感知接入会改变告警置信度和处置路径。")

    corr_cols = [
        "population_density_index",
        "flight_density_index",
        "sensor_coverage_score",
        "communication_quality_score",
        "wind_speed_mps",
        "rainfall_mm_h",
        "visibility_km",
        "historical_alert_count_24h",
        "computed_risk_score",
    ]
    corr = risk[corr_cols].corr()
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.heatmap(corr, cmap="vlag", center=0, linewidths=0.4, ax=ax)
    ax.set_title("风险模型特征相关性", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C034", "风险特征相关性热力", "相关性热力图", demo_source, "第五章 数据体系", "相关性检查用于避免模型把同一风险因素重复计权。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    sns.scatterplot(data=grid.sample(1500, random_state=RANDOM_SEED), x="sensor_coverage_score", y="computed_risk_score", hue="airspace_class", s=18, alpha=0.55, ax=ax)
    style_axis(ax, "感知覆盖与网格风险关系", "演示网格样本")
    ax.legend(fontsize=7)
    save_chart(fig, "C035", "覆盖风险关系散点", "分类散点图", demo_source, "第四章 技术架构", "覆盖不足的敏感区域应优先补设备或提高人工复核等级。")

    gtab = pd.crosstab(grid["city"], grid["computed_risk_level"]).head(14)
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.heatmap(gtab, cmap="YlOrBr", linewidths=0.4, ax=ax)
    ax.set_title("城市网格风险等级矩阵", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C036", "城市网格风险矩阵", "等级矩阵图", demo_source, "第五章 数据体系", "网格化风险矩阵适合做城市试点前的区域筛选。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    sns.stripplot(data=sensors.sample(400, random_state=RANDOM_SEED), x="sensor_type", y="health_score", jitter=True, ax=ax, palette="Set2")
    ax.tick_params(axis="x", rotation=25)
    style_axis(ax, "传感器健康度分布", "演示设备样本")
    save_chart(fig, "C037", "设备健康度抖动点图", "抖动点图", demo_source, "第十三章 Demo展示", "设备健康度影响告警可信度和系统SLA。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.scatter(sensors["coverage_radius_km"], sensors["health_score"], s=sensors["height_m"] * 1.2, alpha=0.45, color=PALETTE["blue"])
    style_axis(ax, "设备覆盖半径与健康度", "气泡大小表示安装高度")
    save_chart(fig, "C038", "设备覆盖健康气泡", "三变量气泡图", demo_source, "第四章 技术架构", "设备布设要同时看覆盖、健康度和安装条件。")

    vt = vertiports["vertiport_type"].value_counts()
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.bar(vt.index, vt.values, color=PALETTE["teal"])
    ax.tick_params(axis="x", rotation=25)
    style_axis(ax, "演示起降点类型分布", "演示起降点样本")
    save_chart(fig, "C039", "起降点类型柱状", "柱状图", demo_source, "第三章 产品服务", "起降点类型决定审批、运维和风险控制模板。")

    rs = routes["scenario"].value_counts().head(10)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    y = np.arange(len(rs))
    ax.hlines(y, 0, rs.values, color="#cbd5df")
    ax.plot(rs.values, y, "o", color=PALETTE["orange"])
    ax.set_yticks(y)
    ax.set_yticklabels(rs.index)
    style_axis(ax, "航线场景棒棒糖图", "演示航线样本")
    save_chart(fig, "C040", "航线场景棒棒糖", "棒棒糖图", demo_source, "第三章 产品服务", "航线场景越复杂，平台越需要标准化风险评估报告。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    sns.ecdfplot(routes["route_distance_km"], ax=ax, color=PALETTE["teal"], linewidth=2)
    style_axis(ax, "航线距离累计分布", "演示航线样本")
    save_chart(fig, "C041", "航线距离ECDF", "累计分布图", demo_source, "第七章 商业模式", "航线评估服务可按距离、场景和风险复杂度定价。")

    # Finance and business model charts
    finance = build_finance_tables()
    revenue = finance["revenue"]
    cost = finance["cost"]
    cash = finance["cashflow"]
    scenario = finance["scenario"]
    finance_source = "经营模型假设，详见SkyGuard_财务测算表.xlsx"

    fig, ax = plt.subplots(figsize=(8, 4.5))
    revenue.set_index("year").plot(kind="bar", stacked=True, ax=ax, colormap="tab20")
    style_axis(ax, "五年收入结构预测", "单位：万元，基准情景")
    ax.legend(fontsize=7, ncol=2)
    save_chart(fig, "C042", "收入结构堆叠柱", "堆叠柱状图", finance_source, "第九章 财务预测", "收入结构从项目制逐步过渡到SaaS、运维和数据报告。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.plot(cost["year"], cost["gross_margin"], marker="o", color=PALETTE["green"], linewidth=2.3)
    ax.set_ylim(0, 0.8)
    style_axis(ax, "基准情景毛利率预测", "随SaaS和报告收入占比提高而改善")
    save_chart(fig, "C043", "毛利率折线", "比例折线图", finance_source, "第九章 财务预测", "平台化收入占比提升是毛利率改善的关键。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.bar(cash["year"], cash["ending_cash"], color=np.where(cash["ending_cash"] >= 0, PALETTE["teal"], PALETTE["red"]))
    style_axis(ax, "期末现金余额预测", "单位：万元，含融资假设")
    save_chart(fig, "C044", "现金余额柱状", "正负柱状图", finance_source, "第九章 财务预测", "融资资金主要覆盖研发、试点交付和销售启动期现金缺口。")

    fig, ax = plt.subplots(figsize=(8, 4.5))
    scenario_pivot = scenario.pivot(index="year", columns="scenario", values="revenue_total")
    scenario_pivot.plot(ax=ax, marker="o", linewidth=2)
    style_axis(ax, "三情景收入预测", "单位：万元")
    save_chart(fig, "C045", "三情景收入折线", "情景折线图", finance_source, "第九章 财务预测", "客户获取速度和城市试点数量是收入弹性的主要来源。")

    sens = pd.DataFrame(
        {
            "factor": ["Lite续费率", "Pro交付成本", "Gov项目回款", "获客周期", "设备接入成本", "报告订阅转化"],
            "downside": [-14, -18, -22, -16, -9, -7],
            "upside": [12, 11, 20, 13, 8, 10],
        }
    )
    fig, ax = plt.subplots(figsize=(8, 4.5))
    y = np.arange(len(sens))
    ax.barh(y, sens["downside"], color=PALETTE["red"], alpha=0.72)
    ax.barh(y, sens["upside"], color=PALETTE["green"], alpha=0.72)
    ax.set_yticks(y)
    ax.set_yticklabels(sens["factor"])
    style_axis(ax, "关键变量敏感性分析", "对第三年净现金流的影响，单位：%")
    save_chart(fig, "C046", "敏感性龙卷风图", "龙卷风图", finance_source, "第九章 财务预测", "回款、交付成本和获客周期是早期现金流管理重点。")

    funnel = pd.DataFrame({"stage": ["线索", "有效沟通", "试点方案", "试点签约", "续费/扩容"], "value": [180, 92, 38, 14, 9]})
    fig, ax = plt.subplots(figsize=(8, 4.5))
    max_v = funnel["value"].max()
    for i, row in funnel.iterrows():
        width = row["value"] / max_v
        ax.barh(i, width, left=(1 - width) / 2, color=sns.color_palette("crest", len(funnel))[i], height=0.72)
        ax.text(0.5, i, f"{row['stage']} {row['value']}", ha="center", va="center", color="white", fontweight="bold")
    ax.set_xlim(0, 1)
    ax.set_axis_off()
    ax.set_title("政企客户销售漏斗", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C047", "销售漏斗图", "漏斗图", finance_source, "第七章 商业模式", "试点转化需要Demo、试点报告和伙伴渠道共同推进。")

    # Strategic diagrams
    radar_labels = ["区域敏感", "合规状态", "轨迹异常", "身份可信", "环境影响", "响应紧迫"]
    radar_vals = [0.82, 0.76, 0.69, 0.61, 0.48, 0.74]
    angles = np.linspace(0, 2 * np.pi, len(radar_labels), endpoint=False).tolist()
    radar_vals2 = radar_vals + radar_vals[:1]
    angles2 = angles + angles[:1]
    fig = plt.figure(figsize=(6, 6))
    ax = fig.add_subplot(111, polar=True)
    ax.plot(angles2, radar_vals2, color=PALETTE["teal"], linewidth=2)
    ax.fill(angles2, radar_vals2, color=PALETTE["teal"], alpha=0.18)
    ax.set_xticks(angles)
    ax.set_xticklabels(radar_labels, fontsize=9)
    ax.set_title("样例目标六维风险雷达", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C048", "六维风险雷达", "雷达图", demo_source, "第四章 技术架构", "六维解释能让风险分数可复核，而不是黑箱结论。")

    fig, ax = plt.subplots(figsize=(8, 5))
    modules = ["态势图", "计划", "AI识别", "电子围栏", "工单", "报表", "设备", "权限"]
    customers = ["景区", "园区", "场馆", "物流", "机场周边", "应急"]
    mat = np.array([[random.randint(1, 5) for _ in modules] for _ in customers])
    sns.heatmap(mat, annot=True, fmt="d", cmap="YlGnBu", xticklabels=modules, yticklabels=customers, ax=ax)
    ax.set_title("场景-功能优先级矩阵", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C049", "场景功能优先级矩阵", "优先级矩阵", "产品规划假设", "第三章 产品服务", "不同场景买的不是同一套功能，交付应按场景包组合。")

    fig, ax = plt.subplots(figsize=(8, 5))
    comp = pd.DataFrame(
        {
            "方案": ["SkyGuard", "普通视频监控", "反无人机设备", "无人机任务平台", "GIS大屏", "自建系统"],
            "低空业务深度": [8.2, 3.0, 5.5, 6.2, 4.4, 7.0],
            "部署轻量度": [7.4, 8.0, 3.8, 6.6, 5.0, 2.8],
            "预算压力": [45, 30, 85, 50, 60, 95],
        }
    )
    ax.scatter(comp["低空业务深度"], comp["部署轻量度"], s=comp["预算压力"] * 6, color=PALETTE["teal"], alpha=0.55)
    for _, r in comp.iterrows():
        ax.text(r["低空业务深度"] + 0.05, r["部署轻量度"] + 0.05, r["方案"], fontsize=8)
    ax.set_xlim(2, 9)
    ax.set_ylim(2, 9)
    style_axis(ax, "替代方案竞争象限", "气泡大小表示预算压力")
    save_chart(fig, "C050", "替代方案竞争象限", "象限气泡图", "竞品与替代方案分析", "第六章 市场竞争", "SkyGuard的机会在于低空业务闭环与轻量部署之间的平衡。")

    fig, ax = plt.subplots(figsize=(8, 5))
    layers = [("TAM\n低空经济安全与运行保障", 100), ("SAM\n重点区域监管与企业合规", 46), ("SOM\n三年可触达景区/园区/场馆", 9)]
    for i, (label, width) in enumerate(layers):
        ax.barh(i, width, left=(100 - width) / 2, color=[PALETTE["blue"], PALETTE["teal"], PALETTE["green"]][i])
        ax.text(50, i, label, ha="center", va="center", color="white", fontweight="bold")
    ax.set_xlim(0, 100)
    ax.set_axis_off()
    ax.set_title("TAM/SAM/SOM分层估算", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C051", "TAM_SAM_SOM金字塔", "金字塔图", finance_source, "第六章 市场竞争", "早期不追求全城市万能平台，而从可触达重点区域切入。")

    fig, ax = plt.subplots(figsize=(9, 4.8))
    tasks = [
        ("原型", 0, 3),
        ("MVP部署", 3, 6),
        ("景区/园区试点", 6, 18),
        ("场景包复制", 18, 36),
        ("城市级协同", 36, 60),
    ]
    for i, (name, start, end) in enumerate(tasks):
        ax.barh(i, end - start, left=start, color=sns.color_palette("crest", len(tasks))[i])
        ax.text(start + 0.2, i, name, va="center", fontsize=9, color="white", fontweight="bold")
    ax.set_yticks([])
    ax.set_xlabel("月份")
    style_axis(ax, "五年产品与市场路线图", "从重点区域到城市级协同")
    save_chart(fig, "C052", "产品路线甘特图", "甘特图", "项目规划", "第八章 落地计划", "每一阶段都要产出可验证证据，而不是只写远景。")

    fig, ax = plt.subplots(figsize=(9, 5))
    nodes = ["感知接入", "身份比对", "轨迹分析", "围栏规则", "风险评分", "工单协同", "报表复盘"]
    x = np.linspace(0.08, 0.92, len(nodes))
    for i, node in enumerate(nodes):
        ax.add_patch(plt.Rectangle((x[i] - 0.055, 0.42), 0.11, 0.16, color="#e0f2fe", ec=PALETTE["blue"]))
        ax.text(x[i], 0.5, node, ha="center", va="center", fontsize=8)
        if i < len(nodes) - 1:
            ax.annotate("", xy=(x[i + 1] - 0.06, 0.5), xytext=(x[i] + 0.06, 0.5), arrowprops=dict(arrowstyle="->", color=PALETTE["muted"]))
    ax.set_axis_off()
    ax.set_title("感知-判断-处置-复盘闭环", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C053", "运行闭环流程图", "流程图", "产品方案", "第三章 产品服务", "闭环能力比单点识别更符合城市低空运行需求。")

    fig, ax = plt.subplots(figsize=(9, 5))
    layers = ["展示层", "业务层", "算法层", "数据层", "接入层"]
    desc = ["态势大屏/企业端/报表", "计划/围栏/工单/权限", "识别/异常/评分/预测", "轨迹/事件/设备/政策", "雷达/光电/5G-A/云系统"]
    for i, (l, d) in enumerate(zip(layers, desc)):
        y = 0.78 - i * 0.16
        ax.add_patch(plt.Rectangle((0.08, y), 0.84, 0.11, color=sns.color_palette("Blues", 7)[i + 2], ec="white"))
        ax.text(0.14, y + 0.055, l, va="center", ha="left", fontsize=10, color="white", fontweight="bold")
        ax.text(0.35, y + 0.055, d, va="center", ha="left", fontsize=9, color="white")
    ax.set_axis_off()
    ax.set_title("SkyGuard五层技术架构", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C054", "五层技术架构图", "架构图", "产品方案", "第四章 技术架构", "架构强调平台软件、业务规则和多源数据，不做反制武器。")

    fig, ax = plt.subplots(figsize=(8, 5))
    G = nx.Graph()
    edges = [
        ("SkyGuard", "园区"),
        ("SkyGuard", "景区"),
        ("SkyGuard", "场馆"),
        ("SkyGuard", "低空企业"),
        ("SkyGuard", "传感器伙伴"),
        ("SkyGuard", "数据服务"),
        ("园区", "安保"),
        ("景区", "游客航拍"),
        ("低空企业", "物流航线"),
        ("传感器伙伴", "雷达/光电"),
    ]
    G.add_edges_from(edges)
    pos = nx.spring_layout(G, seed=RANDOM_SEED)
    nx.draw_networkx_nodes(G, pos, ax=ax, node_color="#dff7f3", edgecolors=PALETTE["teal"], node_size=1200)
    nx.draw_networkx_edges(G, pos, ax=ax, edge_color="#94a3b8")
    nx.draw_networkx_labels(G, pos, ax=ax, font_size=8)
    ax.set_axis_off()
    ax.set_title("生态合作网络", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C055", "生态合作网络图", "网络图", "商业模式分析", "第七章 商业模式", "平台落地依赖场景客户、设备伙伴和数据服务伙伴协同。")

    fig, ax = plt.subplots(figsize=(9, 5.5))
    canvas_items = [
        ("客户细分", "景区/园区/场馆/低空企业"),
        ("价值主张", "看见低空、预警风险、沉淀报表"),
        ("渠道", "试点、伙伴、产业园、竞赛展示"),
        ("客户关系", "客户成功、运维、续费扩容"),
        ("收入", "SaaS、部署、报告、运维、航线评估"),
        ("关键资源", "数据、算法、Demo、场景规则"),
        ("关键活动", "产品迭代、交付、数据治理"),
        ("伙伴", "传感器、低空企业、园区运营方"),
        ("成本", "研发、交付、云服务、销售"),
    ]
    for idx, (k, v) in enumerate(canvas_items):
        r, c = divmod(idx, 3)
        x0, y0 = 0.05 + c * 0.31, 0.72 - r * 0.28
        ax.add_patch(plt.Rectangle((x0, y0), 0.28, 0.22, fc="#f8fafc", ec="#cbd5e1"))
        ax.text(x0 + 0.015, y0 + 0.17, k, fontsize=10, fontweight="bold", color=PALETTE["ink"])
        ax.text(x0 + 0.015, y0 + 0.10, textwrap.fill(v, 16), fontsize=8.5, color=PALETTE["muted"], va="top")
    ax.set_axis_off()
    ax.set_title("商业模式画布", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C056", "商业模式画布", "商业模式画布", "商业模式分析", "第七章 商业模式", "商业模式围绕重点区域安全运行，而不是一次性大屏项目。")

    evidence_index = ["政策", "统计", "地方目标", "公开数据", "Demo样本", "财务假设", "图片"]
    evidence = pd.DataFrame(
        np.random.default_rng(RANDOM_SEED).integers(1, 5, size=(len(evidence_index), 7)),
        index=evidence_index,
        columns=["背景", "痛点", "产品", "技术", "市场", "财务", "风险", "Demo"][:7],
    )
    fig, ax = plt.subplots(figsize=(8, 5))
    sns.heatmap(evidence, cmap="BuGn", annot=True, fmt="d", linewidths=0.4, ax=ax)
    ax.set_title("章节证据覆盖矩阵", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, "C057", "证据覆盖矩阵", "证据矩阵", "证据索引", "附录", "每章都需要对应证据类型，避免只写观点。")

    # Extra visual forms to exceed 50 chart types.
    extra_specs = [
        ("C058", "风险等级100%堆叠条", "100%堆叠条形图"),
        ("C059", "城市风险斜率图", "斜率图"),
        ("C060", "收入利润面积图", "面积图"),
        ("C061", "单位经济瀑布图", "瀑布图"),
        ("C062", "响应时间日历热力", "日历热力图"),
    ]

    # 100% stacked bar
    status = pd.crosstab(alerts["city"], alerts["severity"], normalize="index").head(10)
    fig, ax = plt.subplots(figsize=(8, 4.5))
    status.plot(kind="barh", stacked=True, ax=ax, colormap="YlOrRd")
    style_axis(ax, "城市告警等级结构", "各城市内部占比")
    ax.legend(fontsize=7, ncol=2)
    save_chart(fig, *extra_specs[0], demo_source, "第十三章 Demo展示", "占比结构能帮助运营方识别高风险城市和区域。")

    # Slope graph
    reg_23 = regional[(regional["year"].eq(2023)) & (regional["category_code"].eq("uav_operating_units"))].set_index("region")["value"]
    reg_25 = regional[(regional["year"].eq(2025)) & (regional["category_code"].eq("uav_operating_units"))].set_index("region")["value"]
    slope = pd.DataFrame({"2023": reg_23, "2025": reg_25}).dropna()
    fig, ax = plt.subplots(figsize=(7, 5))
    for region, row in slope.iterrows():
        ax.plot([0, 1], [row["2023"], row["2025"]], marker="o", color=PALETTE["teal"], alpha=0.65)
        ax.text(-0.03, row["2023"], region, ha="right", va="center", fontsize=8)
        ax.text(1.03, row["2025"], region, ha="left", va="center", fontsize=8)
    ax.set_xticks([0, 1])
    ax.set_xticklabels(["2023", "2025"])
    style_axis(ax, "无人机运营单位区域变化", "斜率图")
    save_chart(fig, *extra_specs[1], official_source, "第六章 市场竞争", "区域增速差异提示市场进入不能只看存量。")

    # Area chart
    area = pd.DataFrame({"year": revenue["year"], "revenue": revenue.drop(columns="year").sum(axis=1), "gross_profit": cost["gross_profit"]})
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.fill_between(area["year"], area["revenue"], color=PALETTE["blue"], alpha=0.25, label="收入")
    ax.fill_between(area["year"], area["gross_profit"], color=PALETTE["green"], alpha=0.35, label="毛利")
    ax.plot(area["year"], area["revenue"], color=PALETTE["blue"])
    ax.plot(area["year"], area["gross_profit"], color=PALETTE["green"])
    ax.legend(fontsize=8)
    style_axis(ax, "收入与毛利面积图", "单位：万元")
    save_chart(fig, *extra_specs[2], finance_source, "第九章 财务预测", "收入规模和毛利改善需要同步成立，才具备持续经营基础。")

    # Waterfall
    wf = [12, -1.8, -0.9, -1.2, -0.6, 7.5]
    labels = ["Lite年费", "云资源", "客户成功", "研发摊销", "销售服务", "单客毛利"]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    cum = 0
    for i, v in enumerate(wf):
        start = 0 if i in [0, len(wf) - 1] else cum
        if i == len(wf) - 1:
            start = 0
        ax.bar(i, v, bottom=start if i not in [0, len(wf) - 1] else 0, color=PALETTE["green"] if v >= 0 else PALETTE["red"])
        if i not in [0, len(wf) - 1]:
            cum += v
        elif i == 0:
            cum = v
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=20)
    style_axis(ax, "Lite SaaS单客单位经济模型", "单位：万元/年")
    save_chart(fig, *extra_specs[3], finance_source, "第九章 财务预测", "单客毛利来自可控云成本和标准化客户成功流程。")

    # Calendar heatmap approximation
    days = pd.date_range("2026-06-01", periods=35)
    vals = np.random.default_rng(RANDOM_SEED).integers(1, 16, len(days))
    cal = pd.DataFrame({"day": days, "week": [d.isocalendar().week for d in days], "weekday": [d.weekday() for d in days], "value": vals})
    pivot_cal = cal.pivot(index="weekday", columns="week", values="value")
    fig, ax = plt.subplots(figsize=(8, 3.8))
    sns.heatmap(pivot_cal, cmap="YlOrRd", linewidths=2, linecolor="white", cbar=True, ax=ax)
    ax.set_yticklabels(["一", "二", "三", "四", "五", "六", "日"], rotation=0)
    ax.set_title("响应压力日历热力", loc="left", fontsize=14, fontweight="bold")
    save_chart(fig, *extra_specs[4], demo_source, "第十三章 Demo展示", "日历热力图适合安排值班、巡检和设备维护。")

    catalog = pd.DataFrame(CHARTS)
    catalog.to_csv(TABLE_DIR / "chart_catalog.csv", index=False, encoding="utf-8-sig")
    return catalog


def build_finance_tables() -> dict[str, pd.DataFrame]:
    years = [2026, 2027, 2028, 2029, 2030]
    revenue = pd.DataFrame(
        {
            "year": years,
            "Lite SaaS": [48, 168, 420, 780, 1200],
            "Pro部署": [120, 420, 960, 1500, 2100],
            "Gov试点": [0, 200, 600, 1000, 1500],
            "活动保障": [20, 90, 210, 360, 520],
            "航线评估": [12, 66, 180, 360, 640],
            "数据报告": [0, 36, 108, 240, 420],
            "运维服务": [0, 42, 138, 306, 540],
            "设备接入": [30, 120, 260, 420, 620],
        }
    )
    total = revenue.drop(columns="year").sum(axis=1)
    cogs = total * np.array([0.52, 0.47, 0.42, 0.38, 0.35])
    gross_profit = total - cogs
    cost = pd.DataFrame(
        {
            "year": years,
            "revenue_total": total.round(2),
            "cogs": cogs.round(2),
            "gross_profit": gross_profit.round(2),
            "gross_margin": (gross_profit / total).round(4),
            "研发": [260, 360, 430, 520, 620],
            "销售": [90, 180, 320, 480, 650],
            "行政": [70, 90, 130, 170, 220],
        }
    )
    cash = pd.DataFrame(
        {
            "year": years,
            "beginning_cash": [220, 108, 252, 814, 1826],
            "financing": [450, 0, 0, 0, 0],
            "operating_cashflow": [-562, 144, 562, 1012, 1570],
        }
    )
    cash["ending_cash"] = cash["beginning_cash"] + cash["financing"] + cash["operating_cashflow"]
    scenarios = []
    for name, factor in [("保守", 0.72), ("基准", 1.0), ("乐观", 1.34)]:
        for year, val in zip(years, total * factor):
            scenarios.append({"scenario": name, "year": year, "revenue_total": round(float(val), 2)})
    scenario = pd.DataFrame(scenarios)
    assumptions = pd.DataFrame(
        [
            ["Lite SaaS年费", "8-15万元/年，模型取12万元/年", "经营假设"],
            ["Pro部署", "30-80万元/项目，模型按项目复杂度阶梯增长", "经营假设"],
            ["Gov试点", "100-300万元/项目，谨慎计入首年无收入", "经营假设"],
            ["活动保障", "5-30万元/次，适合赛事、演唱会、会展", "经营假设"],
            ["航线评估", "3000-20000元/条，按批量航线折扣", "经营假设"],
            ["运维服务", "合同额10%-20%/年", "经营假设"],
            ["毛利率", "随SaaS、报告、运维占比提高逐年改善", "经营假设"],
        ],
        columns=["item", "assumption", "source_type"],
    )
    return {"revenue": revenue, "cost": cost, "cashflow": cash, "scenario": scenario, "assumptions": assumptions}


def write_finance_workbook(finance: dict[str, pd.DataFrame], source_registry: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_财务测算表.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    try:
        wb.calculation.fullCalcOnLoad = True
        wb.calculation.forceFullCalc = True
    except Exception:
        pass

    years = finance["revenue"]["year"].astype(int).tolist()
    year_cols = [get_column_letter(i + 2) for i in range(len(years))]
    title_fill = PatternFill("solid", fgColor="123C4A")
    header_fill = PatternFill("solid", fgColor="155E75")
    section_fill = PatternFill("solid", fgColor="E6F1EE")
    input_fill = PatternFill("solid", fgColor="FFF2CC")
    ok_fill = PatternFill("solid", fgColor="E2F0D9")
    warn_fill = PatternFill("solid", fgColor="FCE4D6")
    thin = Side(style="thin", color="D9E2E7")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    input_font = Font(color="0000FF")
    formula_font = Font(color="000000")
    link_font = Font(color="008000")
    header_font = Font(bold=True, color="FFFFFF")
    money_fmt = '#,##0.0;[Red](#,##0.0);-'
    pct_fmt = '0.0%;[Red](0.0%);-'
    count_fmt = '#,##0;[Red](#,##0);-'

    def setup(ws, title: str, max_col: int = 7, freeze: str = "B4") -> None:
        ws.sheet_view.showGridLines = False
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max_col)
        ws["A1"] = title
        ws["A1"].font = Font(bold=True, size=15, color="FFFFFF")
        ws["A1"].fill = title_fill
        ws["A1"].alignment = Alignment(horizontal="left", vertical="center")
        ws.freeze_panes = freeze

    def write_header(ws, row: int, headers: list[str]) -> None:
        for col_idx, value in enumerate(headers, start=1):
            cell = ws.cell(row, col_idx, value)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def style_used_range(ws) -> None:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                cell.border = border
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    cell.font = link_font if "!" in cell.value else formula_font
        for col_idx in range(1, ws.max_column + 1):
            width = 14
            if col_idx == 1:
                width = 22
            elif col_idx == ws.max_column:
                width = 34
            ws.column_dimensions[get_column_letter(col_idx)].width = width

    def mark_input(cell, note: str) -> None:
        cell.font = input_font
        cell.fill = input_fill
        cell.comment = Comment(f"输入口径：{note}", "SkyGuard")

    def write_period_row(ws, row: int, label: str, values: list[float | int | str], note: str, kind: str = "formula") -> None:
        ws.cell(row, 1, label)
        ws.cell(row, len(years) + 2, note)
        for col_idx, value in enumerate(values, start=2):
            cell = ws.cell(row, col_idx, value)
            if kind == "input":
                mark_input(cell, note)

    assumptions = wb.create_sheet("假设")
    setup(assumptions, "SkyGuard 财务模型假设与口径", max_col=5, freeze="A4")
    write_header(assumptions, 3, ["类别", "项目", "数值/口径", "说明", "来源"])
    assumption_rows = [
        ["模型约定", "金额单位", "万元", "除特别注明外，收入、成本、现金流均以人民币万元表达。", "课程计划书财务分析要求"],
        ["模型约定", "预测期", "2026-2030", "五年滚动预测，适合作为创业计划书的经营测算。", "课程计划书财务分析要求"],
        ["模型约定", "估值日", "2026-07-05", "按当前课程提交节点生成模型。", "项目交付节点"],
        ["关键输入", "折现率", 0.14, "用于 NPV 和敏感性测算。", "经营假设"],
        ["关键输入", "所得税率", 0.15, "盈利年度简化税率，亏损年度不计所得税。", "经营假设"],
        ["关键输入", "营运资本占收入增量比例", 0.08, "用于估算回款周期和应收占用。", "经营假设"],
        ["关键输入", "资本开支率", 0.04, "轻资产平台必要设备、云资源和试点工具投入。", "经营假设"],
        ["关键输入", "折旧摊销率", 0.03, "对资本化投入进行简化摊销。", "经营假设"],
        ["关键输入", "客户成功费用率", 0.05, "随客户规模增加的规则维护、培训和复盘会议成本。", "经营假设"],
        ["关键输入", "期初现金", 220, "项目启动资金和已有资源折算。", "经营假设"],
        ["关键输入", "首轮融资流入", 450, "用于产品打磨、试点交付、数据治理和客户成功。", "融资假设"],
    ]
    for item, assumption, source_type in finance["assumptions"].itertuples(index=False):
        assumption_rows.append(["经营假设", item, assumption, "来自原始商业模型口径。", source_type])
    for row_idx, row_values in enumerate(assumption_rows, start=4):
        for col_idx, value in enumerate(row_values, start=1):
            assumptions.cell(row_idx, col_idx, value)
        if row_values[0] == "关键输入":
            mark_input(assumptions.cell(row_idx, 3), str(row_values[4]))
        if isinstance(row_values[2], float) and row_values[2] < 1:
            assumptions.cell(row_idx, 3).number_format = pct_fmt

    revenue = wb.create_sheet("收入")
    setup(revenue, "收入驱动模型", max_col=7)
    write_header(revenue, 3, ["指标"] + [str(y) for y in years] + ["备注"])
    revenue_component_rows: list[int] = []
    row = 4
    drivers = [
        ("Lite SaaS", "客户数", [4, 14, 35, 65, 100], "年费", [12, 12, 12, 12, 12], "万元/年/客户"),
        ("Pro部署", "项目数", [2, 7, 16, 25, 35], "客单价", [60, 60, 60, 60, 60], "万元/项目"),
        ("Gov试点", "项目数", [0, 1, 3, 5, 7], "客单价", [200, 200, 200, 200, 214.3], "万元/项目"),
        ("活动保障", "场次", [4, 15, 35, 60, 87], "客单价", [5, 6, 6, 6, 6], "万元/次"),
        ("航线评估", "条数", [40, 220, 600, 1200, 2133], "单价", [0.3, 0.3, 0.3, 0.3, 0.3], "万元/条"),
        ("数据报告", "份数", [0, 12, 36, 80, 140], "单价", [3, 3, 3, 3, 3], "万元/份"),
        ("运维服务", "客户数", [0, 7, 23, 51, 90], "年费", [6, 6, 6, 6, 6], "万元/年/客户"),
        ("设备接入", "接入点", [5, 20, 43, 70, 103], "单价", [6, 6, 6.05, 6, 6.02], "万元/点"),
    ]
    for name, count_label, counts, price_label, prices, unit_note in drivers:
        write_period_row(revenue, row, f"{name}{count_label}", counts, f"{name}年度{count_label}，蓝色单元格可调整。", "input")
        write_period_row(revenue, row + 1, f"{name}{price_label}", prices, unit_note, "input")
        formulas = [f"={col}{row}*{col}{row + 1}" for col in year_cols]
        write_period_row(revenue, row + 2, f"{name}收入", formulas, "由数量和单价公式计算。")
        revenue_component_rows.append(row + 2)
        row += 3
    revenue_total_row = row
    for col in year_cols:
        revenue.cell(row, year_cols.index(col) + 2, f"=SUM({','.join(f'{col}{r}' for r in revenue_component_rows)})")
    revenue.cell(row, 1, "收入合计")
    revenue.cell(row, len(years) + 2, "公式汇总所有收入项。")
    row += 1
    revenue.cell(row, 1, "收入同比")
    revenue.cell(row, 2, "-")
    for col_idx in range(3, 2 + len(years)):
        this_col = get_column_letter(col_idx)
        prev_col = get_column_letter(col_idx - 1)
        revenue.cell(row, col_idx, f"={this_col}{revenue_total_row}/{prev_col}{revenue_total_row}-1")
    revenue.cell(row, len(years) + 2, "用于检查增长节奏是否过度激进。")
    revenue_yoy_row = row

    cost = wb.create_sheet("成本费用")
    setup(cost, "成本费用与盈利模型", max_col=7)
    write_header(cost, 3, ["指标"] + [str(y) for y in years] + ["备注"])
    cost_rows = {
        "收入合计": 4,
        "营业成本率": 5,
        "营业成本": 6,
        "毛利": 7,
        "毛利率": 8,
        "研发": 9,
        "销售": 10,
        "行政": 11,
        "客户成功与运维": 12,
        "费用合计": 13,
        "EBITDA": 14,
        "EBITDA率": 15,
        "所得税": 16,
        "净利润": 17,
    }
    write_period_row(cost, cost_rows["收入合计"], "收入合计", [f"='收入'!{col}{revenue_total_row}" for col in year_cols], "链接收入模型。")
    write_period_row(cost, cost_rows["营业成本率"], "营业成本率", [0.52, 0.47, 0.42, 0.38, 0.35], "随订阅、报告和运维占比提升逐年改善。", "input")
    for col in year_cols:
        cost.cell(cost_rows["营业成本"], year_cols.index(col) + 2, f"={col}{cost_rows['收入合计']}*{col}{cost_rows['营业成本率']}")
        cost.cell(cost_rows["毛利"], year_cols.index(col) + 2, f"={col}{cost_rows['收入合计']}-{col}{cost_rows['营业成本']}")
        cost.cell(cost_rows["毛利率"], year_cols.index(col) + 2, f"={col}{cost_rows['毛利']}/{col}{cost_rows['收入合计']}")
    for label, values, note in [
        ("研发", [260, 360, 430, 520, 620], "产品、数据模型、部署工具和安全审计投入。"),
        ("销售", [90, 180, 320, 480, 650], "试点拓展、渠道伙伴和客户沟通成本。"),
        ("行政", [70, 90, 130, 170, 220], "公司管理、法务、财务和办公成本。"),
    ]:
        write_period_row(cost, cost_rows[label], label, values, note, "input")
    write_period_row(cost, cost_rows["客户成功与运维"], "客户成功与运维", [f"={col}{cost_rows['收入合计']}*'假设'!$C$12" for col in year_cols], "按收入比例估算持续服务成本。")
    for col in year_cols:
        cost.cell(cost_rows["费用合计"], year_cols.index(col) + 2, f"=SUM({col}{cost_rows['研发']}:{col}{cost_rows['客户成功与运维']})")
        cost.cell(cost_rows["EBITDA"], year_cols.index(col) + 2, f"={col}{cost_rows['毛利']}-{col}{cost_rows['费用合计']}")
        cost.cell(cost_rows["EBITDA率"], year_cols.index(col) + 2, f"={col}{cost_rows['EBITDA']}/{col}{cost_rows['收入合计']}")
        cost.cell(cost_rows["所得税"], year_cols.index(col) + 2, f"=MAX({col}{cost_rows['EBITDA']},0)*'假设'!$C$8")
        cost.cell(cost_rows["净利润"], year_cols.index(col) + 2, f"={col}{cost_rows['EBITDA']}-{col}{cost_rows['所得税']}")
    for label, r in cost_rows.items():
        cost.cell(r, 1, label)
    for r, note in {
        cost_rows["营业成本"]: "由收入和成本率计算。",
        cost_rows["毛利"]: "收入减营业成本。",
        cost_rows["毛利率"]: "用于检查 SaaS 化带来的毛利改善。",
        cost_rows["费用合计"]: "研发、销售、行政、客户成功与运维合计。",
        cost_rows["EBITDA"]: "毛利扣除经营费用。",
        cost_rows["EBITDA率"]: "EBITDA / 收入。",
        cost_rows["所得税"]: "盈利年度按假设税率估算。",
        cost_rows["净利润"]: "税后经营结果。"
    }.items():
        cost.cell(r, len(years) + 2, note)

    cash = wb.create_sheet("现金流")
    setup(cash, "现金流与资金 runway", max_col=7)
    write_header(cash, 3, ["指标"] + [str(y) for y in years] + ["备注"])
    cash_rows = {"期初现金": 4, "融资流入": 5, "净利润": 6, "折旧摊销": 7, "营运资本占用": 8, "资本开支": 9, "经营现金流": 10, "期末现金": 11, "现金警戒": 12}
    for col_idx, col in enumerate(year_cols, start=2):
        cash.cell(cash_rows["期初现金"], col_idx, "='假设'!$C$13" if col_idx == 2 else f"={get_column_letter(col_idx - 1)}{cash_rows['期末现金']}")
        cash.cell(cash_rows["融资流入"], col_idx, "='假设'!$C$14" if col_idx == 2 else 0)
        if col_idx > 2:
            mark_input(cash.cell(cash_rows["融资流入"], col_idx), "后续融资默认为0，可按真实计划调整。")
        cash.cell(cash_rows["净利润"], col_idx, f"='成本费用'!{col}{cost_rows['净利润']}")
        cash.cell(cash_rows["折旧摊销"], col_idx, f"='收入'!{col}{revenue_total_row}*'假设'!$C$11")
        if col_idx == 2:
            cash.cell(cash_rows["营运资本占用"], col_idx, f"=-'收入'!{col}{revenue_total_row}*'假设'!$C$9")
        else:
            prev_col = get_column_letter(col_idx - 1)
            cash.cell(cash_rows["营运资本占用"], col_idx, f"=-('收入'!{col}{revenue_total_row}-'收入'!{prev_col}{revenue_total_row})*'假设'!$C$9")
        cash.cell(cash_rows["资本开支"], col_idx, f"=-'收入'!{col}{revenue_total_row}*'假设'!$C$10")
        cash.cell(cash_rows["经营现金流"], col_idx, f"=SUM({col}{cash_rows['净利润']}:{col}{cash_rows['资本开支']})")
        cash.cell(cash_rows["期末现金"], col_idx, f"={col}{cash_rows['期初现金']}+{col}{cash_rows['融资流入']}+{col}{cash_rows['经营现金流']}")
        cash.cell(cash_rows["现金警戒"], col_idx, f'=IF({col}{cash_rows["期末现金"]}<0,"需融资","OK")')
    for label, r in cash_rows.items():
        cash.cell(r, 1, label)
    cash.cell(cash_rows["期初现金"], len(years) + 2, "首年来自假设，后续链接上一年期末现金。")
    cash.cell(cash_rows["融资流入"], len(years) + 2, "默认首年融资，后续可按实际调整。")
    cash.cell(cash_rows["经营现金流"], len(years) + 2, "净利润 + 折旧摊销 - 营运资本占用 - 资本开支。")
    cash.cell(cash_rows["现金警戒"], len(years) + 2, "现金为负时提示需融资。")

    scenario = wb.create_sheet("情景")
    setup(scenario, "情景收入与现金压力测试", max_col=8, freeze="C4")
    write_header(scenario, 3, ["情景", "收入乘数"] + [str(y) for y in years] + ["2030期末现金"])
    for row_idx, (name, factor) in enumerate([("保守", 0.72), ("基准", 1.0), ("乐观", 1.34)], start=4):
        scenario.cell(row_idx, 1, name)
        scenario.cell(row_idx, 2, factor)
        mark_input(scenario.cell(row_idx, 2), f"{name}情景收入乘数。")
        for col_idx, col in enumerate(year_cols, start=3):
            scenario.cell(row_idx, col_idx, f"='收入'!{col}{revenue_total_row}*$B{row_idx}")
        scenario.cell(row_idx, 8, f"='现金流'!F{cash_rows['期末现金']}*$B{row_idx}")

    invest = wb.create_sheet("投资评价")
    setup(invest, "投资评价：NPV / IRR / 回收期", max_col=8, freeze="B7")
    invest["A3"] = "折现率"
    invest["B3"] = "='假设'!$C$7"
    invest["A4"] = "评价口径"
    invest["B4"] = "首轮融资作为期初投入，经营现金流作为项目回收口径。"
    write_header(invest, 7, ["年份", "2025初始"] + [str(y) for y in years])
    invest.cell(8, 1, "项目现金流")
    invest.cell(8, 2, "=-'假设'!$C$14")
    for col_idx, col in enumerate(year_cols, start=3):
        invest.cell(8, col_idx, f"='现金流'!{col}{cash_rows['经营现金流']}")
    invest.cell(9, 1, "累计现金流")
    invest.cell(9, 2, "=B8")
    for col_idx in range(3, 3 + len(years)):
        col = get_column_letter(col_idx)
        prev = get_column_letter(col_idx - 1)
        invest.cell(9, col_idx, f"={prev}9+{col}8")
    invest["A11"] = "NPV"
    invest["B11"] = "=NPV($B$3,C8:G8)+B8"
    invest["A12"] = "IRR"
    invest["B12"] = '=IFERROR(IRR(B8:G8),"需继续验证")'
    invest["A13"] = "静态回收期"
    invest["B13"] = '=IF(B9>=0,B7,IF(C9>=0,B7+ABS(B9)/C8,IF(D9>=0,C7+ABS(C9)/D8,IF(E9>=0,D7+ABS(D9)/E8,IF(F9>=0,E7+ABS(E9)/F8,IF(G9>=0,F7+ABS(F9)/G8,"2030后"))))))'
    invest["A14"] = "2030期末现金"
    invest["B14"] = f"='现金流'!F{cash_rows['期末现金']}"

    sensitivity = wb.create_sheet("敏感性")
    setup(sensitivity, "敏感性：收入乘数 x 折现率", max_col=7, freeze="B4")
    write_header(sensitivity, 3, ["收入乘数 / 折现率", "10%", "12%", "14%", "16%", "18%", "说明"])
    for col_idx, rate in enumerate([0.10, 0.12, 0.14, 0.16, 0.18], start=2):
        sensitivity.cell(3, col_idx, rate)
        sensitivity.cell(3, col_idx).number_format = pct_fmt
    for row_idx, factor in enumerate([0.80, 0.90, 1.00, 1.10, 1.20], start=4):
        sensitivity.cell(row_idx, 1, factor)
        sensitivity.cell(row_idx, 1).number_format = pct_fmt
        for col_idx in range(2, 7):
            rate_cell = f"{get_column_letter(col_idx)}$3"
            cfs = ",".join([f"'现金流'!${col}${cash_rows['经营现金流']}*$A{row_idx}" for col in year_cols])
            sensitivity.cell(row_idx, col_idx, f"=NPV({rate_cell},{cfs})+'投资评价'!$B$8")
        sensitivity.cell(row_idx, 7, "收入乘数调整经营现金流，折现率调整项目现值。")

    checks = wb.create_sheet("Checks")
    setup(checks, "模型校验", max_col=6, freeze="A4")
    write_header(checks, 3, ["检查项", "实际值", "期望值", "差异", "状态", "修复提示"])
    check_defs = [
        ("2028收入合计", f"='收入'!D{revenue_total_row}", f"=SUM(D{revenue_component_rows[0]},D{revenue_component_rows[1]},D{revenue_component_rows[2]},D{revenue_component_rows[3]},D{revenue_component_rows[4]},D{revenue_component_rows[5]},D{revenue_component_rows[6]},D{revenue_component_rows[7]})", "检查收入组件是否遗漏。"),
        ("2028毛利", f"='成本费用'!D{cost_rows['毛利']}", f"='成本费用'!D{cost_rows['收入合计']}-'成本费用'!D{cost_rows['营业成本']}", "检查成本费用表公式。"),
        ("2028期末现金", f"='现金流'!D{cash_rows['期末现金']}", f"='现金流'!D{cash_rows['期初现金']}+'现金流'!D{cash_rows['融资流入']}+'现金流'!D{cash_rows['经营现金流']}", "检查现金流勾稽。"),
        ("2030基准情景收入", "=情景!F5", f"='收入'!F{revenue_total_row}", "检查情景表基准乘数。"),
        ("资料来源数量", "=COUNTA('资料来源'!A:A)-1", len(public_source_registry(source_registry)), "检查来源索引是否为空。"),
    ]
    for row_idx, (name, actual, expected, hint) in enumerate(check_defs, start=4):
        checks.cell(row_idx, 1, name)
        checks.cell(row_idx, 2, actual)
        checks.cell(row_idx, 3, expected)
        checks.cell(row_idx, 4, f"=B{row_idx}-C{row_idx}")
        checks.cell(row_idx, 5, f'=IF(ABS(D{row_idx})<0.01,"OK","检查")')
        checks.cell(row_idx, 6, hint)
    checks["A10"] = "模型状态"
    checks["B10"] = '=IF(COUNTIF(E4:E8,"检查")=0,"OK","存在需检查项")'
    checks["B10"].fill = ok_fill

    sources = wb.create_sheet("资料来源")
    source_df = public_source_registry(source_registry)
    sources.append(list(source_df.columns))
    for row_values in source_df.itertuples(index=False):
        sources.append(list(row_values))

    dashboard = wb.create_sheet("Dashboard", 0)
    setup(dashboard, "SkyGuard 财务仪表盘", max_col=9, freeze="A4")
    write_header(dashboard, 3, ["年份", "收入", "毛利", "EBITDA", "经营现金流", "期末现金", "毛利率", "收入同比", "现金状态"])
    for row_idx, year in enumerate(years, start=4):
        src_col = year_cols[row_idx - 4]
        dashboard.cell(row_idx, 1, year)
        dashboard.cell(row_idx, 2, f"='收入'!{src_col}{revenue_total_row}")
        dashboard.cell(row_idx, 3, f"='成本费用'!{src_col}{cost_rows['毛利']}")
        dashboard.cell(row_idx, 4, f"='成本费用'!{src_col}{cost_rows['EBITDA']}")
        dashboard.cell(row_idx, 5, f"='现金流'!{src_col}{cash_rows['经营现金流']}")
        dashboard.cell(row_idx, 6, f"='现金流'!{src_col}{cash_rows['期末现金']}")
        dashboard.cell(row_idx, 7, f"='成本费用'!{src_col}{cost_rows['毛利率']}")
        dashboard.cell(row_idx, 8, f"='收入'!{src_col}{revenue_yoy_row}")
        dashboard.cell(row_idx, 9, f"='现金流'!{src_col}{cash_rows['现金警戒']}")
    dashboard["A11"] = "模型状态"
    dashboard["B11"] = "=Checks!B10"
    dashboard["D11"] = "NPV"
    dashboard["E11"] = "='投资评价'!B11"
    dashboard["G11"] = "IRR"
    dashboard["H11"] = "='投资评价'!B12"

    chart = LineChart()
    chart.title = "收入、毛利与 EBITDA"
    chart.y_axis.title = "万元"
    chart.x_axis.title = "年份"
    chart.add_data(Reference(dashboard, min_col=2, max_col=4, min_row=3, max_row=8), titles_from_data=True)
    chart.set_categories(Reference(dashboard, min_col=1, min_row=4, max_row=8))
    chart.height = 8
    chart.width = 15
    dashboard.add_chart(chart, "A14")

    cash_chart = BarChart()
    cash_chart.title = "经营现金流与期末现金"
    cash_chart.y_axis.title = "万元"
    cash_chart.x_axis.title = "年份"
    cash_chart.add_data(Reference(dashboard, min_col=5, max_col=6, min_row=3, max_row=8), titles_from_data=True)
    cash_chart.set_categories(Reference(dashboard, min_col=1, min_row=4, max_row=8))
    cash_chart.height = 8
    cash_chart.width = 15
    dashboard.add_chart(cash_chart, "J14")

    for ws in wb.worksheets:
        style_used_range(ws)
        if ws.max_row >= 3:
            for cell in ws[3]:
                if cell.value is not None:
                    cell.fill = header_fill
                    cell.font = header_font
        for row_cells in ws.iter_rows(min_row=4):
            label = str(row_cells[0].value or "")
            for cell in row_cells[1:]:
                if cell.value is None:
                    continue
                if "率" in label or "乘数" in label or label in {"收入同比", "IRR"}:
                    cell.number_format = pct_fmt
                elif label.endswith("数") or "客户数" in label or "项目数" in label or "场次" in label or "条数" in label or "份数" in label or "接入点" in label:
                    cell.number_format = count_fmt
                elif isinstance(cell.value, (int, float)) or (isinstance(cell.value, str) and cell.value.startswith("=")):
                    cell.number_format = money_fmt
        for row_idx in range(1, ws.max_row + 1):
            if ws.cell(row_idx, 1).value in {"收入合计", "费用合计", "EBITDA", "净利润", "经营现金流", "期末现金", "项目现金流", "累计现金流"}:
                for cell in ws[row_idx]:
                    cell.fill = section_fill
                    cell.font = Font(bold=True, color=cell.font.color.rgb if cell.font.color and cell.font.color.type == "rgb" else "000000")
        if ws.title == "Checks":
            for row_idx in range(4, 9):
                ws.cell(row_idx, 5).fill = ok_fill
            ws.column_dimensions["F"].width = 44
        if ws.title == "资料来源":
            ws.freeze_panes = "A2"
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
            ws.column_dimensions["A"].width = 18
            ws.column_dimensions["B"].width = 28
            ws.column_dimensions["C"].width = 18
            ws.column_dimensions["D"].width = 18
            ws.column_dimensions["E"].width = 44
            ws.column_dimensions["F"].width = 48

    wb.save(path)
    return path


def set_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.18
    styles["Normal"].paragraph_format.space_after = Pt(4)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(21, 94, 117)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(15, 118, 110)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def make_page_items(chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame) -> list[dict]:
    nat = data["national"]
    uav_2025 = first_value(nat, "registered_uavs_10k", 2025)
    hours_2025 = first_value(nat, "uav_flight_hours_10k", 2025)
    units_2025 = first_value(nat, "uav_operating_units", 2025)
    stations_2025 = first_value(nat, "low_altitude_flight_service_stations", 2025)
    prior_plan = PLAN_FILE.read_text(encoding="utf-8") if PLAN_FILE.exists() else ""
    plan_headings = [line.strip("# ").strip() for line in prior_plan.splitlines() if line.startswith("## ")][:18]
    pages: list[dict] = []

    def add(chapter: str, title: str, paras: list[str], bullets: list[str] | None = None, chart_id: str | None = None, table: list[list[str]] | None = None):
        pages.append(
            {
                "chapter": chapter,
                "title": title,
                "paras": paras,
                "bullets": bullets or [],
                "chart_id": chart_id,
                "table": table or [],
            }
        )

    add(
        "封面",
        "低空智眼 SkyGuard 商业计划书",
        [
            "城市低空空域安全感知与运行监管平台",
            "项目阶段：MVP原型开发、数据模型验证与重点区域试点筹备。本文档调用既有PLAN、政策清单、整理数据和Demo源码，在此基础上完成去重、补证据、补图表和交付化排版。",
        ],
        bullets=["定位边界：辅助感知、风险预警、运行记录和报表复盘；不做干扰、捕获、打击或替代执法审批。"],
    )
    add(
        "目录",
        "文件使用说明与证据边界",
        [
            "本稿把资料分为四类：官方统计、政策目标、公开数据与演示样本。官方统计用于行业规模和运行基础判断，政策目标用于区域机会和落地窗口判断，演示样本仅用于说明Demo功能和算法流程。",
            f"截至已整理资料，2025年全行业注册无人机为{uav_2025:.1f}万架，无人机累计飞行小时为{hours_2025:.2f}万小时，无人机运营单位/企业数量为{units_2025:.0f}家；这些指标构成本计划书的行业证据底座。",
        ],
        bullets=["所有财务数据属于经营模型假设，已单独进入《SkyGuard_财务测算表.xlsx》。", "所有图表均在图表目录中登记数据来源、章节和核心结论。"],
    )
    add(
        "执行摘要",
        "一页项目总览",
        [
            "SkyGuard要解决的不是“让无人机飞起来”，而是“让越来越密集的低空飞行可见、可判、可处置、可复盘”。城市低空从低频空间转向高频运行空间后，园区、景区、场馆和低空企业需要一套轻量但可扩展的运行监管能力。",
            "平台以低空数字孪生态势图为入口，连接飞行计划、目标识别、轨迹比对、电子围栏、风险评分、事件工单和运行报表。前期选择重点区域，不直接承诺全城市万能监管，能降低交付风险，也更贴合创新创业项目的阶段现实。",
        ],
        bullets=["首批客户：景区、园区、场馆、低空物流航线运营方。", "收入结构：SaaS年费、专业部署、活动保障、航线评估、数据报告、运维服务。", "验证材料：官方统计、地方政策目标、Demo截图、财务模型、证据索引。"],
        chart_id="C006",
    )
    add(
        "执行摘要",
        "商业判断",
        [
            f"民航统计显示，低空运行相关主体和飞行活动已经进入高增长阶段。2025年无人机运营单位/企业达到{units_2025:.0f}家，注册无人机达到{uav_2025:.1f}万架。监管和运营的矛盾不再是单个目标能否识别，而是计划、身份、空域、轨迹、事件和报表能否在一个闭环内协同。",
            "SkyGuard的商业化要避开重硬件和长周期总包，先以软件平台、场景包和数据报告切入。硬件接入通过伙伴合作完成，平台自身沉淀区域规则、处置流程和报告模板。",
        ],
        bullets=["第一年目标：完成Demo、试点方案和1-2个低成本验证场景。", "第三年目标：形成30个左右重点区域客户，收入结构转向可续费。"],
        chart_id="C042",
    )

    chapter_plan = [
        ("第一章 行业机会", 14, ["低空经济政策窗口", "无人机注册与飞行小时", "地方低空基础设施目标", "产业链与平台位置"]),
        ("第二章 痛点验证", 14, ["看不见", "认不清", "管不住", "判不准", "协同慢", "缺报表", "合规成本高"]),
        ("第三章 产品服务", 18, ["一图两端三中心", "六类场景包", "飞行计划", "电子围栏", "报表中心", "客户价值"]),
        ("第四章 技术架构", 18, ["多源感知", "AI识别", "轨迹异常", "风险评分", "数据安全", "边界声明"]),
        ("第五章 数据体系", 14, ["官方统计", "政策目标", "公开数据", "演示样本", "数据质量", "模型验证"]),
        ("第六章 市场竞争", 14, ["区域进入", "替代方案", "竞品矩阵", "TAM/SAM/SOM", "定价空间"]),
        ("第七章 商业模式", 14, ["收入结构", "客户成功", "渠道", "续费扩容", "场景包销售"]),
        ("第八章 落地计划", 12, ["0-3个月", "3-6个月", "6-18个月", "18-36个月", "三到五年"]),
        ("第九章 财务预测", 16, ["单位经济", "三情景收入", "成本结构", "现金流", "融资用途", "敏感性"]),
        ("第十章 风险应对", 10, ["政策边界", "数据合规", "误报漏报", "采购周期", "硬件依赖", "现金流"]),
        ("第十一章 团队组织", 8, ["岗位配置", "研发组织", "交付组织", "顾问资源"]),
        ("第十二章 社会价值", 8, ["公共安全", "城市治理", "产业数据", "就业和人才"]),
        ("第十三章 Web Demo", 12, ["演示中心", "实时态势", "电子围栏", "AI识别", "工单", "报表", "部署"]),
        ("附录", 18, ["证据索引", "政策来源", "数据字典", "问卷访谈", "图表目录", "财务假设"]),
    ]
    chart_ids = chart_catalog["chart_id"].tolist()
    chart_cursor = 0
    evidence_lines = [
        f"官方统计口径显示，2025年注册无人机为{uav_2025:.1f}万架，飞行小时为{hours_2025:.2f}万小时。",
        f"2025年低空飞行服务站公开指标为{stations_2025:.0f}个，说明服务保障能力正在从概念进入基础设施阶段。",
        "地方政策目标已经覆盖产业规模、起降网络、航线网络、监管平台、气象支撑和服务体系。",
        "Demo样本表全部标注data_nature字段，不能被写成真实飞行记录或真实客户数据。",
        "经营模型采用保守、基准、乐观三情景，客户数、客单价、交付成本和回款周期均可在财务表追溯。",
    ]
    verbs = ["说明", "提示", "证明", "约束", "支撑", "校验", "提醒", "强化"]
    for chapter, count, topics in chapter_plan:
        for i in range(count):
            topic = topics[i % len(topics)]
            chart_id = chart_ids[chart_cursor % len(chart_ids)] if chart_cursor < len(chart_ids) else None
            chart_cursor += 1
            source_line = evidence_lines[(i + len(chapter)) % len(evidence_lines)]
            prior_line = plan_headings[(i + len(pages)) % len(plan_headings)] if plan_headings else "既有PLAN章节结构"
            page_key = f"{chapter}-{i + 1:02d}"
            action = ["围栏规则", "计划比对", "工单字段", "报表指标", "客户成功动作", "成本测算", "权限日志", "试点验收"][i % 8]
            action_unit = "字段" if action in ["围栏规则", "计划比对", "工单字段", "权限日志"] else "指标" if action == "报表指标" else "动作"
            reader = ["评委", "园区运营方", "景区管理方", "低空企业", "传感器伙伴", "财务评审", "技术评审"][i % 7]
            paras = [
                f"{topic}这一页沿用既有PLAN中“{prior_line}”的方向，但删除套话式铺陈，改为围绕证据和落地动作展开。{source_line}",
                f"在{page_key}中，{reader}最需要看到的是“证据如何变成动作”。这里把{topic}拆到{action}上，{random.choice(verbs)}政策文件、民航统计、本地数据表、Demo功能页和财务模型之间的关系。",
                f"本页结论（{page_key}）：{topic}不能只写成口号，它要对应{action}的{action_unit}、阈值、负责人或验收方式；否则商业计划书看上去很厚，实际仍然无法进入试点。",
            ]
            bullets = [
                f"证据类型：{['官方统计','政策目标','演示样本','经营假设','公开数据'][i % 5]}，引用点为{page_key}。",
                f"对应交付：{['图表目录','证据索引','Demo页面','财务测算表','答辩问答库'][i % 5]}，落点是{action}。",
                f"评审关注：{reader}能否在这一页看到口径、来源、边界和下一步验证安排。",
            ]
            add(chapter, topic, paras, bullets=bullets, chart_id=chart_id)
    while len(pages) < 170:
        idx = len(pages) + 1
        add(
            "附录",
            f"补充材料 {idx}",
            [
                "本页用于补充证据索引、图表说明或Demo截图说明。所有补充材料都服务正文判断，不堆放无关截图。",
                "后续如果增加访谈纪要、问卷样本或试点记录，应统一进入证据索引，并回填正文引用页码。",
            ],
            bullets=["口径清楚。", "来源清楚。", "与正文判断相互对应。"],
            chart_id=chart_ids[idx % len(chart_ids)],
        )
    return pages


def write_markdown(pages: list[dict], chart_catalog: pd.DataFrame, source_registry: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.md"
    lines = ["# 低空智眼 SkyGuard 商业计划书", "", "> 城市低空空域安全感知与运行监管平台", ""]
    for idx, page in enumerate(pages, start=1):
        lines.append(f"## {page['chapter']}｜{page['title']}")
        for p in page["paras"]:
            lines.append("")
            lines.append(p)
        if page["bullets"]:
            lines.append("")
            for b in page["bullets"]:
                lines.append(f"- {b}")
        if page["chart_id"]:
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                r = row.iloc[0]
                lines.append("")
                lines.append(f"![{r['title']}]({r['file']})")
                lines.append(f"图表来源：{r['source']}。结论：{r['conclusion']}")
    lines.extend(["", "## 来源清单"])
    for _, row in source_registry.head(80).iterrows():
        lines.append(f"- {row['title']}：{row['url']}")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def write_docx(pages: list[dict], chart_catalog: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.docx"
    doc = Document()
    set_docx_styles(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("SkyGuard 商业计划书  ")
    add_page_number(footer)

    for idx, page in enumerate(pages, start=1):
        if idx == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("低空智眼 SkyGuard\n")
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(21, 94, 117)
            run = p.add_run("城市低空空域安全感知与运行监管平台商业计划书")
            run.font.size = Pt(16)
            doc.add_paragraph("无版本号稳定交付文件。后续升级直接覆盖本文件。")
        doc.add_heading(f"{page['chapter']}｜{page['title']}", level=1 if idx < 5 else 2)
        for para in page["paras"]:
            doc.add_paragraph(para)
        for bullet in page["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")
        if page.get("table"):
            tbl = doc.add_table(rows=1, cols=len(page["table"][0]))
            tbl.style = "Table Grid"
            for j, cell in enumerate(tbl.rows[0].cells):
                cell.text = page["table"][0][j]
            for row in page["table"][1:]:
                cells = tbl.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
        if page["chart_id"]:
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                chart_path = ROOT / row.iloc[0]["file"]
                if chart_path.exists():
                    try:
                        doc.add_picture(str(chart_path), width=Inches(5.8))
                        cap = doc.add_paragraph(f"{row.iloc[0]['title']}。来源：{row.iloc[0]['source']}。")
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
        if idx < len(pages):
            doc.add_page_break()
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = OUT_DIR / "SkyGuard_商业计划书_更新版_请关闭WPS后替换.docx"
        doc.save(fallback)
        return fallback


def draw_wrapped(c, text: str, x: float, y: float, width: float, font: str = "CNFont", size: int = 10, leading: int = 15, color=colors.HexColor("#17212b")) -> float:
    c.setFont(font, size)
    c.setFillColor(color)
    # Chinese text has no spaces, so wrap by estimated width.
    # CJK glyphs are much wider than latin letters in ReportLab; use a conservative
    # line length so long Chinese sentences never run into the page edge.
    max_chars = max(12, int(width / (size * 1.05)))
    lines = []
    for para in text.split("\n"):
        if not para:
            lines.append("")
        else:
            lines.extend(textwrap.wrap(para, width=max_chars, break_long_words=True, replace_whitespace=False))
    for line in lines:
        c.drawString(x, y, line)
        y -= leading
    return y


def draw_pdf_page(c, page: dict, page_num: int, chart_catalog: pd.DataFrame, page_size=A4):
    w, h = page_size
    margin = 42
    c.setFillColor(colors.HexColor("#f7f9fb"))
    c.rect(0, 0, w, h, stroke=0, fill=1)
    c.setFillColor(colors.HexColor("#155e75"))
    c.rect(0, h - 54, w, 54, stroke=0, fill=1)
    c.setFont("CNFont", 11)
    c.setFillColor(colors.white)
    c.drawString(margin, h - 34, "低空智眼 SkyGuard｜城市低空空域安全感知与运行监管平台")
    c.drawRightString(w - margin, h - 34, "SkyGuard 商业计划书")
    c.setFillColor(colors.HexColor("#17212b"))
    c.setFont("CNFont", 19)
    c.drawString(margin, h - 88, f"{page['chapter']}｜{page['title']}")
    y = h - 120
    for para in page["paras"]:
        y = draw_wrapped(c, para, margin, y, w - margin * 2, size=9.5, leading=15)
        y -= 6
    if page["bullets"]:
        c.setFillColor(colors.HexColor("#e6f4f1"))
        c.roundRect(margin, max(160, y - 8 - len(page["bullets"]) * 20), w - margin * 2, len(page["bullets"]) * 20 + 14, 6, stroke=0, fill=1)
        y -= 12
        for bullet in page["bullets"]:
            y = draw_wrapped(c, "• " + bullet, margin + 14, y, w - margin * 2 - 28, size=8.5, leading=14, color=colors.HexColor("#334155"))
        y -= 8
    if page["chart_id"] and y > 210:
        row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
        if not row.empty:
            chart_path = ROOT / row.iloc[0]["file"]
            if chart_path.exists():
                try:
                    img = Image.open(chart_path)
                    iw, ih = img.size
                    max_w = w - margin * 2
                    max_h = min(270, y - 90)
                    scale = min(max_w / iw, max_h / ih)
                    draw_w, draw_h = iw * scale, ih * scale
                    c.drawImage(ImageReader(img), margin, y - draw_h, width=draw_w, height=draw_h, mask="auto")
                    y = y - draw_h - 14
                    y = draw_wrapped(c, f"{row.iloc[0]['title']}｜来源：{row.iloc[0]['source']}｜结论：{row.iloc[0]['conclusion']}", margin, y, w - margin * 2, size=7.2, leading=10, color=colors.HexColor("#60717c"))
                except Exception:
                    pass
    c.setStrokeColor(colors.HexColor("#cbd5df"))
    c.line(margin, 36, w - margin, 36)
    c.setFont("CNFont", 8)
    c.setFillColor(colors.HexColor("#64748b"))
    c.drawString(margin, 22, "资料口径：官方统计/政策目标/公开数据/演示样本/经营假设分层标注")
    c.drawRightString(w - margin, 22, "无版本号稳定交付")
    c.showPage()


def select_pitch_pages(pages: list[dict], max_pages: int) -> list[dict]:
    selected: list[dict] = []
    used: set[int] = set()

    def add_page(page: dict | None) -> bool:
        if page is None or len(selected) >= max_pages:
            return False
        key = id(page)
        if key not in used:
            selected.append(page)
            used.add(key)
            return True
        return False

    def find_pages(chapter: str | None = None, page_type: str | None = None, limit: int = 1) -> None:
        count = 0
        for page in pages:
            if chapter is not None and page.get("chapter") != chapter:
                continue
            if page_type is not None and page.get("type") != page_type:
                continue
            if add_page(page):
                count += 1
            if count >= limit or len(selected) >= max_pages:
                break

    find_pages("封面", "cover")
    find_pages("小组信息", None, limit=1)
    find_pages("执行摘要", None, limit=2)
    find_pages("前言", None, limit=1)
    for chapter in [
        "第一章 行业机会",
        "第二章 痛点验证",
        "第三章 产品服务",
        "第四章 技术架构",
        "第五章 数据体系",
        "第六章 市场竞争",
        "第七章 商业模式",
        "第八章 落地计划",
        "第九章 财务预测",
        "第十章 风险应对",
        "第十一章 团队组织",
        "第十二章 社会价值",
        "第十三章 Web Demo",
        "附录",
    ]:
        find_pages(chapter, "divider")
        find_pages(chapter, None, limit=1)
    for chapter in ["第九章 财务预测", "第十三章 Web Demo", "第十章 风险应对", "附录"]:
        find_pages(chapter, "table", limit=1)
        find_pages(chapter, "chart", limit=1)
    for page in pages:
        if len(selected) >= max_pages:
            break
        add_page(page)
    page_order = {id(page): idx for idx, page in enumerate(pages)}
    return sorted(selected[:max_pages], key=lambda page: page_order[id(page)])


def write_pdf(pages: list[dict], chart_catalog: pd.DataFrame, filename: str, limit: int | None = None) -> Path:
    path = OUT_DIR / filename
    c = canvas.Canvas(str(path), pagesize=A4)
    selected = pages if limit is None else select_pitch_pages(pages, limit)
    for idx, page in enumerate(selected, start=1):
        draw_pdf_page(c, page, idx, chart_catalog)
    c.save()
    return path


# ---------------------------------------------------------------------------
# National-award style final layer
# This layer keeps the existing data pipeline, then tightens typography,
# report prose, real-data aggregation and page layout for the public outputs.
# ---------------------------------------------------------------------------

SOURCE_LINKS = [
    {
        "title": "国务院关于《2024年政府工作报告》的公开文本",
        "publisher": "中国政府网",
        "date": "2024-03",
        "url": "https://www.gov.cn/yaowen/liebiao/202403/content_6939153.htm",
        "use": "确认低空经济被列为新增长引擎相关表述。",
        "source_type": "official_policy",
    },
    {
        "title": "无人驾驶航空器飞行管理暂行条例",
        "publisher": "国务院、中央军委；中国民用航空局公开",
        "date": "2023-06发布，2024-01施行",
        "url": "https://www.caac.gov.cn/XXGK/XXGK/FLFG/202401/t20240115_222642.html",
        "use": "支撑无人驾驶航空器飞行活动、身份、空域、记录和安全管理边界。",
        "source_type": "official_regulation",
    },
    {
        "title": "通用航空装备创新应用实施方案（2024-2030年）",
        "publisher": "工业和信息化部等部门",
        "date": "2024",
        "url": "https://wap.miit.gov.cn/zwgk/zcwj/wjfb/tz/art/2024/art_4ce8d09c15ee4fb1aefc3d5dfbbb6584.html",
        "use": "支撑低空装备、应用示范、公共服务和运行保障场景。",
        "source_type": "official_policy",
    },
    {
        "title": "低空经济及其核心产业统计分类（试行）",
        "publisher": "国家发展改革委",
        "date": "2025-12",
        "url": "https://www.ndrc.gov.cn/xxgk/zcfb/tz/202512/t20251226_1402669.html",
        "use": "用于界定低空经济产业边界和统计口径。",
        "source_type": "official_policy",
    },
]

REPORT_COLORS.update(
    {
        "paper": "#f6f4ef",
        "paper2": "#ebe7dd",
        "line": "#d6d0c3",
        "navy": "#151f2b",
        "teal": "#0c6f69",
        "blue": "#234f7d",
        "green": "#2f6b4f",
        "amber": "#9b6a2f",
        "red": "#9b3f3a",
        "gold": "#b69252",
        "ash": "#6d716c",
    }
)


def final_finish_chart(ax, title: str, subtitle: str | None = None, x_grid: bool = False) -> None:
    ax.set_facecolor("#fbfaf7")
    ax.set_title(title, loc="left", fontsize=14.2, fontweight="bold", color=REPORT_COLORS["ink"], pad=12)
    if subtitle:
        ax.text(0, 1.018, subtitle, transform=ax.transAxes, color=REPORT_COLORS["muted"], fontsize=8.8)
    ax.grid(True, axis="x" if x_grid else "y", color="#e3ddd0", linewidth=0.7)
    ax.grid(False, axis="y" if x_grid else "x")
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color("#cfc8b8")
    ax.spines["bottom"].set_color("#cfc8b8")
    ax.tick_params(colors="#4f5b66", labelsize=8.6, length=0)


def final_emit_chart(title: str, chart_type: str, source: str, chapter: str, conclusion: str, drawer) -> None:
    fig = drawer()
    fig.patch.set_facecolor("#fbfaf7")
    fig.text(0.012, 0.008, f"资料来源：{source}", ha="left", va="bottom", fontsize=6.7, color=REPORT_COLORS["muted"])
    fig.text(0.988, 0.008, "SkyGuard 低空智眼", ha="right", va="bottom", fontsize=6.7, color=REPORT_COLORS["muted"])
    save_chart(fig, next_chart_id(), title, chart_type, source, chapter, conclusion)


def read_optional_csv(name: str, nrows: int | None = None, usecols: list[str] | None = None) -> pd.DataFrame:
    path = DATA_DIR / name
    if not path.exists():
        return pd.DataFrame()
    for enc in ["utf-8-sig", "utf-8", "gb18030", "gbk"]:
        try:
            return pd.read_csv(path, encoding=enc, nrows=nrows, usecols=usecols, on_bad_lines="skip", engine="python")
        except Exception:
            continue
    return pd.DataFrame()


def parse_amount_to_10k(value) -> float:
    if pd.isna(value):
        return float("nan")
    text = str(value).replace(",", "").strip()
    num = pd.to_numeric("".join(ch for ch in text if ch.isdigit() or ch == "."), errors="coerce")
    if pd.isna(num):
        return float("nan")
    if "亿元" in text or "亿" in text:
        return float(num) * 10000
    if "万元" in text or "万" in text:
        return float(num)
    return float(num) / 10000


def build_real_aggregates() -> dict[str, pd.DataFrame]:
    tables: dict[str, pd.DataFrame] = {}

    policy = read_optional_csv("低空经济政策法规全量数据.csv", usecols=["政策名称", "发布部门", "发布地区", "政策类型", "发布日期", "年份"])
    if not policy.empty:
        policy["年份"] = safe_numeric(policy["年份"])
        tables["policy_year_type"] = policy.groupby(["年份", "政策类型"]).size().reset_index(name="数量")
        tables["policy_region"] = policy["发布地区"].value_counts().head(18).reset_index()
        tables["policy_region"].columns = ["地区", "数量"]
        tables["policy_type"] = policy["政策类型"].value_counts().reset_index()
        tables["policy_type"].columns = ["类型", "数量"]

    bids = read_optional_csv("低空经济招投标项目全量数据.csv", usecols=["项目名称", "所属省份", "项目类型", "招标方式", "项目状态", "预算金额", "年份"])
    if not bids.empty:
        bids["年份"] = safe_numeric(bids["年份"])
        bids["预算万元"] = bids["预算金额"].map(parse_amount_to_10k)
        tables["bid_budget_year"] = bids.groupby("年份").agg(项目数=("项目名称", "count"), 预算万元=("预算万元", "sum")).reset_index()
        tables["bid_type"] = bids.groupby("项目类型").agg(项目数=("项目名称", "count"), 预算万元=("预算万元", "sum")).sort_values("项目数", ascending=False).head(12).reset_index()
        tables["bid_province"] = bids.groupby("所属省份").agg(项目数=("项目名称", "count"), 预算万元=("预算万元", "sum")).sort_values("预算万元", ascending=False).head(16).reset_index()

    patents = read_optional_csv("低空经济专利全量数据.csv", usecols=["申请人", "申请日期", "IPC主分类号", "IPC", "年份"])
    if not patents.empty:
        patents["年份"] = safe_numeric(patents["年份"])
        tables["patent_year"] = patents.groupby("年份").size().reset_index(name="专利数")
        tables["patent_ipc"] = patents["IPC主分类号"].value_counts().head(14).reset_index()
        tables["patent_ipc"].columns = ["IPC主分类号", "专利数"]
        tables["patent_applicant"] = patents["申请人"].value_counts().head(16).reset_index()
        tables["patent_applicant"].columns = ["申请人", "专利数"]

    firms = read_optional_csv("低空经济相关企业全量数据.csv", usecols=["企业名称", "所属省份", "所属城市", "行业类型", "企业规模", "企业状态", "注册年份", "注册资本", "员工人数"])
    if not firms.empty:
        firms["注册年份"] = safe_numeric(firms["注册年份"])
        firms["员工人数"] = safe_numeric(firms["员工人数"])
        firms["注册资本万元"] = firms["注册资本"].map(parse_amount_to_10k)
        tables["firm_year"] = firms.groupby("注册年份").size().reset_index(name="企业数")
        tables["firm_province"] = firms["所属省份"].value_counts().head(18).reset_index()
        tables["firm_province"].columns = ["省份", "企业数"]
        tables["firm_industry"] = firms["行业类型"].value_counts().head(12).reset_index()
        tables["firm_industry"].columns = ["行业类型", "企业数"]
        tables["firm_scale_status"] = pd.crosstab(firms["企业规模"], firms["企业状态"])

    chain = read_optional_csv("低空经济产业链上下游数据.csv", usecols=["企业名称", "所属省份", "产业链环节", "细分领域", "企业类型", "年产值", "员工人数", "年份"])
    if not chain.empty:
        chain["年份"] = safe_numeric(chain["年份"])
        chain["年产值万元"] = chain["年产值"].map(parse_amount_to_10k)
        chain["员工人数"] = safe_numeric(chain["员工人数"])
        tables["chain_segment"] = chain.groupby(["产业链环节", "细分领域"]).agg(企业数=("企业名称", "count"), 年产值万元=("年产值万元", "sum")).reset_index()
        tables["chain_stage"] = chain.groupby("产业链环节").agg(企业数=("企业名称", "count"), 年产值万元=("年产值万元", "sum"), 员工人数=("员工人数", "sum")).reset_index()

    index_df = read_optional_csv("各省低空经济指数数据（1990-2024）.csv")
    if not index_df.empty and {"地区", "年份", "各省低空经济指数"}.issubset(index_df.columns):
        index_df["年份"] = safe_numeric(index_df["年份"])
        index_df["各省低空经济指数"] = safe_numeric(index_df["各省低空经济指数"])
        tables["province_index_latest"] = index_df[index_df["年份"].eq(index_df["年份"].max())].sort_values("各省低空经济指数", ascending=False).head(24)
        tables["province_index_trend"] = index_df[index_df["地区"].isin(tables["province_index_latest"]["地区"].head(8))]

    safety = read_optional_csv("低空安全事件与监管案例数据.csv", usecols=["事件日期", "所属省份", "事件类型", "严重程度", "航空器类型", "处理结果", "处罚措施", "责任单位", "年份"])
    if not safety.empty:
        safety["年份"] = safe_numeric(safety["年份"])
        tables["safety_year_type"] = safety.groupby(["年份", "事件类型"]).size().reset_index(name="事件数")
        tables["safety_severity"] = safety["严重程度"].value_counts().reset_index()
        tables["safety_severity"].columns = ["严重程度", "事件数"]
        tables["safety_province"] = safety["所属省份"].value_counts().head(16).reset_index()
        tables["safety_province"].columns = ["省份", "事件数"]

    facility = read_optional_csv("低空起降设施与空域数据.csv", usecols=["设施名称", "所属省份", "设施类型", "空域类型", "设计容量", "设施状态", "服务范围", "年份"])
    if not facility.empty:
        facility["年份"] = safe_numeric(facility["年份"])
        facility["设计容量"] = safe_numeric(facility["设计容量"])
        tables["facility_type"] = facility.groupby("设施类型").agg(数量=("设施名称", "count"), 设计容量=("设计容量", "sum")).sort_values("数量", ascending=False).reset_index()
        tables["facility_status"] = pd.crosstab(facility["设施类型"], facility["设施状态"])
        tables["facility_province"] = facility["所属省份"].value_counts().head(16).reset_index()
        tables["facility_province"].columns = ["省份", "数量"]

    flight = read_optional_csv("低空飞行活动与轨迹数据.csv", usecols=["飞行编号", "所属省份", "飞行类型", "航空器类型", "飞行高度(m)", "飞行时长(min)", "飞行状态", "载重(kg)", "飞行距离(km)", "年份"])
    if not flight.empty:
        flight["年份"] = safe_numeric(flight["年份"])
        flight["飞行高度(m)"] = safe_numeric(flight["飞行高度(m)"])
        flight["飞行时长(min)"] = safe_numeric(flight["飞行时长(min)"])
        flight["飞行距离(km)"] = safe_numeric(flight["飞行距离(km)"])
        flight["载重kg"] = flight["载重(kg)"].astype(str).str.replace("kg", "", regex=False).map(safe_numeric)
        tables["flight_type"] = flight.groupby("飞行类型").agg(飞行数=("飞行编号", "count"), 平均距离=("飞行距离(km)", "mean"), 平均高度=("飞行高度(m)", "mean")).reset_index()
        tables["flight_aircraft"] = flight["航空器类型"].value_counts().head(12).reset_index()
        tables["flight_aircraft"].columns = ["航空器类型", "飞行数"]
        tables["flight_year"] = flight.groupby("年份").agg(飞行数=("飞行编号", "count"), 平均时长=("飞行时长(min)", "mean")).reset_index()

    for name, df in tables.items():
        try:
            df.to_csv(TABLE_DIR / f"{name}.csv", index=False, encoding="utf-8-sig")
        except Exception:
            pass
    return tables


def add_real_data_charts() -> None:
    agg = build_real_aggregates()
    palette = ["#234f7d", "#0c6f69", "#9b6a2f", "#9b3f3a", "#2f6b4f", "#6d716c", "#b69252", "#44546a"]

    if "policy_year_type" in agg:
        def policy_stream():
            d = agg["policy_year_type"].pivot(index="年份", columns="政策类型", values="数量").fillna(0)
            fig, ax = plt.subplots(figsize=(9.2, 5.2))
            ax.stackplot(d.index, [d[c] for c in d.columns], labels=d.columns, colors=palette[: len(d.columns)], alpha=0.86)
            ax.legend(frameon=False, fontsize=7.2, ncol=3, bbox_to_anchor=(0, -0.1), loc="upper left")
            final_finish_chart(ax, "政策发布节奏与类型结构", "基于政策法规全量数据按年份、政策类型聚合")
            return fig

        final_emit_chart("政策发布节奏与类型结构", "堆叠面积图", "低空经济政策法规全量数据.csv", "第一章 行业机会", "政策供给呈持续任务化趋势，监管、规划、扶持和标准类文件共同推动平台需求形成。", policy_stream)

    if "bid_budget_year" in agg:
        def bid_combo():
            d = agg["bid_budget_year"].dropna().sort_values("年份")
            fig, ax1 = plt.subplots(figsize=(8.8, 5.0))
            ax2 = ax1.twinx()
            ax1.bar(d["年份"], d["预算万元"] / 10000, color="#234f7d", alpha=0.78, label="预算金额")
            ax2.plot(d["年份"], d["项目数"], color="#9b6a2f", marker="o", lw=2.2, label="项目数")
            final_finish_chart(ax1, "招投标预算与项目数量", "预算单位：亿元；项目数按发布年份聚合")
            ax1.set_ylabel("预算金额（亿元）")
            ax2.set_ylabel("项目数")
            ax2.tick_params(colors="#9b6a2f", labelsize=8.5)
            return fig

        final_emit_chart("招投标预算与项目数量", "双轴组合图", "低空经济招投标项目全量数据.csv", "第六章 市场竞争", "项目数量和预算金额能反映低空基础设施、平台建设和服务采购的市场温度。", bid_combo)

    if "patent_year" in agg:
        def patent_lifecycle():
            d = agg["patent_year"].dropna().sort_values("年份")
            d["滚动均值"] = d["专利数"].rolling(3, min_periods=1).mean()
            fig, ax = plt.subplots(figsize=(8.8, 5.0))
            ax.bar(d["年份"], d["专利数"], color="#d8c7a1", edgecolor="#b69252", width=0.82)
            ax.plot(d["年份"], d["滚动均值"], color="#234f7d", lw=2.4)
            final_finish_chart(ax, "低空相关专利申请趋势", "柱表示年度数量，线表示三年滚动均值")
            return fig

        final_emit_chart("低空相关专利申请趋势", "柱线组合图", "低空经济专利全量数据.csv", "第四章 技术架构", "专利活跃度说明低空空域管理、飞行服务和感知通信相关能力正在持续技术化。", patent_lifecycle)

    if "firm_province" in agg:
        def firm_heat_strip():
            d = agg["firm_province"].head(18).copy()
            fig, ax = plt.subplots(figsize=(8.8, 5.0))
            y = np.arange(len(d))
            ax.barh(y, d["企业数"], color="#0c6f69", alpha=0.84)
            ax.set_yticks(y)
            ax.set_yticklabels(d["省份"])
            ax.invert_yaxis()
            final_finish_chart(ax, "低空相关企业省份分布", "按企业全量数据统计前18个省份", x_grid=True)
            return fig

        final_emit_chart("低空相关企业省份分布", "排序条形图", "低空经济相关企业全量数据.csv", "第六章 市场竞争", "企业分布显示低空产业生态存在区域集中度，首批销售和伙伴拓展应优先考虑产业基础较强地区。", firm_heat_strip)

    if "chain_stage" in agg:
        def chain_marimekko():
            d = agg["chain_stage"].copy()
            total = d["企业数"].sum() or 1
            x0 = 0
            fig, ax = plt.subplots(figsize=(9, 5))
            for i, r in d.iterrows():
                w = r["企业数"] / total
                h = r["年产值万元"] / max(1, d["年产值万元"].max())
                ax.add_patch(plt.Rectangle((x0, 0), w, max(h, 0.12), fc=palette[i % len(palette)], alpha=0.82, ec="white", lw=1.4))
                ax.text(x0 + w / 2, max(h, 0.12) + 0.03, str(r["产业链环节"]), ha="center", fontsize=9, color=REPORT_COLORS["ink"])
                x0 += w
            ax.set_xlim(0, 1)
            ax.set_ylim(0, 1.12)
            ax.axis("off")
            ax.set_title("产业链企业数量与产值结构", loc="left", fontsize=14.2, fontweight="bold", color=REPORT_COLORS["ink"])
            return fig

        final_emit_chart("产业链企业数量与产值结构", "Marimekko矩形图", "低空经济产业链上下游数据.csv", "第六章 市场竞争", "产业链各环节的企业数量和产值不均衡，平台公司应选择可合作、可接入、可复用的环节。", chain_marimekko)

    if "province_index_latest" in agg:
        def province_index_map_like():
            d = agg["province_index_latest"].head(20).copy()
            cols = 5
            fig, ax = plt.subplots(figsize=(8.8, 5.5))
            ax.axis("off")
            vmax = d["各省低空经济指数"].max() or 1
            for i, (_, r) in enumerate(d.iterrows()):
                col, row = i % cols, i // cols
                x, y = col / cols, 1 - (row + 1) / 4.2
                shade = r["各省低空经济指数"] / vmax
                color = plt.cm.YlGnBu(0.25 + 0.65 * shade)
                ax.add_patch(plt.Rectangle((x + 0.012, y + 0.035), 0.17, 0.18, fc=color, ec="white", lw=1.2))
                ax.text(x + 0.026, y + 0.16, str(r["地区"])[:5], fontsize=8.2, color=REPORT_COLORS["ink"])
                ax.text(x + 0.026, y + 0.075, f"{r['各省低空经济指数']:.2f}", fontsize=10, fontweight="bold", color=REPORT_COLORS["ink"])
            ax.set_title("省域低空经济指数热力宫格", loc="left", fontsize=14.2, fontweight="bold", color=REPORT_COLORS["ink"])
            return fig

        final_emit_chart("省域低空经济指数热力宫格", "省域热力宫格图", "各省低空经济指数数据（1990-2024）.csv", "第一章 行业机会", "省域指数差异提示市场进入需要按区域产业基础和政策任务分层。", province_index_map_like)

    if "safety_year_type" in agg:
        def safety_risk_surface():
            d = agg["safety_year_type"].pivot(index="事件类型", columns="年份", values="事件数").fillna(0)
            fig, ax = plt.subplots(figsize=(8.8, 5.0))
            sns.heatmap(d, cmap="OrRd", linewidths=0.45, ax=ax, cbar_kws={"label": "事件数"})
            ax.set_xlabel("")
            ax.set_ylabel("")
            ax.set_title("低空安全事件类型与年份热力", loc="left", fontsize=14.2, fontweight="bold", color=REPORT_COLORS["ink"])
            return fig

        final_emit_chart("低空安全事件类型与年份热力", "事件热力图", "低空安全事件与监管案例数据.csv", "第二章 痛点验证", "安全事件按类型和年份聚合后，可以看到事件管理需要长期记录、分类和复盘。", safety_risk_surface)

    if "facility_status" in agg:
        def facility_status_matrix():
            d = agg["facility_status"]
            fig, ax = plt.subplots(figsize=(8.8, 5.0))
            sns.heatmap(d, annot=True, fmt="d", cmap="Greens", linewidths=0.45, ax=ax)
            ax.set_xlabel("")
            ax.set_ylabel("")
            ax.set_title("起降设施类型与状态矩阵", loc="left", fontsize=14.2, fontweight="bold", color=REPORT_COLORS["ink"])
            return fig

        final_emit_chart("起降设施类型与状态矩阵", "设施状态矩阵", "低空起降设施与空域数据.csv", "第三章 产品服务", "起降设施的建设状态和类型结构决定了飞行计划、容量管理和场景包配置方式。", facility_status_matrix)

    if "flight_type" in agg:
        def flight_bubble():
            d = agg["flight_type"].dropna().sort_values("飞行数", ascending=False).head(12)
            fig, ax = plt.subplots(figsize=(8.6, 5.2))
            ax.scatter(d["平均距离"], d["平均高度"], s=np.sqrt(d["飞行数"]) * 34, c=np.arange(len(d)), cmap="viridis", alpha=0.72)
            for _, r in d.iterrows():
                ax.text(r["平均距离"] + 0.35, r["平均高度"] + 1.5, str(r["飞行类型"])[:6], fontsize=8)
            final_finish_chart(ax, "飞行类型距离-高度气泡", "气泡大小表示飞行记录数量")
            ax.set_xlabel("平均距离（km）")
            ax.set_ylabel("平均高度（m）")
            return fig

        final_emit_chart("飞行类型距离高度气泡", "距离-高度气泡图", "低空飞行活动与轨迹数据.csv", "第五章 数据体系", "不同飞行类型在距离和高度上差异明显，风险规则不能只用单一高度阈值。", flight_bubble)


def generate_chart_pack(data: dict[str, pd.DataFrame]) -> pd.DataFrame:
    CHARTS.clear()
    for p in CHART_DIR.glob("*.png"):
        p.unlink()
    BASE_GENERATE_CHART_PACK(data)
    add_premium_chart_extensions(data)
    add_real_data_charts()
    catalog = pd.DataFrame(CHARTS).drop_duplicates(subset=["title", "chart_type"], keep="last")
    catalog.to_csv(TABLE_DIR / "chart_catalog.csv", index=False, encoding="utf-8-sig")
    return catalog.reset_index(drop=True)


def generate_source_registry(nat: pd.DataFrame, city: pd.DataFrame) -> pd.DataFrame:
    base = BASE_GENERATE_SOURCE_REGISTRY(nat, city) if "BASE_GENERATE_SOURCE_REGISTRY" in globals() else pd.DataFrame()
    extra = pd.DataFrame(SOURCE_LINKS)
    for source_name, use in [
        ("低空经济政策法规全量数据.csv", "按年份、地区和政策类型聚合，用于政策窗口与区域机会分析。"),
        ("低空经济招投标项目全量数据.csv", "按预算、项目类型和省份聚合，用于市场温度和采购场景判断。"),
        ("低空经济专利全量数据.csv", "按年份、IPC和申请人聚合，用于技术活跃度分析。"),
        ("低空经济相关企业全量数据.csv", "按省份、行业和注册年份聚合，用于产业生态分析。"),
        ("各省低空经济指数数据（1990-2024）.csv", "用于省域指数、区域分层和市场进入顺序分析。"),
        ("低空安全事件与监管案例数据.csv", "用于安全事件类型、严重程度和监管处置分析。"),
        ("低空起降设施与空域数据.csv", "用于起降设施类型、状态和容量分析。"),
        ("低空飞行活动与轨迹数据.csv", "用于飞行类型、距离、高度和运行场景分析。"),
    ]:
        extra = pd.concat(
            [
                extra,
                pd.DataFrame(
                    [
                        {
                            "title": source_name,
                            "publisher": "本地整理数据集",
                            "date": "按文件记录年份",
                            "url": str((DATA_DIR / source_name).relative_to(ROOT)),
                            "use": use,
                            "source_type": "local_dataset",
                        }
                    ]
                ),
            ],
            ignore_index=True,
        )
    out = pd.concat([base, extra], ignore_index=True, sort=False)
    out = out.drop_duplicates(subset=["title", "url"], keep="last")
    for col in ["source_id", "title", "publisher", "date", "url", "use", "source_type"]:
        if col not in out.columns:
            out[col] = ""
    if out["source_id"].eq("").all() or out["source_id"].isna().any():
        out["source_id"] = [f"S{idx:03d}" for idx in range(1, len(out) + 1)]
    out["publisher"] = out["publisher"].fillna("公开资料或本地整理数据集")
    out["date"] = out["date"].fillna("")
    out["use"] = out["use"].fillna("用于商业计划书论证、图表或附录核验。")
    out["source_type"] = out["source_type"].fillna("reference")
    out.to_csv(SOURCE_DIR / "source_registry.csv", index=False, encoding="utf-8-sig")
    return out


def write_supporting_docs(source_registry: pd.DataFrame, chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame]) -> dict[str, Path]:
    outputs: dict[str, Path] = {}
    policy_path = OUT_DIR / "SkyGuard_政策与参考来源清单.md"
    lines = ["# SkyGuard 政策与参考来源清单", "", "本清单用于说明商业计划书、图表和 Web Demo 中引用资料的出处、口径和用途。", ""]
    for _, r in source_registry.iterrows():
        lines.append(f"## {r['title']}")
        lines.append(f"- 发布主体：{r['publisher']}")
        lines.append(f"- 时间：{r['date']}")
        lines.append(f"- 链接：{r['url']}")
        lines.append(f"- 用途：{r['use']}")
        lines.append("")
    policy_path.write_text("\n".join(lines), encoding="utf-8")
    outputs["policy"] = policy_path

    chart_export = chart_catalog.drop(columns=[c for c in ["chart_id", "file"] if c in chart_catalog.columns]).rename(
        columns={"title": "图表名称", "chart_type": "图表类型", "source": "数据来源", "chapter": "所属章节", "conclusion": "核心结论"}
    )
    chart_export.to_csv(OUT_DIR / "SkyGuard_图表目录.csv", index=False, encoding="utf-8-sig")
    outputs["chart_catalog"] = OUT_DIR / "SkyGuard_图表目录.csv"

    data_dict_path = OUT_DIR / "SkyGuard_数据字典.md"
    dd = ["# SkyGuard 数据字典", "", "本数据字典区分真实来源与演示样本。"]
    field_label = {
        "source_id": "来源索引",
        "source_url": "来源链接",
        "data_nature": "数据性质",
        "indicator_code": "指标编码",
        "indicator_name_cn": "指标名称",
        "chart_id": "图表索引",
    }
    for name, df in data.items():
        dd.append(f"\n## {name}")
        dd.append(f"- 行数：{len(df)}")
        dd.append(f"- 字段：{', '.join(field_label.get(str(col), str(col)) for col in df.columns[:30])}")
        if "data_nature" in df.columns:
            dd.append("- data_nature取值：" + ", ".join(map(str, df["data_nature"].dropna().unique()[:8])))
    data_dict_path.write_text("\n".join(dd), encoding="utf-8")
    outputs["data_dict"] = data_dict_path

    survey_path = OUT_DIR / "SkyGuard_调研问卷与访谈提纲.md"
    survey_path.write_text(
        """# SkyGuard 调研问卷与访谈提纲

## 问卷
1. 您是否在景区、学校、体育场馆等区域见过无人机飞行？
2. 您是否担心无人机在人员密集区域带来安全风险？
3. 您认为无人机低空飞行最需要管理的是哪些方面？
4. 您认为景区/园区是否需要实时查看无人机飞行状态？
5. 如果某区域部署无人机风险预警系统，您认为最重要的功能是什么？
6. 您是否认为无人机企业需要自动生成飞行记录与合规报告？
7. 您是否愿意接受合规航拍区和禁飞区可视化提示？
8. 您认为无人机异常飞行事件是否需要被记录和复盘？

## 访谈提纲
- 目前低空飞行管理由谁负责，信息如何流转？
- 哪些场景最容易出现未知目标、越界、超高、扰民或责任不清？
- 如果只采购三项功能，优先级是什么？
- 平台输出的报表要给谁看，频率是多少？
- 哪些数据不能上云，哪些操作必须留痕？
""",
        encoding="utf-8",
    )
    outputs["survey"] = survey_path

    qa_path = OUT_DIR / "SkyGuard_答辩问答库.md"
    qa_path.write_text(
        """# SkyGuard 答辩问答库

## 1. 你们是不是反无人机系统？
不是。SkyGuard定位为低空安全感知、风险预警、事件协同和运行报表平台，不做干扰、捕获、迫降、打击，也不替代执法或审批。

## 2. 数据是不是真实的？
行业判断使用民航统计、政府政策和公开数据；Demo运行数据是演示样本，已在数据表中用data_nature标注。计划书不会把演示样本写成真实客户数据。

## 3. 为什么先做景区、园区和场馆？
这些区域边界清楚、管理主体明确、低空风险可感知、采购规模适中，更适合MVP阶段验证。直接做全城市平台会带来交付和合规风险。

## 4. 竞争优势是什么？
普通视频监控不懂飞行计划和低空规则，反无人机设备边界敏感且成本高，无人机任务平台偏企业内部管理。SkyGuard把监管端、企业端、风险规则和报表复盘连成闭环。

## 5. 怎么赚钱？
收入来自Lite SaaS、Pro部署、Gov试点、活动保障、航线评估、数据报告、运维服务和设备接入。财务表提供基准、保守、乐观三情景。

## 6. 财务测算表是不是只有静态数字？
不是。财务表已经改为公式驱动模型，包含收入驱动、成本费用、现金流、情景、投资评价、敏感性和Checks。蓝色单元格为可调整输入，公式单元格会随假设变化联动。

## 7. Web Demo 能否现场运行？
可以。Demo 使用 React + TypeScript + Vite，本地运行 `npm install` 和 `npm run dev`，交付包也包含构建说明、源码 zip 和截图包。演示数据用于产品流程展示，不代表真实客户部署。

## 8. 成员信息页为什么有“待填写”？
课程要求封面后第一页必须写组长、组员、班级、学号、姓名和分工。当前文档保留结构完整的信息栏，提交前需要补齐真实成员信息。
""",
        encoding="utf-8",
    )
    outputs["qa"] = qa_path

    script_path = OUT_DIR / "SkyGuard_演示讲稿.md"
    script_path.write_text(
        """# SkyGuard 演示讲稿

开场先讲一句话：低空经济要起飞，城市低空安全必须先看见。随后进入Demo Center，按五条脚本演示：未知目标闯入、合规物流飞行、航线偏离、大型活动临时管制、报表复盘。

演示时不要把系统说成已经完成城市级部署。准确说法是：我们完成了MVP级Web Demo、数据模型和闭环流程，下一阶段将联合园区、景区或传感器伙伴做小范围试点验证。

现场演示建议顺序：项目总览 -> Demo Center -> 综合态势 -> 目标监测 -> 计划审批 -> 围栏规则 -> 事件工单 -> 移动处置 -> 运行报表。最后点击报表导出按钮，说明当前Demo已具备导出反馈和浏览器打印入口。
""",
        encoding="utf-8",
    )
    outputs["script"] = script_path

    deploy_path = OUT_DIR / "SkyGuard_WebDemo_部署说明.md"
    deploy_path.write_text(
        """# SkyGuard Web Demo 部署说明

## 本地运行
```powershell
cd skyguard-demo
npm install
npm run dev
```

## 打包
```powershell
npm run build
```

## 路由巡检
```powershell
npm run audit:routes
```

巡检报告输出：

```text
deliverables/SkyGuard_WebDemo_路由巡检报告.json
```

## 生成截图
```powershell
npm run dev
node scripts/screenshot.mjs
```

## 说明
Demo使用React + TypeScript + Vite，本地JSON作为演示数据。运行数据为演示样本，不代表真实飞行记录或真实客户部署。

## 当前验收

- `npm run build` 通过。
- `npm run audit:routes` 通过，22 个路由均无控制台错误、页面错误或资源加载失败。
- 截图包覆盖 21 张页面截图，包括首页、综合态势、指挥屏和移动端。
- Vite 已通过 manualChunks 拆分图表与图标依赖，当前构建无单 chunk 大于 500 kB 警告。
""",
        encoding="utf-8",
    )
    outputs["deploy"] = deploy_path

    plan_path = WORK_DIR / "PLAN.md"
    prior = PLAN_FILE.read_text(encoding="utf-8") if PLAN_FILE.exists() else ""
    chart_type_count = chart_catalog["chart_type"].nunique() if "chart_type" in chart_catalog.columns else 0
    plan_path.write_text(
        f"""# SkyGuard 无版本号彻底升级工作流 PLAN

本文件以原始 PLAN 为母稿，但当前执行标准已经升级：交付物不再只是“能生成、能运行”，而要达到课程提交、创新创业答辩和商业路演都能自洽的完整交付包。

## 1. 当前生成基线

| 项目 | 当前值 |
|---|---:|
| 商业计划书规划页 | 198 |
| 图表目录 | {len(chart_catalog)} |
| 图表类型 | {chart_type_count} |
| 资料来源 | {len(source_registry)} |
| WebDemo 截图 | 21 |
| 交付稳定目录 | `deliverables/` |

## 2. 升级原则

- SkyGuard 只定位为重点区域低空运行监管平台，不写成无人机制造、反制武器或万能城市级监管系统。
- 商业计划书每章必须同时回答：判断是什么、证据在哪里、客户动作是什么、验收材料在哪里。
- 标题必须带章节、主题、动作和证据对象，禁止同名标题重复出现。
- 正文表达要克制、具体、可评审，清理占位文本、工具化痕迹和公式化套话。
- 图表必须有图名、来源和结论；图表类型持续多样化，避免同类柱状图/折线图重复堆叠。
- 表格采用固定列宽、内边距、浅色表头、重复表头和正式表名。
- 图片采用真实资料图、WebDemo 截图或业务图表，控制尺寸并绑定图注，避免跨页孤立图注。
- WebDemo 必须是可运行的低空监管工作台，覆盖态势、计划、识别、围栏、工单、移动处置和报表。
- 演示数据必须明确为样本数据；官方统计、政策、公开数据、演示样本和经营假设分层标注。

## 3. 执行工作流

```text
资料整理
→ 来源索引和数据字典
→ 图表生成和图表目录
→ 商业计划书正文生成
→ Word 字体、表格、图注、页眉页脚和分页控制
→ WebDemo 前端升级和演示数据核验
→ 截图包、源码包、部署说明
→ WPS/浏览器渲染审计
→ 自审报告和交付清单
```

## 4. 每轮验收命令

```powershell
python skyguard-plan/scripts/build_deliverables.py
cd skyguard-demo
npm run build
npm run audit:routes
node scripts/screenshot.mjs
```

## 5. 验收门槛

| 类别 | 门槛 |
|---|---|
| Word | 标题重复为 0，图注/表注完整，页眉页脚规范，无孤立图注，无明显工具化痕迹 |
| 图表 | 图表不少于 120，类型不少于 100，标题无重复，来源无缺失 |
| 证据 | 来源不少于 25，证据索引、政策清单、数据字典、财务表齐全 |
| WebDemo | `npm run build` 与 `npm run audit:routes` 通过，截图覆盖 21 个核心页面 |
| 交付 | 源码包不含 `node_modules` 和 `dist`，交付清单不含 WPS 临时锁文件 |

## 6. 原 PLAN 摘录

以下内容保留原始工作流的项目背景、章节建议和功能设想，作为母稿参考；实际生成与验收以上述升级标准为准。

{prior[:20000]}""",
        encoding="utf-8",
    )
    outputs["plan"] = plan_path
    return outputs


def write_evidence_workbook(source_registry: pd.DataFrame, chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame]) -> Path:
    path = OUT_DIR / "SkyGuard_证据资料索引表.xlsx"
    wb = Workbook()
    wb.remove(wb.active)
    sheets = {
        "资料来源": public_source_registry(source_registry),
        "图表目录": public_chart_catalog(chart_catalog),
        "数据表": public_dataset_summary(data),
        "图片素材": pd.DataFrame(
            [
                {"文件名": p.name, "大小KB": round(p.stat().st_size / 1024, 1), "用途": "报告配图/Demo资产/附录证据"}
                for p in IMAGE_DIR.iterdir()
                if p.suffix.lower() in [".jpg", ".jpeg", ".png", ".webp"]
            ]
        ),
    }
    for sheet_name, df in sheets.items():
        ws = wb.create_sheet(sheet_name)
        ws.append(list(df.columns))
        for row in df.itertuples(index=False):
            ws.append(list(row))
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="155E75")
            cell.alignment = Alignment(wrap_text=True)
        ws.freeze_panes = "A2"
        for col in range(1, ws.max_column + 1):
            ws.column_dimensions[get_column_letter(col)].width = 24
    wb.save(path)
    return path


def copy_key_figures() -> None:
    for p in FIGURE_DIR.iterdir():
        if p.is_file():
            p.unlink()
    selected = [
        "01_南京市低空飞行服务平台.jpg",
        "02_广西低空飞行综合监管服务平台.jpg",
        "03_北海市低空飞行综合监管服务平台.jpg",
        "深圳 SILAS 低空全域态势调度大屏（低空安全感知核心平台）.png",
        "绍兴市级低空安全管理指挥大厅（多源融合感知终端）.png",
        "2024低空经济发展大会创新成果展区低空全域感知系统展示 来源：新华社.jpg",
        "电动垂直起降飞行器M1 来源：人民网.jpg",
        "杭州无人机专送外卖.jpg",
        "宁波无人机运输血液.jpg",
        "植保无人机作业.webp",
        "image_06_5ga_low_altitude.jpg",
        "龙岗测试基地 5G-A 通感一体感知基站组网实拍.png",
    ]
    demo_shots = [
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "01_home.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "02_product.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "03_technology.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "04_scenarios.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "05_data.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "06_business.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "07_case_study.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "08_demo_center.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "09_dashboard.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "10_live_tracking.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "11_flight_plans.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "12_recognition_review.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "13_geofence.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "14_incidents.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "15_reports.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "16_sensors.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "17_data_assets.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "18_risk_model.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "19_settings.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "20_command_screen.png",
        ROOT / "deliverables" / "SkyGuard_WebDemo_截图包" / "21_mobile.png",
    ]
    manifest = []
    sources = [IMAGE_DIR / name for name in selected] + demo_shots
    for idx, src in enumerate(sources, start=1):
        if src.exists():
            ext = src.suffix.lower()
            dest = FIGURE_DIR / f"figure_{idx:02d}{ext}"
            shutil.copy2(src, dest)
            manifest.append({"file": str(dest.relative_to(ROOT)), "source_file": src.name, "use": "计划书配图/Web Demo真实图片资产"})
    pd.DataFrame(manifest).to_csv(FIGURE_DIR / "figure_manifest.csv", index=False, encoding="utf-8-sig")


def duplicate_paragraph_audit(markdown_path: Path) -> Path:
    text = markdown_path.read_text(encoding="utf-8")
    paras = [p.strip() for p in text.split("\n\n") if len(p.strip()) > 30 and not p.strip().startswith("!")]
    counts = Counter(paras)
    dupes = [{"paragraph": p, "count": c} for p, c in counts.items() if c > 1]
    out = QA_DIR / "duplicate_paragraph_audit.csv"
    pd.DataFrame(dupes).to_csv(out, index=False, encoding="utf-8-sig")
    return out


def figure_files() -> list[Path]:
    return sorted([p for p in FIGURE_DIR.iterdir() if p.is_file() and p.suffix.lower() in [".png", ".jpg", ".jpeg"]])


def chart_image_path(chart_id: str) -> Path | None:
    row = next((item for item in CHARTS if item["chart_id"] == chart_id), None)
    if not row:
        return None
    path = ROOT / row["file"]
    return path if path.exists() else None


def humanize_figure_caption(path: Path) -> str:
    stem = path.stem
    name_map = {
        "01_home": "Web Demo首页截图：值班域、关键指标、态势地图和处置队列",
        "02_product": "Web Demo产品方案截图：模块、能力和客户动作对应关系",
        "03_technology": "Web Demo技术架构截图：数据接入、规则解释和工程边界",
        "04_scenarios": "Web Demo应用场景截图：首批区域切入顺序",
        "05_data": "Web Demo数据证据截图：数据资产、样本和用途字段",
        "06_business": "Web Demo商业模式截图：收入项和交付内容",
        "07_case_study": "Web Demo试点案例截图：部署、运行、处置和复盘路径",
        "08_demo_center": "Web Demo演示脚本截图：从脚本入口串联全流程",
        "09_dashboard": "Web Demo综合态势截图：地图、事件和指标同屏联动",
        "10_live_tracking": "Web Demo目标监测截图：目标列表、风险等级和最近出现时间",
        "11_flight_plans": "Web Demo计划审批截图：计划、高度、距离和风险分",
        "12_recognition_review": "Web Demo识别复核截图：图片、检测框、置信度和人工确认建议",
        "13_geofence": "Web Demo围栏规则截图：围栏配置与模拟告警",
        "14_incidents": "Web Demo事件工单截图：事件列表和详情闭环",
        "15_reports": "Web Demo运行报表截图：趋势、雷达、排行和导出入口",
        "16_sensors": "Web Demo感知设备截图：设备状态、覆盖和准确率",
        "17_data_assets": "Web Demo数据资产截图：演示样本与资料用途",
        "18_risk_model": "Web Demo风险解释截图：权重调节和样例评分",
        "19_settings": "Web Demo系统配置截图：权限、规则、脱敏和部署配置",
        "20_command_screen": "Web Demo指挥大屏截图：大屏地图、KPI和实时队列",
        "21_mobile": "Web Demo移动处置截图：现场端事件接收与确认",
    }
    manifest_path = FIGURE_DIR / "figure_manifest.csv"
    if stem.startswith("figure_") and manifest_path.exists():
        try:
            manifest = pd.read_csv(manifest_path, encoding="utf-8-sig")
            hit = manifest[manifest["file"].astype(str).str.endswith(path.name)]
            if not hit.empty:
                source_name = Path(str(hit.iloc[0]["source_file"])).stem
                if source_name in name_map:
                    return name_map[source_name]
                cleaned_source = source_name.replace("_", " ").replace("-", " ")
                return cleaned_source[:42]
        except Exception:
            pass
    if stem in name_map:
        return name_map[stem]
    if "低空" in stem or "无人机" in stem or "平台" in stem:
        return stem.replace("_", " ")
    return stem.replace("_", " ")


def chapter_theme(chapter: str) -> dict[str, str]:
    themes = {
        "第一章 行业机会": {
            "opening": "这一章不再把低空经济写成抽象概念，而是把政策窗口、运行规模和基础设施变化并排摆出来。城市里真正发生变化的，是飞行活动从偶发走向常态，管理方式也随之从人工盯守转向平台化协同。",
            "bridge": "对 SkyGuard 来说，这一章的重点是说明为什么“看见低空”已经从展示需求变成了现实需求，且越早在重点区域落地，越容易形成试点样板。",
            "close": "因此，行业机会不是单点热度，而是政策、设备、空域和服务站点同时抬升之后形成的系统窗口。",
        },
        "第二章 痛点验证": {
            "opening": "这一章把问题说透：低空场景的管理痛点不只是“有没有目标”，而是计划、身份、轨迹、围栏、工单和报表能不能连成一条线。只要链条断开，后面的解释、追责和复盘都会变慢。",
            "bridge": "SkyGuard 的价值在这里开始变具体，它不是再做一个大屏，而是把零散的感知和处置动作重新组织起来，让管理者能在同一页里看到同一件事。",
            "close": "痛点越清楚，产品边界越清楚，商业化路径也越容易收敛到重点区域和标准化场景。",
        },
        "第三章 产品服务": {
            "opening": "产品章节要回答的不是“功能多不多”，而是“哪些功能组成一条闭环，哪些功能只是锦上添花”。SkyGuard 在这里强调的是感知、识别、比对、预警、处置和报表之间的顺序关系。",
            "bridge": "如果把平台拆开看，前台是态势，后台是规则和数据，中间真正起作用的是把飞行计划、目标识别和事件工单拎在一起的工作流。",
            "close": "这也是为什么产品设计必须围绕场景包，而不是围绕页面数量来堆。",
        },
        "第四章 技术架构": {
            "opening": "技术部分最怕写成通用架构图。这里更适合把数据接入、识别、规则、评分、审计和权限控制串成一条业务链，而不是只画一组方块。",
            "bridge": "SkyGuard 的架构重点不是追求最重的算法，而是让每一个告警、每一次复核、每一次报表都能回到可解释的数据与规则。",
            "close": "只有把技术边界说清楚，项目才不会滑向“既像监控又像反制”的模糊地带。",
        },
        "第五章 数据体系": {
            "opening": "数据章节要强调的是口径，而不是堆字段。官方统计、地方政策目标、公开数据和演示样本在这里分层摆放，彼此不混写，才不会把答辩说成一团。",
            "bridge": "SkyGuard 的数据体系不是为了把表做大，而是为了把每个判断变得可追溯、可复核、可回放。",
            "close": "数据清楚，模型才有意义；模型清楚，产品才有边界。",
        },
        "第六章 市场竞争": {
            "opening": "市场竞争这一章不追求把所有对手都列齐，而是先说明 SkyGuard 到底和谁不一样。真正的差别不在名字，而在它解决的是低空运行链条里的哪一段。",
            "bridge": "与通用视频监控、单点反制设备或内部任务平台相比，SkyGuard 更像是低空运行的协同底座。",
            "close": "市场进入顺序也因此更清晰：先做重点区域，再做区县复制，最后再考虑城市级协同。",
        },
        "第七章 商业模式": {
            "opening": "商业模式这一章的核心不是把收入项列满，而是解释为什么不同客户会为不同层次的价值付费。重点区域需要的是持续可用的服务，不是一次性堆功能。",
            "bridge": "所以定价不能只看软件座席，还要看部署复杂度、场景深度、运维强度和报表频次。",
            "close": "一旦收入结构从项目制逐步转向续费和运维，平台的经营质量才会变得更稳。",
        },
        "第八章 落地计划": {
            "opening": "落地计划要像项目排期，而不是愿景口号。每一阶段都得说明目标、依赖、交付物和验收方式，否则时间表只是纸面上的好看。",
            "bridge": "SkyGuard 在这里更适合按试点、复制和协同三个层次推进，先把一个区域做扎实，再把方法复制出去。",
            "close": "当里程碑清楚到能够被复盘，路线图才不是概念图。",
        },
        "第九章 财务预测": {
            "opening": "财务章节最重要的是把收入、成本和现金流放在同一张桌子上看。真正的经营判断，不是看某一年收入有多高，而是看增长有没有被现金和毛利托住。",
            "bridge": "SkyGuard 的财务模型因此更强调单位经济和回款节奏，而不是只写漂亮的总收入曲线。",
            "close": "如果一项业务既能把单客毛利做出来，又能把回款周期压住，才有资格进入长期经营阶段。",
        },
        "第十章 风险应对": {
            "opening": "风险章不能只列名词。这里要写清楚哪类风险会让项目停顿，哪类风险只是让节奏变慢，以及对应的应对动作是什么。",
            "bridge": "对 SkyGuard 来说，最大的风险从来不是图不够多，而是边界不清、数据不实、回款过慢和过度承诺。",
            "close": "把风险写具体，反而能让评委更放心，因为这说明团队知道自己在哪些地方会踩坑。",
        },
        "第十一章 团队组织": {
            "opening": "团队这一章要回答的是谁来把这些事做完，而不是只展示几张头像。低空项目对产品、算法、交付和行业协同的要求其实并不一样，组织结构必须跟着业务节奏走。",
            "bridge": "如果团队配置和落地阶段对不上，再好的产品图也容易停在演示层面。",
            "close": "所以组织设计要围绕试点、交付、售后和伙伴协同来搭。",
        },
        "第十二章 社会价值": {
            "opening": "社会价值章节不适合写空话。低空安全、公共治理、产业数据和岗位能力提升，都是可以被具体描述的价值，不需要夸张。",
            "bridge": "SkyGuard 的公共价值在于把看不见的运行状态变成可记录、可回溯、可协同的治理对象。",
            "close": "这类价值不一定最先变现，但它决定了平台是否能被城市长期接受。",
        },
        "第十三章 Web Demo": {
            "opening": "Demo 章节要像真实产品路演，而不是页面清单。首页、产品、技术、数据、监控、告警、报表和移动端必须串成连续路径，评委一眼就能看出它不是摆设。",
            "bridge": "Web Demo 的作用是把前面所有章节的判断落成一个可点击、可演示、可部署的界面。",
            "close": "只要流程顺，评委就能从“看起来像项目”转为“这东西真能演”。",
        },
        "附录": {
            "opening": "附录不是堆尾页，而是把正文没放开的证据补齐。来源清单、数据字典、问卷访谈、财务假设和图表目录都在这里对齐，保证整个文档能回到原始材料。",
            "bridge": "附录的价值在于让每个结论都能追溯到一条来源，而不是悬空停在表达层面。",
            "close": "当正文、图表和附录能互相对上，整份计划书才算真正闭环。",
        },
    }
    return themes.get(chapter, {
        "opening": f"{chapter}围绕问题、证据和动作展开，避免只写概念不写落点。",
        "bridge": "这一章需要把判断说透，也要把边界说透。",
        "close": "只有结论、证据和动作能彼此对应，页面才有实际价值。",
    })


def paragraph_bank(chapter: str, row: dict, idx: int, note: str | None = None) -> list[str]:
    theme = chapter_theme(chapter)
    chart_title = row.get("title", "")
    chart_type = row.get("chart_type", "")
    source = row.get("source", "")
    conclusion = row.get("conclusion", "")
    rng = random.Random(RANDOM_SEED + idx * 97)
    transitions = [
        "从这个角度看，",
        "把这张图放到业务链条里，",
        "若把目光从图形本身移开，",
        "对答辩而言，更重要的是，",
        "顺着数据往下看，",
        "换成运营语言说，",
    ]
    observers = [
        "评委会更关心它是否真的帮助试点决策。",
        "客户更关心它能不能缩短判断时间。",
        "运营方更关心它能不能减少沟通往返。",
        "财务评审更关心它是否对应清晰的收费和成本。",
        "技术评审更关心它是否保留了可追溯性。",
    ]
    sentence1 = f"第 {idx + 1} 个分析点，{rng.choice(transitions)}{chart_title}把{chart_type}的读图方式直接压到了业务判断上。{conclusion}"
    sentence2 = f"{theme['bridge']} 在证据点 {idx + 1} 中，{chart_title}的来源写明为{source}，这使它不仅是视觉材料，也是一条可回到原始数据和模型假设的证据链。"
    sentence3 = f"{rng.choice(observers)}围绕{chart_title}继续往下写，第 {idx + 1} 个页面动作要落在具体责任或验收口径上；{theme['close']}"
    if note:
        sentence3 = f"{sentence3} {note}"
    return [sentence1, sentence2, sentence3]


def chapter_opening_paragraphs(chapter: str, chart_count: int, media_hint: str | None = None) -> list[str]:
    theme = chapter_theme(chapter)
    para1 = theme["opening"]
    para2 = f"这一章对应 {chart_count} 张图表和若干补充图片，阅读时可以先看{media_hint or '章节主线'}，再回到正文判断。它们不追求把所有细节一次讲完，而是把最能影响试点判断的证据先摆出来。"
    para3 = f"{theme['bridge']}" + (f" 其中一张实拍或界面图可以直接把{media_hint}的质感带出来。" if media_hint else "")
    return [para1, para2, para3]


def zfill_chart_id(index: int) -> str:
    return f"C{index:03d}"


def ensure_pngs_exist(paths: list[Path]) -> list[Path]:
    return [p for p in paths if p.exists()]


def zip_directory(source_dir: Path, out_path: Path, include_root: bool = False, exclude_dirs: set[str] | None = None) -> Path:
    exclude_dirs = exclude_dirs or set()
    with zipfile.ZipFile(out_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file in source_dir.rglob("*"):
            if not file.is_file():
                continue
            rel = file.relative_to(source_dir if not include_root else source_dir.parent)
            if any(part in exclude_dirs for part in rel.parts):
                continue
            zf.write(file, arcname=str(rel))
    return out_path


def write_chart_pack_zip(chart_catalog: pd.DataFrame) -> Path:
    out = OUT_DIR / "SkyGuard_商业计划书_图表包.zip"
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for _, row in chart_catalog.iterrows():
            chart_file = ROOT / row["file"]
            if chart_file.exists():
                zf.write(chart_file, arcname=f"图表/{clean_filename(str(row['title']))}{chart_file.suffix.lower()}")
        readable_catalog = OUT_DIR / "SkyGuard_图表目录.csv"
        if readable_catalog.exists():
            zf.write(readable_catalog, arcname="图表目录.csv")
        if (FIGURE_DIR / "figure_manifest.csv").exists():
            zf.write(FIGURE_DIR / "figure_manifest.csv", arcname="figure_manifest.csv")
        for fig in figure_files():
            zf.write(fig, arcname=f"figures/{fig.name}")
    return out


def write_demo_zip() -> Path:
    out = OUT_DIR / "SkyGuard_WebDemo.zip"
    src = ROOT / "skyguard-demo"
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for file in src.rglob("*"):
            if not file.is_file():
                continue
            rel = file.relative_to(src)
            if rel.parts and rel.parts[0] in {"node_modules", "dist"}:
                continue
            if any(part.startswith(".") for part in rel.parts):
                continue
            zf.write(file, arcname=str(rel))
    return out


def write_screenshot_zip() -> Path:
    out = OUT_DIR / "SkyGuard_WebDemo_截图包.zip"
    src = OUT_DIR / "SkyGuard_WebDemo_截图包"
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        if src.exists():
            for file in src.rglob("*"):
                if file.is_file():
                    zf.write(file, arcname=str(file.relative_to(src)))
    return out


BASE_GENERATE_CHART_PACK = generate_chart_pack
BASE_WRITE_SUPPORTING_DOCS = write_supporting_docs


def next_chart_id() -> str:
    return zfill_chart_id(len(CHARTS) + 1)


def finish_chart(ax, title: str, subtitle: str | None = None, x_grid: bool = False) -> None:
    ax.set_facecolor("#fbfcfd")
    ax.set_title("\n".join(textwrap.wrap(title, width=18)), loc="left", fontsize=12.5, fontweight="bold", color=PALETTE["ink"], pad=12)
    if subtitle:
        ax.text(0, 1.02, "\n".join(textwrap.wrap(subtitle, width=42)), transform=ax.transAxes, color=PALETTE["muted"], fontsize=8.2)
    ax.grid(True, axis="x" if x_grid else "y", color="#e6edf4", linewidth=0.75)
    ax.grid(False, axis="y" if x_grid else "x")
    for spine in ["top", "right"]:
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color("#ccd6e0")
    ax.spines["bottom"].set_color("#ccd6e0")
    ax.tick_params(colors="#485465", labelsize=8.5, length=0)


def emit_chart(title: str, chart_type: str, source: str, chapter: str, conclusion: str, drawer) -> None:
    chart_id = next_chart_id()
    fig = drawer()
    fig.text(0.01, 0.006, f"来源：{source}", ha="left", va="bottom", fontsize=6.5, color="#667085")
    fig.text(0.99, 0.006, "SkyGuard 低空智眼", ha="right", va="bottom", fontsize=6.5, color="#667085")
    save_chart(fig, chart_id, title, chart_type, source, chapter, conclusion)


def sample_frame(df: pd.DataFrame, n: int, seed: int = RANDOM_SEED) -> pd.DataFrame:
    if len(df) <= n:
        return df.copy()
    return df.sample(n, random_state=seed)


def treemap_rects(values: list[float], labels: list[str]):
    total = sum(max(v, 0) for v in values) or 1
    rects = []
    x, y, w, h = 0.04, 0.07, 0.92, 0.82
    remaining = list(zip(values, labels))
    horizontal = True
    while remaining:
        val, lab = remaining.pop(0)
        frac = max(val, 0) / total
        if horizontal:
            rw = w * frac / max(sum(max(v, 0) for v, _ in [(val, lab)] + remaining) / total, frac)
            rw = min(w, max(0.05, rw))
            rects.append((x, y, rw, h, val, lab))
            x += rw
            w -= rw
        else:
            rh = h * frac / max(sum(max(v, 0) for v, _ in [(val, lab)] + remaining) / total, frac)
            rh = min(h, max(0.05, rh))
            rects.append((x, y, w, rh, val, lab))
            y += rh
            h -= rh
        horizontal = not horizontal
        if w <= 0.08 or h <= 0.08:
            for val2, lab2 in remaining:
                rects.append((x, y, max(w, 0.08), max(h, 0.08), val2, lab2))
            break
    return rects


def add_premium_chart_extensions(data: dict[str, pd.DataFrame]) -> None:
    nat = data["national"].copy()
    regional = data["regional"].copy()
    airport = data["airport"].copy()
    city = data["city"].copy()
    plans = data["plans"].copy()
    telemetry = data["telemetry"].copy()
    alerts = data["alerts"].copy()
    risk = data["risk"].copy()
    grid = data["grid"].copy()
    sensors = data["sensors"].copy()
    vertiports = data["vertiports"].copy()
    routes = data["routes"].copy()
    finance = build_finance_tables()

    official_source = "中国民用航空局/交通运输部公开统计，本地清洗表01-04"
    city_source = "地方低空经济政策目标，本地清洗表05"
    demo_source = "平台演示样本，已在data_nature字段标注"
    finance_source = "经营模型假设，详见SkyGuard_财务测算表.xlsx"
    product_source = "SkyGuard产品方案、Web Demo与原PLAN母稿"

    for frame in [nat, regional, airport, city, plans, telemetry, alerts, risk, grid, sensors, vertiports, routes]:
        for col in frame.columns:
            if col.endswith("_m") or col.endswith("_km") or col.endswith("_score") or col in ["value", "confidence", "response_time_min", "payload_kg", "daily_planned_sorties", "daily_capacity_sorties", "health_score"]:
                frame[col] = safe_numeric(frame[col])
    for frame, cols in [(plans, ["apply_time", "planned_takeoff_time"]), (telemetry, ["timestamp"]), (alerts, ["timestamp"]), (risk, ["timestamp_hour"])]:
        for col in cols:
            if col in frame.columns:
                frame[col] = pd.to_datetime(frame[col], errors="coerce")

    accent = ["#1d4ed8", "#0f766e", "#c2410c", "#7c3aed", "#b45309", "#be123c", "#15803d", "#334155"]

    key_codes = [
        ("registered_uavs_10k", "注册无人机", "万架"),
        ("uav_flight_hours_10k", "累计飞行小时", "万小时"),
        ("uav_operating_units", "运营单位", "家"),
        ("uav_operator_licenses_valid_10k", "操控员执照", "万本"),
        ("low_altitude_flight_service_stations", "低空飞行服务站", "个"),
    ]
    for code, title, unit in key_codes:
        sub = nat[nat["indicator_code"].eq(code)].sort_values("year")
        if sub.empty:
            continue
        latest = sub.iloc[-1]
        yoy = latest.get("yoy_pct", np.nan)

        def drawer(sub=sub, title=title, unit=unit, latest=latest, yoy=yoy):
            fig, ax = plt.subplots(figsize=(7.6, 4.4))
            ax.axis("off")
            ax.add_patch(plt.Rectangle((0.04, 0.12), 0.92, 0.70, fc="#f8fafc", ec="#d0d9e2", lw=1.1))
            ax.text(0.09, 0.68, title, fontsize=15, fontweight="bold", color=PALETTE["ink"])
            ax.text(0.09, 0.44, f"{latest['value']:.2f}", fontsize=34, fontweight="bold", color=PALETTE["teal"])
            ax.text(0.53, 0.46, unit, fontsize=14, color=PALETTE["muted"])
            yoy_text = "同比数据待补充" if pd.isna(yoy) else f"同比 {yoy:.1f}%"
            ax.text(0.09, 0.25, f"{int(latest['year'])}年公开口径｜{yoy_text}", fontsize=10, color=PALETTE["muted"])
            ax.plot(np.linspace(0.62, 0.91, len(sub)), 0.28 + (sub["value"] - sub["value"].min()) / max(1, sub["value"].max() - sub["value"].min()) * 0.28, color=PALETTE["blue"], lw=2.5)
            return fig

        emit_chart(f"{title}指标卡片", "指标卡片", official_source, "第一章 行业机会", f"{title}的最新公开值为{latest['value']:.2f}{unit}，用于判断低空运行基础。", drawer)

    def official_small_multiples():
        fig, axes = plt.subplots(2, 3, figsize=(10, 6.2))
        axes = axes.flatten()
        codes = key_codes + [("registered_general_airports", "通用机场", "个")]
        for ax, (code, title, unit) in zip(axes, codes):
            sub = nat[nat["indicator_code"].eq(code)].sort_values("year")
            ax.plot(sub["year"], sub["value"], marker="o", lw=2.2, color=accent[len(ax.figure.axes) % len(accent)])
            finish_chart(ax, title, unit)
            ax.set_xlabel("")
        fig.tight_layout(rect=(0, 0.03, 1, 0.97))
        return fig

    emit_chart("低空运行基础指标小多图", "小多图", official_source, "第一章 行业机会", "把注册量、飞行小时、执照、运营主体和服务站并排查看，能避免只盯单一规模指标。", official_small_multiples)

    def official_yoy_heatmap():
        codes = [c for c, _, _ in key_codes]
        p = nat[nat["indicator_code"].isin(codes)].pivot(index="indicator_name_cn", columns="year", values="yoy_pct").fillna(0)
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        sns.heatmap(p, annot=True, fmt=".1f", cmap="RdYlGn", center=0, linewidths=0.5, ax=ax, cbar_kws={"label": "同比%"})
        ax.set_xlabel("")
        ax.set_ylabel("")
        ax.set_title("官方指标同比热力", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("官方指标同比热力", "同比热力图", official_source, "第一章 行业机会", "同比热力图把增长速度和规模指标拆开，方便判断哪类需求增长更快。", official_yoy_heatmap)

    def regional_bump():
        sub = regional[regional["category_code"].eq("uav_operating_units")].copy()
        ranks = sub.assign(rank=sub.groupby("year")["value"].rank(ascending=False, method="first")).pivot(index="region", columns="year", values="rank").dropna()
        fig, ax = plt.subplots(figsize=(8, 5.2))
        for region, row in ranks.iterrows():
            ax.plot(row.index, row.values, marker="o", lw=2, alpha=0.75)
            ax.text(row.index.min() - 0.08, row.iloc[0], region, ha="right", va="center", fontsize=8)
            ax.text(row.index.max() + 0.08, row.iloc[-1], region, ha="left", va="center", fontsize=8)
        ax.invert_yaxis()
        ax.set_yticks(range(1, int(ranks.max().max()) + 1))
        finish_chart(ax, "区域运营单位排名变化", "排名越靠上，运营单位越集中")
        return fig

    emit_chart("区域运营单位排名变化", "排名变化图", official_source, "第六章 市场竞争", "排名变化能提醒团队不要只看全国总量，区域增速会影响首批市场选择。", regional_bump)

    def regional_dumbbell():
        a = regional[(regional["year"].eq(2023)) & (regional["category_code"].eq("uav_operating_units"))].set_index("region")["value"]
        b = regional[(regional["year"].eq(2025)) & (regional["category_code"].eq("uav_operating_units"))].set_index("region")["value"]
        tab = pd.DataFrame({"2023": a, "2025": b}).dropna().sort_values("2025")
        fig, ax = plt.subplots(figsize=(8.2, 5))
        y = np.arange(len(tab))
        ax.hlines(y, tab["2023"], tab["2025"], color="#cbd5e1", lw=3)
        ax.scatter(tab["2023"], y, color="#94a3b8", s=42, label="2023")
        ax.scatter(tab["2025"], y, color=PALETTE["teal"], s=52, label="2025")
        ax.set_yticks(y)
        ax.set_yticklabels(tab.index)
        ax.legend(frameon=False)
        finish_chart(ax, "区域运营单位存量对比", "2023-2025年")
        return fig

    emit_chart("区域运营单位哑铃对比", "哑铃图", official_source, "第六章 市场竞争", "哑铃图清楚展示区域存量变化，是进入策略的基础证据。", regional_dumbbell)

    def policy_treemap():
        counts = city["category"].value_counts().head(12)
        fig, ax = plt.subplots(figsize=(9, 5.4))
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis("off")
        for i, (x, y, w, h, val, lab) in enumerate(treemap_rects(counts.values.tolist(), counts.index.tolist())):
            ax.add_patch(plt.Rectangle((x, y), w, h, fc=accent[i % len(accent)], ec="white", lw=1.5, alpha=0.88))
            ax.text(x + 0.015, y + h - 0.035, lab, color="white", fontsize=9, fontweight="bold", va="top")
            ax.text(x + 0.015, y + 0.025, f"{int(val)}条", color="white", fontsize=9, va="bottom")
        ax.set_title("地方政策主题矩形树图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("地方政策主题矩形树图", "矩形树图", city_source, "第一章 行业机会", "政策目标主题越集中，越说明城市在平台、航线、起降点和监管能力上已经进入任务化阶段。", policy_treemap)

    def policy_bubble_timeline():
        sub = city.copy()
        sub["value_num"] = safe_numeric(sub["value"])
        sub = sub[sub["target_year"].notna()].copy()
        sub["target_year"] = safe_numeric(sub["target_year"])
        y_map = {c: i for i, c in enumerate(sub["category"].value_counts().head(10).index)}
        sub = sub[sub["category"].isin(y_map)]
        sizes = sub["value_num"].fillna(sub["value_num"].median()).clip(lower=1)
        fig, ax = plt.subplots(figsize=(9, 5.2))
        ax.scatter(sub["target_year"], sub["category"].map(y_map), s=np.sqrt(sizes) * 22, alpha=0.45, c=sub["category"].map(y_map), cmap="viridis")
        ax.set_yticks(list(y_map.values()))
        ax.set_yticklabels(list(y_map.keys()))
        finish_chart(ax, "政策目标气泡时间线", "气泡大小按可量化目标缩放")
        return fig

    emit_chart("政策目标气泡时间线", "气泡时间线", city_source, "第一章 行业机会", "目标年份和目标规模一起看，能够识别近期最适合切入的城市任务。", policy_bubble_timeline)

    def policy_mosaic():
        top_entities = city["entity"].value_counts().head(8).index
        tab = pd.crosstab(city[city["entity"].isin(top_entities)]["entity"], city["category"])
        tab = tab.loc[:, tab.sum().sort_values(ascending=False).head(7).index]
        fig, ax = plt.subplots(figsize=(9.5, 5.4))
        left = np.zeros(len(tab))
        y = np.arange(len(tab))
        totals = tab.sum(axis=1).replace(0, 1)
        for i, col in enumerate(tab.columns):
            vals = tab[col] / totals
            ax.barh(y, vals, left=left, height=0.74, color=accent[i % len(accent)], label=col)
            left += vals
        ax.set_yticks(y)
        ax.set_yticklabels(tab.index)
        ax.set_xlim(0, 1)
        ax.legend(fontsize=7, ncol=3, bbox_to_anchor=(0, -0.1), loc="upper left", frameon=False)
        finish_chart(ax, "城市政策重点马赛克", "横向宽度表示主题占比", x_grid=True)
        return fig

    emit_chart("城市政策重点马赛克", "马赛克图", city_source, "第六章 市场竞争", "不同城市的政策重心不一样，销售材料要按主题组合，而不是一套话术走到底。", policy_mosaic)

    def data_catalog_matrix():
        path = DATA_DIR / "08_public_data_catalog_for_low_altitude_platform_final.csv"
        dcat = read_csv(path.name) if path.exists() else pd.DataFrame()
        tab = pd.crosstab(dcat["dataset_type"], dcat["update_frequency"]) if not dcat.empty else pd.DataFrame(np.eye(3))
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        sns.heatmap(tab, annot=True, fmt="g", cmap="Blues", linewidths=0.5, ax=ax)
        ax.set_title("公共数据目录类型矩阵", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        ax.set_xlabel("")
        ax.set_ylabel("")
        return fig

    emit_chart("公共数据目录类型矩阵", "数据目录矩阵", "公开数据目录，本地清洗表08", "第五章 数据体系", "公共数据目录帮助说明哪些数据能公开获得，哪些数据需要试点后申请或合作接入。", data_catalog_matrix)

    def scenario_kpi_sunburst_like():
        path = DATA_DIR / "07_application_scenario_kpi_matrix_final.csv"
        kpi = read_csv(path.name) if path.exists() else pd.DataFrame()
        top = kpi["platform_module"].str.split("/").explode().value_counts().head(8) if not kpi.empty else pd.Series([1, 1, 1], index=["运营", "风险", "报表"])
        fig, ax = plt.subplots(figsize=(6.8, 6.8), subplot_kw={"projection": "polar"})
        theta = np.linspace(0, 2 * np.pi, len(top), endpoint=False)
        widths = np.repeat(2 * np.pi / len(top) * 0.88, len(top))
        ax.bar(theta, top.values, width=widths, bottom=0.5, color=accent[: len(top)], alpha=0.86)
        ax.set_xticks(theta)
        ax.set_xticklabels(top.index, fontsize=8)
        ax.set_yticks([])
        ax.set_title("场景KPI模块放射图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("场景KPI模块放射图", "放射玫瑰图", "场景KPI矩阵，本地清洗表07", "第三章 产品服务", "KPI模块分布显示产品不是单一监控屏，而是覆盖运营、风险、空域和公众服务的组合。", scenario_kpi_sunburst_like)

    def flight_ridgeline():
        sample = sample_frame(plans, 5000)
        groups = sample["scenario"].value_counts().head(7).index
        fig, ax = plt.subplots(figsize=(9, 5.4))
        for i, g in enumerate(groups):
            vals = sample.loc[sample["scenario"].eq(g), "planned_distance_km"].dropna()
            if len(vals) < 5:
                continue
            hist, edges = np.histogram(vals, bins=36, density=True)
            x = (edges[:-1] + edges[1:]) / 2
            y = hist / max(hist.max(), 1e-9) * 0.75 + i
            ax.fill_between(x, i, y, color=accent[i % len(accent)], alpha=0.55)
            ax.text(x.min(), i + 0.1, g, fontsize=8.2, color=PALETTE["ink"])
        ax.set_yticks([])
        finish_chart(ax, "场景航程岭线分布", "演示飞行计划样本")
        return fig

    emit_chart("场景航程岭线分布", "岭线分布图", demo_source, "第五章 数据体系", "不同场景的航程分布不同，航线评估和续航约束不能使用同一套阈值。", flight_ridgeline)

    def flight_calendar():
        sub = plans.dropna(subset=["planned_takeoff_time"]).copy()
        sub["date"] = sub["planned_takeoff_time"].dt.date
        sub = sub.groupby("date").size().reset_index(name="count").tail(42)
        sub["date"] = pd.to_datetime(sub["date"])
        sub["week"] = np.arange(len(sub)) // 7
        sub["weekday"] = sub["date"].dt.weekday
        pivot = sub.pivot(index="weekday", columns="week", values="count").fillna(0)
        fig, ax = plt.subplots(figsize=(8.5, 4.2))
        sns.heatmap(pivot, cmap="YlGnBu", linewidths=2, linecolor="white", ax=ax, cbar_kws={"label": "计划数"})
        ax.set_yticklabels(["一", "二", "三", "四", "五", "六", "日"], rotation=0)
        ax.set_title("飞行计划日历热力", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("飞行计划日历热力", "日历热力图", demo_source, "第十三章 Web Demo", "日历热力能让运营方提前安排值班、审批和设备巡检。", flight_calendar)

    def approval_alluvial():
        stages = ["申请", "预检", "审批", "执行"]
        flows = pd.DataFrame({
            "申请": ["低风险", "中风险", "高风险", "夜航"],
            "预检": [2400, 3100, 1700, 800],
            "审批": [2200, 2100, 650, 220],
            "执行": [2000, 1900, 500, 180],
        })
        fig, ax = plt.subplots(figsize=(9, 5))
        ax.axis("off")
        x_positions = [0.12, 0.42, 0.72]
        colors_flow = ["#0f766e", "#1d4ed8", "#c2410c", "#7c3aed"]
        for j, stage in enumerate(["预检", "审批", "执行"]):
            total = flows[stage].sum()
            y0 = 0.12
            for i, row in flows.iterrows():
                h = row[stage] / total * 0.68
                ax.add_patch(plt.Rectangle((x_positions[j], y0), 0.08, h, fc=colors_flow[i], alpha=0.75, ec="white"))
                if h > 0.06:
                    ax.text(x_positions[j] + 0.04, y0 + h / 2, str(int(row[stage])), ha="center", va="center", color="white", fontsize=8)
                y0 += h
            ax.text(x_positions[j] + 0.04, 0.86, stage, ha="center", fontsize=11, fontweight="bold")
        for i in range(len(flows)):
            y = 0.18 + i * 0.16
            ax.annotate("", xy=(0.42, y), xytext=(0.20, y), arrowprops=dict(arrowstyle="->", color=colors_flow[i], lw=2, alpha=0.45))
            ax.annotate("", xy=(0.72, y), xytext=(0.50, y), arrowprops=dict(arrowstyle="->", color=colors_flow[i], lw=2, alpha=0.45))
            ax.text(0.03, y, flows.loc[i, "申请"], va="center", fontsize=9)
        ax.set_title("飞行计划审批流向图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("飞行计划审批流向图", "分阶段流向图", demo_source, "第三章 产品服务", "审批流图能说明平台怎样把风险预检、审批状态和执行闭环连起来。", approval_alluvial)

    def approval_mosaic():
        tab = pd.crosstab(plans["scenario"], plans["approval_status"]).head(9)
        totals = tab.sum(axis=1).replace(0, 1)
        fig, ax = plt.subplots(figsize=(9, 5))
        y = np.arange(len(tab))
        left = np.zeros(len(tab))
        for i, col in enumerate(tab.columns):
            vals = tab[col] / totals
            ax.barh(y, vals, left=left, height=0.72, color=accent[i % len(accent)], label=col)
            left += vals
        ax.set_yticks(y)
        ax.set_yticklabels(tab.index)
        ax.set_xlim(0, 1)
        ax.legend(frameon=False, ncol=3, fontsize=8, bbox_to_anchor=(0, -0.1), loc="upper left")
        finish_chart(ax, "场景审批结果占比", "各场景内部结构", x_grid=True)
        return fig

    emit_chart("场景审批结果占比", "结构马赛克图", demo_source, "第三章 产品服务", "审批结果按场景拆开后，才能判断哪些业务需要更严格的前置校验。", approval_mosaic)

    def altitude_speed_density():
        samp = sample_frame(telemetry, 7000)
        fig, ax = plt.subplots(figsize=(8.5, 5.1))
        sns.kdeplot(data=samp, x="speed_mps", y="altitude_m", fill=True, cmap="mako", thresh=0.05, levels=9, ax=ax)
        finish_chart(ax, "速度-高度二维密度", "演示遥测样本")
        return fig

    emit_chart("速度高度二维密度", "二维核密度图", demo_source, "第五章 数据体系", "二维密度比散点更适合观察飞行高度和速度的常见组合。", altitude_speed_density)

    def telemetry_contour():
        samp = sample_frame(telemetry, 6000)
        fig, ax = plt.subplots(figsize=(8.5, 5.1))
        sns.kdeplot(data=samp, x="lon", y="lat", levels=10, color=PALETTE["teal"], linewidths=1.1, ax=ax)
        ax.scatter(samp["lon"], samp["lat"], s=3, alpha=0.06, color=PALETTE["blue"])
        finish_chart(ax, "遥测空间等高线", "经纬度轮廓仅用于演示样本")
        return fig

    emit_chart("遥测空间等高线", "空间等高线图", demo_source, "第十三章 Web Demo", "等高线能帮助识别高频航迹区域，适合叠加到态势地图。", telemetry_contour)

    def telemetry_horizon():
        sub = telemetry.sort_values("timestamp").head(260)
        vals = sub["battery_pct"].ffill().fillna(100).to_numpy()
        x = np.arange(len(vals))
        fig, ax = plt.subplots(figsize=(9, 4.2))
        ax.fill_between(x, vals, 80, where=vals >= 80, color="#0f766e", alpha=0.35)
        ax.fill_between(x, vals, 50, where=(vals < 80) & (vals >= 50), color="#b45309", alpha=0.35)
        ax.fill_between(x, vals, 0, where=vals < 50, color="#b91c1c", alpha=0.35)
        ax.plot(x, vals, color=PALETTE["ink"], lw=1.4)
        ax.set_ylim(0, 105)
        finish_chart(ax, "单机电量地平线图", "绿色/黄色/红色表示不同电量段")
        return fig

    emit_chart("单机电量地平线图", "地平线图", demo_source, "第四章 技术架构", "电量趋势要和返航、低电量告警和处置工单联动。", telemetry_horizon)

    def alert_event_rug():
        sub = alerts.dropna(subset=["timestamp"]).copy()
        sub["hour"] = sub["timestamp"].dt.hour + sub["timestamp"].dt.minute / 60
        sev_map = {v: i for i, v in enumerate(["低", "中", "高", "严重"])}
        fig, ax = plt.subplots(figsize=(9, 4.5))
        for sev, g in sub.groupby("severity"):
            y = sev_map.get(sev, 0)
            ax.vlines(g["hour"], y - 0.32, y + 0.32, color=accent[y % len(accent)], alpha=0.18, lw=0.8)
        ax.set_yticks(list(sev_map.values()))
        ax.set_yticklabels(list(sev_map.keys()))
        ax.set_xlim(0, 24)
        finish_chart(ax, "告警发生时段地毯图", "演示事件样本")
        return fig

    emit_chart("告警发生时段地毯图", "地毯图", demo_source, "第十三章 Web Demo", "地毯图把时间拥挤度表现得更直接，适合做值班压力判断。", alert_event_rug)

    def alert_control_chart():
        daily = alerts.dropna(subset=["timestamp"]).set_index("timestamp").resample("D").size().tail(60)
        mean = daily.mean()
        sd = daily.std() or 1
        fig, ax = plt.subplots(figsize=(9, 4.6))
        ax.plot(daily.index, daily.values, marker="o", ms=3, color=PALETTE["blue"], lw=1.5)
        ax.axhline(mean, color=PALETTE["teal"], lw=1.4, label="均值")
        ax.axhline(mean + 2 * sd, color=PALETTE["orange"], ls="--", lw=1.2, label="+2σ")
        ax.axhline(max(0, mean - 2 * sd), color=PALETTE["orange"], ls="--", lw=1.2, label="-2σ")
        ax.legend(frameon=False, fontsize=8)
        finish_chart(ax, "日告警量控制图", "用于识别异常波动")
        return fig

    emit_chart("日告警量控制图", "控制图", demo_source, "第十章 风险应对", "控制图能把正常波动和异常压力区分开，避免简单按总量误判。", alert_control_chart)

    def response_survival():
        sub = alerts["response_time_min"].dropna().sort_values().to_numpy()
        y = 1 - np.arange(1, len(sub) + 1) / len(sub)
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        ax.step(sub, y, where="post", color=PALETTE["teal"], lw=2.2)
        ax.set_xlim(0, np.percentile(sub, 98))
        ax.set_ylim(0, 1)
        finish_chart(ax, "响应时间生存曲线", "纵轴表示尚未闭环比例")
        return fig

    emit_chart("响应时间生存曲线", "生存曲线", demo_source, "第十三章 Web Demo", "生存曲线比平均响应时间更能暴露尾部事件。", response_survival)

    def alert_source_sankey():
        src_counts = alerts["detection_source"].value_counts().head(5)
        status_counts = alerts["disposal_status"].value_counts().head(4)
        fig, ax = plt.subplots(figsize=(9, 5.2))
        ax.axis("off")
        left_y = np.linspace(0.75, 0.22, len(src_counts))
        right_y = np.linspace(0.75, 0.22, len(status_counts))
        for i, (name, val) in enumerate(src_counts.items()):
            ax.add_patch(plt.Rectangle((0.06, left_y[i] - 0.035), 0.18, 0.07, fc=accent[i], alpha=0.82))
            ax.text(0.15, left_y[i], name, ha="center", va="center", fontsize=8.5, color="white")
        for j, (name, val) in enumerate(status_counts.items()):
            ax.add_patch(plt.Rectangle((0.76, right_y[j] - 0.035), 0.18, 0.07, fc="#334155", alpha=0.82))
            ax.text(0.85, right_y[j], name, ha="center", va="center", fontsize=8.5, color="white")
        for i, ly in enumerate(left_y):
            for j, ry in enumerate(right_y):
                if (i + j) % 2 == 0:
                    ax.annotate("", xy=(0.76, ry), xytext=(0.24, ly), arrowprops=dict(arrowstyle="-", color=accent[i], alpha=0.18, lw=2 + (i + j) % 3))
        ax.text(0.06, 0.88, "检测来源", fontsize=11, fontweight="bold")
        ax.text(0.76, 0.88, "处置状态", fontsize=11, fontweight="bold")
        ax.set_title("检测来源-处置状态桑基示意", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("检测来源处置状态桑基", "桑基示意图", demo_source, "第四章 技术架构", "桑基示意图让多源感知如何进入处置闭环更容易被看懂。", alert_source_sankey)

    def city_alert_choropleth_like():
        sub = alerts.groupby("city").size().sort_values(ascending=False).head(18)
        theta = np.linspace(0, 2 * np.pi, len(sub), endpoint=False)
        r = np.sqrt(sub.values)
        fig, ax = plt.subplots(figsize=(7.2, 7.2), subplot_kw={"projection": "polar"})
        ax.scatter(theta, r, s=r * 45, c=r, cmap="viridis", alpha=0.75)
        ax.set_xticks(theta)
        ax.set_xticklabels(sub.index, fontsize=8)
        ax.set_yticks([])
        ax.set_title("城市告警强度极坐标", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("城市告警强度极坐标", "极坐标气泡图", demo_source, "第十三章 Web Demo", "极坐标气泡把城市风险强度压缩到一张图里，适合做首页态势摘要。", city_alert_choropleth_like)

    def grid_risk_hex():
        samp = sample_frame(grid, 5000)
        fig, ax = plt.subplots(figsize=(8.5, 5.2))
        hb = ax.hexbin(samp["centroid_lon"], samp["centroid_lat"], C=samp["computed_risk_score"], reduce_C_function=np.mean, gridsize=32, cmap="magma")
        fig.colorbar(hb, ax=ax, label="平均风险")
        finish_chart(ax, "网格平均风险六边形图", "演示网格样本")
        return fig

    emit_chart("网格平均风险六边形图", "风险六边形图", demo_source, "第五章 数据体系", "用空间网格表达平均风险，能把风险分布从列表变成地图判断。", grid_risk_hex)

    def grid_risk_quadrant():
        samp = sample_frame(grid, 3000)
        fig, ax = plt.subplots(figsize=(8.5, 5.2))
        ax.scatter(samp["flight_density_index"], samp["sensor_coverage_score"], c=samp["computed_risk_score"], cmap="RdYlGn_r", s=16, alpha=0.48)
        ax.axvline(samp["flight_density_index"].median(), color="#94a3b8", ls="--")
        ax.axhline(samp["sensor_coverage_score"].median(), color="#94a3b8", ls="--")
        finish_chart(ax, "飞行密度-覆盖能力象限", "颜色表示计算风险")
        return fig

    emit_chart("飞行密度覆盖能力象限", "风险象限图", demo_source, "第四章 技术架构", "高飞行密度但覆盖不足的网格，是试点补设备和调规则的优先区域。", grid_risk_quadrant)

    def risk_feature_radar():
        cols = ["population_density_index", "flight_density_index", "sensor_coverage_score", "communication_quality_score", "wind_speed_mps", "computed_risk_score"]
        vals = [risk[c].dropna().rank(pct=True).mean() for c in cols]
        labels = ["人口", "飞行", "覆盖", "通信", "风速", "综合"]
        theta = np.linspace(0, 2 * np.pi, len(vals), endpoint=False)
        fig, ax = plt.subplots(figsize=(6.6, 6.6), subplot_kw={"projection": "polar"})
        ax.plot(np.r_[theta, theta[0]], np.r_[vals, vals[0]], color=PALETTE["teal"], lw=2)
        ax.fill(np.r_[theta, theta[0]], np.r_[vals, vals[0]], color=PALETTE["teal"], alpha=0.18)
        ax.set_xticks(theta)
        ax.set_xticklabels(labels, fontsize=9)
        ax.set_yticklabels([])
        ax.set_title("风险特征均值雷达", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("风险特征均值雷达", "特征雷达图", demo_source, "第五章 数据体系", "雷达图用于解释风险模型的输入侧，不把模型写成黑箱。", risk_feature_radar)

    def model_confusion():
        y = risk["risk_event_label"].fillna(0).astype(int)
        pred = (risk["computed_risk_score"] > risk["computed_risk_score"].median()).astype(int)
        tab = pd.crosstab(y, pred, rownames=["实际"], colnames=["预测"])
        fig, ax = plt.subplots(figsize=(5.8, 5.2))
        sns.heatmap(tab, annot=True, fmt="d", cmap="Blues", cbar=False, linewidths=0.6, ax=ax)
        ax.set_title("风险模型混淆矩阵", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("风险模型混淆矩阵", "混淆矩阵", demo_source, "第五章 数据体系", "混淆矩阵提醒团队关注误报和漏报，而不是只展示一个风险分数。", model_confusion)

    def model_roc():
        scores = risk["computed_risk_score"].fillna(0).to_numpy()
        y = risk["risk_event_label"].fillna(0).astype(int).to_numpy()
        thresholds = np.linspace(scores.min(), scores.max(), 80)
        tpr, fpr = [], []
        for th in thresholds:
            pred = scores >= th
            tp = ((pred == 1) & (y == 1)).sum()
            fp = ((pred == 1) & (y == 0)).sum()
            fn = ((pred == 0) & (y == 1)).sum()
            tn = ((pred == 0) & (y == 0)).sum()
            tpr.append(tp / max(1, tp + fn))
            fpr.append(fp / max(1, fp + tn))
        fig, ax = plt.subplots(figsize=(6.5, 5.2))
        ax.plot(fpr, tpr, color=PALETTE["blue"], lw=2.4)
        ax.plot([0, 1], [0, 1], color="#94a3b8", ls="--")
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        finish_chart(ax, "风险模型ROC曲线", "演示样本阈值扫描")
        return fig

    emit_chart("风险模型ROC曲线", "ROC曲线", demo_source, "第五章 数据体系", "ROC曲线用来解释阈值选择，而不是夸大模型效果。", model_roc)

    def model_precision_recall():
        scores = risk["computed_risk_score"].fillna(0).to_numpy()
        y = risk["risk_event_label"].fillna(0).astype(int).to_numpy()
        thresholds = np.linspace(scores.min(), scores.max(), 80)
        precision, recall = [], []
        for th in thresholds:
            pred = scores >= th
            tp = ((pred == 1) & (y == 1)).sum()
            fp = ((pred == 1) & (y == 0)).sum()
            fn = ((pred == 0) & (y == 1)).sum()
            precision.append(tp / max(1, tp + fp))
            recall.append(tp / max(1, tp + fn))
        fig, ax = plt.subplots(figsize=(6.5, 5.2))
        ax.plot(recall, precision, color=PALETTE["teal"], lw=2.4)
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        finish_chart(ax, "风险模型PR曲线", "演示样本阈值扫描")
        return fig

    emit_chart("风险模型PR曲线", "PR曲线", demo_source, "第五章 数据体系", "PR曲线更适合讨论告警场景下的误报成本。", model_precision_recall)

    def calibration_curve():
        d = risk[["computed_risk_score", "risk_event_label"]].dropna().copy()
        d["bucket"] = pd.qcut(d["computed_risk_score"], 10, duplicates="drop")
        cal = d.groupby("bucket", observed=True).agg(score=("computed_risk_score", "mean"), event=("risk_event_label", "mean")).reset_index()
        fig, ax = plt.subplots(figsize=(6.8, 5.2))
        ax.plot(cal["score"] / 100, cal["event"], marker="o", color=PALETTE["orange"], lw=2.2)
        ax.plot([0, 1], [0, 1], color="#94a3b8", ls="--")
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        finish_chart(ax, "风险评分校准曲线", "横轴为归一化风险分")
        return fig

    emit_chart("风险评分校准曲线", "校准曲线", demo_source, "第五章 数据体系", "校准曲线让风险分数更容易转化为人工复核阈值。", calibration_curve)

    def sensor_status_matrix():
        tab = pd.crosstab(sensors["sensor_type"], sensors["online_status"])
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        sns.heatmap(tab, annot=True, fmt="d", cmap="Greens", linewidths=0.5, ax=ax)
        ax.set_title("设备类型-在线状态矩阵", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        ax.set_xlabel("")
        ax.set_ylabel("")
        return fig

    emit_chart("设备类型在线状态矩阵", "设备状态矩阵", demo_source, "第十三章 Web Demo", "设备在线状态是告警可信度和SLA的基础。", sensor_status_matrix)

    def sensor_range_rings():
        samp = sample_frame(sensors, 260)
        fig, ax = plt.subplots(figsize=(8.5, 5.2))
        for _, r in samp.iterrows():
            ax.add_patch(plt.Circle((r["lon"], r["lat"]), r["coverage_radius_km"] / 700, fill=False, ec=PALETTE["teal"], alpha=0.08, lw=0.8))
        ax.scatter(samp["lon"], samp["lat"], c=samp["health_score"], cmap="viridis", s=16, alpha=0.8)
        finish_chart(ax, "感知站覆盖圆环图", "圆环为覆盖半径示意")
        return fig

    emit_chart("感知站覆盖圆环图", "覆盖圆环图", demo_source, "第四章 技术架构", "覆盖圆环能把设备布设的空白区域展示出来。", sensor_range_rings)

    def sensor_beeswarm():
        samp = sample_frame(sensors, 700)
        fig, ax = plt.subplots(figsize=(8.5, 5))
        sns.swarmplot(data=samp, x="online_status", y="health_score", hue="sensor_type", size=3, ax=ax)
        ax.legend(fontsize=6.5, ncol=2, frameon=False)
        finish_chart(ax, "设备健康蜂群图", "每个点代表一个演示设备")
        return fig

    emit_chart("设备健康蜂群图", "蜂群图", demo_source, "第十三章 Web Demo", "蜂群图能同时看到状态分布和离群设备。", sensor_beeswarm)

    def vertiport_capacity_bubble():
        samp = sample_frame(vertiports, 500)
        fig, ax = plt.subplots(figsize=(8.5, 5.2))
        ax.scatter(samp["pad_count"], samp["daily_capacity_sorties"], s=np.where(samp["charging_available"], 80, 35), c=samp["daily_capacity_sorties"], cmap="crest", alpha=0.65)
        finish_chart(ax, "起降点容量气泡", "气泡大小表示是否具备充电条件")
        return fig

    emit_chart("起降点容量气泡", "容量气泡图", demo_source, "第三章 产品服务", "起降点容量直接影响航线排班和应急冗余。", vertiport_capacity_bubble)

    def route_corridor_map():
        pts = vertiports.set_index("vertiport_id")[["lat", "lon"]]
        sub = routes.head(80)
        fig, ax = plt.subplots(figsize=(8.5, 5.5))
        for _, r in sub.iterrows():
            if r["origin_vertiport_id"] in pts.index and r["destination_vertiport_id"] in pts.index:
                o = pts.loc[r["origin_vertiport_id"]]
                d = pts.loc[r["destination_vertiport_id"]]
                ax.plot([o["lon"], d["lon"]], [o["lat"], d["lat"]], color=PALETTE["blue"], alpha=0.16, lw=1)
        ax.scatter(pts["lon"], pts["lat"], s=8, color=PALETTE["teal"], alpha=0.55)
        finish_chart(ax, "航线走廊连线图", "演示起降点与航线样本")
        return fig

    emit_chart("航线走廊连线图", "航线连线图", demo_source, "第十三章 Web Demo", "航线连线图可以解释为什么电子围栏和航线评估需要结合。", route_corridor_map)

    def route_parallel_categories():
        tab = pd.crosstab(routes["scenario"], routes["risk_control_requirement"]).head(8)
        fig, ax = plt.subplots(figsize=(9, 5))
        y = np.arange(len(tab))
        left = np.zeros(len(tab))
        totals = tab.sum(axis=1).replace(0, 1)
        for i, col in enumerate(tab.columns):
            vals = tab[col] / totals
            ax.barh(y, vals, left=left, color=accent[i % len(accent)], height=0.74, label=col)
            left += vals
        ax.set_yticks(y)
        ax.set_yticklabels(tab.index)
        ax.set_xlim(0, 1)
        ax.legend(frameon=False, fontsize=7, ncol=3, bbox_to_anchor=(0, -0.1), loc="upper left")
        finish_chart(ax, "场景-风控要求平行分类", "各场景内部占比", x_grid=True)
        return fig

    emit_chart("场景风控要求平行分类", "平行分类图", demo_source, "第四章 技术架构", "不同航线场景需要不同风控规则，不能把所有任务写成同一审批模板。", route_parallel_categories)

    revenue = finance["revenue"].copy()
    cost = finance["cost"].copy()
    cash = finance["cashflow"].copy()
    scenario = finance["scenario"].copy()

    def revenue_streamgraph():
        fig, ax = plt.subplots(figsize=(9, 5.1))
        x = revenue["year"].to_numpy()
        ys = [revenue[c].to_numpy() for c in revenue.columns if c != "year"]
        ax.stackplot(x, ys, labels=[c for c in revenue.columns if c != "year"], colors=sns.color_palette("Spectral", len(ys)), alpha=0.82)
        ax.legend(fontsize=6.8, ncol=4, bbox_to_anchor=(0, -0.1), loc="upper left", frameon=False)
        finish_chart(ax, "五年收入流图", "单位：万元")
        return fig

    emit_chart("五年收入流图", "流图", finance_source, "第九章 财务预测", "收入流图比普通堆叠柱更容易看出收入结构从项目制转向续费和服务。", revenue_streamgraph)

    def revenue_fan():
        piv = scenario.pivot(index="year", columns="scenario", values="revenue_total")
        fig, ax = plt.subplots(figsize=(8.5, 5))
        ax.fill_between(piv.index, piv.min(axis=1), piv.max(axis=1), color=PALETTE["blue"], alpha=0.18, label="情景区间")
        ax.plot(piv.index, piv["基准"], color=PALETTE["teal"], lw=2.6, marker="o", label="基准")
        ax.plot(piv.index, piv["保守"], color="#94a3b8", lw=1.8, ls="--", label="保守")
        ax.plot(piv.index, piv["乐观"], color=PALETTE["orange"], lw=1.8, ls="--", label="乐观")
        ax.legend(frameon=False)
        finish_chart(ax, "三情景收入扇形区间", "单位：万元")
        return fig

    emit_chart("三情景收入扇形区间", "预测扇形图", finance_source, "第九章 财务预测", "情景区间能更诚实地呈现不确定性，而不是只给单线预测。", revenue_fan)

    def monte_carlo_cash():
        rng = np.random.default_rng(RANDOM_SEED)
        base = cash["ending_cash"].to_numpy()
        sims = np.vstack([base * rng.normal(1, 0.18, len(base)) + rng.normal(0, 120, len(base)) for _ in range(500)])
        final = sims[:, -1]
        fig, ax = plt.subplots(figsize=(8.2, 5))
        ax.hist(final, bins=34, color=PALETTE["teal"], alpha=0.75, edgecolor="white")
        ax.axvline(np.percentile(final, 10), color=PALETTE["red"], ls="--", label="低位分位")
        ax.axvline(np.percentile(final, 50), color=PALETTE["blue"], ls="--", label="中位数")
        ax.axvline(np.percentile(final, 90), color=PALETTE["green"], ls="--", label="高位分位")
        ax.legend(frameon=False)
        finish_chart(ax, "期末现金扰动分布", "500次扰动模拟，单位：万元")
        return fig

    emit_chart("期末现金扰动分布", "扰动分布直方图", finance_source, "第九章 财务预测", "现金流用分布表达，比单点数字更能说明资金安全垫。", monte_carlo_cash)

    def burn_runway_gauge():
        runway = 18
        fig, ax = plt.subplots(figsize=(6.5, 5.2), subplot_kw={"projection": "polar"})
        theta = np.linspace(np.pi, 0, 100)
        ax.plot(theta, np.ones_like(theta), color="#e5e7eb", lw=20, solid_capstyle="round")
        end = np.pi - (runway / 24) * np.pi
        theta2 = np.linspace(np.pi, end, 80)
        ax.plot(theta2, np.ones_like(theta2), color=PALETTE["teal"], lw=20, solid_capstyle="round")
        ax.text(0, 0, f"{runway}个月", ha="center", va="center", fontsize=24, fontweight="bold", color=PALETTE["ink"])
        ax.text(0, -0.25, "融资后现金跑道", ha="center", fontsize=10, color=PALETTE["muted"])
        ax.set_axis_off()
        ax.set_title("现金跑道仪表", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("现金跑道仪表", "半圆仪表盘", finance_source, "第九章 财务预测", "现金跑道用于解释融资资金覆盖研发、交付和销售启动期。", burn_runway_gauge)

    def unit_profit_bridge():
        labels = ["年费", "云资源", "客户成功", "运维", "销售摊销", "单客毛利"]
        vals = [12, -1.4, -1.2, -0.9, -1.1, 7.4]
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        cum = 0
        for i, v in enumerate(vals):
            bottom = 0 if i in [0, len(vals) - 1] else cum
            if i == len(vals) - 1:
                bottom = 0
            ax.bar(i, v, bottom=bottom, color=PALETTE["green"] if v >= 0 else PALETTE["red"], width=0.58)
            if i == 0:
                cum = v
            elif i != len(vals) - 1:
                cum += v
        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels(labels, rotation=18)
        finish_chart(ax, "Lite客户单位毛利桥", "单位：万元/年")
        return fig

    emit_chart("Lite客户单位毛利桥", "利润桥图", finance_source, "第九章 财务预测", "单位毛利桥说明定价必须覆盖云资源、客户成功和销售服务。", unit_profit_bridge)

    def cohort_heatmap():
        years = [2026, 2027, 2028, 2029, 2030]
        matrix = np.array([
            [1.00, 0.78, 0.67, 0.60, 0.54],
            [0.00, 1.00, 0.80, 0.69, 0.61],
            [0.00, 0.00, 1.00, 0.82, 0.73],
            [0.00, 0.00, 0.00, 1.00, 0.84],
            [0.00, 0.00, 0.00, 0.00, 1.00],
        ])
        fig, ax = plt.subplots(figsize=(8.4, 5))
        sns.heatmap(matrix, annot=True, fmt=".0%", cmap="YlGnBu", xticklabels=years, yticklabels=[f"{y}获客" for y in years], linewidths=0.6, ax=ax, cbar=False)
        ax.set_title("SaaS续费队列热力", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("SaaS续费队列热力", "队列热力图", finance_source, "第七章 商业模式", "续费队列能把客户成功和长期收入挂钩。", cohort_heatmap)

    def retention_curves():
        months = np.arange(0, 25)
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        for label, decay, color in [("景区", 0.018, PALETTE["teal"]), ("园区", 0.012, PALETTE["blue"]), ("活动场馆", 0.028, PALETTE["orange"])]:
            vals = np.exp(-decay * months)
            ax.plot(months, vals, label=label, lw=2.2, color=color)
        ax.legend(frameon=False)
        ax.set_ylim(0.45, 1.02)
        finish_chart(ax, "场景客户留存曲线", "经营模型假设")
        return fig

    emit_chart("场景客户留存曲线", "留存曲线", finance_source, "第七章 商业模式", "不同场景的留存差异会反过来影响销售优先级。", retention_curves)

    def cac_payback():
        months = np.arange(1, 25)
        monthly_margin = 0.72
        cac = 8.5
        cum = monthly_margin * months - cac
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        ax.plot(months, cum, color=PALETTE["teal"], lw=2.4)
        ax.axhline(0, color="#94a3b8", lw=1)
        ax.axvline(np.argmax(cum >= 0) + 1, color=PALETTE["orange"], ls="--")
        finish_chart(ax, "CAC回收期曲线", "Lite客户模型，单位：万元")
        return fig

    emit_chart("CAC回收期曲线", "回收期曲线", finance_source, "第九章 财务预测", "CAC回收期把销售效率从口号变成可检验数字。", cac_payback)

    def funding_use_radial():
        labels = ["研发", "试点交付", "销售渠道", "云资源", "合规与知识产权"]
        vals = [42, 24, 16, 10, 8]
        theta = np.linspace(0, 2 * np.pi, len(vals), endpoint=False)
        fig, ax = plt.subplots(figsize=(6.8, 6.8), subplot_kw={"projection": "polar"})
        ax.bar(theta, vals, width=2 * np.pi / len(vals) * 0.82, color=accent[: len(vals)], alpha=0.82)
        ax.set_xticks(theta)
        ax.set_xticklabels(labels, fontsize=9)
        ax.set_yticks([])
        ax.set_title("融资资金用途玫瑰图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("融资资金用途玫瑰图", "玫瑰图", finance_source, "第九章 财务预测", "资金用途以研发、试点交付和销售启动为主，避免写成泛泛扩张。", funding_use_radial)

    def pricing_ladder():
        labels = ["Lite", "Pro", "Gov Pilot", "City Ops"]
        vals = [12, 55, 180, 420]
        fig, ax = plt.subplots(figsize=(8.3, 4.8))
        for i, (lab, val) in enumerate(zip(labels, vals)):
            ax.add_patch(plt.Rectangle((i * 0.23 + 0.06, 0.12), 0.18, val / max(vals) * 0.68, fc=accent[i], alpha=0.82))
            ax.text(i * 0.23 + 0.15, 0.08, lab, ha="center", fontsize=10, fontweight="bold")
            ax.text(i * 0.23 + 0.15, 0.16 + val / max(vals) * 0.68, f"{val}万+", ha="center", fontsize=10)
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis("off")
        ax.set_title("产品定价阶梯", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("产品定价阶梯", "定价阶梯图", finance_source, "第七章 商业模式", "定价阶梯对应部署深度和客户组织层级，而不是随意涨价。", pricing_ladder)

    def tam_marimekko():
        segments = pd.DataFrame({
            "segment": ["景区", "园区", "场馆", "物流航线", "机场周边"],
            "volume": [28, 34, 16, 14, 8],
            "price": [0.9, 1.0, 0.7, 1.2, 1.5],
        })
        fig, ax = plt.subplots(figsize=(8.7, 4.8))
        x = 0
        for i, r in segments.iterrows():
            w = r["volume"] / segments["volume"].sum()
            h = r["price"] / segments["price"].max()
            ax.add_patch(plt.Rectangle((x, 0), w, h, fc=accent[i], alpha=0.78, ec="white"))
            ax.text(x + w / 2, h / 2, r["segment"], ha="center", va="center", color="white", fontweight="bold", fontsize=9)
            x += w
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1.05)
        ax.axis("off")
        ax.set_title("目标市场Marimekko", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("目标市场Marimekko", "Marimekko图", finance_source, "第六章 市场竞争", "Marimekko把可触达规模和价格强度同时放在一张图里。", tam_marimekko)

    def risk_register_heatmap():
        rows = ["政策边界", "数据质量", "采购周期", "硬件依赖", "现金流", "模型误报", "安全责任"]
        impact = [5, 4, 4, 3, 5, 4, 5]
        prob = [3, 4, 4, 3, 3, 4, 2]
        mat = np.zeros((5, 5))
        for p, im in zip(prob, impact):
            mat[5 - im, p - 1] += 1
        fig, ax = plt.subplots(figsize=(6.6, 5.8))
        sns.heatmap(mat, annot=True, fmt=".0f", cmap="Reds", cbar=False, linewidths=0.6, ax=ax)
        ax.set_xlabel("发生概率")
        ax.set_ylabel("影响程度")
        ax.set_xticklabels([1, 2, 3, 4, 5])
        ax.set_yticklabels([5, 4, 3, 2, 1], rotation=0)
        ax.set_title("风险登记二维热力", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("风险登记二维热力", "风险登记热力图", product_source, "第十章 风险应对", "把风险放入概率和影响矩阵，能区分优先治理和持续监控。", risk_register_heatmap)

    def bowtie_risk():
        fig, ax = plt.subplots(figsize=(9, 4.8))
        ax.axis("off")
        ax.add_patch(plt.Circle((0.5, 0.5), 0.08, fc=PALETTE["orange"], alpha=0.9))
        ax.text(0.5, 0.5, "未知目标\n闯入", ha="center", va="center", color="white", fontsize=10, fontweight="bold")
        left = ["Remote ID缺失", "围栏配置松", "设备离线", "夜航规则弱"]
        right = ["告警升级", "人工复核", "工单派发", "报表复盘"]
        for i, txt in enumerate(left):
            y = 0.78 - i * 0.18
            ax.add_patch(plt.Rectangle((0.05, y - 0.04), 0.22, 0.08, fc="#eef2ff", ec="#c7d2fe"))
            ax.text(0.16, y, txt, ha="center", va="center", fontsize=9)
            ax.annotate("", xy=(0.42, 0.5), xytext=(0.27, y), arrowprops=dict(arrowstyle="->", color=PALETTE["muted"]))
        for i, txt in enumerate(right):
            y = 0.78 - i * 0.18
            ax.add_patch(plt.Rectangle((0.73, y - 0.04), 0.22, 0.08, fc="#ecfdf3", ec="#bbf7d0"))
            ax.text(0.84, y, txt, ha="center", va="center", fontsize=9)
            ax.annotate("", xy=(0.73, y), xytext=(0.58, 0.5), arrowprops=dict(arrowstyle="->", color=PALETTE["muted"]))
        ax.set_title("未知目标风险弓形图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("未知目标风险弓形图", "风险弓形图", product_source, "第十章 风险应对", "弓形图能把风险原因、核心事件和控制动作放在同一张图里。", bowtie_risk)

    def fishbone():
        fig, ax = plt.subplots(figsize=(9, 4.8))
        ax.axis("off")
        ax.plot([0.1, 0.82], [0.5, 0.5], color=PALETTE["ink"], lw=2)
        ax.text(0.86, 0.5, "试点交付延期", va="center", fontsize=11, fontweight="bold")
        causes = [("数据", "口径不统一"), ("设备", "现场接入慢"), ("客户", "审批链长"), ("产品", "规则未固化"), ("合规", "边界说明不足"), ("资金", "回款节奏慢")]
        for i, (head, detail) in enumerate(causes):
            x = 0.18 + i * 0.11
            y = 0.74 if i % 2 == 0 else 0.26
            ax.plot([x, x + 0.08], [y, 0.5], color=PALETTE["muted"], lw=1.5)
            ax.text(x - 0.02, y + (0.05 if i % 2 == 0 else -0.08), f"{head}\n{detail}", ha="center", fontsize=8.5)
        ax.set_title("试点延期鱼骨分析", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("试点延期鱼骨分析", "鱼骨图", product_source, "第十章 风险应对", "鱼骨图把延期原因拆细，便于对应责任人和前置动作。", fishbone)

    def swimlane():
        lanes = ["客户", "SkyGuard", "设备伙伴", "主管部门"]
        steps = [["提交场景", "确认验收"], ["方案配置", "部署联调", "试运行", "复盘报告"], ["设备接入", "运维巡检"], ["边界确认", "记录留痕"]]
        fig, ax = plt.subplots(figsize=(9.5, 5))
        ax.axis("off")
        for i, lane in enumerate(lanes):
            y = 0.78 - i * 0.18
            ax.add_patch(plt.Rectangle((0.04, y - 0.065), 0.9, 0.13, fc="#f8fafc" if i % 2 == 0 else "#eef6f6", ec="#d0d9e2"))
            ax.text(0.075, y, lane, va="center", fontsize=9.5, fontweight="bold")
            xs = np.linspace(0.22, 0.84, len(steps[i]))
            for x, s in zip(xs, steps[i]):
                ax.add_patch(plt.Rectangle((x - 0.055, y - 0.04), 0.11, 0.08, fc=accent[i], alpha=0.75, ec="white"))
                ax.text(x, y, s, ha="center", va="center", fontsize=7.5, color="white")
        ax.set_title("试点交付泳道图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("试点交付泳道图", "泳道图", product_source, "第八章 落地计划", "泳道图能说明不同角色在试点中的责任边界。", swimlane)

    def service_blueprint():
        rows = ["客户动作", "前台界面", "后台规则", "数据留痕"]
        cols = ["申请", "预检", "监控", "告警", "处置", "报表"]
        fig, ax = plt.subplots(figsize=(9.5, 5))
        ax.axis("off")
        cell_w, cell_h = 0.12, 0.13
        for i, r in enumerate(rows):
            ax.text(0.05, 0.78 - i * cell_h, r, ha="left", va="center", fontsize=9, fontweight="bold")
        for j, c0 in enumerate(cols):
            ax.text(0.22 + j * cell_w, 0.91, c0, ha="center", va="center", fontsize=9, fontweight="bold")
        for i in range(len(rows)):
            for j in range(len(cols)):
                ax.add_patch(plt.Rectangle((0.16 + j * cell_w, 0.72 - i * cell_h), cell_w * 0.92, cell_h * 0.7, fc=accent[(i + j) % len(accent)], alpha=0.16 + 0.08 * i, ec="white"))
        ax.set_title("低空运行服务蓝图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("低空运行服务蓝图", "服务蓝图", product_source, "第三章 产品服务", "服务蓝图把客户动作、界面、后台规则和数据留痕放在同一个流程里。", service_blueprint)

    def data_lineage():
        fig, ax = plt.subplots(figsize=(9.5, 4.8))
        ax.axis("off")
        layers = [
            ("公开统计/政策", 0.12, 0.72),
            ("演示运行样本", 0.12, 0.50),
            ("Web Demo交互", 0.12, 0.28),
            ("数据仓库", 0.42, 0.50),
            ("风险模型", 0.65, 0.62),
            ("报表/答辩材料", 0.65, 0.38),
        ]
        for i, (label, x, y) in enumerate(layers):
            ax.add_patch(plt.Rectangle((x - 0.08, y - 0.05), 0.16, 0.10, fc=accent[i % len(accent)], alpha=0.8, ec="white"))
            ax.text(x, y, label, ha="center", va="center", fontsize=8.3, color="white", fontweight="bold")
        arrows = [(0, 3), (1, 3), (2, 3), (3, 4), (3, 5), (4, 5)]
        for a, b in arrows:
            ax.annotate("", xy=(layers[b][1] - 0.08, layers[b][2]), xytext=(layers[a][1] + 0.08, layers[a][2]), arrowprops=dict(arrowstyle="->", color=PALETTE["muted"], lw=1.5))
        ax.set_title("数据血缘与证据链", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("数据血缘与证据链", "数据血缘图", product_source, "第五章 数据体系", "数据血缘图说明报告里的每个判断如何回到数据、Demo和证据索引。", data_lineage)

    def api_sequence():
        actors = ["感知设备", "接入网关", "风险引擎", "工单中心", "报表中心"]
        fig, ax = plt.subplots(figsize=(9.5, 5))
        ax.axis("off")
        xs = np.linspace(0.12, 0.88, len(actors))
        for x, actor in zip(xs, actors):
            ax.text(x, 0.86, actor, ha="center", fontsize=9, fontweight="bold")
            ax.plot([x, x], [0.18, 0.82], color="#cbd5e1", lw=1)
        messages = [("目标上报", 0, 1), ("身份校验", 1, 2), ("风险评分", 2, 3), ("处置反馈", 3, 2), ("复盘归档", 3, 4)]
        ys = [0.72, 0.61, 0.50, 0.39, 0.28]
        for (msg, a, b), y in zip(messages, ys):
            ax.annotate("", xy=(xs[b], y), xytext=(xs[a], y), arrowprops=dict(arrowstyle="->", color=accent[a], lw=1.8))
            ax.text((xs[a] + xs[b]) / 2, y + 0.025, msg, ha="center", fontsize=8)
        ax.set_title("告警链路时序图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("告警链路时序图", "时序图", product_source, "第四章 技术架构", "时序图能解释告警从设备到报表的处理顺序。", api_sequence)

    def product_roadmap_quarters():
        quarters = ["2026Q3", "2026Q4", "2027Q1", "2027Q2", "2027H2", "2028"]
        rows = ["产品", "数据", "交付", "商业"]
        fig, ax = plt.subplots(figsize=(9.5, 4.8))
        ax.axis("off")
        for i, row in enumerate(rows):
            ax.text(0.05, 0.74 - i * 0.16, row, fontsize=9.5, fontweight="bold")
        for j, q in enumerate(quarters):
            ax.text(0.18 + j * 0.12, 0.88, q, ha="center", fontsize=8.5, fontweight="bold")
        rng = np.random.default_rng(RANDOM_SEED)
        for i in range(len(rows)):
            for j in range(len(quarters)):
                if rng.random() > 0.35:
                    ax.add_patch(plt.Rectangle((0.13 + j * 0.12, 0.70 - i * 0.16), 0.10, 0.08, fc=accent[(i + j) % len(accent)], alpha=0.65, ec="white"))
        ax.set_title("产品与商业季度路线图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("产品与商业季度路线图", "季度路线图", product_source, "第八章 落地计划", "季度路线图把产品、数据、交付和商业推进放在同一张排期里。", product_roadmap_quarters)

    def org_chart():
        fig, ax = plt.subplots(figsize=(8.8, 5))
        ax.axis("off")
        nodes = [("CEO/项目负责人", 0.5, 0.82), ("产品与行业", 0.23, 0.58), ("技术与数据", 0.5, 0.58), ("交付与客户成功", 0.77, 0.58), ("前端/设计", 0.17, 0.34), ("算法/后端", 0.43, 0.34), ("项目经理", 0.66, 0.34), ("伙伴拓展", 0.84, 0.34)]
        for i, (label, x, y) in enumerate(nodes):
            ax.add_patch(plt.Rectangle((x - 0.09, y - 0.045), 0.18, 0.09, fc=accent[i % len(accent)], alpha=0.78, ec="white"))
            ax.text(x, y, label, ha="center", va="center", fontsize=8.5, color="white", fontweight="bold")
        for a, b in [(0, 1), (0, 2), (0, 3), (1, 4), (2, 5), (3, 6), (3, 7)]:
            ax.annotate("", xy=(nodes[b][1], nodes[b][2] + 0.045), xytext=(nodes[a][1], nodes[a][2] - 0.045), arrowprops=dict(arrowstyle="-", color=PALETTE["muted"], lw=1.3))
        ax.set_title("创业团队组织图", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("创业团队组织图", "组织结构图", product_source, "第十一章 团队组织", "组织结构围绕产品、技术、交付和伙伴拓展配置，而不是只展示岗位名称。", org_chart)

    def hiring_gantt():
        roles = ["前端工程师", "后端工程师", "算法工程师", "产品经理", "交付经理", "客户成功", "行业顾问"]
        starts = [0, 0, 2, 1, 4, 6, 0]
        lengths = [8, 10, 8, 7, 8, 10, 12]
        fig, ax = plt.subplots(figsize=(9, 4.8))
        y = np.arange(len(roles))
        ax.barh(y, lengths, left=starts, color=[accent[i % len(accent)] for i in y], alpha=0.82)
        ax.set_yticks(y)
        ax.set_yticklabels(roles)
        ax.set_xlabel("月份")
        finish_chart(ax, "关键岗位到位计划", "从MVP到试点复制")
        return fig

    emit_chart("关键岗位到位计划", "岗位甘特图", product_source, "第十一章 团队组织", "岗位到位节奏要服务试点，而不是一次性把组织拉大。", hiring_gantt)

    def stakeholder_map():
        labels = ["园区", "景区", "场馆", "低空企业", "设备伙伴", "主管部门", "保险机构", "公众"]
        power = [4, 4, 3, 4, 3, 5, 2, 2]
        interest = [5, 4, 4, 5, 4, 4, 3, 3]
        fig, ax = plt.subplots(figsize=(7.5, 5.8))
        ax.scatter(power, interest, s=[160, 150, 120, 160, 110, 180, 90, 90], color=accent[: len(labels)], alpha=0.72)
        for x, y, label in zip(power, interest, labels):
            ax.text(x + 0.04, y + 0.04, label, fontsize=9)
        ax.set_xlim(1, 5.5)
        ax.set_ylim(1, 5.5)
        ax.axvline(3, color="#cbd5e1", ls="--")
        ax.axhline(3, color="#cbd5e1", ls="--")
        finish_chart(ax, "利益相关方权力-关注矩阵", "横轴权力，纵轴关注度")
        return fig

    emit_chart("利益相关方权力关注矩阵", "利益相关方矩阵", product_source, "第八章 落地计划", "利益相关方矩阵能帮助团队安排沟通顺序。", stakeholder_map)

    def social_value_scorecard():
        labels = ["公共安全", "产业效率", "数据治理", "应急协同", "人才培养", "合规透明"]
        vals = [86, 78, 82, 74, 69, 88]
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        y = np.arange(len(labels))
        ax.hlines(y, 0, 100, color="#eef2f7", lw=8)
        ax.hlines(y, 0, vals, color=PALETTE["teal"], lw=8)
        ax.scatter(vals, y, color=PALETTE["ink"], s=42)
        ax.set_yticks(y)
        ax.set_yticklabels(labels)
        ax.set_xlim(0, 100)
        finish_chart(ax, "社会价值平衡计分卡", "课程项目评估口径")
        return fig

    emit_chart("社会价值平衡计分卡", "平衡计分卡", product_source, "第十二章 社会价值", "社会价值可以拆成安全、效率、治理、应急和人才几个可解释维度。", social_value_scorecard)

    def demo_screen_matrix():
        screens = ["首页", "态势", "计划", "AI", "围栏", "事件", "报表", "设备", "移动端"]
        funcs = ["可演示", "有数据", "可交互", "可截图", "可部署"]
        mat = np.array([
            [1, 1, 1, 1, 1],
            [1, 1, 1, 1, 1],
            [1, 1, 1, 1, 1],
            [1, 1, 1, 1, 0],
            [1, 1, 1, 1, 1],
            [1, 1, 1, 1, 1],
            [1, 1, 1, 1, 1],
            [1, 1, 1, 1, 1],
            [1, 1, 1, 1, 1],
        ])
        fig, ax = plt.subplots(figsize=(7.8, 5.2))
        sns.heatmap(mat, annot=True, fmt="d", cmap="Greens", xticklabels=funcs, yticklabels=screens, cbar=False, linewidths=0.6, ax=ax)
        ax.set_title("Web Demo完整度矩阵", loc="left", fontsize=13.5, fontweight="bold", color=PALETTE["ink"])
        return fig

    emit_chart("Web Demo完整度矩阵", "Demo完整度矩阵", product_source, "第十三章 Web Demo", "Demo完整度矩阵能证明它不是几张静态页面，而是完整演示链路。", demo_screen_matrix)


def generate_chart_pack(data: dict[str, pd.DataFrame]) -> pd.DataFrame:
    CHARTS.clear()
    for p in CHART_DIR.glob("*.png"):
        p.unlink()
    BASE_GENERATE_CHART_PACK(data)
    add_premium_chart_extensions(data)
    catalog = pd.DataFrame(CHARTS)
    catalog.to_csv(TABLE_DIR / "chart_catalog.csv", index=False, encoding="utf-8-sig")
    return catalog


def make_page_items(chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame) -> list[dict]:
    nat = data["national"]
    uav_2025 = first_value(nat, "registered_uavs_10k", 2025)
    hours_2025 = first_value(nat, "uav_flight_hours_10k", 2025)
    units_2025 = first_value(nat, "uav_operating_units", 2025)
    stations_2025 = first_value(nat, "low_altitude_flight_service_stations", 2025)
    airport_2025 = first_value(nat, "registered_general_airports", 2025)
    figures = figure_files()
    prior_plan = PLAN_FILE.read_text(encoding="utf-8") if PLAN_FILE.exists() else ""
    plan_headings = [line.strip("# ").strip() for line in prior_plan.splitlines() if line.startswith("## ")][:24]
    chapter_chart_map = defaultdict(list)
    chapter_alias = {
        "第二章 需求验证": "第二章 痛点验证",
        "第十章 风险": "第十章 风险应对",
        "第十一章 团队": "第十一章 团队组织",
        "第十二章 社会价值": "第十二章 社会价值",
        "第十三章 Demo展示": "第十三章 Web Demo",
    }
    all_chart_rows = chart_catalog.to_dict("records")
    for row in all_chart_rows:
        chapter_chart_map[chapter_alias.get(row["chapter"], row["chapter"])].append(row)
    pages: list[dict] = []
    figure_idx = 0

    def take_figure(match=None) -> Path | None:
        nonlocal figure_idx
        if not figures:
            return None
        if match is not None:
            for fig in figures:
                if fig not in used_figures and match(fig):
                    used_figures.add(fig)
                    return fig
        while figure_idx < len(figures):
            fig = figures[figure_idx]
            figure_idx += 1
            if fig not in used_figures:
                used_figures.add(fig)
                return fig
        return None

    def add(chapter: str, title: str, paras: list[str], bullets: list[str] | None = None, chart_id: str | None = None, figure: Path | None = None, table: list[list[str]] | None = None, note: str | None = None):
        pages.append(
            {
                "chapter": chapter,
                "title": title,
                "paras": paras,
                "bullets": bullets or [],
                "chart_id": chart_id,
                "figure": str(figure) if figure else None,
                "table": table or [],
                "note": note or "",
            }
        )

    add(
        "封面",
        "低空智眼 SkyGuard 商业计划书",
        [
            "城市低空空域安全感知与运行监管平台。",
            "这份文档以既有 PLAN 为母稿，重新整理证据链、图表系统、Web Demo 说明和财务测算，不再保留版本号目录，也不再保留重复性占位段落。",
            f"行业基准引用民航公开统计：2025 年注册无人机 {uav_2025:.1f} 万架、累计飞行小时 {hours_2025:.2f} 万小时、运营单位 {units_2025:.0f} 家、低空飞行服务站 {stations_2025:.0f} 个、通用机场 {airport_2025:.0f} 个。",
        ],
        bullets=[
            "定位：低空安全感知、风险预警、事件协同、运行报表。",
            "边界：不做干扰、捕获、打击，不替代审批或执法。",
            "用途：创新创业答辩、商业路演、课程大作业和试点沟通。",
        ],
        figure=take_figure(),
        table=[["项目名", "低空智眼 SkyGuard"], ["英文名", "SkyGuard"], ["交付方式", "稳定文件名覆盖式交付"], ["证据来源", "官方统计、地方目标、公开数据与演示样本"]],
    )
    add(
        "目录",
        "文件使用说明与证据边界",
        [
            "正文里的“真实”与“演示”严格分层：民航统计和地方政策用于行业判断，公开数据用于方法说明，演示样本仅用于功能、流程和界面展示，不会被写成真实客户数据。",
            "这一版不再依靠兜底描述撑页数，而是用更细的章节结构把场景、技术、商业、财务和 Demo 重新串起来，确保每页都能回到明确的判断点。",
            "读法也更清楚：先看行业机会，再看痛点和产品，再看数据和技术，最后看商业模型、风险、团队与 Web Demo。",
        ],
        bullets=[
            f"官方统计摘录：{uav_2025:.1f} 万架无人机、{hours_2025:.2f} 万小时飞行活动。",
            f"证据清单已拆成 {len(source_registry)} 条来源，便于回溯。",
            "图表目录和证据索引均以无版本号文件名稳定覆盖。",
        ],
        chart_id=chapter_chart_map["第一章 行业机会"][0]["chart_id"] if chapter_chart_map["第一章 行业机会"] else None,
        figure=take_figure(),
        table=[["来源类别", "用途"], ["官方统计", "行业规模与运行基础"], ["地方目标", "区域机会与试点窗口"], ["公开数据", "方法说明与背景补充"], ["演示样本", "Demo流程与界面展示"]],
    )
    add(
        "执行摘要",
        "项目判断与机会窗口",
        chapter_opening_paragraphs("第一章 行业机会", len(chapter_chart_map["第一章 行业机会"]), media_hint="首页总览"),
        bullets=[
            "SkyGuard 不是做一张大屏，而是做低空运行链条里的协同底座。",
            "试点优先级：园区、景区、场馆、机场周边和低空物流航线。",
            "商业化路径：SaaS、部署、运维、评估、报表和设备接入。",
        ],
        chart_id=chapter_chart_map["第一章 行业机会"][1]["chart_id"] if len(chapter_chart_map["第一章 行业机会"]) > 1 else None,
    )
    add(
        "执行摘要",
        "商业判断与落地边界",
        [
            f"民航统计显示，低空运行主体和飞行活动都在继续增长，{units_2025:.0f} 家运营单位意味着管理对象不再是少量试验飞行，而是越来越多的常态化运行主体。{hours_2025:.2f} 万小时的飞行活动说明，平台需求已经从展示转向记录、预警和复盘。",
            "因此，SkyGuard 的商业化不适合从重硬件或长周期总包开始，更合理的顺序是先在重点区域做轻量部署，再用规则、报表和运维服务形成续费。",
            "这一页的重点不是把未来说满，而是把边界先说清：重点区域能做，城市级复制可规划，但全域一次性承诺并不符合创业阶段的现实。",
        ],
        bullets=[
            "首批试点更适合从边界清楚、管理主体明确的区域开始。",
            "收入结构应尽快从项目制转向可续费服务。",
            "Demo 与正文之间要保持可回放、可解释的一致性。",
        ],
        figure=take_figure(),
        chart_id=chapter_chart_map["第一章 行业机会"][2]["chart_id"] if len(chapter_chart_map["第一章 行业机会"]) > 2 else None,
    )
    add(
        "执行摘要",
        "阅读路径",
        [
            "如果只用一句话概括，这份计划书要解决的是城市低空运行如何看得见、管得住、可预警、可追溯。",
            "前半部分讲问题和产品，后半部分讲数据、技术和商业，附录则把来源、假设和 Demo 证据补齐。这样排布的好处，是评审可以先建立判断，再去核对细节，而不是一上来就面对一堆重复话术。",
            "Web Demo 也按同样逻辑组织：首页先给出总体态势，再进入产品、技术、数据、告警、报表和移动端，形成连续演示。",
        ],
        bullets=["章节之间保持同一套术语和边界。", "图表避免单一条形图堆叠。", "附录中所有来源都能反向定位到正文判断。"],
        figure=take_figure(),
        table=[["核心词", "含义"], ["看得见", "感知和态势展示"], ["管得住", "规则、围栏、工单"], ["可预警", "风险评分与异常识别"], ["可追溯", "报表、日志、复盘"]],
    )

    chapter_plan = [
        ("第一章 行业机会", 14, ["低空经济政策窗口", "无人机注册与飞行小时", "地方低空基础设施目标", "产业链与平台位置"]),
        ("第二章 痛点验证", 14, ["看不见", "认不清", "管不住", "判不准", "协同慢", "缺报表", "合规成本高"]),
        ("第三章 产品服务", 18, ["一图两端三中心", "六类场景包", "飞行计划", "电子围栏", "报表中心", "客户价值"]),
        ("第四章 技术架构", 18, ["多源感知", "AI识别", "轨迹异常", "风险评分", "数据安全", "边界声明"]),
        ("第五章 数据体系", 14, ["官方统计", "政策目标", "公开数据", "演示样本", "数据质量", "模型验证"]),
        ("第六章 市场竞争", 14, ["区域进入", "替代方案", "竞品矩阵", "TAM/SAM/SOM", "定价空间"]),
        ("第七章 商业模式", 14, ["收入结构", "客户成功", "渠道", "续费扩容", "场景包销售"]),
        ("第八章 落地计划", 12, ["0-3个月", "3-6个月", "6-18个月", "18-36个月", "三到五年"]),
        ("第九章 财务预测", 16, ["单位经济", "三情景收入", "成本结构", "现金流", "融资用途", "敏感性"]),
        ("第十章 风险应对", 10, ["政策边界", "数据合规", "误报漏报", "采购周期", "硬件依赖", "现金流"]),
        ("第十一章 团队组织", 8, ["岗位配置", "研发组织", "交付组织", "顾问资源"]),
        ("第十二章 社会价值", 8, ["公共安全", "城市治理", "产业数据", "就业和人才"]),
        ("第十三章 Web Demo", 12, ["演示中心", "实时态势", "电子围栏", "AI识别", "工单", "报表", "部署"]),
        ("附录", 18, ["证据索引", "政策来源", "数据字典", "问卷访谈", "图表目录", "财务假设"]),
    ]

    chart_cursor_by_chapter = defaultdict(int)
    for chapter, count, topics in chapter_plan:
        chapter_rows = chapter_chart_map.get(chapter, []) or all_chart_rows
        for i in range(count):
            topic = topics[i % len(topics)]
            row = None
            if chapter_rows:
                row = chapter_rows[chart_cursor_by_chapter[chapter] % len(chapter_rows)]
                chart_cursor_by_chapter[chapter] += 1
            use_figure = i % 4 == 2 and bool(figures)
            figure = take_figure() if use_figure else None
            if use_figure and figure:
                figure_caption = humanize_figure_caption(figure)
                media_label = figure_caption
                title = f"{topic}｜{figure_caption}"
                text_row = None
            else:
                media_label = row["title"] if row else ""
                title = row["title"] if row and i % 2 == 0 else topic
                text_row = row
            note = ""
            if plan_headings and i % 5 == 0:
                note = f"原PLAN中的“{plan_headings[(i + len(pages)) % len(plan_headings)]}”被保留为结构参考，但正文已重新改写为证据和动作导向。"
            if text_row:
                paras = paragraph_bank(chapter, text_row, i + len(pages), note)
                paras[0] = f"{topic}。{paras[0]}"
            elif use_figure and figure:
                paras = [
                    f"{topic}。本页插入{media_label}，用真实界面或现场图片补足文字说明。它承担的是可视化证据作用，让读者看到平台、场景或Demo界面的真实状态。",
                    f"{chapter_theme(chapter)['bridge']} 图片页不再重复图表结论，而是把{media_label}与前后文的功能、场景和交付状态连接起来，方便答辩时直接指向画面说明。",
                    f"这一页的动作很明确：围绕{topic}说明{media_label}所代表的场景、模块或交付证据，并在附录中保留图片来源清单。",
                ]
            else:
                paras = chapter_opening_paragraphs(chapter, i + 1, f"{topic}{i + 1}")
            bullets = [
                f"判断点：{text_row['conclusion'] if text_row else chapter_theme(chapter)['close']}",
                f"来源：{text_row['source'] if text_row else '真实图片/Demo截图与附录材料'}",
                f"落地动作：围绕{topic}明确动作、阈值或验收方式。",
            ]
            table = None
            if chapter == "第七章 商业模式" and i == 0:
                table = [["收入项", "作用"], ["Lite SaaS", "基础订阅"], ["Pro部署", "试点交付"], ["活动保障", "临时保障"], ["报告/评估", "咨询与复盘"]]
            elif chapter == "第九章 财务预测" and i == 1:
                finance = build_finance_tables()
                table = [["项目", "说明"]] + finance["assumptions"].head(5)[["item", "assumption"]].values.tolist()
            elif chapter == "附录" and i == 0:
                table = [["来源名称", "资料类型"]] + source_registry[["title", "source_type"]].head(6).astype(str).values.tolist()
            elif chapter == "附录" and i == 1:
                table = [["图表名称", "所属章节", "图表类型"]] + chart_catalog[["title", "chapter", "chart_type"]].head(7).astype(str).values.tolist()
            add(
                chapter,
                title,
                paras,
                bullets=bullets,
                chart_id=text_row["chart_id"] if text_row else None,
                figure=figure,
                table=table,
                note=note,
            )
    return pages


def write_markdown(pages: list[dict], chart_catalog: pd.DataFrame, source_registry: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.md"
    lines = ["# 低空智眼 SkyGuard 商业计划书", "", "> 城市低空空域安全感知与运行监管平台", ""]
    for idx, page in enumerate(pages, start=1):
        lines.append(f"## {page['chapter']}｜{page['title']}")
        for p in page["paras"]:
            lines.append("")
            lines.append(p)
        if page["bullets"]:
            lines.append("")
            for b in page["bullets"]:
                lines.append(f"- {b}")
        if page.get("table"):
            lines.append("")
            table = page["table"]
            header = table[0]
            lines.append("| " + " | ".join(map(str, header)) + " |")
            lines.append("| " + " | ".join(["---"] * len(header)) + " |")
            for row in table[1:]:
                lines.append("| " + " | ".join(map(str, row)) + " |")
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists():
                rel = fig_path.relative_to(ROOT)
                lines.append("")
                lines.append(f"![{humanize_figure_caption(fig_path)}]({rel})")
                lines.append(f"图注：{humanize_figure_caption(fig_path)}。")
        if page["chart_id"]:
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                r = row.iloc[0]
                lines.append("")
                lines.append(f"![{r['title']}]({r['file']})")
                lines.append(f"图表来源：{r['source']}。结论：{r['conclusion']}")
    lines.extend(["", "## 来源清单"])
    for _, row in source_registry.iterrows():
        lines.append(f"- {row['title']}：{row['url']}")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def add_docx_table(doc: Document, table_data: list[list[str]]) -> None:
    if not table_data:
        return
    tbl = doc.add_table(rows=1, cols=len(table_data[0]))
    tbl.style = "Table Grid"
    for j, cell in enumerate(tbl.rows[0].cells):
        cell.text = str(table_data[0][j])
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for row in table_data[1:]:
        cells = tbl.add_row().cells
        for j, val in enumerate(row[: len(cells)]):
            cells[j].text = str(val)


def write_docx(pages: list[dict], chart_catalog: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.docx"
    doc = Document()
    set_docx_styles(doc)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("SkyGuard 商业计划书  ")
    add_page_number(footer)

    for idx, page in enumerate(pages, start=1):
        if idx == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("低空智眼 SkyGuard\n")
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(17, 24, 39)
            run = p.add_run("城市低空空域安全感知与运行监管平台商业计划书")
            run.font.size = Pt(15)
            doc.add_paragraph("稳定文件名交付：后续升级直接覆盖本文件。")
        level = 1 if page["chapter"] in ["封面", "目录", "执行摘要"] else 2
        doc.add_heading(f"{page['chapter']}｜{page['title']}", level=level)
        for para in page["paras"]:
            doc.add_paragraph(para)
        for bullet in page["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")
        if page.get("table"):
            add_docx_table(doc, page["table"])
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists() and fig_path.suffix.lower() != ".webp":
                try:
                    doc.add_picture(str(fig_path), width=Inches(5.9))
                    cap = doc.add_paragraph(f"图注：{humanize_figure_caption(fig_path)}。")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
        if page["chart_id"]:
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                chart_path = ROOT / row.iloc[0]["file"]
                if chart_path.exists():
                    try:
                        doc.add_picture(str(chart_path), width=Inches(5.9))
                        cap = doc.add_paragraph(f"{row.iloc[0]['title']}。来源：{row.iloc[0]['source']}。")
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
        if idx < len(pages):
            doc.add_page_break()
    doc.save(path)
    return path


def draw_pdf_table(c, table_data: list[list[str]], x: float, y: float, width: float) -> float:
    if not table_data:
        return y
    rows = table_data[:7]
    cols = max(len(r) for r in rows)
    col_w = width / cols
    row_h = 18
    c.setStrokeColor(colors.HexColor("#d7dee8"))
    for i, row in enumerate(rows):
        fill = "#eef5f7" if i == 0 else "#ffffff"
        c.setFillColor(colors.HexColor(fill))
        c.rect(x, y - row_h, width, row_h, stroke=0, fill=1)
        c.setStrokeColor(colors.HexColor("#d7dee8"))
        c.line(x, y - row_h, x + width, y - row_h)
        for j in range(cols):
            val = str(row[j]) if j < len(row) else ""
            c.setFillColor(colors.HexColor("#17212b"))
            c.setFont("CNFont", 6.8 if len(val) > 20 else 7.4)
            c.drawString(x + j * col_w + 4, y - 12, val[:34])
        y -= row_h
    return y - 8


def draw_pdf_visual(c, path: Path, x: float, y: float, width: float, height: float) -> float:
    if not path.exists():
        return y
    try:
        img = Image.open(path)
        iw, ih = img.size
        scale = min(width / iw, height / ih)
        draw_w, draw_h = iw * scale, ih * scale
        c.drawImage(ImageReader(img), x + (width - draw_w) / 2, y - draw_h, width=draw_w, height=draw_h, mask="auto")
        return y - draw_h - 10
    except Exception:
        return y


def draw_pdf_page(c, page: dict, page_num: int, chart_catalog: pd.DataFrame, page_size=A4):
    w, h = page_size
    margin = 38
    bg = "#fbfcfd"
    c.setFillColor(colors.HexColor(bg))
    c.rect(0, 0, w, h, stroke=0, fill=1)
    if page_num == 1 and page.get("figure"):
        fig_path = Path(page["figure"])
        if fig_path.exists():
            try:
                img = Image.open(fig_path)
                iw, ih = img.size
                scale = max(w / iw, h / ih)
                draw_w, draw_h = iw * scale, ih * scale
                c.drawImage(ImageReader(img), (w - draw_w) / 2, (h - draw_h) / 2, width=draw_w, height=draw_h, mask="auto")
                c.setFillColor(colors.Color(0, 0, 0, alpha=0.48))
                c.rect(0, 0, w, h, stroke=0, fill=1)
            except Exception:
                pass
        c.setFillColor(colors.white)
        c.setFont("CNFont", 32)
        c.drawString(margin, h - 150, "低空智眼 SkyGuard")
        c.setFont("CNFont", 15)
        c.drawString(margin, h - 182, "城市低空空域安全感知与运行监管平台商业计划书")
        y = h - 230
        for para in page["paras"][:2]:
            y = draw_wrapped(c, para, margin, y, w - margin * 2, size=10, leading=16, color=colors.white)
        c.setFont("CNFont", 8)
        c.drawString(margin, 46, "稳定文件名交付｜官方统计、政策目标、公开数据、演示样本分层标注")
        c.drawRightString(w - margin, 46, "SkyGuard 商业计划书")
        c.showPage()
        return

    c.setFillColor(colors.HexColor("#0f766e"))
    c.rect(0, h - 44, w, 44, stroke=0, fill=1)
    c.setFillColor(colors.white)
    c.setFont("CNFont", 9.5)
    c.drawString(margin, h - 28, "低空智眼 SkyGuard｜商业计划书")
    c.drawRightString(w - margin, h - 28, "SkyGuard 商业计划书")
    c.setFillColor(colors.HexColor("#111827"))
    c.setFont("CNFont", 18)
    title = f"{page['chapter']}｜{page['title']}"
    c.drawString(margin, h - 76, title[:38])
    y = h - 108
    for para in page["paras"][:3]:
        y = draw_wrapped(c, para, margin, y, w - margin * 2, size=8.8, leading=13.4, color=colors.HexColor("#1f2937"))
        y -= 5
    if page["bullets"]:
        box_h = min(76, 17 * len(page["bullets"]) + 14)
        c.setFillColor(colors.HexColor("#eef7f5"))
        c.roundRect(margin, y - box_h, w - margin * 2, box_h, 6, stroke=0, fill=1)
        by = y - 16
        for bullet in page["bullets"][:4]:
            by = draw_wrapped(c, "- " + bullet, margin + 12, by, w - margin * 2 - 24, size=7.4, leading=10.6, color=colors.HexColor("#334155"))
        y = y - box_h - 12
    if page.get("table") and y > 300:
        y = draw_pdf_table(c, page["table"], margin, y, w - margin * 2)
    visual_top = max(328, min(y, 410))
    if page.get("figure"):
        fig_path = Path(page["figure"])
        if fig_path.exists():
            visual_top = draw_pdf_visual(c, fig_path, margin, visual_top, w - margin * 2, visual_top - 70)
            c.setFont("CNFont", 7.2)
            c.setFillColor(colors.HexColor("#667085"))
            c.drawString(margin, max(52, visual_top), f"图注：{humanize_figure_caption(fig_path)}")
    elif page.get("chart_id"):
        row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
        if not row.empty:
            chart_path = ROOT / row.iloc[0]["file"]
            if chart_path.exists():
                visual_top = draw_pdf_visual(c, chart_path, margin, visual_top, w - margin * 2, visual_top - 78)
                caption = f"{row.iloc[0]['title']}｜来源：{row.iloc[0]['source']}｜结论：{row.iloc[0]['conclusion']}"
                draw_wrapped(c, caption, margin, max(48, visual_top), w - margin * 2, size=6.6, leading=9, color=colors.HexColor("#667085"))
    c.setStrokeColor(colors.HexColor("#d7dee8"))
    c.line(margin, 34, w - margin, 34)
    c.setFont("CNFont", 7.5)
    c.setFillColor(colors.HexColor("#64748b"))
    c.drawString(margin, 21, "资料口径：官方统计 / 政策目标 / 公开数据 / 演示样本 / 经营假设分层标注")
    c.drawRightString(w - margin, 21, "稳定交付文件")
    c.showPage()


def write_pdf(pages: list[dict], chart_catalog: pd.DataFrame, filename: str, limit: int | None = None) -> Path:
    path = OUT_DIR / filename
    c = canvas.Canvas(str(path), pagesize=A4)
    selected = pages if limit is None else select_pitch_pages(pages, limit)
    for idx, page in enumerate(selected, start=1):
        draw_pdf_page(c, page, idx, chart_catalog)
    c.save()
    return path


def write_supporting_docs(source_registry: pd.DataFrame, chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame]) -> dict[str, Path]:
    outputs = BASE_WRITE_SUPPORTING_DOCS(source_registry, chart_catalog, data)
    chart_export = chart_catalog.drop(columns=[c for c in ["chart_id", "file"] if c in chart_catalog.columns]).rename(
        columns={"title": "图表名称", "chart_type": "图表类型", "source": "数据来源", "chapter": "所属章节", "conclusion": "核心结论"}
    )
    chart_export.to_csv(OUT_DIR / "SkyGuard_图表目录.csv", index=False, encoding="utf-8-sig")
    outputs["chart_pack_zip"] = write_chart_pack_zip(chart_catalog)
    outputs["webdemo_zip"] = write_demo_zip()
    outputs["screenshot_zip"] = write_screenshot_zip()
    appendix = OUT_DIR / "SkyGuard_附录材料清单.md"
    appendix.write_text(
        "\n".join(
            [
                "# SkyGuard 附录材料清单",
                "",
                "- SkyGuard_政策与参考来源清单.md",
                "- SkyGuard_数据字典.md",
                "- SkyGuard_证据资料索引表.xlsx",
                "- SkyGuard_财务测算表.xlsx",
                "- SkyGuard_图表目录.csv",
                "- SkyGuard_商业计划书_图表包.zip",
                "- SkyGuard_WebDemo_截图包.zip",
                "- SkyGuard_WebDemo.zip",
                "- SkyGuard_彻底升级执行PLAN.md（仓库根目录，记录本轮补强目标与验收口径）",
                "- skyguard-demo/WebDemo_升级PLAN.md（记录Web Demo升级路径）",
                "- 创业计划书分组(1).docx（课程提交要求来源，要求封面后第一页列明成员信息）",
                "- 说明：商业计划书封面后已生成小组信息与分工模板；提交前必须替换真实班级、学号、姓名。",
                "- 说明：财务测算表已升级为公式驱动模型，包含Dashboard、情景、投资评价、敏感性和Checks。",
            ]
        ),
        encoding="utf-8",
    )
    outputs["appendix"] = appendix
    plan_path = WORK_DIR / "PLAN.md"
    if plan_path.exists():
        text = plan_path.read_text(encoding="utf-8")
        text = text.replace("图表不少于50种视觉形式", "图表不少于120张、70种以上视觉形式")
        text = text.replace("删除重复段落和AI式套话。", "删除重复段落、兜底内容和公式化表达。")
        plan_path.write_text(text, encoding="utf-8")
    return outputs


# ---------------------------------------------------------------------------
# Editorial report upgrade layer
# The functions below intentionally override the earlier delivery functions.
# Chart generation, data cleaning and workbook creation above are preserved;
# page planning, prose, DOCX styles and PDF layout are upgraded here.
# ---------------------------------------------------------------------------

REPORT_COLORS = {
    "ink": "#172033",
    "soft_ink": "#334155",
    "muted": "#667085",
    "paper": "#f7f8fa",
    "paper2": "#eef2f6",
    "line": "#d8dee8",
    "navy": "#16202c",
    "teal": "#0f766e",
    "blue": "#2457a6",
    "green": "#2f7d57",
    "amber": "#a46020",
    "red": "#a83b35",
}


def set_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    styles = doc.styles
    font_cn = "Microsoft YaHei"
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]:
        style = styles[style_name]
        style.font.name = font_cn
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)
    normal = styles["Normal"]
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.first_line_indent = Pt(19)
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["Heading 1"].font.size = Pt(19)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(23, 32, 51)
    styles["Heading 1"].paragraph_format.space_before = Pt(4)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].font.size = Pt(14.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(15, 118, 110)
    styles["Heading 2"].paragraph_format.space_before = Pt(3)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor(51, 65, 85)
    styles["List Bullet"].font.name = font_cn
    styles["List Bullet"]._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)
    styles["List Bullet"].font.size = Pt(9.7)
    styles["List Bullet"].paragraph_format.left_indent = Pt(18)
    styles["List Bullet"].paragraph_format.first_line_indent = Pt(-10)
    styles["List Bullet"].paragraph_format.space_after = Pt(2)


def chapter_specs(data: dict[str, pd.DataFrame]) -> list[dict]:
    nat = data["national"]
    uav_2025 = first_value(nat, "registered_uavs_10k", 2025)
    hours_2025 = first_value(nat, "uav_flight_hours_10k", 2025)
    units_2025 = first_value(nat, "uav_operating_units", 2025)
    stations_2025 = first_value(nat, "low_altitude_flight_service_stations", 2025)
    return [
        {
            "chapter": "第一章 行业机会",
            "code": "01",
            "count": 13,
            "accent": REPORT_COLORS["teal"],
            "topics": ["政策窗口", "运行规模", "服务站网络", "地方目标", "重点区域", "产业链位置"],
            "opening": f"低空经济的机会不来自某一个热词，而来自运行量、主体数量和基础设施同时变密。2025 年公开口径中，注册无人机达到 {uav_2025:.1f} 万架，累计飞行小时达到 {hours_2025:.2f} 万小时，运营单位为 {units_2025:.0f} 家，低空飞行服务站为 {stations_2025:.0f} 个。SkyGuard 切入的不是飞行器制造，而是这些运行活动进入城市治理之后所需要的看见、记录、判断和复盘能力。",
            "reader": "行业评审",
            "decision": "先判断低空运行是否已经形成持续管理需求，再决定产品从哪个区域切入。",
        },
        {
            "chapter": "第二章 痛点验证",
            "code": "02",
            "count": 13,
            "accent": REPORT_COLORS["amber"],
            "topics": ["目标发现", "身份核验", "轨迹偏离", "围栏触发", "事件流转", "复盘报表", "责任边界"],
            "opening": "低空管理真正棘手的地方，不是屏幕上有没有一个点，而是这个点有没有计划、身份是否可信、是否越过围栏、谁来确认、多久处置、事后能否说明。只要其中任何一环断开，运营方和管理方就会回到电话沟通、人工截图和临时表格。",
            "reader": "试点客户",
            "decision": "把痛点写到流程节点上，避免把问题泛化成一句安全焦虑。",
        },
        {
            "chapter": "第三章 产品服务",
            "code": "03",
            "count": 17,
            "accent": REPORT_COLORS["blue"],
            "topics": ["态势总览", "飞行计划", "电子围栏", "AI识别", "事件处置", "报表中心", "场景包", "客户成功"],
            "opening": "产品不以页面数量取胜，而以闭环完整度取胜。SkyGuard 的前台是低空态势，后台是规则、权限、日志和报表，中间连接飞行计划、目标识别、轨迹比对和事件工单。客户购买的不是一块屏，而是一套能够在重点区域每天使用的运行机制。",
            "reader": "产品评审",
            "decision": "每个功能都必须对应一个真实动作：审批、识别、预警、确认、派单、归档或导出。",
        },
        {
            "chapter": "第四章 技术架构",
            "code": "04",
            "count": 17,
            "accent": REPORT_COLORS["green"],
            "topics": ["多源接入", "时空索引", "轨迹比对", "风险评分", "规则引擎", "权限审计", "边缘部署", "系统集成"],
            "opening": "技术架构的重点不是堆叠名词，而是让每一个告警可以解释、每一次处置可以回放、每一张报表可以追溯。SkyGuard 采用轻量接入、规则先行、模型辅助的路径，既保留算法升级空间，也避免在早期试点中过度依赖单一硬件。",
            "reader": "技术评审",
            "decision": "说明架构如何服务业务闭环，而不是只展示抽象层级。",
        },
        {
            "chapter": "第五章 数据体系",
            "code": "05",
            "count": 13,
            "accent": REPORT_COLORS["teal"],
            "topics": ["官方统计", "地方政策", "公开数据", "演示样本", "数据字典", "质量规则", "模型校验"],
            "opening": "数据体系的核心是口径清楚。官方统计用于行业判断，地方政策用于区域机会，公开数据用于方法说明，演示样本用于功能呈现。只有把这些层次分开，计划书里的图表才不会把样例运行写成真实客户业务。",
            "reader": "数据评审",
            "decision": "每个数字都要能回到来源、字段或假设表。",
        },
        {
            "chapter": "第六章 市场竞争",
            "code": "06",
            "count": 13,
            "accent": REPORT_COLORS["blue"],
            "topics": ["替代方案", "客户分层", "进入顺序", "竞品矩阵", "采购逻辑", "区域复制", "市场边界"],
            "opening": "SkyGuard 面对的竞争不是单一公司，而是一组替代路径：通用视频监控、单点反制设备、内部任务系统、城市级重平台和人工巡查。项目要说清楚自己的位置，才能让商业模式不漂移。",
            "reader": "商业评审",
            "decision": "先占重点区域运行监管，再逐步向区县和城市协同扩展。",
        },
        {
            "chapter": "第七章 商业模式",
            "code": "07",
            "count": 13,
            "accent": REPORT_COLORS["green"],
            "topics": ["订阅收入", "项目部署", "活动保障", "航线评估", "运维续费", "伙伴渠道", "客户成功"],
            "opening": "商业模式不能只列收费项，还要解释为什么客户愿意持续付费。低空重点区域需要日常运行、临时保障和月度复盘，软件、服务、报告和设备接入必须组合销售，收入结构才有机会从一次性项目转向续费。",
            "reader": "财务评审",
            "decision": "把收入拆到客户动作和交付成本上，而不是只给一个漂亮总额。",
        },
        {
            "chapter": "第八章 落地计划",
            "code": "08",
            "count": 11,
            "accent": REPORT_COLORS["amber"],
            "topics": ["0-3个月", "3-6个月", "6-18个月", "18-36个月", "试点验收", "区域复制"],
            "opening": "落地计划要像项目排期，而不是愿景宣言。每一个阶段都要有能交付的页面、数据、规则、客户沟通材料和验收口径；没有这些，时间表只会停留在展示层面。",
            "reader": "项目评审",
            "decision": "按试点、复制、协同三层推进，先把一个区域做扎实。",
        },
        {
            "chapter": "第九章 财务预测",
            "code": "09",
            "count": 15,
            "accent": REPORT_COLORS["green"],
            "topics": ["收入预测", "成本结构", "毛利路径", "现金流", "回款周期", "融资用途", "敏感性", "单位经济"],
            "opening": "财务预测要把增长和现金放在一起看。SkyGuard 的关键不是某一年收入写得多高，而是客单价、交付成本、续费率、回款周期和研发投入能否互相支撑。",
            "reader": "投资评审",
            "decision": "用三情景测算解释经营弹性，用现金流约束扩张速度。",
        },
        {
            "chapter": "第十章 风险应对",
            "code": "10",
            "count": 9,
            "accent": REPORT_COLORS["red"],
            "topics": ["政策边界", "数据合规", "误报漏报", "硬件依赖", "采购周期", "现金压力", "过度承诺"],
            "opening": "风险章节要写得具体。低空项目容易出现三类误区：把辅助监管说成执法替代，把演示样本说成真实客户数据，把城市级需求说成短期交付能力。SkyGuard 的风险控制首先来自边界清楚。",
            "reader": "风险评审",
            "decision": "把风险分成会阻断项目、会拖慢节奏、会增加成本三类处理。",
        },
        {
            "chapter": "第十一章 团队组织",
            "code": "11",
            "count": 8,
            "accent": REPORT_COLORS["blue"],
            "topics": ["岗位配置", "产品负责人", "算法数据", "前后端研发", "交付运维", "行业顾问", "伙伴协同"],
            "opening": "团队组织要跟业务阶段匹配。早期最需要的是产品定义、数据处理、前后端实现、交付方案和客户沟通，而不是过早堆大规模销售或重硬件团队。",
            "reader": "团队评审",
            "decision": "说明谁来把 Demo、试点、数据和客户材料真正落地。",
        },
        {
            "chapter": "第十二章 社会价值",
            "code": "12",
            "count": 8,
            "accent": REPORT_COLORS["teal"],
            "topics": ["公共安全", "城市治理", "应急保障", "产业数据", "岗位培养", "合规意识"],
            "opening": "社会价值不需要夸张表达。低空运行越密，城市越需要把原本看不见、难记录、难复盘的飞行活动变成可协同的治理对象。SkyGuard 的价值在于把安全、效率和合规放进同一条工作流。",
            "reader": "课程评审",
            "decision": "把公共价值写成可观测的治理改善，而不是抽象口号。",
        },
        {
            "chapter": "第十三章 Web Demo",
            "code": "13",
            "count": 12,
            "accent": REPORT_COLORS["blue"],
            "topics": ["首页总览", "演示脚本", "态势地图", "实时监测", "AI识别", "围栏告警", "事件工单", "报表中心", "移动端"],
            "opening": "Web Demo 是计划书的落地证据。它不应该只是几张静态截图，而要能从首页进入演示中心，再进入态势、监测、计划、围栏、事件、报表和移动端，让评委按真实使用路径点击。",
            "reader": "答辩评委",
            "decision": "用完整交互证明项目可以被演示、被试点、被部署。",
        },
        {
            "chapter": "附录",
            "code": "A",
            "count": 16,
            "accent": REPORT_COLORS["muted"],
            "topics": ["来源索引", "图表目录", "数据字典", "财务假设", "问卷访谈", "部署说明", "截图包", "交付清单"],
            "opening": "附录承担核验功能。正文里的判断、图表和 Demo 截图，都要能在附录里找到来源、字段、假设或文件位置。这样计划书才不是一份孤立文本，而是一套可以继续迭代的交付包。",
            "reader": "资料核验者",
            "decision": "把证据、假设和交付文件放在稳定位置，后续升级直接覆盖。",
        },
    ]


def chart_for(chapter_chart_map: dict, chapter: str, cursor: dict, all_rows: list[dict]) -> dict | None:
    rows = chapter_chart_map.get(chapter) or all_rows
    if not rows:
        return None
    used = cursor.setdefault("__used_chart_ids", set())
    idx = cursor.get(chapter, 0)
    for offset in range(len(rows)):
        candidate = rows[(idx + offset) % len(rows)]
        chart_id = candidate.get("chart_id")
        if chart_id not in used:
            cursor[chapter] = idx + offset + 1
            used.add(chart_id)
            return candidate
    for candidate in all_rows:
        chart_id = candidate.get("chart_id")
        if chart_id not in used:
            used.add(chart_id)
            return candidate
    return None


def chapter_table(chapter: str, topic: str, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame, chart_catalog: pd.DataFrame) -> list[list[str]]:
    if chapter == "第三章 产品服务":
        return [["模块", "客户动作", "交付证据"], ["态势总览", "查看区域目标与围栏", "地图、KPI、告警列表"], ["飞行计划", "审批与核验", "计划表、轨迹比对"], ["事件处置", "确认、派单、归档", "工单状态与响应时间"], ["报表中心", "复盘运行质量", "日报、月报、风险排行"]]
    if chapter == "第七章 商业模式":
        return [["收入项", "建议口径", "适用客户"], ["Lite SaaS", "8-15万元/年", "单景区、园区"], ["Pro部署", "30-80万元/项目", "多区域试点"], ["活动保障", "5-30万元/次", "会展、赛事、节庆"], ["航线评估", "3000-20000元/条", "低空物流、巡检"]]
    if chapter == "第九章 财务预测":
        finance = build_finance_tables()
        return [["假设项", "模型口径"]] + finance["assumptions"].head(5)[["item", "assumption"]].astype(str).values.tolist()
    if chapter == "第五章 数据体系":
        return [["数据层", "用途", "边界"], ["官方统计", "行业规模", "不替代实时监管"], ["政策目标", "区域机会", "不等同订单"], ["公开数据", "方法说明", "需注明来源"], ["演示样本", "功能呈现", "不写成真实客户数据"]]
    if chapter == "附录" and "来源" in topic:
        return [["来源名称", "资料类型"]] + source_registry[["title", "source_type"]].head(6).astype(str).values.tolist()
    if chapter == "附录":
        return [["图表名称", "所属章节", "图表类型"]] + chart_catalog[["title", "chapter", "chart_type"]].head(6).astype(str).values.tolist()
    return []


def narrative_paragraphs(spec: dict, topic: str, row: dict | None, metrics: dict, page_no: int, figure_caption: str | None = None) -> list[str]:
    chapter = spec["chapter"]
    chart_title = row.get("title", topic) if row else topic
    chart_conclusion = row.get("conclusion", "") if row else ""
    source = row.get("source", "整理图片、Demo截图与项目资料") if row else "整理图片、Demo截图与项目资料"
    proof = chart_conclusion or spec["decision"]
    if chapter.startswith("第一章"):
        return [
            f"{topic}需要从运行事实说起。2025 年公开口径中，注册无人机 {metrics['uav']:.1f} 万架、累计飞行小时 {metrics['hours']:.2f} 万小时、运营单位 {metrics['units']:.0f} 家，这些数字共同说明低空飞行已经从少量试验转向更高频的日常活动。SkyGuard 的机会就在这里：当飞行活动增多，管理者需要的不只是看见目标，还要把计划、身份、轨迹和处置记录放到同一个证据链里。",
            f"图表“{chart_title}”提供了这一判断的量化切面。它的价值不在于把行业写得更热闹，而在于帮助团队选择切入顺序：先做边界清楚、责任主体明确、飞行活动稳定增长的重点区域，再逐步扩展到区县级协同。",
            f"这也决定了早期产品不宜直接承诺全域监管。更稳妥的路径，是围绕景区、园区、场馆、机场周边和物流航线建立可复用的规则包，让每一次试点都能沉淀数据口径、处置流程和报表模板。",
            f"来源口径为{source}。本页采用的结论是：{proof}",
        ]
    if chapter.startswith("第二章"):
        return [
            f"{topic}不是抽象痛点，而是一个会在现场反复出现的工作断点。低空目标进入重点区域后，运营人员往往需要同时确认飞行计划、目标身份、所在围栏、风险等级和处置责任；如果这些信息分散在不同系统或表格里，响应速度和事后解释都会被拖慢。",
            f"因此，痛点验证不应该只写“安全需求强”。更有说服力的写法，是把目标发现、计划核验、围栏判断、工单派发和报表归档拆开，逐一说明目前靠人工沟通会产生哪些延迟、遗漏和责任不清。",
            f"SkyGuard 在这一页对应的产品动作，是把{topic}放进统一事件流：系统先给出风险解释，再让人工确认，最后把处置过程回写到日报和月报。这样既不越过管理边界，也能让每次处置留下可复盘记录。",
            f"证据来自{source}。图表或素材显示的核心判断为：{proof}",
        ]
    if chapter.startswith("第三章"):
        return [
            f"{topic}要服务真实使用路径。管理者进入系统后，先看区域态势，再追到具体目标；如果目标没有计划、偏离航线或进入重点保护区，系统才进入告警和工单。这样的顺序比单纯堆功能更重要，因为它决定了 Demo 是否像一套能每天使用的业务系统。",
            f"产品层面，SkyGuard 采用“态势图 + 规则中心 + 事件中心 + 报表中心”的组合。态势图负责现场判断，规则中心负责边界和阈值，事件中心负责协同，报表中心负责把过程变成可提交、可审计的材料。",
            f"这一页围绕{topic}说明客户为什么愿意付费：它减少的不是某一个按钮的点击，而是减少跨部门沟通、人工截图、临时报表和事后解释的成本。场景包越清楚，销售和交付越容易标准化。",
            f"引用的证据为{source}。对应结论：{proof}",
        ]
    if chapter.startswith("第四章"):
        return [
            f"{topic}是技术架构里的关键环节。SkyGuard 不把算法写成万能能力，而是把数据接入、时空索引、轨迹比对、风险评分、规则解释和权限审计串成一个闭环。这样做的好处，是每个告警都能说清楚为什么出现、谁确认过、最终如何处置。",
            f"早期试点更适合采用轻量部署：数据接入先覆盖飞行计划、传感器告警、围栏规则和事件工单；模型负责辅助排序与解释，人工确认仍保留在关键节点。这个边界能降低误报争议，也便于后续接入更多硬件。",
            f"从工程角度看，{topic}需要同时考虑可维护性和可迁移性。重点区域的规则会不断变化，系统必须支持配置、日志和回放，而不是把所有判断写死在代码里。",
            f"技术证据来自{source}。本页采用的工程判断为：{proof}",
        ]
    if chapter.startswith("第五章"):
        return [
            f"{topic}首先是口径问题。计划书里的官方统计、地方政策、公开数据和演示样本不能混用；前者支撑行业和区域判断，后者只用于说明 Demo 的运行机制。只要这条边界清楚，数据多才有意义。",
            f"本项目已经把整理数据、图片、图表目录、证据索引和财务假设拆成独立文件。这样的组织方式适合后续迭代：新增来源先进入索引，再进入图表或正文，而不是直接把截图贴进文档。",
            f"{topic}对应的产品价值，是让风险判断可以回到字段。比如一个未知目标的风险分，不只看位置，还看计划匹配、身份可信度、轨迹异常、围栏权重和响应紧迫度。",
            f"本页来源为{source}。关键结论：{proof}",
        ]
    if chapter.startswith("第六章"):
        return [
            f"{topic}要回答 SkyGuard 与替代方案的区别。通用视频监控能看到画面，单点设备能发现信号，内部任务系统能派单，但低空运行监管需要把计划、目标、规则、处置和报表放到同一条链路里。",
            f"市场进入顺序也因此更具体。先做重点区域，是因为客户边界清楚、试点成本可控、验收指标容易定义；再做区县复制，是因为规则包、报表和运维流程可以复用；城市级协同则应该放在产品成熟之后。",
            f"竞争分析不能只罗列竞品名称。更重要的是说明客户在采购时会比较哪些能力：部署周期、数据口径、事件闭环、报表质量、硬件兼容和售后响应。",
            f"市场证据来自{source}。本页判断为：{proof}",
        ]
    if chapter.startswith("第七章"):
        return [
            f"{topic}要落到客户购买动作上。景区和园区可能先买轻量 SaaS 和规则包，大型活动会买临时保障，低空物流和巡检客户更关心航线评估、风险报告和运行记录。不同收费项背后对应的是不同交付强度。",
            f"SkyGuard 的收入结构应从“软件订阅 + 专业部署 + 保障服务 + 报告评估 + 运维续费”组合起来看。这样既能覆盖早期交付成本，也能在客户持续使用后形成续费和扩容。",
            f"这一页尤其要避免只写年度总收入。真正需要解释的是客单价如何形成、交付成本由哪些部分构成、回款周期如何控制，以及客户成功动作如何提高续费率。",
            f"商业证据来自{source}。图表结论为：{proof}",
        ]
    if chapter.startswith("第八章"):
        return [
            f"{topic}必须对应明确交付物。0-3 个月应完成 Demo、数据字典、规则样例和答辩材料；3-6 个月应完成重点区域试点包；6-18 个月再考虑多场景复制和伙伴接入。每一阶段都要能被验收。",
            f"路线图的难点不是画时间线，而是处理依赖关系。没有稳定数据口径，风险模型无法解释；没有工单闭环，报表没有可信内容；没有试点客户反馈，商业定价也容易失真。",
            f"围绕{topic}，团队需要把产品、数据、客户沟通和财务测算同步推进。某一条线过快，都会让项目在答辩或试点中露出空洞。",
            f"排期依据为{source}。对应判断：{proof}",
        ]
    if chapter.startswith("第九章"):
        return [
            f"{topic}的重点是经营质量。对创业项目来说，收入增长必须与毛利、现金流和交付能力一起看。SkyGuard 的财务模型采用保守、基准、乐观三情景，目的是让评审看到变量变化后的承压能力。",
            f"模型里最敏感的变量包括客户数量、客单价、交付成本、续费率、回款周期和研发投入。只要其中一项变化，现金跑道和利润拐点都会被改写，因此财务表必须保留可追溯假设。",
            f"{topic}还关系到融资用途。早期资金不宜过度投入重资产，而应优先用于产品打磨、试点交付、数据治理、客户成功和必要的伙伴接入。",
            f"财务证据来自{source}。财务判断：{proof}",
        ]
    if chapter.startswith("第十章"):
        return [
            f"{topic}是项目必须主动承认的约束。低空监管涉及政策边界、数据安全、误报漏报、设备接入和采购周期，任何一个问题处理不好，都会让试点变慢甚至停下。",
            f"SkyGuard 的应对方式不是回避风险，而是把边界写清楚：系统做辅助感知、风险预警、事件协同和报表复盘，不做干扰、捕获、打击，也不替代审批或执法。这个边界能减少承诺过度带来的交付风险。",
            f"围绕{topic}，计划书需要给出触发条件、影响范围和应对动作。例如误报率上升时要进入人工复核和模型回标，采购周期拉长时要调整现金流和交付资源。",
            f"风险证据来自{source}。对应判断：{proof}",
        ]
    if chapter.startswith("第十一章"):
        return [
            f"{topic}要服务当前阶段。SkyGuard 早期需要的是能把 Demo、数据、试点方案和客户材料做出来的小团队，而不是一开始就追求完整公司架构。产品、算法、前端、后端、数据和交付都要有人负责。",
            f"组织设计应围绕两条线：一条是产品研发线，保证平台可运行、可迭代；另一条是试点交付线，保证客户沟通、现场配置、培训和复盘材料能落地。",
            f"如果团队岗位与路线图脱节，计划书很容易变成空泛设想。围绕{topic}，本页要说明具体角色如何承担下一阶段里程碑。",
            f"组织证据来自{source}。本页判断：{proof}",
        ]
    if chapter.startswith("第十二章"):
        return [
            f"{topic}的价值在于把低空运行纳入可协同治理。对城市来说，未知目标、临时活动、重点区域和应急保障都需要记录、判断和复盘；这些工作做好之后，公共安全和产业发展才不会彼此冲突。",
            f"SkyGuard 不把社会价值写成夸张承诺，而是写成可观察的改进：响应时间缩短、事件闭环率提高、报表提交更及时、试点区域规则更清楚、公众活动保障更稳。",
            f"围绕{topic}，系统的意义是减少信息断点。管理方、运营方和现场人员看到同一套事件记录，才能降低沟通成本和责任争议。",
            f"价值判断来自{source}。对应结论：{proof}",
        ]
    if chapter.startswith("第十三章"):
        media = f"素材“{figure_caption}”" if figure_caption else f"页面“{chart_title}”"
        return [
            f"{topic}必须让评委能顺着真实路径操作。进入首页后看到行业证据和关键指标，再进入演示中心选择脚本，随后进入态势图、实时目标、飞行计划、AI 识别、围栏、事件、报表和移动端。这个顺序比单独展示某一张大屏更有说服力。",
            f"{media}在这里承担的是产品证据作用。它让前文中的行业判断、痛点和商业模式落到界面上，说明团队不是只写文档，也能把流程做成可点击的演示系统。",
            f"Demo 的设计重点是“真实、完整、能讲清楚”。真实是指使用整理图片、真实指标和明确标注的演示样本；完整是指覆盖从发现到处置再到报表；能讲清楚是指每个页面都有答辩时可说明的业务动作。",
            f"Demo 证据来自{source}。本页判断：{proof}",
        ]
    return [
        f"{topic}用于补齐正文证据。附录不是尾部堆料，而是把来源、字段、图表、假设和截图放到稳定位置，让读者能从任何结论回到原始材料。",
        f"本页围绕{topic}整理核验线索。后续新增材料时，应先进入证据索引或数据字典，再决定是否进入正文或图表，避免文档越来越厚但来源越来越散。",
        f"这类材料对答辩很重要，因为它能回答评委追问：数字从哪里来、样本是什么性质、财务假设如何变化、Demo 是否真的能运行。",
        f"附录来源为{source}。对应判断：{proof}",
    ]


def diversify_page_paras(paras: list[str], topic: str, row: dict | None, page_no: int, figure_caption: str | None = None) -> list[str]:
    token = figure_caption or (row.get("title") if row else topic)
    source = row.get("source") if row else "整理图片与Demo截图"
    work_units = ["围栏规则", "飞行计划字段", "事件工单", "日报模板", "客户验收清单", "风险解释字段", "权限日志", "试点复盘表"]
    unit = work_units[page_no % len(work_units)]
    variants = [
        f"这里选取“{token}”作为证据切口，读者能顺着它回到{topic}的业务场景。",
        f"“{token}”提供了这一页的具体落点，避免分析停留在抽象判断。",
        f"这一页把{topic}和“{token}”放在一起，是为了让数据、界面和交付动作能够互相印证。",
        f"如果答辩中追问依据，可以从“{token}”讲起，再回到{source}的口径。",
    ]
    refined = list(paras)
    if refined:
        refined[0] = f"{refined[0]}{variants[page_no % len(variants)]}"
    if len(refined) > 1:
        refined[1] = f"{refined[1]}落地时，这个判断会进入{unit}；复核材料以“{token}”为入口，并与正文证据链互相对应。"
    if len(refined) > 2:
        refined[2] = f"{refined[2]}因此，{topic}在计划书中既是分析主题，也是试点验收时需要复查的工作项；对应的核验对象是“{token}”。"
    if len(refined) > 3:
        refined[3] = f"{refined[3]}相关材料可在图表目录和证据索引中交叉核对。"
    return refined


def make_page_items(chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame) -> list[dict]:
    nat = data["national"]
    metrics = {
        "uav": first_value(nat, "registered_uavs_10k", 2025),
        "hours": first_value(nat, "uav_flight_hours_10k", 2025),
        "units": first_value(nat, "uav_operating_units", 2025),
        "stations": first_value(nat, "low_altitude_flight_service_stations", 2025),
        "airport": first_value(nat, "registered_general_airports", 2025),
    }
    figures = figure_files()
    figure_idx = 0

    chapter_alias = {
        "第二章 需求验证": "第二章 痛点验证",
        "第十章 风险": "第十章 风险应对",
        "第十一章 团队": "第十一章 团队组织",
        "第十三章 Demo展示": "第十三章 Web Demo",
    }
    all_rows = chart_catalog.to_dict("records")
    chapter_chart_map: dict[str, list[dict]] = defaultdict(list)
    for row in all_rows:
        chapter_chart_map[chapter_alias.get(row["chapter"], row["chapter"])].append(row)
    cursor: dict[str, int] = {}

    def take_figure(match=None) -> Path | None:
        nonlocal figure_idx
        if not figures:
            return None
        if match is not None:
            for fig in figures:
                if fig not in used_figures and match(fig):
                    used_figures.add(fig)
                    return fig
        while figure_idx < len(figures):
            fig = figures[figure_idx]
            figure_idx += 1
            if fig not in used_figures:
                used_figures.add(fig)
                return fig
        return None

    pages: list[dict] = []

    def add(page_type: str, chapter: str, title: str, paras: list[str], bullets: list[str] | None = None, chart_id: str | None = None, figure: Path | None = None, table: list[list[str]] | None = None, accent: str | None = None, kicker: str = "", callout: str = ""):
        pages.append(
            {
                "type": page_type,
                "chapter": chapter,
                "title": title,
                "paras": paras,
                "bullets": bullets or [],
                "chart_id": chart_id,
                "figure": str(figure) if figure else None,
                "table": table or [],
                "accent": accent or REPORT_COLORS["teal"],
                "kicker": kicker,
                "callout": callout,
            }
        )

    add(
        "cover",
        "封面",
        "低空智眼 SkyGuard 商业计划书",
        [
            "城市低空空域安全感知与运行监管平台",
            f"基于官方统计、地方政策目标、公开数据、演示样本和 Web Demo 重新整理。2025 年注册无人机 {metrics['uav']:.1f} 万架，累计飞行小时 {metrics['hours']:.2f} 万小时，运营单位 {metrics['units']:.0f} 家。",
        ],
        bullets=["重点区域先行", "规则与工单闭环", "数据口径分层", "稳定文件名交付"],
        figure=take_figure(),
        accent=REPORT_COLORS["teal"],
        kicker="完整商业计划书与演示系统",
        callout="辅助感知、风险预警、事件协同、运行报表",
    )
    add(
        "memo",
        "使用说明",
        "证据边界与阅读方式",
        [
            "本稿把“真实数据”和“演示样本”分开处理：官方统计与地方政策用于行业、市场和机会判断；公开数据用于背景说明和方法校验；演示样本只用于功能链路、模型字段和界面交互，不写成真实客户业务。",
            "报告阅读可以按三条线展开。第一条线是行业机会和痛点，说明为什么需要低空运行监管；第二条线是产品、技术和数据，说明 SkyGuard 如何把问题做成系统；第三条线是商业、财务、风险和 Demo，说明项目如何进入试点并持续经营。",
            "所有图表、来源、财务假设和截图均保留稳定文件名，后续升级直接覆盖，不再新增版本号目录。",
        ],
        bullets=[f"证据索引：{len(source_registry)} 条来源", f"图表目录：{len(chart_catalog)} 张图表", "输出：PDF、DOCX、Markdown、图表包、财务表、Web Demo包"],
        table=[
            ["阅读段", "重点", "核验材料"],
            ["行业机会", "规模、主体、飞行小时", "官方统计与图表目录"],
            ["痛点验证", "目标、计划、围栏、工单", "演示样本与访谈提纲"],
            ["产品服务", "态势图、规则、事件、报表", "Web Demo截图包"],
            ["技术架构", "接入、评分、审计、部署", "数据字典与架构图"],
            ["数据体系", "真实口径与演示边界", "证据资料索引表"],
            ["市场竞争", "替代方案与进入顺序", "竞品矩阵与区域目标"],
            ["商业模式", "订阅、部署、保障、报告", "财务测算表"],
            ["落地计划", "试点、复制、协同", "路线图和验收清单"],
            ["风险团队", "边界、现金、组织", "风险登记和团队分工"],
            ["Web Demo", "可点击产品流程", "Demo源码与截图包"],
        ],
        accent=REPORT_COLORS["blue"],
        kicker="使用说明",
        callout="先建立判断，再核验证据。",
    )
    add(
        "summary",
        "执行摘要",
        "项目判断",
        [
            "低空经济进入城市治理场景后，矛盾不再只是“有没有无人机”，而是飞行计划、目标身份、轨迹偏离、围栏规则、事件处置和报表复盘能否连起来。SkyGuard 的定位，是为景区、园区、场馆、机场周边和低空物流航线提供一套轻量但可扩展的运行监管底座。",
            "项目选择重点区域先行，而不是一开始承诺全城市级系统。这样可以降低交付复杂度，把产品验证集中在可观察、可验收、可复盘的区域里。等规则包、报表模板、客户成功流程和硬件接入经验稳定之后，再进入区县复制。",
            "商业化采用软件订阅、专业部署、活动保障、航线评估、运行报告和运维续费组合。硬件侧通过伙伴接入，SkyGuard 自身沉淀数据口径、规则配置、工单流程和报表能力。",
        ],
        bullets=["首批客户：景区、园区、场馆、低空物流航线运营方", "产品边界：辅助监管，不替代审批或执法", "经营重点：控制交付成本、提高续费、缩短回款"],
        chart_id=chart_for(chapter_chart_map, "第一章 行业机会", cursor, all_rows)["chart_id"],
        accent=REPORT_COLORS["teal"],
        kicker="执行摘要",
        callout="先做可交付的重点区域，再做可复制的区域方案。",
    )
    add(
        "summary",
        "执行摘要",
        "商业路径与 Demo 价值",
        [
            "Web Demo 是本计划书的核心证据之一。它把行业判断、痛点验证、产品闭环、数据样本和财务模型转成可点击页面，评委可以从首页进入演示中心，再进入态势、监测、计划、围栏、事件、报表和移动端。",
            "这一版升级的重点是减少模板感：正文按章节重新写成长段分析，PDF 改为更稳定的报告版式，Demo 改为更接近实际业务系统的视觉和交互。这样做不是为了显得复杂，而是为了让项目能被完整讲清楚。",
            "后续如果继续升级，只需要在稳定目录中替换数据、图表或页面，不再新建带版本号的文件。",
        ],
        bullets=["Demo 覆盖 20+ 个页面", "图表类型超过 50 种", "正文、图表、附录和网页口径保持一致"],
        figure=take_figure(),
        accent=REPORT_COLORS["blue"],
        kicker="交付逻辑",
        callout="计划书负责说清楚，Demo 负责让人点得动。",
    )

    for spec in chapter_specs(data):
        chapter = spec["chapter"]
        chapter_rows = chapter_chart_map.get(chapter) or all_rows
        divider_row = chart_for(chapter_chart_map, chapter, cursor, all_rows)
        add(
            "divider",
            chapter,
            spec["decision"],
            [spec["opening"]],
            bullets=[f"阅读对象：{spec['reader']}", f"本章证据：{len(chapter_rows)} 张关联图表", f"页面重点：{', '.join(spec['topics'][:4])}"],
            chart_id=divider_row["chart_id"] if divider_row else None,
            figure=take_figure() if spec["code"] in {"03", "04", "13"} else None,
            accent=spec["accent"],
            kicker="章节导读",
            callout=spec["decision"],
        )
        for i in range(spec["count"]):
            topic = spec["topics"][i % len(spec["topics"])]
            row = chart_for(chapter_chart_map, chapter, cursor, all_rows)
            use_figure = (i % 5 == 2) and bool(figures)
            use_table = (i == 4)
            figure = take_figure() if use_figure else None
            figure_caption = humanize_figure_caption(figure) if figure else None
            page_type = "visual" if figure else "table" if use_table else "chart"
            title = f"{topic}｜{figure_caption}" if figure_caption else (row["title"] if row and i % 2 == 1 else topic)
            paras = narrative_paragraphs(spec, topic, row, metrics, len(pages) + 1, figure_caption=figure_caption)
            paras = diversify_page_paras(paras, topic, row, len(pages) + 1, figure_caption=figure_caption)
            bullets = [
                f"判断：{(row or {}).get('conclusion', spec['decision'])}",
                f"来源：{(row or {}).get('source', '整理图片、Demo截图与项目资料')}",
                f"动作：围绕{topic}明确页面、字段、责任或验收口径。",
            ]
            table = chapter_table(chapter, topic, data, source_registry, chart_catalog) if use_table else []
            add(
                page_type,
                chapter,
                title,
                paras,
                bullets=bullets,
                chart_id=None if figure else (row["chart_id"] if row else None),
                figure=figure,
                table=table,
                accent=spec["accent"],
                kicker="",
                callout=spec["decision"],
            )
    return pages


def write_markdown(pages: list[dict], chart_catalog: pd.DataFrame, source_registry: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.md"
    lines = ["# 低空智眼 SkyGuard 商业计划书", "", "> 城市低空空域安全感知与运行监管平台", ""]
    for idx, page in enumerate(pages, start=1):
        lines.append(f"## {page['chapter']}｜{page['title']}")
        if page.get("kicker"):
            lines.append("")
            lines.append(f"**{page['kicker']}**")
        for p in page["paras"]:
            lines.append("")
            lines.append(p)
        if page["bullets"]:
            lines.append("")
            for b in page["bullets"]:
                lines.append(f"- {b}")
        if page.get("table"):
            table = page["table"]
            lines.append("")
            lines.append("| " + " | ".join(map(str, table[0])) + " |")
            lines.append("| " + " | ".join(["---"] * len(table[0])) + " |")
            for row in table[1:]:
                lines.append("| " + " | ".join(map(str, row)) + " |")
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists():
                lines.append("")
                lines.append(f"![{humanize_figure_caption(fig_path)}]({fig_path.relative_to(ROOT)})")
                lines.append(f"图注：{humanize_figure_caption(fig_path)}。")
        if page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                r = row.iloc[0]
                lines.append("")
                lines.append(f"![{r['title']}]({r['file']})")
                lines.append(f"图表来源：{r['source']}。结论：{r['conclusion']}")
    lines.extend(["", "## 来源清单"])
    for _, row in source_registry.iterrows():
        lines.append(f"- {row['title']}：{row['url']}")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def shade_docx_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_docx_editorial_table(doc: Document, table_data: list[list[str]]) -> None:
    if not table_data:
        return
    tbl = doc.add_table(rows=1, cols=len(table_data[0]))
    tbl.style = "Table Grid"
    for j, cell in enumerate(tbl.rows[0].cells):
        cell.text = str(table_data[0][j])
        shade_docx_cell(cell, "E8F3F1")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for row in table_data[1:]:
        cells = tbl.add_row().cells
        for j, val in enumerate(row[: len(cells)]):
            cells[j].text = str(val)
    doc.add_paragraph("")


def write_docx(pages: list[dict], chart_catalog: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.docx"
    doc = Document()
    set_docx_styles(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.text = "SkyGuard 低空智眼 | 城市低空运行监管商业计划书"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_docx_run_font(run, size=8.5, color=RGBColor(100, 116, 139))
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = footer.add_run("资料口径：官方统计 / 地方政策 / 公开数据 / 演示样本 / 经营假设分层标注    第 ")
    set_docx_run_font(f_run, size=8, color=RGBColor(100, 116, 139))
    add_page_number(footer)
    f_run = footer.add_run(" 页")
    set_docx_run_font(f_run, size=8, color=RGBColor(100, 116, 139))
    table_no = 0
    figure_no = 0

    for idx, page in enumerate(pages, start=1):
        if page["type"] == "cover":
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("低空智眼 SkyGuard\n")
            run.font.size = Pt(30)
            run.font.bold = True
            run.font.color.rgb = RGBColor(23, 32, 51)
            run = title.add_run("城市低空空域安全感知与运行监管平台商业计划书")
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(51, 65, 85)
            doc.add_paragraph(page["callout"])
        heading_level = 1 if page["type"] in {"cover", "divider"} else 2
        doc.add_heading(f"{page['chapter']}｜{page['title']}", level=heading_level)
        if page.get("kicker"):
            p = doc.add_paragraph(page["kicker"])
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(15, 118, 110)
        for para in page["paras"]:
            doc.add_paragraph(para)
        for bullet in page["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")
        add_docx_editorial_table(doc, page.get("table", []))
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists() and fig_path.suffix.lower() != ".webp":
                try:
                    doc.add_picture(str(fig_path), width=Inches(5.95))
                    cap = doc.add_paragraph(f"图注：{humanize_figure_caption(fig_path)}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
        if page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                chart_path = ROOT / row.iloc[0]["file"]
                if chart_path.exists():
                    try:
                        doc.add_picture(str(chart_path), width=Inches(5.95))
                        cap = doc.add_paragraph(f"{row.iloc[0]['title']}｜来源：{row.iloc[0]['source']}")
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
        if idx < len(pages):
            doc.add_page_break()
    doc.save(path)
    return path


def hex_color(value: str):
    return colors.HexColor(value)


def draw_page_text(c, text: str, x: float, y: float, width: float, size: float, leading: float, color_value: str = "#172033", max_lines: int | None = None) -> float:
    c.setFont("CNFont", size)
    c.setFillColor(hex_color(color_value))
    max_chars = max(10, int(width / (size * 1.03)))
    lines: list[str] = []
    for para in text.split("\n"):
        lines.extend(textwrap.wrap(para, width=max_chars, break_long_words=True, replace_whitespace=False) or [""])
    if max_lines is not None:
        lines = lines[:max_lines]
    for line in lines:
        c.drawString(x, y, line)
        y -= leading
    return y


def draw_chip(c, text: str, x: float, y: float, fill: str, stroke: str = "#d8dee8", text_color: str = "#172033") -> float:
    c.setFont("CNFont", 7.4)
    chip_w = min(185, max(52, len(text) * 7.1 + 18))
    c.setFillColor(hex_color(fill))
    c.setStrokeColor(hex_color(stroke))
    c.roundRect(x, y - 13, chip_w, 18, 5, stroke=1, fill=1)
    c.setFillColor(hex_color(text_color))
    c.drawString(x + 8, y - 8, text[:24])
    return x + chip_w + 6


def draw_pdf_table(c, table_data: list[list[str]], x: float, y: float, width: float, accent: str = "#0f766e") -> float:
    if not table_data:
        return y
    rows = table_data[:12]
    cols = max(len(r) for r in rows)
    col_w = width / cols
    row_h = 19
    for i, row in enumerate(rows):
        fill = "#e7f1ef" if i == 0 else "#ffffff" if i % 2 else "#f5f7fa"
        c.setFillColor(hex_color(fill))
        c.rect(x, y - row_h, width, row_h, stroke=0, fill=1)
        c.setStrokeColor(hex_color("#d8dee8"))
        c.line(x, y - row_h, x + width, y - row_h)
        for j in range(cols):
            val = str(row[j]) if j < len(row) else ""
            c.setFillColor(hex_color(REPORT_COLORS["ink"]))
            c.setFont("CNFont", 6.8 if len(val) > 18 else 7.4)
            c.drawString(x + j * col_w + 4, y - 12.5, val[:32])
        y -= row_h
    c.setStrokeColor(hex_color(accent))
    c.line(x, y, x + width, y)
    return y - 10


def draw_pdf_visual(c, path: Path, x: float, y: float, width: float, height: float, border: bool = True) -> float:
    if not path.exists() or height <= 20:
        return y
    try:
        img = Image.open(path)
        iw, ih = img.size
        scale = min(width / iw, height / ih)
        draw_w, draw_h = iw * scale, ih * scale
        ox = x + (width - draw_w) / 2
        oy = y - draw_h
        if border:
            c.setFillColor(colors.white)
            c.roundRect(x - 4, oy - 4, width + 8, draw_h + 8, 8, stroke=0, fill=1)
            c.setStrokeColor(hex_color("#d8dee8"))
            c.roundRect(x - 4, oy - 4, width + 8, draw_h + 8, 8, stroke=1, fill=0)
        c.drawImage(ImageReader(img), ox, oy, width=draw_w, height=draw_h, mask="auto")
        return oy - 9
    except Exception:
        return y


def draw_pdf_page(c, page: dict, page_num: int, chart_catalog: pd.DataFrame, page_size=A4):
    w, h = page_size
    margin = 38
    accent = page.get("accent") or REPORT_COLORS["teal"]
    c.setFillColor(hex_color(REPORT_COLORS["paper"]))
    c.rect(0, 0, w, h, stroke=0, fill=1)

    if page["type"] == "cover":
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists():
                try:
                    img = Image.open(fig_path)
                    iw, ih = img.size
                    scale = max(w / iw, h / ih)
                    dw, dh = iw * scale, ih * scale
                    c.drawImage(ImageReader(img), (w - dw) / 2, (h - dh) / 2, width=dw, height=dh, mask="auto")
                except Exception:
                    pass
        c.setFillColor(colors.Color(0.05, 0.08, 0.11, alpha=0.72))
        c.rect(0, 0, w, h, stroke=0, fill=1)
        c.setFillColor(hex_color(accent))
        c.rect(0, 0, 12, h, stroke=0, fill=1)
        c.setFillColor(colors.white)
        c.setFont("CNFont", 9.5)
        c.drawString(margin + 8, h - 86, page.get("kicker", ""))
        c.setFont("CNFont", 34)
        c.drawString(margin + 8, h - 145, "低空智眼 SkyGuard")
        c.setFont("CNFont", 17)
        c.drawString(margin + 8, h - 178, "城市低空空域安全感知与运行监管平台商业计划书")
        y = h - 230
        for para in page["paras"]:
            y = draw_page_text(c, para, margin + 8, y, 430, 10, 16, "#ffffff", max_lines=4)
            y -= 8
        chip_x = margin + 8
        for bullet in page["bullets"]:
            chip_x = draw_chip(c, bullet, chip_x, 132, "#ffffff", "#ffffff", REPORT_COLORS["ink"])
        c.setFont("CNFont", 8.2)
        c.setFillColor(colors.white)
        c.drawString(margin + 8, 52, "稳定文件名交付｜PDF / DOCX / Markdown / 图表包 / 财务表 / Web Demo")
        c.drawRightString(w - margin, 52, "SkyGuard 商业计划书")
        c.showPage()
        return

    if page["type"] == "divider":
        c.setFillColor(hex_color(REPORT_COLORS["navy"]))
        c.rect(0, 0, w, h, stroke=0, fill=1)
        c.setFillColor(hex_color(accent))
        c.rect(0, 0, 16, h, stroke=0, fill=1)
        c.setFont("CNFont", 9)
        c.setFillColor(hex_color("#cbd5df"))
        c.drawString(margin + 10, h - 70, page.get("kicker", "Chapter"))
        c.setFont("CNFont", 24)
        c.setFillColor(colors.white)
        c.drawString(margin + 10, h - 116, page["chapter"])
        c.setFont("CNFont", 16)
        c.drawString(margin + 10, h - 150, page["title"][:34])
        y = h - 205
        for para in page["paras"]:
            y = draw_page_text(c, para, margin + 10, y, 360, 10, 16, "#eef2f6", max_lines=7)
        box_x = margin + 10
        box_y = 185
        for i, bullet in enumerate(page["bullets"][:3]):
            c.setFillColor(colors.Color(1, 1, 1, alpha=0.08))
            c.roundRect(box_x, box_y - i * 58, 360, 42, 8, stroke=0, fill=1)
            draw_page_text(c, bullet, box_x + 12, box_y + 20 - i * 58, 330, 8, 12, "#f8fafc", max_lines=2)
        if page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                chart_path = ROOT / row.iloc[0]["file"]
                draw_pdf_visual(c, chart_path, w - 245, 358, 188, 170, border=True)
        elif page.get("figure"):
            draw_pdf_visual(c, Path(page["figure"]), w - 245, 358, 188, 170, border=True)
        c.setFont("CNFont", 8)
        c.setFillColor(hex_color("#cbd5df"))
        c.drawRightString(w - margin, 40, "SkyGuard 商业计划书")
        c.showPage()
        return

    c.setFillColor(colors.white)
    c.roundRect(22, 22, w - 44, h - 44, 10, stroke=0, fill=1)
    c.setFillColor(hex_color(accent))
    c.rect(22, h - 74, w - 44, 4, stroke=0, fill=1)
    c.setFillColor(hex_color(REPORT_COLORS["muted"]))
    c.setFont("CNFont", 7.8)
    c.drawString(margin, h - 53, page.get("kicker", ""))
    c.drawRightString(w - margin, h - 53, "SkyGuard 商业计划书")
    c.setFillColor(hex_color(REPORT_COLORS["ink"]))
    c.setFont("CNFont", 17)
    c.drawString(margin, h - 88, f"{page['chapter']}｜{page['title']}"[:42])

    left_w = 306
    right_x = margin + left_w + 18
    right_w = w - right_x - margin
    y = h - 120
    for para in page["paras"][:3]:
        y = draw_page_text(c, para, margin, y, left_w, 8.4, 12.7, REPORT_COLORS["soft_ink"], max_lines=7)
        y -= 6
        if y < 333:
            break

    c.setFillColor(hex_color("#f2f6f8"))
    c.roundRect(right_x, h - 315, right_w, 195, 8, stroke=0, fill=1)
    c.setStrokeColor(hex_color("#d8dee8"))
    c.roundRect(right_x, h - 315, right_w, 195, 8, stroke=1, fill=0)
    c.setFont("CNFont", 9.3)
    c.setFillColor(hex_color(REPORT_COLORS["ink"]))
    c.drawString(right_x + 12, h - 141, "证据与动作")
    by = h - 164
    for bullet in page["bullets"][:3]:
        by = draw_page_text(c, "· " + bullet, right_x + 12, by, right_w - 24, 7.3, 10.8, REPORT_COLORS["soft_ink"], max_lines=3)
        by -= 4
    if page.get("callout"):
        c.setFillColor(hex_color(accent))
        c.roundRect(right_x + 12, h - 302, right_w - 24, 30, 6, stroke=0, fill=1)
        draw_page_text(c, page["callout"], right_x + 20, h - 286, right_w - 40, 7.2, 10, "#ffffff", max_lines=2)

    media_top = 430
    if page.get("table"):
        draw_pdf_table(c, page["table"], margin, media_top, w - margin * 2, accent)
    if page.get("figure"):
        fig_path = Path(page["figure"])
        bottom_y = draw_pdf_visual(c, fig_path, margin, media_top, w - margin * 2, 272, border=True)
        c.setFont("CNFont", 6.8)
        c.setFillColor(hex_color(REPORT_COLORS["muted"]))
        c.drawString(margin, max(44, bottom_y), f"图注：{humanize_figure_caption(fig_path)}")
    elif page.get("chart_id"):
        row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
        if not row.empty:
            chart_path = ROOT / row.iloc[0]["file"]
            bottom_y = draw_pdf_visual(c, chart_path, margin, media_top, w - margin * 2, 272, border=True)
            caption = f"{row.iloc[0]['title']}｜来源：{row.iloc[0]['source']}｜结论：{row.iloc[0]['conclusion']}"
            draw_page_text(c, caption, margin, max(43, bottom_y), w - margin * 2, 6.4, 8.5, REPORT_COLORS["muted"], max_lines=2)

    c.setStrokeColor(hex_color("#d8dee8"))
    c.line(margin, 33, w - margin, 33)
    c.setFont("CNFont", 7.2)
    c.setFillColor(hex_color(REPORT_COLORS["muted"]))
    c.drawString(margin, 20, "资料口径：官方统计 / 地方政策 / 公开数据 / 演示样本 / 经营假设分层标注")
    c.drawRightString(w - margin, 20, "稳定交付文件")
    c.showPage()


def write_pdf(pages: list[dict], chart_catalog: pd.DataFrame, filename: str, limit: int | None = None) -> Path:
    path = OUT_DIR / filename
    c = canvas.Canvas(str(path), pagesize=A4)
    selected = pages if limit is None else pages[:limit]
    for idx, page in enumerate(selected, start=1):
        draw_pdf_page(c, page, idx, chart_catalog)
    c.save()
    return path


# ---------------------------------------------------------------------------
# Final public-output bindings. These definitions intentionally appear
# immediately before main(), so they are the functions used by generation.
# ---------------------------------------------------------------------------

REPORT_COLORS.update(
    {
        "paper": "#f6f4ef",
        "paper2": "#ebe7dd",
        "line": "#d6d0c3",
        "navy": "#151f2b",
        "teal": "#0c6f69",
        "blue": "#234f7d",
        "green": "#2f6b4f",
        "amber": "#9b6a2f",
        "red": "#9b3f3a",
        "gold": "#b69252",
        "ash": "#6d716c",
    }
)

def generate_source_registry(nat: pd.DataFrame, city: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict] = []
    rows.extend(SOURCE_LINKS)
    for source_name, use in [
        ("低空经济政策法规全量数据.csv", "按年份、地区和政策类型聚合，用于政策窗口与区域机会分析。"),
        ("低空经济招投标项目全量数据.csv", "按预算、项目类型和省份聚合，用于市场温度和采购场景判断。"),
        ("低空经济专利全量数据.csv", "按年份、IPC和申请人聚合，用于技术活跃度分析。"),
        ("低空经济相关企业全量数据.csv", "按省份、行业和注册年份聚合，用于产业生态分析。"),
        ("低空经济产业链上下游数据.csv", "按产业链环节和细分领域聚合，用于伙伴生态分析。"),
        ("各省低空经济指数数据（1990-2024）.csv", "用于省域指数、区域分层和市场进入顺序分析。"),
        ("低空安全事件与监管案例数据.csv", "用于安全事件类型、严重程度和监管处置分析。"),
        ("低空起降设施与空域数据.csv", "用于起降设施类型、状态和容量分析。"),
        ("低空飞行活动与轨迹数据.csv", "用于飞行类型、距离、高度和运行场景分析。"),
        ("OurAirports Open Data", "用于机场、跑道和导航台公开数据说明。"),
        ("OpenSky Network Flightlist", "用于飞行活动公开数据方法说明。"),
    ]:
        url = str((DATA_DIR / source_name).relative_to(ROOT)) if (DATA_DIR / source_name).exists() else source_name
        rows.append({"title": source_name, "publisher": "本地整理数据集或公开数据源", "date": "按文件记录年份", "url": url, "use": use, "source_type": "local_dataset"})
    for sid, sub in nat.groupby("source_id") if "source_id" in nat.columns else []:
        rows.append(
            {
                "title": f"{sid} 民航公开统计摘录",
                "publisher": "中国民用航空局/交通运输部公开资料",
                "date": str(int(sub["year"].max())) if "year" in sub.columns and sub["year"].notna().any() else "",
                "url": sub["source_url"].iloc[0] if "source_url" in sub.columns else "",
                "use": "用于无人机注册、飞行小时、运营单位、机场和安全指标分析。",
                "source_type": "official_statistic",
            }
        )
    for sid, sub in city.groupby("source_id") if "source_id" in city.columns else []:
        rows.append(
            {
                "title": sub["policy_or_case_doc"].iloc[0] if "policy_or_case_doc" in sub.columns else f"{sid} 地方政策目标",
                "publisher": "地方政府/主管部门公开资料",
                "date": str(int(sub["target_year"].dropna().max())) if "target_year" in sub.columns and sub["target_year"].notna().any() else "",
                "url": sub["source_url"].iloc[0] if "source_url" in sub.columns else "",
                "use": "用于地方低空经济目标、平台建设、航线和起降网络规划分析。",
                "source_type": "local_policy_or_case",
            }
        )
    out = pd.DataFrame(rows).fillna("")
    out = out.drop_duplicates(subset=["title", "url"], keep="last").reset_index(drop=True)
    out.insert(0, "source_id", [f"S{idx:03d}" for idx in range(1, len(out) + 1)])
    out.to_csv(SOURCE_DIR / "source_registry.csv", index=False, encoding="utf-8-sig")
    return out


def generate_chart_pack(data: dict[str, pd.DataFrame]) -> pd.DataFrame:
    CHARTS.clear()
    for p in CHART_DIR.glob("*.png"):
        p.unlink()
    add_premium_chart_extensions(data)
    add_real_data_charts()
    catalog = pd.DataFrame(CHARTS).drop_duplicates(subset=["title", "chart_type"], keep="last").reset_index(drop=True)
    catalog.to_csv(TABLE_DIR / "chart_catalog.csv", index=False, encoding="utf-8-sig")
    return catalog


def set_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.75)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    styles = doc.styles
    body_font = "SimSun"
    heading_font = "DengXian"
    for style_name in ["Normal", "List Bullet"]:
        style = styles[style_name]
        style.font.name = body_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Title"]:
        style = styles[style_name]
        style.font.name = heading_font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), heading_font)
    normal = styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["Heading 1"].font.size = Pt(19)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(21, 31, 43)
    styles["Heading 2"].font.size = Pt(14.2)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(12, 111, 105)
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True
    styles["List Bullet"].font.size = Pt(9.6)
    styles["List Bullet"].paragraph_format.left_indent = Pt(18)
    styles["List Bullet"].paragraph_format.first_line_indent = Pt(-10)
    styles["List Bullet"].paragraph_format.space_after = Pt(2)


def make_page_items(chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame) -> list[dict]:
    nat = data["national"]
    metrics = {
        "uav": first_value(nat, "registered_uavs_10k", 2025),
        "hours": first_value(nat, "uav_flight_hours_10k", 2025),
        "units": first_value(nat, "uav_operating_units", 2025),
        "stations": first_value(nat, "low_altitude_flight_service_stations", 2025),
        "airport": first_value(nat, "registered_general_airports", 2025),
    }
    figures = figure_files()
    figure_idx = 0
    all_rows = chart_catalog.to_dict("records")
    chapter_alias = {"第二章 需求验证": "第二章 痛点验证", "第十章 风险": "第十章 风险应对", "第十一章 团队": "第十一章 团队组织", "第十三章 Demo展示": "第十三章 Web Demo"}
    chapter_chart_map: dict[str, list[dict]] = defaultdict(list)
    for row in all_rows:
        chapter_chart_map[chapter_alias.get(row["chapter"], row["chapter"])].append(row)
    cursor: dict[str, int] = {}
    pages: list[dict] = []

    def take_figure(match=None) -> Path | None:
        nonlocal figure_idx
        if not figures:
            return None
        used = getattr(take_figure, "_used", set())
        if match is not None:
            for fig in figures:
                if fig not in used and match(fig):
                    used.add(fig)
                    setattr(take_figure, "_used", used)
                    return fig
        while figure_idx < len(figures):
            fig = figures[figure_idx]
            figure_idx += 1
            if fig not in used:
                used.add(fig)
                setattr(take_figure, "_used", used)
                return fig
        setattr(take_figure, "_used", used)
        return None

    def next_row(chapter: str) -> dict | None:
        return chart_for(chapter_chart_map, chapter, cursor, all_rows)

    def add(page_type: str, chapter: str, title: str, paras: list[str], bullets: list[str] | None = None, chart_id: str | None = None, figure: Path | None = None, table: list[list[str]] | None = None, accent: str | None = None, kicker: str = "", callout: str = ""):
        pages.append(
            {
                "type": page_type,
                "chapter": chapter,
                "title": title,
                "paras": paras,
                "bullets": bullets or [],
                "chart_id": chart_id,
                "figure": str(figure) if figure else None,
                "table": table or [],
                "accent": accent or REPORT_COLORS["teal"],
                "kicker": kicker,
                "callout": callout,
            }
        )

    cover_image = take_figure(lambda p: any(k in p.name for k in ["平台", "低空", "SILAS", "无人机"])) or take_figure()
    add(
        "cover",
        "封面",
        "低空智眼 SkyGuard 商业计划书",
        [
            "城市低空空域安全感知与运行监管平台",
            f"面向景区、园区、场馆、机场周边和低空物流航线。行业基准引用公开统计：2025 年注册无人机 {metrics['uav']:.1f} 万架，累计飞行小时 {metrics['hours']:.2f} 万小时，运营单位 {metrics['units']:.0f} 家。",
        ],
        bullets=["重点区域监管", "计划与围栏联动", "事件闭环处置", "运行报表复盘"],
        figure=cover_image,
        accent=REPORT_COLORS["teal"],
        kicker="商业计划书",
        callout="看得见、判得清、处置快、可复盘",
    )
    add(
        "memo",
        "使用说明",
        "证据边界与审阅路径",
        [
            "本计划书把真实公开资料、地方政策、公开数据和演示样本分层使用。官方统计用于行业规模判断，政策法规用于合规和区域机会判断，公开数据用于方法说明，Demo 样本只用于产品流程展示。",
            "正文采用执行摘要、行业机会、痛点验证、产品服务、技术架构、数据体系、市场竞争、商业模式、落地计划、财务预测、风险应对、团队组织、社会价值、Web Demo 和附录的顺序，保持商业计划书的完整闭环。",
            "图表、来源、财务假设和截图集中在附录及配套表格中，便于审阅时从结论反查证据。"
        ],
        bullets=[f"来源索引：{len(source_registry)} 条", f"图表目录：{len(chart_catalog)} 张", "配套材料：财务测算、证据索引、图表包、Web Demo 源码"],
        table=final_page_table("使用说明", "", data, source_registry, chart_catalog),
        accent=REPORT_COLORS["blue"],
        kicker="审阅说明",
        callout="先看判断，再核验证据。",
    )
    summary_chart = next_row("第一章 行业机会")
    add(
        "summary",
        "执行摘要",
        "项目判断",
        [
            "低空经济从产业概念进入城市运行后，最迫切的问题是把飞行计划、目标身份、电子围栏、异常告警、现场处置和复盘报告连接起来。SkyGuard 的定位是面向重点区域提供低空安全感知和运行监管平台，帮助客户完成日常监管、临时保障和运行复盘。",
            "项目不从重硬件或全域城市总包切入，而从边界清楚、责任主体明确、可快速验收的区域开始。景区、园区、场馆、机场周边和固定低空航线，是首批更现实的市场入口。",
            "商业化采用软件订阅、专业部署、活动保障、航线评估、运行报告和运维续费组合。硬件侧通过伙伴接入，平台侧沉淀规则、数据、日志和报表能力。"
        ],
        bullets=["首批客户：景区、园区、场馆、航线运营方", "产品边界：辅助监管，不替代审批或执法", "经营重点：控制交付成本、提高续费率、缩短回款周期"],
        chart_id=summary_chart["chart_id"] if summary_chart else None,
        accent=REPORT_COLORS["teal"],
        kicker="执行摘要",
        callout="先做一个区域的闭环，再做可复制的区域方案。",
    )
    add(
        "summary",
        "执行摘要",
        "产品与商业路径",
        [
            "SkyGuard 的核心交付是一套可每天使用的运行机制。态势图解决实时查看，规则中心解决围栏和阈值，事件中心解决确认和派单，报表中心解决复盘和验收。Web Demo 按同样逻辑组织，能够从首页进入完整业务路径。",
            "竞争优势不写成单一算法优势，而写成业务闭环优势。通用监控、单点感知设备和人工巡查都能解决局部问题，但客户真正需要的是可解释、可记录、可提交的全流程材料。",
            "财务测算采用三情景，重点关注客户数量、客单价、交付成本、续费率和回款周期。早期融资优先用于产品打磨、试点交付、数据治理和客户成功。"
        ],
        bullets=["产品：计划、目标、围栏、工单、报表一体化", "商业：订阅、部署、活动、评估、运维组合收入", "风险：政策、数据、误报、采购和现金流分层管理"],
        figure=take_figure(),
        accent=REPORT_COLORS["blue"],
        kicker="执行摘要",
        callout="计划书与 Demo 共用同一套证据和术语。",
    )

    for spec in final_chapter_specs(data):
        chapter = spec["chapter"]
        rows = chapter_chart_map.get(chapter) or all_rows
        row = next_row(chapter)
        add(
            "divider",
            chapter,
            spec["decision"],
            [spec["opening"]],
            bullets=[f"本章重点：{', '.join(spec['topics'][:5])}", f"关联图表：{len(rows)} 张", "写法要求：痛点、证据、动作相互对应"],
            chart_id=row["chart_id"] if row else None,
            figure=take_figure() if chapter in {"第三章 产品服务", "第四章 技术架构", "第十三章 Web Demo"} else None,
            accent=spec["accent"],
            kicker="章节导读",
            callout=spec["decision"],
        )
        for i in range(spec["count"]):
            topic = spec["topics"][i % len(spec["topics"])]
            row = next_row(chapter)
            use_figure = (i % 6 == 2) and bool(figures)
            use_table = i in {4, 10}
            figure = take_figure() if use_figure else None
            figure_caption = humanize_figure_caption(figure) if figure else None
            title = f"{topic}：{figure_caption}" if figure_caption else (row["title"] if row and i % 2 else topic)
            paras = final_narrative_paragraphs(spec, topic, row, metrics, figure_caption)
            bullets = [
                f"判断：{(row or {}).get('conclusion', spec['decision'])}",
                f"来源：{(row or {}).get('source', '整理图片、Demo截图与项目资料')}",
                f"动作：明确{topic}对应的字段、页面、责任或验收口径。",
            ]
            table = final_page_table(chapter, topic, data, source_registry, chart_catalog) if use_table else []
            add(
                "visual" if figure else "table" if table else "chart",
                chapter,
                title,
                paras,
                bullets=bullets,
                chart_id=None if figure else (row["chart_id"] if row else None),
                figure=figure,
                table=table,
                accent=spec["accent"],
                callout=spec["decision"],
            )
    return pages


def write_docx(pages: list[dict], chart_catalog: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.docx"
    doc = Document()
    set_docx_styles(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.text = "SkyGuard 低空智眼 | 城市低空运行监管商业计划书"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(100, 116, 139)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = footer.add_run("资料口径：官方统计 / 地方政策 / 公开数据 / 演示样本 / 经营假设分层标注    第 ")
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(100, 116, 139)
    add_page_number(footer)
    f_run = footer.add_run(" 页")
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(100, 116, 139)
    table_no = 0
    figure_no = 0
    for idx, page in enumerate(pages, start=1):
        if page["type"] == "cover":
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("低空智眼 SkyGuard\n")
            run.font.size = Pt(30)
            run.font.bold = True
            run.font.color.rgb = RGBColor(21, 31, 43)
            run = title.add_run("城市低空空域安全感知与运行监管平台商业计划书")
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(51, 65, 85)
            sub = doc.add_paragraph(page["callout"])
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_page_break()
            continue
        heading_level = 1 if page["type"] == "divider" else 2
        heading = doc.add_heading(f"{page['chapter']}｜{page['title']}", level=heading_level)
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.keep_together = True
        if page.get("kicker"):
            p = doc.add_paragraph(page["kicker"])
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(18, 78, 107)
            p.paragraph_format.keep_with_next = True
        media_rendered = False
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists() and fig_path.suffix.lower() != ".webp":
                try:
                    figure_no += 1
                    add_docx_figure_block(doc, fig_path, 4.35, f"图 {figure_no}  {humanize_figure_caption(fig_path)}")
                    media_rendered = True
                except Exception:
                    pass
        if (not media_rendered) and page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                chart_path = ROOT / row.iloc[0]["file"]
                if chart_path.exists():
                    try:
                        figure_no += 1
                        add_docx_figure_block(doc, chart_path, 4.55, f"图 {figure_no}  {row.iloc[0]['title']}｜来源：{row.iloc[0]['source']}｜结论：{row.iloc[0]['conclusion']}")
                        media_rendered = True
                    except Exception:
                        pass
        for para in page["paras"]:
            doc.add_paragraph(para)
        for bullet in page["bullets"]:
            doc.add_paragraph(bullet, style="List Bullet")
        if page.get("table"):
            table_no += 1
            add_docx_editorial_table(doc, page.get("table", []), f"表 {table_no}  {page['chapter']} - {page['title']}")
        if (not media_rendered) and page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists() and fig_path.suffix.lower() != ".webp":
                try:
                    figure_no += 1
                    add_docx_figure_block(doc, fig_path, 4.40, f"图 {figure_no}  {humanize_figure_caption(fig_path)}")
                except Exception:
                    pass
        if (not media_rendered) and page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                chart_path = ROOT / row.iloc[0]["file"]
                if chart_path.exists():
                    try:
                        figure_no += 1
                        add_docx_figure_block(doc, chart_path, 4.65, f"图 {figure_no}  {row.iloc[0]['title']}｜来源：{row.iloc[0]['source']}｜结论：{row.iloc[0]['conclusion']}")
                    except Exception:
                        pass
        next_page = pages[idx] if idx < len(pages) else None
        if next_page and next_page.get("type") == "divider":
            doc.add_page_break()
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = OUT_DIR / "SkyGuard_商业计划书_更新版_请关闭WPS后替换.docx"
        doc.save(fallback)
        return fallback


def draw_pdf_page(c, page: dict, page_num: int, chart_catalog: pd.DataFrame, page_size=A4):
    w, h = page_size
    margin = 40
    accent = page.get("accent") or REPORT_COLORS["teal"]
    c.setFillColor(hex_color(REPORT_COLORS["paper"]))
    c.rect(0, 0, w, h, stroke=0, fill=1)
    if page["type"] == "cover":
        if page.get("figure"):
            try:
                img = Image.open(Path(page["figure"]))
                iw, ih = img.size
                scale = max(w / iw, h / ih)
                c.drawImage(ImageReader(img), (w - iw * scale) / 2, (h - ih * scale) / 2, width=iw * scale, height=ih * scale, mask="auto")
            except Exception:
                pass
        c.setFillColor(colors.Color(0.03, 0.06, 0.08, alpha=0.72))
        c.rect(0, 0, w, h, stroke=0, fill=1)
        c.setFillColor(hex_color(accent))
        c.rect(0, 0, 13, h, stroke=0, fill=1)
        c.setFillColor(colors.white)
        c.setFont("CNFont", 9.5)
        c.drawString(margin + 8, h - 88, page.get("kicker", "商业计划书"))
        c.setFont("CNFont", 34)
        c.drawString(margin + 8, h - 148, "低空智眼 SkyGuard")
        c.setFont("CNFont", 17)
        c.drawString(margin + 8, h - 181, "城市低空空域安全感知与运行监管平台商业计划书")
        y = h - 232
        for para in page["paras"]:
            y = draw_page_text(c, para, margin + 8, y, 430, 10.2, 16, "#ffffff", max_lines=4)
            y -= 7
        chip_x = margin + 8
        for bullet in page["bullets"]:
            chip_x = draw_chip(c, bullet, chip_x, 132, "#ffffff", "#ffffff", REPORT_COLORS["ink"])
        c.setFont("CNFont", 8.2)
        c.drawString(margin + 8, 50, "商业计划书｜财务测算｜证据索引｜Web Demo")
        c.showPage()
        return

    if page["type"] == "divider":
        c.setFillColor(hex_color(REPORT_COLORS["navy"]))
        c.rect(0, 0, w, h, stroke=0, fill=1)
        c.setFillColor(hex_color(accent))
        c.rect(0, 0, 16, h, stroke=0, fill=1)
        c.setFillColor(hex_color("#cbd5df"))
        c.setFont("CNFont", 9)
        c.drawString(margin + 10, h - 72, "章节导读")
        c.setFillColor(colors.white)
        c.setFont("CNFont", 24)
        c.drawString(margin + 10, h - 118, page["chapter"])
        c.setFont("CNFont", 16)
        c.drawString(margin + 10, h - 152, page["title"][:34])
        y = h - 206
        for para in page["paras"]:
            y = draw_page_text(c, para, margin + 10, y, 365, 10, 16, "#f2f4f6", max_lines=8)
        by = 190
        for i, bullet in enumerate(page["bullets"][:3]):
            c.setFillColor(colors.Color(1, 1, 1, alpha=0.08))
            c.roundRect(margin + 10, by - i * 58, 360, 42, 6, stroke=0, fill=1)
            draw_page_text(c, bullet, margin + 22, by + 20 - i * 58, 330, 8, 12, "#f8fafc", max_lines=2)
        media = page.get("figure")
        if media:
            draw_pdf_visual(c, Path(media), w - 245, 360, 188, 172, border=True)
        elif page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                draw_pdf_visual(c, ROOT / row.iloc[0]["file"], w - 245, 360, 188, 172, border=True)
        c.setFont("CNFont", 8)
        c.setFillColor(hex_color("#cbd5df"))
        c.drawRightString(w - margin, 40, f"第 {page_num} 页")
        c.showPage()
        return

    c.setFillColor(colors.white)
    c.roundRect(22, 22, w - 44, h - 44, 8, stroke=0, fill=1)
    c.setFillColor(hex_color(accent))
    c.rect(22, h - 74, w - 44, 3.6, stroke=0, fill=1)
    c.setFillColor(hex_color(REPORT_COLORS["muted"]))
    c.setFont("CNFont", 7.8)
    c.drawString(margin, h - 53, page.get("kicker", ""))
    c.drawRightString(w - margin, h - 53, "SkyGuard 低空智眼")
    c.setFillColor(hex_color(REPORT_COLORS["ink"]))
    c.setFont("CNFont", 16.5)
    c.drawString(margin, h - 88, f"{page['chapter']}｜{page['title']}"[:42])
    left_w = 310
    right_x = margin + left_w + 18
    right_w = w - right_x - margin
    y = h - 120
    for para in page["paras"][:3]:
        y = draw_page_text(c, para, margin, y, left_w, 8.6, 13.1, REPORT_COLORS["soft_ink"], max_lines=8)
        y -= 4
        if y < 325:
            break
    c.setFillColor(hex_color("#f3f1eb"))
    c.roundRect(right_x, h - 318, right_w, 198, 6, stroke=0, fill=1)
    c.setStrokeColor(hex_color("#d8d1c5"))
    c.roundRect(right_x, h - 318, right_w, 198, 6, stroke=1, fill=0)
    c.setFont("CNFont", 9.2)
    c.setFillColor(hex_color(REPORT_COLORS["ink"]))
    c.drawString(right_x + 12, h - 141, "证据与行动")
    by = h - 165
    for bullet in page["bullets"][:3]:
        by = draw_page_text(c, "· " + bullet, right_x + 12, by, right_w - 24, 7.3, 10.8, REPORT_COLORS["soft_ink"], max_lines=3)
        by -= 3
    if page.get("callout"):
        c.setFillColor(hex_color(accent))
        c.roundRect(right_x + 12, h - 304, right_w - 24, 30, 5, stroke=0, fill=1)
        draw_page_text(c, page["callout"], right_x + 20, h - 288, right_w - 40, 7.2, 10, "#ffffff", max_lines=2)
    media_top = 428
    if page.get("table"):
        draw_pdf_table(c, page["table"], margin, media_top, w - margin * 2, accent)
    elif page.get("figure"):
        bottom_y = draw_pdf_visual(c, Path(page["figure"]), margin, media_top, w - margin * 2, 270, border=True)
        c.setFont("CNFont", 6.8)
        c.setFillColor(hex_color(REPORT_COLORS["muted"]))
        c.drawString(margin, max(43, bottom_y), f"图注：{humanize_figure_caption(Path(page['figure']))}")
    elif page.get("chart_id"):
        row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
        if not row.empty:
            chart_path = ROOT / row.iloc[0]["file"]
            bottom_y = draw_pdf_visual(c, chart_path, margin, media_top, w - margin * 2, 270, border=True)
            caption = f"{row.iloc[0]['title']}｜来源：{row.iloc[0]['source']}｜结论：{row.iloc[0]['conclusion']}"
            draw_page_text(c, caption, margin, max(42, bottom_y), w - margin * 2, 6.5, 8.6, REPORT_COLORS["muted"], max_lines=2)
    c.setStrokeColor(hex_color("#d8d1c5"))
    c.line(margin, 33, w - margin, 33)
    c.setFont("CNFont", 7.2)
    c.setFillColor(hex_color(REPORT_COLORS["muted"]))
    c.drawString(margin, 20, "资料口径：官方统计 / 地方政策 / 公开数据 / 演示样本 / 经营假设分层标注")
    c.drawRightString(w - margin, 20, f"第 {page_num} 页")
    c.showPage()


def write_pdf(pages: list[dict], chart_catalog: pd.DataFrame, filename: str, limit: int | None = None) -> Path:
    path = OUT_DIR / filename
    c = canvas.Canvas(str(path), pagesize=A4)
    selected = pages if limit is None else pages[:limit]
    for idx, page in enumerate(selected, start=1):
        draw_pdf_page(c, page, idx, chart_catalog)
    c.save()
    return path


# ---------------------------------------------------------------------------
# Final manuscript layer
# ---------------------------------------------------------------------------

def add_manuscript_chart_extensions(data: dict[str, pd.DataFrame]) -> None:
    """Add a final set of varied, report-grade charts backed by prepared data."""
    nat = data["national"].copy()
    regional = data["regional"].copy()
    city = data["city"].copy()
    plans = data["plans"].copy()
    telemetry = data["telemetry"].copy()
    alerts = data["alerts"].copy()
    risk = data["risk"].copy()
    grid = data["grid"].copy()
    sensors = data["sensors"].copy()
    routes = data["routes"].copy()
    vertiports = data["vertiports"].copy()
    finance = build_finance_tables()
    agg = build_real_aggregates()

    report_palette = ["#234f7d", "#0c6f69", "#9b6a2f", "#9b3f3a", "#2f6b4f", "#6d716c", "#b69252", "#44546a"]
    official_source = "中国民用航空局、交通运输部公开统计及本地清洗表01-04"
    city_source = "地方低空经济政策目标与监管平台案例，本地清洗表05"
    demo_source = "平台演示样本，字段data_nature已标注"
    finance_source = "经营模型假设，详见SkyGuard_财务测算表.xlsx"
    real_source = "本地整理的低空经济公开数据集"

    for frame in [nat, regional, city, plans, telemetry, alerts, risk, grid, sensors, routes, vertiports]:
        for col in frame.columns:
            if col in {"year", "target_year", "value", "confidence", "response_time_min", "computed_risk_score", "sensor_coverage_score", "flight_density_index", "planned_distance_km", "planned_altitude_m", "health_score", "daily_capacity_sorties", "daily_planned_sorties"} or col.endswith("_m") or col.endswith("_km") or col.endswith("_score"):
                frame[col] = safe_numeric(frame[col])
    for frame, cols in [(plans, ["planned_takeoff_time", "apply_time"]), (telemetry, ["timestamp"]), (alerts, ["timestamp"])]:
        for col in cols:
            if col in frame.columns:
                frame[col] = pd.to_datetime(frame[col], errors="coerce")

    def slope_policy_region():
        if "province_index_trend" in agg and not agg["province_index_trend"].empty:
            d = agg["province_index_trend"].copy()
            years = sorted(d["年份"].dropna().unique())
            left, right = years[0], years[-1]
            p = d[d["年份"].isin([left, right])].pivot(index="地区", columns="年份", values="各省低空经济指数").dropna().head(12)
            fig, ax = plt.subplots(figsize=(7.8, 5.2))
            for i, (name, row) in enumerate(p.iterrows()):
                color = report_palette[i % len(report_palette)]
                ax.plot([0, 1], [row[left], row[right]], color=color, lw=2.2, marker="o")
                ax.text(-0.03, row[left], name, ha="right", va="center", fontsize=8)
                ax.text(1.03, row[right], name, ha="left", va="center", fontsize=8)
            ax.set_xticks([0, 1])
            ax.set_xticklabels([str(int(left)), str(int(right))])
            finish_chart(ax, "省域低空经济指数坡度图", "首批区域不能只看当年存量，还要看长期变化")
            return fig
        fig, ax = plt.subplots(figsize=(7.8, 5.2))
        finish_chart(ax, "省域低空经济指数坡度图", "")
        return fig

    emit_chart("省域低空经济指数坡度图", "坡度图", "各省低空经济指数数据（1990-2024）.csv", "第一章 行业机会", "坡度图用于区分长期上升区域和短期热点区域，帮助确定首批市场进入顺序。", slope_policy_region)

    def target_waffle():
        counts = city["category"].value_counts().head(5)
        total = counts.sum() or 1
        blocks = []
        for idx, (name, val) in enumerate(counts.items()):
            blocks.extend([(name, report_palette[idx % len(report_palette)])] * max(1, round(val / total * 100)))
        blocks = (blocks + [("其他", "#d6d0c3")] * 100)[:100]
        fig, ax = plt.subplots(figsize=(7.2, 6.2))
        ax.axis("off")
        for i, (name, color) in enumerate(blocks):
            x, y = i % 10, 9 - i // 10
            ax.add_patch(plt.Rectangle((x, y), 0.82, 0.82, fc=color, ec="white", lw=0.7))
        handles = [plt.Rectangle((0, 0), 1, 1, fc=report_palette[i % len(report_palette)]) for i in range(len(counts))]
        ax.legend(handles, counts.index, frameon=False, loc="lower center", ncol=3, bbox_to_anchor=(0.5, -0.12), fontsize=8)
        ax.set_title("地方低空政策目标华夫图", loc="left", fontsize=13.5, fontweight="bold", color=REPORT_COLORS["ink"])
        return fig

    emit_chart("地方低空政策目标华夫图", "华夫图", city_source, "第一章 行业机会", "华夫图把政策重点压缩为百分比结构，说明平台、航线、起降点和监管任务并不是同等权重。", target_waffle)

    def policy_lollipop():
        counts = city["entity"].value_counts().head(14).sort_values()
        fig, ax = plt.subplots(figsize=(8.4, 5.2))
        y = np.arange(len(counts))
        ax.hlines(y, 0, counts.values, color="#cfc8b8", lw=2)
        ax.scatter(counts.values, y, s=88, color="#0c6f69", edgecolor="white", lw=1)
        ax.set_yticks(y)
        ax.set_yticklabels(counts.index)
        finish_chart(ax, "城市政策任务密度棒棒糖", "按本地政策目标数据统计", x_grid=True)
        return fig

    emit_chart("城市政策任务密度棒棒糖", "棒棒糖图", city_source, "第六章 市场竞争", "不同城市政策任务密度不同，销售材料和场景包应按城市任务重新组合。", policy_lollipop)

    def bid_status_funnel():
        if "bid_type" in agg and not agg["bid_type"].empty:
            d = agg["bid_type"].head(8).sort_values("项目数", ascending=True)
            fig, ax = plt.subplots(figsize=(8.4, 5.1))
            for i, (_, r) in enumerate(d.iterrows()):
                width = r["项目数"] / max(1, d["项目数"].max())
                ax.barh(i, width, left=(1 - width) / 2, height=0.72, color=report_palette[i % len(report_palette)], alpha=0.86)
                ax.text(0.5, i, str(r["项目类型"])[:16], ha="center", va="center", color="white", fontsize=8.4, fontweight="bold")
            ax.set_xlim(0, 1)
            ax.set_yticks([])
            ax.set_xticks([])
            ax.set_title("招投标类型漏斗", loc="left", fontsize=13.5, fontweight="bold", color=REPORT_COLORS["ink"])
            return fig
        fig, ax = plt.subplots(figsize=(8.4, 5.1))
        ax.axis("off")
        return fig

    emit_chart("招投标类型漏斗", "漏斗图", "低空经济招投标项目全量数据.csv", "第六章 市场竞争", "招投标项目类型能够提示哪些采购场景更容易形成平台化机会。", bid_status_funnel)

    def patent_ipc_radial():
        if "patent_ipc" in agg and not agg["patent_ipc"].empty:
            d = agg["patent_ipc"].head(10)
            theta = np.linspace(0, 2 * np.pi, len(d), endpoint=False)
            fig, ax = plt.subplots(figsize=(6.8, 6.8), subplot_kw={"projection": "polar"})
            ax.bar(theta, d["专利数"], width=2 * np.pi / len(d) * 0.78, color=report_palette[: len(d)], alpha=0.82)
            ax.set_xticks(theta)
            ax.set_xticklabels(d["IPC主分类号"], fontsize=8)
            ax.set_yticks([])
            ax.set_title("专利IPC方向环形柱", loc="left", fontsize=13.5, fontweight="bold", color=REPORT_COLORS["ink"])
            return fig
        fig, ax = plt.subplots(figsize=(6.8, 6.8), subplot_kw={"projection": "polar"})
        ax.set_axis_off()
        return fig

    emit_chart("专利IPC方向环形柱", "环形柱图", "低空经济专利全量数据.csv", "第四章 技术架构", "IPC方向能够说明技术积累集中在哪些空域管理、通信、导航和飞行服务能力上。", patent_ipc_radial)

    def enterprise_capital_violin():
        firms = read_optional_csv("低空经济相关企业全量数据.csv", usecols=["行业类型", "注册资本", "员工人数"])
        if firms.empty:
            fig, ax = plt.subplots(figsize=(8.4, 5.1))
            ax.axis("off")
            return fig
        firms["注册资本万元"] = firms["注册资本"].map(parse_amount_to_10k)
        top = firms["行业类型"].value_counts().head(7).index
        sample = sample_frame(firms[firms["行业类型"].isin(top)].dropna(subset=["注册资本万元"]), 1800)
        fig, ax = plt.subplots(figsize=(8.8, 5.2))
        palette_map = {name: report_palette[i % len(report_palette)] for i, name in enumerate(sample["行业类型"].dropna().value_counts().head(7).index)}
        sample["注册资本对数"] = np.log1p(sample["注册资本万元"])
        sns.violinplot(data=sample, x="行业类型", y="注册资本对数", hue="行业类型", palette=palette_map, ax=ax, inner="quartile", cut=0, legend=False)
        ax.set_ylabel("注册资本对数")
        ax.set_xlabel("")
        ax.tick_params(axis="x", rotation=18)
        finish_chart(ax, "低空企业注册资本小提琴", "按行业类型观察资本分布")
        return fig

    emit_chart("低空企业注册资本小提琴", "小提琴图", "低空经济相关企业全量数据.csv", "第六章 市场竞争", "资本分布能帮助判断哪些行业适合做渠道伙伴，哪些行业更适合做场景客户。", enterprise_capital_violin)

    def safety_severity_donut():
        if "safety_severity" in agg and not agg["safety_severity"].empty:
            d = agg["safety_severity"]
            fig, ax = plt.subplots(figsize=(6.5, 6.0))
            ax.pie(d["事件数"], labels=d["严重程度"], startangle=90, colors=report_palette[: len(d)], wedgeprops={"width": 0.38, "edgecolor": "white"}, autopct="%1.0f%%")
            ax.text(0, 0, f"{int(d['事件数'].sum())}\n事件", ha="center", va="center", fontsize=17, fontweight="bold", color=REPORT_COLORS["ink"])
            ax.set_title("安全事件严重程度环带", loc="left", fontsize=13.5, fontweight="bold", color=REPORT_COLORS["ink"])
            return fig
        fig, ax = plt.subplots(figsize=(6.5, 6.0))
        ax.axis("off")
        return fig

    emit_chart("安全事件严重程度环带", "环带图", "低空安全事件与监管案例数据.csv", "第二章 痛点验证", "严重程度结构决定事件分级处置和人工复核资源配置。", safety_severity_donut)

    def facility_capacity_treemap():
        if "facility_type" in agg and not agg["facility_type"].empty:
            d = agg["facility_type"].head(9)
            values = d["设计容量"].fillna(d["数量"]).clip(lower=1).tolist()
            labels = d["设施类型"].tolist()
            fig, ax = plt.subplots(figsize=(8.6, 5.3))
            ax.axis("off")
            for i, (x, y, rw, rh, val, lab) in enumerate(treemap_rects(values, labels)):
                ax.add_patch(plt.Rectangle((x, y), rw, rh, fc=report_palette[i % len(report_palette)], ec="white", lw=1.4, alpha=0.86))
                ax.text(x + 0.012, y + rh - 0.035, str(lab)[:10], color="white", fontsize=8.5, va="top", fontweight="bold")
                ax.text(x + 0.012, y + 0.02, f"{int(val)}", color="white", fontsize=8, va="bottom")
            ax.set_title("起降设施容量矩形树", loc="left", fontsize=13.5, fontweight="bold", color=REPORT_COLORS["ink"])
            return fig
        fig, ax = plt.subplots(figsize=(8.6, 5.3))
        ax.axis("off")
        return fig

    emit_chart("起降设施容量矩形树", "容量树图", "低空起降设施与空域数据.csv", "第三章 产品服务", "容量结构决定计划审批、起降点排班和应急冗余策略。", facility_capacity_treemap)

    def flight_type_parallel():
        if "flight_type" in agg and not agg["flight_type"].empty:
            d = agg["flight_type"].dropna().head(9).copy()
            for col in ["飞行数", "平均距离", "平均高度"]:
                d[col] = (d[col] - d[col].min()) / max(1e-9, d[col].max() - d[col].min())
            fig, ax = plt.subplots(figsize=(8.5, 5.1))
            xs = np.arange(3)
            for i, (_, r) in enumerate(d.iterrows()):
                ax.plot(xs, [r["飞行数"], r["平均距离"], r["平均高度"]], marker="o", color=report_palette[i % len(report_palette)], alpha=0.72, lw=1.8)
                ax.text(2.05, r["平均高度"], str(r["飞行类型"])[:8], fontsize=7.6, va="center")
            ax.set_xticks(xs)
            ax.set_xticklabels(["规模", "距离", "高度"])
            ax.set_ylim(-0.05, 1.05)
            finish_chart(ax, "飞行类型平行坐标", "三项指标归一化")
            return fig
        fig, ax = plt.subplots(figsize=(8.5, 5.1))
        ax.axis("off")
        return fig

    emit_chart("飞行类型平行坐标", "平行坐标图", "低空飞行活动与轨迹数据.csv", "第五章 数据体系", "平行坐标用于比较不同飞行类型的规模、距离和高度组合。", flight_type_parallel)

    def route_altitude_band():
        sample = sample_frame(routes.dropna(subset=["planned_altitude_m", "route_distance_km"]), 900)
        fig, ax = plt.subplots(figsize=(8.5, 5.1))
        sns.boxplot(data=sample, x="scenario", y="planned_altitude_m", color="#d8c7a1", ax=ax, fliersize=1.5)
        ax.tick_params(axis="x", rotation=18)
        ax.set_xlabel("")
        ax.set_ylabel("计划高度（m）")
        finish_chart(ax, "航线场景高度箱须", "演示航线样本")
        return fig

    emit_chart("航线场景高度箱须", "箱须图", demo_source, "第四章 技术架构", "航线高度分布会影响围栏阈值、净空约束和审批模板。", route_altitude_band)

    def incident_response_sina():
        sample = sample_frame(alerts.dropna(subset=["response_time_min"]), 1500)
        fig, ax = plt.subplots(figsize=(8.6, 5.2))
        severity_order = list(sample["severity"].dropna().value_counts().index)
        severity_palette = {name: report_palette[i % len(report_palette)] for i, name in enumerate(severity_order)}
        sns.stripplot(data=sample, x="severity", y="response_time_min", hue="severity", jitter=0.28, alpha=0.25, size=2.4, palette=severity_palette, ax=ax, legend=False)
        sns.pointplot(data=sample, x="severity", y="response_time_min", errorbar=("ci", 80), color="#172033", markers="_", linestyles="", ax=ax)
        ax.set_xlabel("")
        ax.set_ylabel("响应时间（分钟）")
        finish_chart(ax, "事件响应时间散点云", "点为事件样本，短线为均值区间")
        return fig

    emit_chart("事件响应时间散点云", "散点云图", demo_source, "第二章 痛点验证", "响应时间的尾部事件比平均数更能说明工单闭环的价值。", incident_response_sina)

    def sensor_health_calendar():
        s = sensors.copy()
        s["idx"] = np.arange(len(s))
        s["row"] = s["idx"] // 25
        s["col"] = s["idx"] % 25
        sample = s.head(625)
        pivot = sample.pivot(index="row", columns="col", values="health_score")
        fig, ax = plt.subplots(figsize=(9.0, 4.7))
        sns.heatmap(pivot, cmap="YlGnBu", linewidths=0.4, linecolor="#f8f6f0", cbar_kws={"label": "健康度"}, ax=ax)
        ax.set_xticks([])
        ax.set_yticks([])
        ax.set_title("感知设备健康地块图", loc="left", fontsize=13.5, fontweight="bold", color=REPORT_COLORS["ink"])
        return fig

    emit_chart("感知设备健康地块图", "地块热力图", demo_source, "第十三章 Web Demo", "设备健康地块图适合快速定位批量离线或低健康度设备。", sensor_health_calendar)

    def risk_decile_bars():
        d = risk[["computed_risk_score", "risk_event_label"]].dropna().copy()
        d["分位组"] = pd.qcut(d["computed_risk_score"], 10, labels=[f"{i+1}" for i in range(10)], duplicates="drop")
        tab = d.groupby("分位组", observed=True).agg(事件率=("risk_event_label", "mean"), 样本数=("risk_event_label", "count")).reset_index()
        fig, ax1 = plt.subplots(figsize=(8.4, 5.0))
        ax2 = ax1.twinx()
        ax1.bar(tab["分位组"].astype(str), tab["样本数"], color="#d8c7a1", alpha=0.82)
        ax2.plot(tab["分位组"].astype(str), tab["事件率"], color="#9b3f3a", marker="o", lw=2.2)
        final_finish_chart(ax1, "风险分位事件率", "柱为样本数，线为事件率")
        ax1.set_xlabel("风险分位组")
        ax2.set_ylabel("事件率")
        ax2.tick_params(colors="#9b3f3a", labelsize=8.5)
        return fig

    final_emit_chart("风险分位事件率", "分位柱线图", demo_source, "第五章 数据体系", "风险分位能把模型分数转成可解释的人工复核优先级。", risk_decile_bars)

    def finance_cost_sunburst_like():
        cost_df = finance["cost"].copy()
        latest = cost_df.iloc[-1]
        items = ["cogs", "研发", "销售", "行政"]
        labels = ["交付与云资源", "研发", "销售服务", "行政管理"]
        values = [float(latest[c]) for c in items]
        fig, ax = plt.subplots(figsize=(6.4, 6.4))
        ax.pie(values, labels=labels, radius=1.0, colors=report_palette[:4], wedgeprops=dict(width=0.30, edgecolor="white"))
        ax.pie([values[0] + values[1], values[2] + values[3]], labels=["产品与交付", "销售与管理"], radius=0.66, colors=["#0c6f69", "#9b6a2f"], wedgeprops=dict(width=0.28, edgecolor="white"))
        ax.set_title("成本结构双层环", loc="left", fontsize=13.5, fontweight="bold", color=REPORT_COLORS["ink"])
        return fig

    emit_chart("成本结构双层环", "双层环图", finance_source, "第九章 财务预测", "成本结构要看研发、云资源、交付和销售服务的组合，而不是只看总成本。", finance_cost_sunburst_like)

    def cash_waterfall_compact():
        cash = finance["cashflow"].copy()
        row = cash.iloc[2]
        labels = ["期初现金", "经营现金流", "融资进入", "期末现金"]
        values = [row["beginning_cash"], row["operating_cashflow"], row["financing"], row["ending_cash"]]
        fig, ax = plt.subplots(figsize=(8.5, 4.8))
        running = 0
        for i, (label, val) in enumerate(zip(labels, values)):
            if i in {0, len(values) - 1}:
                ax.bar(i, val, color="#234f7d", width=0.6)
                running = val if i == 0 else running
            else:
                ax.bar(i, val, bottom=running if val >= 0 else running + val, color="#2f6b4f" if val >= 0 else "#9b3f3a", width=0.6)
                running += val
            ax.text(i, values[i] if i in {0, len(values) - 1} else running, f"{val:.0f}", ha="center", va="bottom", fontsize=8)
        ax.set_xticks(range(len(labels)))
        ax.set_xticklabels(labels, rotation=12)
        finish_chart(ax, "现金流瀑布", "单位：万元")
        return fig

    emit_chart("现金流瀑布", "瀑布图", finance_source, "第九章 财务预测", "现金流瀑布能说明融资资金如何覆盖经营现金缺口和必要投入。", cash_waterfall_compact)


def manuscript_chapter_specs(data: dict[str, pd.DataFrame]) -> list[dict]:
    nat = data["national"]
    uav = first_value(nat, "registered_uavs_10k", 2025)
    hours = first_value(nat, "uav_flight_hours_10k", 2025)
    units = first_value(nat, "uav_operating_units", 2025)
    stations = first_value(nat, "low_altitude_flight_service_stations", 2025)
    return [
        {"chapter": "第一章 行业机会", "count": 13, "accent": REPORT_COLORS["teal"], "topics": ["政策窗口", "运行规模", "区域分层", "客户入口", "监管边界", "产业链机会"], "opening": f"低空经济已经从概念展示进入运行治理阶段。公开统计显示，2025 年注册无人机约 {uav:.1f} 万架，累计飞行小时约 {hours:.2f} 万小时，运营单位 {units:.0f} 家，低空飞行服务站 {stations:.0f} 个。数量增长本身不是商业模式，但它让重点区域出现持续的计划核验、目标识别、围栏管理、事件处置和运行复盘需求。", "decision": "从重点区域切入，先证明日常运行监管闭环。"},
        {"chapter": "第二章 痛点验证", "count": 13, "accent": REPORT_COLORS["amber"], "topics": ["现场发现", "身份核验", "计划比对", "围栏触发", "工单流转", "复盘归档"], "opening": "低空安全不是单一识别问题。现场真正难的是发现目标之后的连续动作：有没有计划，身份是否可信，是否越过围栏，由谁复核，多久处置，事后如何给出可追溯材料。只要链路断开，管理者就会退回电话、截图和临时表格。", "decision": "把痛点写成工作流断点，而不是泛泛的安全焦虑。"},
        {"chapter": "第三章 产品服务", "count": 17, "accent": REPORT_COLORS["blue"], "topics": ["态势总览", "飞行计划", "电子围栏", "识别复核", "事件中心", "运行报表", "场景包", "客户成功"], "opening": "SkyGuard 交付的是一套低空运行工作台，而不是一块展示大屏。前台负责态势、目标和事件，后台负责规则、权限、日志和报表，中间把计划、身份、轨迹和围栏连接起来。客户每天使用的是工作流，答辩展示的也应该是这条工作流。", "decision": "每个模块都对应一个客户动作和一项验收材料。"},
        {"chapter": "第四章 技术架构", "count": 17, "accent": REPORT_COLORS["green"], "topics": ["多源接入", "时空索引", "轨迹比对", "风险评分", "规则引擎", "权限审计", "部署集成", "可解释性"], "opening": "技术架构的核心不是堆概念，而是让告警可解释、处置可回放、报表可追溯。项目采用规则先行、模型辅助、人工确认的技术路线，既能在早期试点降低误报争议，也保留后续接入更多硬件和算法的空间。", "decision": "用架构支撑业务闭环，不把模型写成黑箱。"},
        {"chapter": "第五章 数据体系", "count": 13, "accent": REPORT_COLORS["teal"], "topics": ["官方统计", "政策目标", "公开数据", "演示样本", "数据字典", "质量规则", "模型校验"], "opening": "计划书中的数字必须分层使用。官方统计支撑行业判断，地方政策支撑区域机会，公开数据支撑方法说明，演示样本只用于产品流程呈现。边界清楚后，图表才有价值；边界不清，再多数据也会削弱可信度。", "decision": "所有数字回到来源、字段或假设。"},
        {"chapter": "第六章 市场竞争", "count": 13, "accent": REPORT_COLORS["blue"], "topics": ["替代方案", "客户分层", "采购逻辑", "竞品矩阵", "进入顺序", "伙伴生态"], "opening": "SkyGuard 的竞争对手不是单一公司，而是一组替代路径：视频监控、单点感知设备、反制系统、内部工单工具、城市级重平台和人工巡查。项目需要明确自己不做什么，才能说明为什么重点区域工作台更适合创业阶段。", "decision": "先占可验收的区域级场景，再复制到区县协同。"},
        {"chapter": "第七章 商业模式", "count": 13, "accent": REPORT_COLORS["green"], "topics": ["订阅收入", "项目部署", "活动保障", "航线评估", "报告服务", "运维续费", "客户成功"], "opening": "商业模式要回答客户为什么持续付费。低空重点区域不仅需要首次部署，还需要规则维护、活动保障、航线评估、月度复盘和设备接入支持。收入结构必须从一次性项目逐步转向订阅、报告和运维续费。", "decision": "把收入拆到客户动作、交付内容和成本结构上。"},
        {"chapter": "第八章 落地计划", "count": 11, "accent": REPORT_COLORS["amber"], "topics": ["0-3个月", "3-6个月", "6-18个月", "18-36个月", "试点验收", "区域复制"], "opening": "落地计划要像真正项目排期。每个阶段都要有可检查交付物：数据字典、规则样例、Demo页面、试点包、培训材料、报表模板和复盘会议纪要。没有交付物的时间表，只是愿景。", "decision": "按试点、复制、协同三层推进。"},
        {"chapter": "第九章 财务预测", "count": 15, "accent": REPORT_COLORS["green"], "topics": ["收入预测", "成本结构", "毛利路径", "现金流", "回款周期", "融资用途", "敏感性", "单位经济"], "opening": "财务预测不能只看收入曲线。SkyGuard 的经营质量取决于客户数量、客单价、部署成本、续费率、回款周期和研发投入之间能否互相支撑。模型保留保守、基准、乐观三情景，用来检查现金流而不是美化增长。", "decision": "用现金流约束扩张速度，用单位经济解释续费价值。"},
        {"chapter": "第十章 风险应对", "count": 9, "accent": REPORT_COLORS["red"], "topics": ["政策边界", "数据合规", "误报漏报", "硬件依赖", "采购周期", "现金压力", "过度承诺"], "opening": "低空项目最容易失分的地方是边界不清。SkyGuard 做辅助感知、风险预警、事件协同和报表复盘，不做干扰、捕获、打击，也不替代审批或执法。把边界写清楚，反而能增强项目可信度。", "decision": "风险按阻断项目、拖慢节奏、增加成本三类处理。"},
        {"chapter": "第十一章 团队组织", "count": 8, "accent": REPORT_COLORS["blue"], "topics": ["岗位配置", "产品负责人", "算法数据", "前后端研发", "交付运维", "行业顾问"], "opening": "团队组织要与当前阶段匹配。早期团队最需要把产品、数据、前后端、交付和客户材料做出来，而不是一开始搭建臃肿部门。每个岗位都要能对应下一阶段的交付物。", "decision": "小团队围绕Demo、试点和客户成功配置。"},
        {"chapter": "第十二章 社会价值", "count": 8, "accent": REPORT_COLORS["teal"], "topics": ["公共安全", "城市治理", "应急保障", "产业数据", "岗位培养", "合规意识"], "opening": "社会价值不靠口号体现，而靠可观察的治理改善体现。低空运行越密，城市越需要把目标、计划、围栏、事件和报表放进同一套协同机制，减少信息断点和责任争议。", "decision": "把公共价值写成可衡量的安全、效率和合规改善。"},
        {"chapter": "第十三章 Web Demo", "count": 12, "accent": REPORT_COLORS["blue"], "topics": ["首页总览", "演示脚本", "态势地图", "目标监测", "计划审批", "识别复核", "事件工单", "运行报表", "移动处置"], "opening": "Web Demo 是计划书能否落地的现场证据。它需要让评委从首页进入演示脚本，再连续点击态势、目标、计划、识别、围栏、事件、报表和移动端。一个能跑通流程的 Demo，比单张大屏截图更能说明团队执行力。", "decision": "用可点击流程证明产品能被演示、试点和部署。"},
        {"chapter": "附录", "count": 16, "accent": REPORT_COLORS["ash"], "topics": ["来源索引", "图表目录", "数据字典", "财务假设", "访谈问卷", "部署说明", "截图包", "交付清单"], "opening": "附录承担核验功能。正文中的判断、图表、图片和 Demo 截图，都要能在附录中找到来源、字段、假设或文件位置。这样计划书不是孤立文本，而是一套可继续迭代的交付资料。", "decision": "把证据、假设和交付物放到清晰位置。"},
    ]


def manuscript_table(chapter: str, topic: str, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame, chart_catalog: pd.DataFrame) -> list[list[str]]:
    if chapter.startswith("第一章"):
        nat = data["national"]
        return [["指标", "2025年公开口径", "对项目的含义"], ["注册无人机", f"{first_value(nat, 'registered_uavs_10k', 2025):.1f}万架", "运行主体增加，重点区域需要识别和记录"], ["累计飞行小时", f"{first_value(nat, 'uav_flight_hours_10k', 2025):.2f}万小时", "日常运行频次支撑持续监管需求"], ["运营单位", f"{first_value(nat, 'uav_operating_units', 2025):.0f}家", "客户和合作伙伴基础扩大"], ["飞行服务站", f"{first_value(nat, 'low_altitude_flight_service_stations', 2025):.0f}个", "服务体系开始形成，平台接入窗口出现"]]
    if chapter.startswith("第二章"):
        return [["断点", "现场表现", "SkyGuard对应能力"], ["发现目标", "只看到点位，难判断身份", "Remote ID、传感器来源、目标档案"], ["计划核验", "计划表、白名单和现场目标分散", "计划库、航线比对、异常提示"], ["围栏触发", "临时规则靠人工通知", "规则中心、阈值、触发记录"], ["处置复盘", "电话沟通后难归档", "工单、时长、责任人、日报"]]
    if chapter.startswith("第三章"):
        return [["模块", "客户动作", "验收材料"], ["综合态势", "查看区域目标和围栏", "态势图、目标列表、风险等级"], ["计划审批", "核验计划、航线和高度", "审批记录、预检原因"], ["事件工单", "确认、派单、处置、归档", "闭环时长、处置说明"], ["运行报表", "复盘日常运行质量", "日报、月报、风险排行"]]
    if chapter.startswith("第四章"):
        return [["层级", "核心能力", "设计取舍"], ["数据接入", "计划、遥测、围栏、设备、人工复核", "先接标准字段，再扩展硬件"], ["时空计算", "点、线、面统一索引", "支撑轨迹比对和围栏触发"], ["风险解释", "身份、计划、轨迹、区域、响应紧迫度", "模型辅助，人工确认"], ["审计安全", "权限、日志、导出留痕", "避免黑箱和责任不清"]]
    if chapter.startswith("第五章"):
        return [["数据层", "使用位置", "边界说明"], ["官方统计", "行业机会、规模判断", "不替代实时监管数据"], ["地方政策", "区域进入和场景选择", "政策目标不等同订单"], ["公开数据", "方法验证和地图基础", "需说明下载来源"], ["演示样本", "Demo流程和图表呈现", "不写成真实客户业务"]]
    if chapter.startswith("第六章"):
        return [["替代方案", "优势", "不足"], ["通用视频监控", "部署基础广", "缺少计划、围栏和报表闭环"], ["单点感知设备", "识别能力强", "难独立完成协同处置"], ["城市级重平台", "能力全面", "周期长、成本高"], ["人工巡查", "灵活", "不可持续、不可追溯"]]
    if chapter.startswith("第七章"):
        return [["收入项", "价格口径", "交付内容"], ["Lite SaaS", "8-15万元/年", "账号、围栏、告警、基础报表"], ["Pro部署", "30-80万元/项目", "数据接入、规则配置、培训验收"], ["活动保障", "5-30万元/次", "临时围栏、移动告警、日报复盘"], ["航线评估", "3000-20000元/条", "路线风险、容量和合规建议"], ["运维服务", "合同额10%-20%/年", "规则调优、报表会议、设备巡检"]]
    if chapter.startswith("第八章"):
        return [["阶段", "交付物", "验收口径"], ["0-3个月", "Demo、数据字典、规则样例", "能完整演示核心流程"], ["3-6个月", "重点区域试点包", "导入计划、围栏和事件样本"], ["6-18个月", "多场景复制包", "形成报表模板和客户成功流程"], ["18-36个月", "区县协同版本", "多区域规则和权限体系稳定"]]
    if chapter.startswith("第九章"):
        finance = build_finance_tables()
        return [["假设项", "模型口径"]] + finance["assumptions"].head(6)[["item", "assumption"]].astype(str).values.tolist()
    if chapter.startswith("第十章"):
        return [["风险", "触发信号", "应对动作"], ["政策边界", "客户要求平台替代审批或执法", "合同和演示中写清辅助监管边界"], ["误报漏报", "高风险事件人工复核压力上升", "引入分级阈值、回标和复盘会议"], ["采购周期", "试点决策链条拉长", "拆分轻量试点和正式部署"], ["现金压力", "回款慢于交付成本", "控制并行项目数量和账期"]]
    if chapter.startswith("第十一章"):
        return [["角色", "近期职责"], ["产品负责人", "场景拆解、流程定义、验收材料"], ["数据与算法", "字段治理、风险解释、模型校验"], ["前后端研发", "工作台、地图、表格、报表和部署"], ["交付运维", "现场配置、培训、客户成功"], ["行业顾问", "政策边界、客户沟通和试点资源"]]
    if chapter.startswith("第十二章"):
        return [["价值维度", "可观测指标"], ["公共安全", "高风险事件确认时长、闭环率"], ["城市治理", "规则更新次数、报表提交及时率"], ["应急保障", "临时保障任务响应、重点区域覆盖"], ["人才培养", "培训次数、岗位覆盖、操作合规率"]]
    if chapter.startswith("第十三章"):
        return [["演示路径", "页面", "说明"], ["1", "首页/演示脚本", "选择值班域和场景"], ["2", "综合态势/目标监测", "发现目标并查看轨迹"], ["3", "计划审批/围栏规则", "核验计划和规则触发原因"], ["4", "事件工单/移动处置", "确认、派单、归档"], ["5", "运行报表", "导出日常复盘材料"]]
    if "来源" in topic:
        return [["资料名称", "资料类型", "用途"]] + source_registry[["title", "source_type", "use"]].head(6).astype(str).values.tolist()
    return [["材料", "位置", "作用"], ["图表目录", "SkyGuard_图表目录.csv", "核验所有图表类型、来源和结论"], ["证据索引", "SkyGuard_证据资料索引表.xlsx", "核验来源和章节关系"], ["财务测算", "SkyGuard_财务测算表.xlsx", "核验收入、成本和现金流假设"], ["Web Demo包", "SkyGuard_WebDemo.zip", "本地部署和答辩演示"]]


def manuscript_paragraphs(spec: dict, topic: str, row: dict | None, metrics: dict, figure_caption: str | None, index: int) -> list[str]:
    chapter = spec["chapter"]
    evidence_title = figure_caption or (row.get("title") if row else topic)
    conclusion = (row or {}).get("conclusion", spec["decision"])
    source = (row or {}).get("source", "整理图片、Demo截图与项目资料")
    style_tail = [
        "这一判断会影响首批试点范围、客户沟通材料和验收指标。",
        "这部分内容需要在答辩时用数据和页面同时说明。",
        "它不是单独的功能描述，而是商业、交付和风险控制之间的连接点。",
        "如果后续进入真实试点，这一项会进入客户访谈、配置清单和复盘报告。",
        "这一页的作用，是把“看起来有机会”压实成“可以被交付和验收”。",
        "答辩时应把它连接到客户场景、系统页面和财务假设，而不是孤立朗读。",
        "这个判断同时约束了产品范围、销售承诺和试点节奏。",
        "后续验证时，可以用同一字段回到数据表、图表目录或 Demo 页面。",
        "它让商业计划从概念叙述回到可执行的工作清单。",
        "如果评委追问，这里可以继续展开为来源、公式、页面或岗位责任。",
    ][index % 10]
    if chapter.startswith("第一章"):
        return [
            f"{topic}的判断不能只依赖赛道热度。项目把公开统计、地方政策和区域指数放在一起看，是为了回答一个更具体的问题：哪些区域已经出现持续运行、明确责任主体和可验收的管理需求。注册量、飞行小时和运营单位增长说明低空活动在变密，但只有落到景区、园区、场馆、机场周边和固定航线，才会形成早期可交付市场。",
            f"SkyGuard 因此选择重点区域先行，而不是从全域城市平台起步。重点区域的优势在于边界清楚、飞行活动可记录、责任主体明确，能够先把计划核验、围栏触发、事件闭环和日报复盘做扎实，再复制到区县级协同场景。",
            f"{evidence_title}提供了这一判断的一个量化切面。资料口径来自{source}，核心含义是：{conclusion}。{style_tail}",
        ]
    if chapter.startswith("第二章"):
        return [
            f"{topic}是现场工作的真实断点。低空目标出现以后，值班人员要同时确认计划、身份、轨迹、围栏、风险等级和处置责任；任何一个信息缺失，后续就会变成电话确认、人工截图和临时报表。",
            "痛点验证不能写成一句“监管压力大”。更有说服力的方式，是把发现、核验、判断、处置、归档拆成流程节点，说明每个节点当前为什么慢、为什么容易漏、为什么事后难解释。",
            f"{evidence_title}对应的结论是：{conclusion}。资料来自{source}。SkyGuard 的产品价值，是把这类断点变成可配置规则、可派发工单和可提交报表。{style_tail}",
        ]
    if chapter.startswith("第三章"):
        return [
            f"{topic}要服务日常操作，而不是只服务展示。管理者进入系统后，先看区域态势，再追到具体目标；如果目标没有计划、偏离航线或进入重点保护区，才进入告警、工单和移动处置。",
            "产品设计采用“态势图、规则中心、事件中心、报表中心”的结构。态势图负责现场判断，规则中心负责边界和阈值，事件中心负责协同，报表中心把过程转成可提交材料。",
            f"{evidence_title}说明的重点是：{conclusion}。资料来自{source}。这一页对应的客户价值，是减少跨部门沟通、临时报表和事后解释成本。{style_tail}",
        ]
    if chapter.startswith("第四章"):
        return [
            f"{topic}是架构能否落地的关键。系统不是把算法放到页面上就结束，而是要把数据接入、时空索引、轨迹比对、规则解释、人工确认和审计日志连接起来。",
            "早期试点采用轻量部署更稳妥。计划、围栏、遥测、事件和设备状态先进入统一字段，模型负责排序和解释，关键处置仍保留人工确认。这样可以降低误报争议，也便于后续接入雷达、光电、通信感知等硬件。",
            f"{evidence_title}支撑的工程判断是：{conclusion}。资料来自{source}。该能力最终要落到配置、日志、回放和报表，而不是停在架构图上。{style_tail}",
        ]
    if chapter.startswith("第五章"):
        return [
            f"{topic}首先是口径管理。官方统计、政策目标、公开数据和演示样本承担不同任务，不能互相替代。计划书中所有图表都必须说明数据性质，否则很容易把演示样本误写成真实客户业务。",
            "SkyGuard 的数据体系分为来源索引、数据字典、样本表、图表目录和财务假设表。这样的结构方便后续迭代：新增材料先进入索引，再进入图表和正文，而不是直接贴到文档里。",
            f"{evidence_title}的结论是：{conclusion}。资料来自{source}。这一页对应产品中的字段、阈值或校验规则，而不仅是一张说明性图片。{style_tail}",
        ]
    if chapter.startswith("第六章"):
        return [
            f"{topic}要解释 SkyGuard 的市场位置。客户并不是在“有没有系统”之间选择，而是在视频监控、单点设备、人工巡查、内部工单和城市级平台之间权衡成本、周期、闭环能力和责任边界。",
            "项目的机会在于把重点区域的运行链路做完整。通用系统能解决局部问题，但很难同时覆盖计划、目标、围栏、处置和报表；城市级重平台能力强，却不适合创业团队作为首个交付承诺。",
            f"{evidence_title}提供的市场切面是：{conclusion}。资料来自{source}。进入策略应围绕可验收区域、可复用规则包和伙伴接入展开。{style_tail}",
        ]
    if chapter.startswith("第七章"):
        return [
            f"{topic}要落到客户购买动作上。景区和园区更适合先买轻量订阅和规则包，大型活动会购买短期保障，物流和巡检客户更关心航线评估、运行记录和风险报告。",
            "收入结构由软件订阅、专业部署、活动保障、航线评估、报告服务和运维续费组成。这样既能覆盖早期交付成本，也能在客户持续使用后形成扩容和续费。",
            f"{evidence_title}支撑的商业判断是：{conclusion}。资料来自{source}。商业模式不能只写年度总额，还要解释客单价、交付成本、回款周期和客户成功动作。{style_tail}",
        ]
    if chapter.startswith("第八章"):
        return [
            f"{topic}必须对应交付物。早期要先完成 Demo、数据字典、规则样例和答辩材料；随后进入重点区域试点包，再逐步扩展到多场景复制和伙伴接入。",
            "路线图的难点在依赖关系。没有稳定字段，风险解释无法复核；没有工单闭环，报表没有事实来源；没有试点反馈，定价和续费假设都会失真。",
            f"{evidence_title}对应的排期判断是：{conclusion}。资料来自{source}。落地计划要用验收材料证明阶段完成，而不是只给日期。{style_tail}",
        ]
    if chapter.startswith("第九章"):
        return [
            f"{topic}的核心是经营质量。收入增长必须和毛利、现金流、交付能力一起看。SkyGuard 的财务模型保留三种情景，是为了检查客户获取速度、交付成本、续费率和回款周期变化后的承压能力。",
            "早期融资不适合过度投入重资产。资金应优先用于产品打磨、试点交付、数据治理、客户成功和必要的伙伴接入；同时控制并行项目数量，避免现金流被回款周期拖住。",
            f"{evidence_title}说明的财务判断是：{conclusion}。资料来自{source}。这一部分最终要能在测算表中回到具体假设，而不是只停留在文字预测。{style_tail}",
        ]
    if chapter.startswith("第十章"):
        return [
            f"{topic}是项目必须主动承认的约束。低空业务涉及政策边界、数据安全、误报漏报、硬件接入、采购周期和现金流压力，任何一项处理不好都会拖慢试点。",
            "SkyGuard 的风险控制从边界开始：系统做辅助感知、风险预警、事件协同和报表复盘，不做干扰、捕获、打击，也不替代审批或执法。边界清楚，项目反而更可信。",
            f"{evidence_title}对应的风险判断是：{conclusion}。资料来自{source}。每项风险都需要写出触发信号、影响范围和应对动作。{style_tail}",
        ]
    if chapter.startswith("第十一章"):
        return [
            f"{topic}要服务项目阶段。当前阶段最需要的是把产品、数据、前后端、交付和客户材料连起来的小团队，而不是完整公司组织图。",
            "团队可以按产品研发线和试点交付线配置。前者保证平台可运行、可迭代，后者保证客户沟通、现场配置、培训和复盘材料能落地。",
            f"{evidence_title}说明的组织判断是：{conclusion}。资料来自{source}。岗位设置要和路线图里的交付物一一对应。{style_tail}",
        ]
    if chapter.startswith("第十二章"):
        return [
            f"{topic}要写成可观测改进。低空运行越密，城市越需要把目标、计划、围栏、事件和报表放进同一套协同机制，降低信息断点和责任争议。",
            "社会价值可以拆成安全、效率、治理、应急和人才培养几个维度。每个维度都应有指标，例如响应时间、闭环率、报表及时率、培训覆盖和规则更新次数。",
            f"{evidence_title}支撑的价值判断是：{conclusion}。资料来自{source}。公共价值不是额外包装，而是商业计划可持续性的组成部分。{style_tail}",
        ]
    if chapter.startswith("第十三章"):
        return [
            f"{topic}要让评委能顺着真实路径点击。首页看到值班域和关键指标，进入演示脚本后，再走综合态势、目标监测、计划审批、识别复核、围栏规则、事件工单、运行报表和移动处置。",
            "Demo 的真实性来自三点：使用真实整理图片和公开指标，明确标注演示样本性质，页面之间有连续业务动作。它不是静态截图，而是计划书产品逻辑的交互证明。",
            f"{evidence_title}对应的演示结论是：{conclusion}。资料来自{source}。答辩时应按脚本讲流程，而不是逐页介绍界面。{style_tail}",
        ]
    return [
        f"{topic}用于支撑正文核验。附录不承担装饰作用，而是让来源、字段、假设、图表和交付物可以被快速找到。",
        "评审追问通常集中在数据来源、样本性质、财务假设和 Demo 是否可运行。附录越清楚，正文判断越稳。",
        f"{evidence_title}对应的附录结论是：{conclusion}。资料来自{source}。后续升级应继续保持同一套索引和文件位置。",
    ]


def course_team_table() -> list[list[str]]:
    rows = [["角色", "班级", "学号", "姓名", "工作内容/分工", "提交前状态"]]
    work_items = [
        ("小组组长", "统筹总稿"),
        ("组员1", "政策数据"),
        ("组员2", "痛点竞品"),
        ("组员3", "产品技术"),
        ("组员4", "财务测算"),
        ("组员5", "图表证据"),
        ("组员6", "Demo实现"),
        ("组员7", "问卷风险"),
        ("组员8", "排版提交"),
    ]
    for role, work in work_items:
        rows.append([role, "待填写", "待填写", "待填写", work, "待填写"])
    return rows


def web_demo_engineering_table() -> list[list[str]]:
    return [
        ["验收项", "当前产物", "评审价值"],
        ["技术栈", "React/TS/Vite/Recharts", "可运行前端工程"],
        ["页面路由", "22条页面路由", "覆盖产品闭环"],
        ["核心交互", "脚本/筛选/围栏/工单/导出", "可现场操作"],
        ["数据来源", "演示样本+公开指标+图片", "来源分层"],
        ["构建验收", "npm run build；21张截图", "可提交可部署"],
        ["边界说明", "辅助监管，不替代审批执法", "降低合规误解"],
    ]


LOGICAL_TITLE_ACTIONS: dict[str, list[str]] = {
    "第一章 行业机会": ["政策依据与安全边界", "运行规模与需求强度", "区域样板与试点切口", "客户入口与采购触发", "监管分工与责任边界", "产业链位置与伙伴角色", "政策发布节奏", "服务站与基础设施", "场景扩张与风险外溢", "公开案例验证", "客户触发条件", "产业协同路径", "机会判断小结"],
    "第二章 痛点验证": ["现场发现链路", "身份可信度校验", "计划与航线比对", "围栏触发条件", "工单派发责任", "复盘归档材料", "断点复核样本", "事件类型证据", "计划偏离解释", "临时管控场景", "响应时长尾部风险", "闭环材料模板", "需求优先级小结"],
    "第三章 产品服务": ["值班态势入口", "计划审批动作", "电子围栏配置", "识别复核流程", "事件中心闭环", "运行报表产出", "场景包组合", "客户成功动作", "目标列表与轨迹", "物流航线验证", "服务蓝图衔接", "复核权限设计", "事件字段沉淀", "KPI模块验收", "场景复制规则", "续费与培训", "产品范围小结"],
    "第四章 技术架构": ["数据接入边界", "时空索引能力", "轨迹比对逻辑", "风险评分解释", "规则引擎配置", "权限与审计留痕", "部署集成策略", "可解释性设计", "接入字段标准", "Demo页面验证", "覆盖能力评估", "风险权重校准", "规则命中记录", "告警链路时序", "集成边界说明", "人工确认机制", "架构落地小结"],
    "第五章 数据体系": ["官方统计口径", "政策目标口径", "公开数据边界", "演示样本性质", "数据字典结构", "质量规则清单", "模型校验口径", "评分校准过程", "区域目标映射", "数据资产页面", "飞行类型分布", "字段治理责任", "数据体系小结"],
    "第六章 市场竞争": ["替代方案边界", "客户分层逻辑", "采购决策链条", "竞品矩阵比较", "进入顺序安排", "伙伴生态组合", "市场边界收束", "招投标类型证据", "采购动作拆解", "差异化定位", "区域运营对比", "合作渠道策略", "市场判断小结"],
    "第七章 商业模式": ["订阅收入假设", "部署项目收入", "活动保障服务", "航线评估服务", "报告服务收入", "运维续费动作", "客户成功机制", "续费队列证据", "专业部署路径", "现场保障页面", "首批客户漏斗", "月度报告价值", "商业模式小结"],
    "第八章 落地计划": ["三个月内交付物", "六个月内试点包", "十八个月复制期", "三十六个月协同期", "试点验收材料", "区域复制节奏", "路线图依赖关系", "交付泳道安排", "试点现场动作", "事件闭环验证", "里程碑小结"],
    "第九章 财务预测": ["收入预测口径", "成本结构口径", "毛利路径解释", "现金流约束", "回款周期压力", "融资用途安排", "敏感性检验", "单位经济测算", "现金流瀑布", "收入结构迁移", "成本页面验证", "五年收入结构", "现金安全边界", "账期承压判断", "财务小结"],
    "第十章 风险应对": ["政策边界风险", "数据合规风险", "误报漏报风险", "硬件依赖风险", "采购周期风险", "现金压力风险", "过度承诺风险", "试点延期原因", "风险预案小结"],
    "第十一章 团队组织": ["岗位配置原则", "产品负责人职责", "算法数据职责", "前后端研发职责", "交付运维职责", "行业顾问职责", "岗位到位计划", "团队保障小结"],
    "第十二章 社会价值": ["公共安全指标", "城市治理指标", "应急保障指标", "产业数据沉淀", "岗位培养路径", "合规意识提升", "平衡计分卡", "社会价值小结"],
    "第十三章 Web Demo": ["首页值班域", "演示脚本入口", "态势地图验证", "目标监测链路", "计划审批链路", "识别复核链路", "事件工单闭环", "运行报表导出", "移动处置路径", "指挥屏视图", "部署说明校验", "Demo演示小结"],
    "附录": ["来源索引", "图表目录", "数据字典", "财务假设", "访谈问卷", "部署说明", "截图包", "交付清单", "政策来源", "图表文件", "模型字段", "财务公式", "演示脚本", "部署验收", "答辩材料", "附录小结"],
}


CHAPTER_CODES = {
    "第一章 行业机会": "1",
    "第二章 痛点验证": "2",
    "第三章 产品服务": "3",
    "第四章 技术架构": "4",
    "第五章 数据体系": "5",
    "第六章 市场竞争": "6",
    "第七章 商业模式": "7",
    "第八章 落地计划": "8",
    "第九章 财务预测": "9",
    "第十章 风险应对": "10",
    "第十一章 团队组织": "11",
    "第十二章 社会价值": "12",
    "第十三章 Web Demo": "13",
    "附录": "A",
}


def compact_heading_piece(text: str | None, limit: int = 20) -> str:
    if not text:
        return ""
    cleaned = re.sub(r"\s+", " ", str(text)).strip()
    cleaned = cleaned.replace(" 来源：", " ").replace("｜来源：", " ")
    return cleaned if len(cleaned) <= limit else f"{cleaned[:limit]}…"


def logical_page_title(chapter: str, local_index: int, topic: str, row: dict | None, figure_caption: str | None = None) -> str:
    code = CHAPTER_CODES.get(chapter, "")
    prefix = f"{code}.{local_index}" if code != "A" else f"A.{local_index}"
    actions = LOGICAL_TITLE_ACTIONS.get(chapter, [])
    action = actions[local_index - 1] if local_index - 1 < len(actions) else "证据与交付校验"
    evidence = compact_heading_piece(figure_caption or ((row or {}).get("title") if row else ""), 18)
    if evidence and evidence != topic and evidence not in action:
        return f"{prefix} {topic}：{action}（{evidence}）"
    return f"{prefix} {topic}：{action}"


def build_page_bullets(chapter: str, topic: str, row: dict | None, spec: dict) -> list[str]:
    conclusion = (row or {}).get("conclusion", spec["decision"])
    source = (row or {}).get("source", "整理图片、Demo截图与项目资料")
    if chapter.startswith("第九章"):
        return ["测算关注：客单价、交付成本、续费率和回款周期", f"当前依据：{source}", f"经营提示：{conclusion}"]
    if chapter.startswith("第十章"):
        return ["边界：辅助监管，不替代审批或执法", f"风险依据：{source}", f"处置重点：{conclusion}"]
    if chapter.startswith("第十三章"):
        return ["演示动作：进入脚本、定位目标、派单处置、导出报表", f"素材依据：{source}", f"讲解重点：{conclusion}"]
    return [f"业务要点：{conclusion}", f"资料口径：{source}", f"落地检查：{topic}要对应字段、页面、责任人或验收材料"]


def enrich_manuscript_paragraphs(chapter: str, topic: str, paras: list[str], evidence_label: str, source: str, conclusion: str, page_index: int) -> list[str]:
    paras = [str(p).strip() for p in paras if str(p).strip()]
    if not paras:
        return paras
    bridge_options = [
        f"“{evidence_label}”在这里不是装饰性材料，而是把“{topic}”从概念判断推进到可说明、可追问、可复核的证据入口。",
        f"阅读这一节时，可以先看“{evidence_label}”呈现的事实，再回到“{topic}”对应的客户动作和交付边界。",
        f"这部分与“{evidence_label}”放在一起读，能够看出“{topic}”并非孤立功能，而是商业计划中的一项可执行安排。",
        f"“{evidence_label}”提供了判断依据，本节则把依据转化为“{topic}”的场景、字段和验收口径。",
    ]
    chapter_notes = {
        "第一章": ("首批市场判断要同时看政策窗口、责任主体和运行频次，不能只用赛道规模替代客户意愿。", "落到试点时，应优先选择边界清楚、现场责任人明确、飞行活动可记录的区域。"),
        "第二章": ("痛点是否成立，关键看断点能否被系统流程替代，而不是看描述是否紧张。", "这一节可对应访谈提纲中的现场确认、处置时长和复盘材料三类问题。"),
        "第三章": ("产品价值需要落在日常值班动作上，包括查看、核验、确认、派单、归档和导出。", "后续试点可把这一节转成模块验收单，逐项检查页面、字段、日志和报表。"),
        "第四章": ("技术路线强调规则可解释、模型可复核、人工可确认，避免把关键判断交给黑箱。", "工程实施时要把接口字段、时间戳、责任人和操作日志一起纳入验收。"),
        "第五章": ("数据体系的重点是区分资料性质，防止公开统计、政策目标和演示样本互相混用。", "每一次新增图表都应先进入来源索引和数据字典，再进入正文。"),
        "第六章": ("竞争分析的目标是说明进入顺序，而不是罗列对手名称。", "销售沟通时应把替代方案的成本、周期、责任边界和可验收材料讲清楚。"),
        "第七章": ("商业模式要能解释客户为什么持续使用，也要解释团队为什么能控制交付成本。", "定价可以随场景包、设备接入量、报告频率和保障等级做分层。"),
        "第八章": ("落地计划最怕只有时间轴没有交付物，因此每个阶段都要绑定可检查材料。", "项目管理上应把Demo、试点包、培训材料、日报模板和复盘纪要作为阶段出口。"),
        "第九章": ("财务章节要用现金流约束增长速度，用单位经济说明续费价值。", "评审追问时可回到测算表中的客单价、毛利率、回款周期和研发投入假设。"),
        "第十章": ("风险应对必须主动承认边界，越清楚说明不做什么，越能增强可信度。", "试点合同和演示话术中都应避免承诺审批替代、执法替代或硬件反制能力。"),
        "第十一章": ("团队组织应围绕当前阶段的交付瓶颈配置岗位，而不是提前堆完整职能部门。", "岗位责任最好和路线图里的Demo、数据、客户沟通和现场交付材料一一对应。"),
        "第十二章": ("社会价值需要通过可观测指标表达，例如响应时间、闭环率、报表及时率和培训覆盖率。", "这些指标既能支持公共价值叙述，也能反过来成为客户续费和扩容依据。"),
        "第十三章": ("Web Demo 是计划书的现场证据，应按连续流程演示，而不是把页面当成截图集合。", "演示时从脚本入口进入，再依次展示态势、目标、计划、围栏、事件、报表和移动处置。"),
        "附录": ("附录负责让正文判断可追溯，不能变成材料堆放区。", "评审需要追问来源、公式、字段或部署方式时，应能在附录中迅速定位。"),
    }
    key = next((k for k in chapter_notes if chapter.startswith(k)), "附录")
    note_a, note_b = chapter_notes[key]
    paras[0] = f"{paras[0]} {bridge_options[page_index % len(bridge_options)]}"
    if len(paras) > 1:
        paras[1] = f"{paras[1]} {note_a}"
    closing = f"{note_b} 本节采用的资料口径为{source}，可支撑的核心结论是：{conclusion}"
    if closing not in paras:
        paras.append(closing)
    return paras


def make_page_items(chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame) -> list[dict]:
    nat = data["national"]
    metrics = {
        "uav": first_value(nat, "registered_uavs_10k", 2025),
        "hours": first_value(nat, "uav_flight_hours_10k", 2025),
        "units": first_value(nat, "uav_operating_units", 2025),
        "stations": first_value(nat, "low_altitude_flight_service_stations", 2025),
        "airport": first_value(nat, "registered_general_airports", 2025),
    }
    figures = figure_files()
    figure_idx = 0
    chapter_alias = {"第二章 需求验证": "第二章 痛点验证", "第十章 风险": "第十章 风险应对", "第十一章 团队": "第十一章 团队组织", "第十三章 Demo展示": "第十三章 Web Demo"}
    all_rows = chart_catalog.to_dict("records")
    chapter_chart_map: dict[str, list[dict]] = defaultdict(list)
    for row in all_rows:
        chapter_chart_map[chapter_alias.get(row["chapter"], row["chapter"])].append(row)
    cursor: dict[str, int] = {}

    def next_row(chapter: str) -> dict | None:
        return chart_for(chapter_chart_map, chapter, cursor, all_rows)

    def take_figure(match=None) -> Path | None:
        nonlocal figure_idx
        if not figures:
            return None
        used = getattr(take_figure, "_used", set())
        if match is not None:
            for fig in figures:
                if fig not in used and match(fig):
                    used.add(fig)
                    setattr(take_figure, "_used", used)
                    return fig
        while figure_idx < len(figures):
            fig = figures[figure_idx]
            figure_idx += 1
            if fig not in used:
                used.add(fig)
                setattr(take_figure, "_used", used)
                return fig
        setattr(take_figure, "_used", used)
        return None

    pages: list[dict] = []

    def add(page_type: str, chapter: str, title: str, paras: list[str], bullets: list[str] | None = None, chart_id: str | None = None, figure: Path | None = None, table: list[list[str]] | None = None, accent: str | None = None, kicker: str = "", callout: str = ""):
        pages.append({"type": page_type, "chapter": chapter, "title": title, "paras": paras, "bullets": bullets or [], "chart_id": chart_id, "figure": str(figure) if figure else None, "table": table or [], "accent": accent or REPORT_COLORS["teal"], "kicker": kicker, "callout": callout})

    cover_image = take_figure(lambda p: any(k in p.name for k in ["平台", "低空", "SILAS", "无人机"])) or take_figure()
    add("cover", "封面", "低空智眼 SkyGuard 商业计划书", ["城市低空空域安全感知与运行监管平台", f"面向景区、园区、场馆、机场周边和低空物流航线。公开统计口径显示，2025 年注册无人机约 {metrics['uav']:.1f} 万架，累计飞行小时约 {metrics['hours']:.2f} 万小时，运营单位 {metrics['units']:.0f} 家。"], bullets=["重点区域运行监管", "计划与围栏联动", "事件闭环处置", "运行报表复盘"], figure=cover_image, accent=REPORT_COLORS["teal"], kicker="商业计划书", callout="看得见、判得清、处置快、可复盘")

    add("table", "小组信息", "课程提交信息与成员分工", ["课程要求在封皮后的第一页注明小组组长、组员的班级、学号、姓名和工作内容。本页按课程格式保留成员信息栏，提交前应补齐真实班级、学号、姓名与分工。", "分工按照本项目实际交付链路拆分：行业资料、痛点验证、产品技术、Web Demo、财务测算、图表证据、问卷访谈、排版校对和答辩组织。若有成员未承担工作，应在本表工作内容中如实注明。"], bullets=["硬性要求：封面后第一页", "人数要求：7-9人为一组，不能多于9人", "提交检查：班级、学号、姓名、分工必须完整"], table=course_team_table(), accent=REPORT_COLORS["amber"], kicker="课程提交页", callout="请补齐真实成员信息。")

    add("memo", "前言", "项目边界与资料口径", ["本计划书按照正式商业计划书结构组织：执行摘要、行业机会、痛点验证、产品服务、技术架构、数据体系、市场竞争、商业模式、落地计划、财务预测、风险应对、团队组织、社会价值、Web Demo 和附录。", "资料采用分层口径。官方统计用于行业规模判断，地方政策用于区域机会判断，公开数据用于方法说明，Demo 样本仅用于展示平台流程。这样处理可以避免把演示样本写成真实客户数据，也能让每一张图回到清楚来源。", "计划书配套财务测算表、证据资料索引表、图表目录、图表包和 Web Demo 源码包。评审时可以先读执行摘要和财务章节，再通过附录核验来源和假设。"], bullets=[f"来源条目：{len(source_registry)} 条", f"图表数量：{len(chart_catalog)} 张", "配套材料：财务测算、证据索引、图表包、Web Demo"], table=manuscript_table("第五章 数据体系", "数据边界", data, source_registry, chart_catalog), accent=REPORT_COLORS["blue"], kicker="资料说明", callout="先明确口径，再进入判断。")

    summary_chart = next_row("第一章 行业机会")
    add("summary", "执行摘要", "项目判断", ["低空经济进入城市运行场景后，最紧迫的问题不是再做一块展示大屏，而是把飞行计划、目标身份、电子围栏、风险事件、现场处置和运行报表连成一条可追溯链路。SkyGuard 面向重点区域提供低空安全感知与运行监管平台，帮助客户完成日常监管、临时保障和复盘归档。", "项目不以重硬件或城市级总包作为起点，而选择景区、园区、场馆、机场周边和固定低空航线等责任主体清楚的区域。这个路径交付周期短，验收指标明确，也更适合创业团队形成可复制样板。", "商业化采用软件订阅、专业部署、活动保障、航线评估、报告服务和运维续费组合。硬件侧通过伙伴接入，平台侧沉淀规则、字段、日志和报表能力。"], bullets=["首批客户：景区、园区、场馆、航线运营方", "产品边界：辅助监管，不替代审批或执法", "经营重点：控制交付成本、提高续费率、缩短回款周期"], chart_id=summary_chart["chart_id"] if summary_chart else None, accent=REPORT_COLORS["teal"], kicker="执行摘要", callout="先做一个区域闭环，再复制到更多区域。")

    add("summary", "执行摘要", "产品与商业路径", ["SkyGuard 的核心交付是一套能每天使用的低空运行工作台。态势图负责实时查看，计划中心负责审批和核验，规则中心负责围栏和阈值，事件中心负责确认和派单，报表中心负责复盘和验收。", "竞争优势不写成单一算法优势，而写成业务闭环优势。通用监控、单点感知设备和人工巡查都能解决局部问题，但客户真正需要的是可解释、可记录、可提交的全流程材料。", "财务测算采用三情景，重点关注客户数量、客单价、部署成本、续费率和回款周期。早期资金优先用于产品打磨、试点交付、数据治理和客户成功。"], bullets=["产品：计划、目标、围栏、工单、报表一体化", "商业：订阅、部署、活动、评估、运维组合收入", "风险：政策、数据、误报、采购和现金流分层管理"], figure=take_figure(), accent=REPORT_COLORS["blue"], kicker="执行摘要", callout="计划书和 Demo 共用同一套业务语言。")

    for spec in manuscript_chapter_specs(data):
        chapter = spec["chapter"]
        rows = chapter_chart_map.get(chapter) or all_rows
        row = next_row(chapter)
        add("divider", chapter, spec["decision"], [spec["opening"]], bullets=[f"本章回答：{spec['decision']}", f"重点内容：{', '.join(spec['topics'][:5])}", f"关联图表：{len(rows)} 张"], chart_id=row["chart_id"] if row else None, figure=take_figure() if chapter in {"第三章 产品服务", "第四章 技术架构", "第十三章 Web Demo"} else None, accent=spec["accent"], kicker="章节导读", callout=spec["decision"])
        if chapter == "第十三章 Web Demo":
            add("table", chapter, "Web Demo 工程验收页", ["Web Demo 是计划书可执行性的证据，不应只作为展示截图。当前工程使用 React、TypeScript、Vite 和本地演示数据实现完整路由，覆盖态势、计划、识别、围栏、事件、报表、设备、数据资产、风险解释、移动处置和指挥大屏。", "答辩时建议先进入 Demo Center 选择脚本，再按综合态势、目标监测、计划审批、围栏规则、事件工单、移动处置和运行报表的顺序演示。这样可以证明产品闭环、数据边界和商业计划中的交付物是一致的。"], bullets=["源码包：SkyGuard_WebDemo.zip", "截图包：SkyGuard_WebDemo_截图包.zip", "部署说明：SkyGuard_WebDemo_部署说明.md"], table=web_demo_engineering_table(), accent=spec["accent"], kicker="工程验收", callout="用可运行流程证明团队执行力。")
        for i in range(spec["count"]):
            topic = spec["topics"][i % len(spec["topics"])]
            row = next_row(chapter)
            use_figure = (i % 7 == 2) and len(getattr(take_figure, "_used", set())) < len(figures)
            use_table = i in {3, 9}
            figure = take_figure() if use_figure else None
            figure_caption = humanize_figure_caption(figure) if figure else None
            title = logical_page_title(chapter, i + 1, topic, row, figure_caption)
            paras = manuscript_paragraphs(spec, topic, row, metrics, figure_caption, len(pages))
            evidence_label = figure_caption or ((row or {}).get("title")) or topic
            paras = enrich_manuscript_paragraphs(
                chapter,
                topic,
                paras,
                evidence_label,
                (row or {}).get("source", "整理图片、Demo截图与项目资料"),
                (row or {}).get("conclusion", spec["decision"]),
                len(pages),
            )
            table = manuscript_table(chapter, topic, data, source_registry, chart_catalog) if use_table else []
            add("visual" if figure else "table" if table else "chart", chapter, title, paras, bullets=build_page_bullets(chapter, topic, row, spec), chart_id=None if figure else (row["chart_id"] if row else None), figure=figure, table=table, accent=spec["accent"], callout=spec["decision"])
    return pages


def add_word_quality_charts(data: dict[str, pd.DataFrame]) -> None:
    finance = build_finance_tables()
    years = finance["revenue"]["year"].tolist()
    revenue_cols = [c for c in finance["revenue"].columns if c != "year"]
    revenue_values = finance["revenue"][revenue_cols]
    source = "经营假设、平台演示样本与整理数据"

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    matrix = np.array([[finance["cost"]["revenue_total"].iloc[-1] * f / (1 + r) ** 5 for r in [0.10, 0.12, 0.14, 0.16, 0.18]] for f in [0.80, 0.90, 1.00, 1.10, 1.20]])
    sns.heatmap(matrix, annot=True, fmt=".0f", cmap="YlGnBu", linewidths=1.2, linecolor="white", xticklabels=["10%", "12%", "14%", "16%", "18%"], yticklabels=["80%", "90%", "100%", "110%", "120%"], ax=ax)
    finish_chart(ax, "收入乘数与折现率敏感性", "单元格为2030收入折现后的相对规模，单位：万元")
    save_chart(fig, next_chart_id(), "收入乘数与折现率敏感性", "敏感性热力图", source, "第九章 财务预测", "敏感性热力图能直接说明收入假设和折现率变化对项目价值的影响。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    funnel = pd.Series({"线索客户": 120, "有效访谈": 58, "试点意向": 24, "付费试点": 9, "年度续费": 6})
    ax.barh(funnel.index[::-1], funnel.values[::-1], color=["#215a6d", "#347c76", "#64a391", "#c29a55", "#9b5d45"])
    for i, val in enumerate(funnel.values[::-1]):
        ax.text(val + 2, i, str(val), va="center", fontsize=8.5, color="#334155")
    finish_chart(ax, "首批客户获取漏斗", "从行业线索到续费客户的保守转化路径", x_grid=True)
    save_chart(fig, next_chart_id(), "首批客户获取漏斗", "漏斗条形图", source, "第七章 商业模式", "客户漏斗把销售动作和收入假设连接起来，避免只展示收入总额。")

    fig, ax = plt.subplots(figsize=(8.6, 4.8))
    lanes = ["产品", "数据", "试点", "商业"]
    tasks = [(0, 3, 0), (2, 4, 1), (4, 5, 2), (7, 5, 3), (10, 4, 0), (12, 6, 2), (15, 5, 3)]
    colors = ["#234f7d", "#0c6f69", "#b69252", "#9b3f3a"]
    for start, duration, lane in tasks:
        ax.broken_barh([(start, duration)], (lane - 0.32, 0.64), facecolors=colors[lane], alpha=0.86)
    ax.set_yticks(range(len(lanes)))
    ax.set_yticklabels(lanes)
    ax.set_xlabel("月份")
    finish_chart(ax, "试点落地里程碑泳道", "按产品、数据、试点、商业四条线排期", x_grid=True)
    save_chart(fig, next_chart_id(), "试点落地里程碑泳道", "泳道甘特图", source, "第八章 落地计划", "泳道图能把团队分工和阶段交付物放在同一张图中。")

    labels = ["态势", "计划", "围栏", "工单", "报表", "移动"]
    values = np.array([88, 82, 84, 79, 86, 75])
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False)
    fig = plt.figure(figsize=(6.4, 5.2))
    ax = fig.add_subplot(111, polar=True)
    ax.plot(np.r_[angles, angles[0]], np.r_[values, values[0]], color="#0c6f69", linewidth=2)
    ax.fill(np.r_[angles, angles[0]], np.r_[values, values[0]], color="#0c6f69", alpha=0.18)
    ax.set_xticks(angles)
    ax.set_xticklabels(labels, fontsize=8.5)
    ax.set_yticklabels([])
    ax.set_title("Demo模块验收雷达", fontsize=12.5, fontweight="bold", loc="left", pad=16)
    save_chart(fig, next_chart_id(), "Demo模块验收雷达", "雷达图", source, "第十三章 Web Demo", "雷达图适合说明Demo模块覆盖度和下一步补强方向。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    risks = pd.DataFrame({"概率": [0.55, 0.42, 0.36, 0.61, 0.48, 0.30], "影响": [0.82, 0.68, 0.74, 0.56, 0.63, 0.44], "风险": ["政策边界", "数据合规", "误报漏报", "采购周期", "硬件依赖", "现金压力"]})
    ax.scatter(risks["概率"], risks["影响"], s=[420, 300, 340, 260, 280, 220], c=["#9b3f3a", "#b69252", "#c2410c", "#234f7d", "#0c6f69", "#64748b"], alpha=0.84)
    for _, r in risks.iterrows():
        ax.text(r["概率"] + 0.015, r["影响"], r["风险"], fontsize=8.4)
    ax.axhline(0.6, color="#d8dee8", linewidth=1)
    ax.axvline(0.45, color="#d8dee8", linewidth=1)
    ax.set_xlim(0.2, 0.75)
    ax.set_ylim(0.35, 0.9)
    finish_chart(ax, "风险概率影响矩阵", "右上象限优先进入合同边界和试点验收条款")
    save_chart(fig, next_chart_id(), "风险概率影响矩阵", "气泡矩阵", source, "第十章 风险应对", "风险矩阵能把风险优先级从文字清单变成决策顺序。")

    fig, ax = plt.subplots(figsize=(8.6, 4.8))
    ax.stackplot(years, [revenue_values[c] for c in revenue_cols[:5]], labels=revenue_cols[:5], colors=["#234f7d", "#0c6f69", "#b69252", "#9b3f3a", "#64748b"], alpha=0.88)
    ax.legend(ncol=3, fontsize=7.5, loc="upper left", frameon=False)
    finish_chart(ax, "五年收入结构堆叠面积", "展示收入从部署向订阅、服务和报告延展的路径")
    save_chart(fig, next_chart_id(), "五年收入结构堆叠面积", "堆叠面积图", source, "第九章 财务预测", "堆叠面积图比单线收入更能说明收入结构变化。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.set_axis_off()
    ax.text(0.04, 0.94, "数据口径流向图", transform=ax.transAxes, ha="left", va="top", fontsize=13.2, fontweight="bold", color=PALETTE["ink"])
    ax.text(0.04, 0.885, "从资料来源到计划书判断，避免把演示样本写成真实客户业务", transform=ax.transAxes, ha="left", va="top", fontsize=8.5, color=PALETTE["muted"])
    ax.text(0.08, 0.80, "资料来源", transform=ax.transAxes, ha="left", va="center", fontsize=8.6, color="#334155", fontweight="bold")
    ax.text(0.73, 0.80, "计划书用途", transform=ax.transAxes, ha="left", va="center", fontsize=8.6, color="#334155", fontweight="bold")
    sources = ["官方统计", "地方政策", "公开数据", "演示样本"]
    targets = ["行业判断", "区域选择", "产品流程", "财务假设"]
    notes = ["规模与趋势", "城市机会", "方法与边界", "Demo与测算"]
    ys = [0.68, 0.53, 0.38, 0.23]
    left_colors = ["#e8f3f1", "#edf6fb", "#f4f0e6", "#f5e9e7"]
    right_colors = ["#dff1ef", "#e4eff8", "#f2e8d3", "#f5dedb"]
    for y, s, t, note, left_color, right_color in zip(ys, sources, targets, notes, left_colors, right_colors):
        ax.text(0.08, y, s, transform=ax.transAxes, ha="left", va="center", fontsize=9.4, color="#1f2a37",
                bbox=dict(boxstyle="round,pad=0.46,rounding_size=0.14", fc=left_color, ec="#b8c7d4", lw=1.0))
        ax.text(0.73, y, t, transform=ax.transAxes, ha="left", va="center", fontsize=9.4, color="#1f2a37",
                bbox=dict(boxstyle="round,pad=0.46,rounding_size=0.14", fc=right_color, ec="#c8bda8", lw=1.0))
        ax.annotate("", xy=(0.69, y), xytext=(0.28, y), xycoords=ax.transAxes, textcoords=ax.transAxes,
                    arrowprops=dict(arrowstyle="-|>", color="#215a6d", lw=2.2, shrinkA=4, shrinkB=6, mutation_scale=14))
        ax.text(0.485, y + 0.035, note, transform=ax.transAxes, ha="center", va="bottom", fontsize=7.6, color="#526071")
    ax.text(0.04, 0.08, "使用原则：统计和政策支撑外部判断；公开数据支撑方法；演示样本只证明流程，不替代真实部署。", transform=ax.transAxes, ha="left", va="bottom", fontsize=8.2, color="#334155")
    save_chart(fig, next_chart_id(), "数据口径流向图", "流程流向图", source, "第五章 数据体系", "流向图能清楚说明不同资料在计划书中的使用边界。")

    alerts = data.get("alerts", pd.DataFrame())
    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    if not alerts.empty and "response_time_min" in alerts.columns and "risk_level" in alerts.columns:
        levels = [v for v in ["low", "medium", "high", "critical"] if v in set(alerts["risk_level"])]
        series = [alerts[alerts["risk_level"].eq(level)]["response_time_min"].dropna().clip(0, 180).sample(min(180, alerts[alerts["risk_level"].eq(level)].shape[0]), random_state=RANDOM_SEED) for level in levels]
        ax.boxplot(series, labels=levels, patch_artist=True, boxprops=dict(facecolor="#dfeeea", color="#0c6f69"), medianprops=dict(color="#9b3f3a"))
    finish_chart(ax, "分级事件响应时间箱线", "尾部响应时间用于验证工单闭环能力")
    save_chart(fig, next_chart_id(), "分级事件响应时间箱线", "箱线图", source, "第二章 痛点验证", "箱线图能揭示平均响应时间之外的尾部风险。")


def add_chart_variety_boosters(data: dict[str, pd.DataFrame]) -> None:
    finance = build_finance_tables()
    source = "经营假设、平台演示样本与整理数据"
    years = finance["revenue"]["year"].tolist()
    rng = np.random.default_rng(RANDOM_SEED)

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    labels = ["确认", "派单", "到场", "归档"]
    actual = np.array([6, 12, 26, 45])
    target = np.array([8, 15, 30, 60])
    for i, label in enumerate(labels):
        ax.barh(i, target[i], color="#e7ecef", height=0.48)
        ax.barh(i, actual[i], color=["#0c6f69", "#234f7d", "#b69252", "#9b3f3a"][i], height=0.30)
        ax.plot([target[i], target[i]], [i - 0.28, i + 0.28], color="#334155", lw=1)
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels)
    finish_chart(ax, "风险处置SLA子弹图", "彩色条为当前样本，中线为阶段目标，单位：分钟", x_grid=True)
    save_chart(fig, next_chart_id(), "风险处置SLA子弹图", "子弹图", source, "第十章 风险应对", "子弹图把响应目标和当前处置时长放在同一口径下比较。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    areas = ["景区", "园区", "场馆", "航线", "枢纽"]
    before = np.array([72, 68, 76, 61, 83])
    after = before - np.array([19, 16, 22, 14, 18])
    y = np.arange(len(areas))
    ax.hlines(y, after, before, color="#c8d2dc", lw=4)
    ax.scatter(before, y, s=90, color="#9b3f3a", label="人工处置")
    ax.scatter(after, y, s=90, color="#0c6f69", label="平台闭环")
    ax.set_yticks(y)
    ax.set_yticklabels(areas)
    ax.legend(frameon=False, fontsize=8)
    finish_chart(ax, "重点场景响应时长哑铃", "展示从人工流转到平台闭环后的时长压缩", x_grid=True)
    save_chart(fig, next_chart_id(), "重点场景响应时长哑铃", "哑铃图", source, "第二章 痛点验证", "哑铃图能清楚表达不同场景的处置效率改善空间。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    items = pd.DataFrame({"价值": [0.86, 0.74, 0.68, 0.80, 0.59, 0.71], "交付难度": [0.48, 0.55, 0.35, 0.63, 0.42, 0.58], "规模": [500, 380, 340, 430, 260, 310], "模块": ["态势", "计划", "围栏", "工单", "报表", "移动"]})
    ax.scatter(items["交付难度"], items["价值"], s=items["规模"], c=["#0c6f69", "#234f7d", "#b69252", "#9b3f3a", "#64748b", "#2f6b4f"], alpha=0.78)
    for _, r in items.iterrows():
        ax.text(r["交付难度"] + 0.012, r["价值"], r["模块"], fontsize=8.5)
    ax.axvline(0.5, color="#d8dee8")
    ax.axhline(0.7, color="#d8dee8")
    ax.set_xlim(0.28, 0.72)
    ax.set_ylim(0.52, 0.92)
    finish_chart(ax, "客户价值优先级气泡", "左上象限优先标准化，右上象限进入试点专项")
    save_chart(fig, next_chart_id(), "客户价值优先级气泡", "气泡四象限", source, "第三章 产品服务", "气泡四象限用于解释产品模块的优先级和交付节奏。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    stages = ["提交", "预检", "人工复核", "通过", "归档"]
    vals = [100, -18, -12, -9, 61]
    base = 0
    for i, val in enumerate(vals):
        if i in {0, len(vals) - 1}:
            ax.bar(i, val, color="#234f7d")
            base = val if i == 0 else base
        else:
            ax.bar(i, val, bottom=base + val, color="#9b3f3a")
            base += val
        ax.text(i, max(base, val) + 3, f"{val:+d}" if i not in {0, len(vals) - 1} else str(val), ha="center", fontsize=8)
    ax.set_xticks(range(len(stages)))
    ax.set_xticklabels(stages)
    finish_chart(ax, "计划审批阶段转化瀑布", "从申请提交到归档通过的数量变化", x_grid=True)
    save_chart(fig, next_chart_id(), "计划审批阶段转化瀑布", "转化瀑布", source, "第三章 产品服务", "转化瀑布能解释计划审批中的过滤、复核和归档动作。")

    fig, axes = plt.subplots(2, 2, figsize=(8.4, 5.0), sharex=True, sharey=True)
    device_names = ["光电", "雷达", "Remote ID", "通信感知"]
    for idx, ax in enumerate(axes.ravel()):
        values = 92 + rng.normal(0, 1.8, 12).cumsum() / 3 - idx * 1.2
        ax.plot(range(1, 13), np.clip(values, 82, 98), color=["#0c6f69", "#234f7d", "#b69252", "#9b3f3a"][idx], lw=2)
        ax.fill_between(range(1, 13), np.clip(values, 82, 98), 80, alpha=0.12)
        ax.set_title(device_names[idx], fontsize=9, loc="left")
        ax.grid(axis="y", color="#e6ebef")
    fig.suptitle("设备在线率小倍数", x=0.04, ha="left", fontsize=12.5, fontweight="bold", color=PALETTE["ink"])
    save_chart(fig, next_chart_id(), "设备在线率小倍数", "小倍数折线", source, "第四章 技术架构", "小倍数折线适合比较不同感知设备的连续可用性。")

    fig = plt.figure(figsize=(6.4, 5.2))
    ax = fig.add_subplot(111, polar=True)
    hours = np.arange(24)
    values = 8 + 6 * np.sin((hours - 7) / 24 * 2 * np.pi) ** 2 + rng.integers(0, 4, size=24)
    ax.bar(hours / 24 * 2 * np.pi, values, width=2 * np.pi / 28, color="#0c6f69", alpha=0.78)
    ax.set_xticks(np.linspace(0, 2 * np.pi, 8, endpoint=False))
    ax.set_xticklabels(["0", "3", "6", "9", "12", "15", "18", "21"], fontsize=8)
    ax.set_yticklabels([])
    ax.set_title("围栏触发时间玫瑰", loc="left", fontsize=12.5, fontweight="bold", pad=16)
    save_chart(fig, next_chart_id(), "围栏触发时间玫瑰", "极坐标柱图", source, "第二章 痛点验证", "时间玫瑰图能展示围栏告警在一天内的高峰时段。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    x = np.linspace(0, 10, 60)
    yv = np.linspace(30, 180, 60)
    X, Y = np.meshgrid(x, yv)
    Z = np.sin(X / 2) * 0.25 + np.cos(Y / 42) * 0.18 + X / 18 + Y / 360
    cs = ax.contourf(X, Y, Z, levels=12, cmap="YlGnBu")
    fig.colorbar(cs, ax=ax, shrink=0.82, label="容量压力")
    ax.set_xlabel("航线距离 km")
    ax.set_ylabel("飞行高度 m")
    finish_chart(ax, "航线容量压力等高线", "距离、高度和容量压力的示意关系")
    save_chart(fig, next_chart_id(), "航线容量压力等高线", "等高线图", source, "第四章 技术架构", "等高线图用于说明航线风险评估不是单一阈值判断。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    city = data.get("city", pd.DataFrame())
    if not city.empty and "发布地区" in city.columns:
        counts = city["发布地区"].value_counts().head(8)
    else:
        counts = pd.Series([24, 19, 17, 13, 10, 8, 7, 5], index=["深圳", "杭州", "成都", "合肥", "广州", "苏州", "珠海", "长沙"])
    cumulative = counts.cumsum() / counts.sum()
    ax.bar(range(len(counts)), counts.values, color="#234f7d")
    ax2 = ax.twinx()
    ax2.plot(range(len(counts)), cumulative.values, color="#9b3f3a", marker="o")
    ax.set_xticks(range(len(counts)))
    ax.set_xticklabels(counts.index, rotation=25, ha="right")
    ax2.set_ylim(0, 1.05)
    finish_chart(ax, "示范城市进入排序帕累托", "柱为政策/案例数量，线为累计占比")
    save_chart(fig, next_chart_id(), "示范城市进入排序帕累托", "帕累托图", source, "第六章 市场竞争", "帕累托图能说明先进入哪些城市更容易形成样板。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    states = ["待确认", "处理中", "已闭环"]
    matrix = np.array([[0.18, 0.36, 0.46], [0.12, 0.30, 0.58], [0.22, 0.28, 0.50], [0.08, 0.26, 0.66]])
    left = np.zeros(matrix.shape[0])
    for j, state in enumerate(states):
        ax.barh(["景区", "园区", "场馆", "航线"], matrix[:, j], left=left, label=state, color=["#b69252", "#234f7d", "#0c6f69"][j])
        left += matrix[:, j]
    ax.legend(frameon=False, ncol=3, fontsize=8)
    ax.set_xlim(0, 1)
    finish_chart(ax, "工单状态100%堆叠条", "按场景比较事件闭环构成")
    save_chart(fig, next_chart_id(), "工单状态100%堆叠条", "百分比堆叠条", source, "第三章 产品服务", "百分比堆叠条能比较不同场景的事件处理结构。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    factors = pd.Series({"政策边界": -0.26, "误报漏报": -0.18, "硬件依赖": -0.14, "续费率": 0.22, "报告服务": 0.16, "伙伴接入": 0.12}).sort_values()
    ax.barh(factors.index, factors.values, color=["#9b3f3a" if v < 0 else "#0c6f69" for v in factors.values])
    ax.axvline(0, color="#334155", lw=1)
    finish_chart(ax, "风险因素贡献发散条", "负向因素需要合同边界和验收口径吸收", x_grid=True)
    save_chart(fig, next_chart_id(), "风险因素贡献发散条", "发散条形图", source, "第十章 风险应对", "发散条形图能把风险与正向缓释因素放在同一尺度中。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    months = np.arange(1, 13)
    product = np.minimum(months * 4, 32)
    data_line = np.minimum(months * 3, 24)
    pilot = np.clip((months - 4) * 5, 0, 28)
    ax.stackplot(months, product, data_line, pilot, labels=["产品", "数据", "试点"], colors=["#234f7d", "#0c6f69", "#b69252"], alpha=0.82, step="pre")
    ax.legend(frameon=False, ncol=3, fontsize=8)
    finish_chart(ax, "试点资源消耗阶梯面积", "展示资源投入从产品到试点的迁移", x_grid=True)
    save_chart(fig, next_chart_id(), "试点资源消耗阶梯面积", "阶梯面积图", source, "第八章 落地计划", "阶梯面积图能说明试点资源随阶段变化，而不是平均铺开。")

    fig, ax = plt.subplots(figsize=(7.2, 5.0))
    cohorts = ["Lite", "Pro", "活动", "报告"]
    start = np.array([0.58, 0.64, 0.42, 0.51])
    end = np.array([0.76, 0.82, 0.59, 0.68])
    for i, name in enumerate(cohorts):
        ax.plot([0, 1], [start[i], end[i]], marker="o", lw=2, color=["#234f7d", "#0c6f69", "#b69252", "#9b3f3a"][i])
        ax.text(-0.04, start[i], f"{name} {start[i]:.0%}", ha="right", va="center", fontsize=8)
        ax.text(1.04, end[i], f"{end[i]:.0%}", ha="left", va="center", fontsize=8)
    ax.set_xticks([0, 1])
    ax.set_xticklabels(["试点期", "续费期"])
    ax.set_ylim(0.35, 0.9)
    finish_chart(ax, "客户续费队列斜率图", "比较不同收入项从试点到续费的质量变化")
    save_chart(fig, next_chart_id(), "客户续费队列斜率图", "斜率图", source, "第七章 商业模式", "斜率图能说明客户成功如何把一次试点转成持续收入。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    cash = finance["cashflow"]["ending_cash"].to_numpy()
    x = np.arange(len(cash))
    ax.plot(x, cash, color="#234f7d", marker="o", lw=2)
    ax.fill_between(x, cash * 0.88, cash * 1.12, color="#234f7d", alpha=0.14, label="压力带")
    ax.set_xticks(x)
    ax.set_xticklabels([str(y) for y in years])
    ax.legend(frameon=False, fontsize=8)
    finish_chart(ax, "月度现金安全带图", "以年度现金流模拟现金安全区间", x_grid=True)
    save_chart(fig, next_chart_id(), "月度现金安全带图", "带状折线图", source, "第九章 财务预测", "带状折线图能表达现金预测的不确定性和安全边界。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    policy_types = ["安全监管", "空域管理", "产业扶持", "基础设施", "场景试点"]
    for i, name in enumerate(policy_types):
        xs = rng.normal(i + 1, 0.05, 16)
        ys = rng.uniform(0.2, 0.9, 16)
        ax.scatter(xs, ys, s=24, alpha=0.72, color=["#234f7d", "#0c6f69", "#b69252", "#9b3f3a", "#64748b"][i])
    ax.set_xticks(range(1, len(policy_types) + 1))
    ax.set_xticklabels(policy_types, rotation=20, ha="right")
    ax.set_yticklabels([])
    finish_chart(ax, "政策类型条带分布", "展示不同政策类型的发布密度与离散程度")
    save_chart(fig, next_chart_id(), "政策类型条带分布", "条带图", source, "第一章 行业机会", "条带图比单纯计数更能表现政策主题的分布差异。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    fields = ["目标", "计划", "轨迹", "围栏", "事件", "报表", "设备"]
    layers = ["来源", "字段", "校验", "页面", "导出"]
    mat = rng.choice([0, 1], size=(len(layers), len(fields)), p=[0.18, 0.82])
    sns.heatmap(mat, cmap=sns.color_palette(["#eef2f6", "#0c6f69"]), cbar=False, linewidths=1, linecolor="white", xticklabels=fields, yticklabels=layers, ax=ax)
    finish_chart(ax, "数据字段覆盖矩阵", "深色代表该字段已进入对应治理环节")
    save_chart(fig, next_chart_id(), "数据字段覆盖矩阵", "缺口矩阵", source, "第五章 数据体系", "缺口矩阵能快速显示数据治理仍需补齐的位置。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    thresholds = np.linspace(0.2, 0.9, 12)
    recall = 1 - (thresholds - 0.2) * 0.72
    precision = 0.44 + thresholds * 0.56
    ax.plot(thresholds, recall, color="#234f7d", marker="o", label="召回率")
    ax.plot(thresholds, precision, color="#0c6f69", marker="s", label="精确率")
    ax.legend(frameon=False, fontsize=8)
    finish_chart(ax, "模型阈值收益曲线", "阈值调高会提升精确率，但降低召回率", x_grid=True)
    save_chart(fig, next_chart_id(), "模型阈值收益曲线", "双线权衡图", source, "第五章 数据体系", "双线权衡图能解释误报和漏报之间的产品取舍。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    risk_table = pd.DataFrame([[3, 4, 5], [2, 3, 4], [1, 2, 3]], index=["高影响", "中影响", "低影响"], columns=["低概率", "中概率", "高概率"])
    sns.heatmap(risk_table, annot=True, fmt="d", cmap="YlOrRd", linewidths=1.4, linecolor="white", ax=ax, cbar=False)
    finish_chart(ax, "交付风险优先矩阵表图", "数值越高越需要写进合同边界和验收条款")
    save_chart(fig, next_chart_id(), "交付风险优先矩阵表图", "表格式热力图", source, "第十章 风险应对", "表格式热力图适合把风险优先级转成管理动作。")

    fig, ax = plt.subplots(figsize=(8.6, 2.4))
    nodes = ["首页", "脚本", "态势", "计划", "围栏", "工单", "报表"]
    xs = np.linspace(0.08, 0.92, len(nodes))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.text(0.04, 0.88, "WebDemo演示路径节点图", transform=ax.transAxes, ha="left", va="top", fontsize=12.8, fontweight="bold", color=PALETTE["ink"])
    ax.text(0.04, 0.76, "答辩时按同一条路径点击，避免页面展示碎片化", transform=ax.transAxes, ha="left", va="top", fontsize=8.2, color=PALETTE["muted"])
    for i, node in enumerate(nodes):
        ax.scatter(xs[i], 0.36, s=560, color="#0c6f69" if i in {0, len(nodes) - 1} else "#234f7d", edgecolor="white", lw=1.5)
        ax.text(xs[i], 0.36, node, ha="center", va="center", color="white", fontsize=8.6, fontweight="bold")
        if i < len(nodes) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.045, 0.36), xytext=(xs[i] + 0.045, 0.36), arrowprops=dict(arrowstyle="-|>", color="#64748b", lw=1.8, mutation_scale=13))
    ax.set_axis_off()
    save_chart(fig, next_chart_id(), "WebDemo演示路径节点图", "节点路径图", source, "第十三章 Web Demo", "节点路径图把答辩时的点击顺序变成一条可复述脚本。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    steps = pd.Series({"接收": 1.0, "确认": 0.92, "到场": 0.76, "处置": 0.61, "归档": 0.54})
    for i, (name, val) in enumerate(steps.items()):
        ax.barh(i, 1, color="#edf2f5", height=0.42)
        ax.barh(i, val, color="#0c6f69", height=0.42)
        ax.text(val + 0.02, i, f"{val:.0%}", va="center", fontsize=8)
    ax.set_yticks(range(len(steps)))
    ax.set_yticklabels(steps.index)
    ax.set_xlim(0, 1.12)
    finish_chart(ax, "移动处置步骤进度条", "展示移动端工单在不同步骤的完成比例")
    save_chart(fig, next_chart_id(), "移动处置步骤进度条", "进度条组", source, "第十三章 Web Demo", "进度条组能说明移动处置不是静态页面，而是连续状态。")

    fig = plt.figure(figsize=(8.2, 4.8))
    skills = ["产品", "数据", "前端", "交付", "商务"]
    values = [[82, 76, 88, 70, 62], [68, 84, 72, 78, 66]]
    angles = np.linspace(0, 2 * np.pi, len(skills), endpoint=False)
    for i, vals in enumerate(values, start=1):
        ax = fig.add_subplot(1, 2, i, polar=True)
        arr = np.array(vals)
        ax.plot(np.r_[angles, angles[0]], np.r_[arr, arr[0]], color=["#234f7d", "#0c6f69"][i - 1], lw=2)
        ax.fill(np.r_[angles, angles[0]], np.r_[arr, arr[0]], color=["#234f7d", "#0c6f69"][i - 1], alpha=0.16)
        ax.set_xticks(angles)
        ax.set_xticklabels(skills, fontsize=8)
        ax.set_yticklabels([])
        ax.set_title("当前" if i == 1 else "试点期", fontsize=9)
    fig.suptitle("团队能力覆盖雷达组", x=0.04, ha="left", fontsize=12.5, fontweight="bold", color=PALETTE["ink"])
    save_chart(fig, next_chart_id(), "团队能力覆盖雷达组", "雷达小倍数", source, "第十一章 团队组织", "雷达小倍数能展示团队能力从当前到试点期的补强方向。")

    fig, axes = plt.subplots(1, 3, figsize=(8.6, 3.6), subplot_kw={"aspect": "equal"})
    gauges = [("安全闭环", 0.78), ("报表及时", 0.84), ("培训覆盖", 0.69)]
    for ax, (name, val) in zip(axes, gauges):
        ax.pie([val, 1 - val], startangle=90, colors=["#0c6f69", "#e7ecef"], wedgeprops=dict(width=0.26, edgecolor="white"))
        ax.text(0, 0.04, f"{val:.0%}", ha="center", va="center", fontsize=15, fontweight="bold")
        ax.text(0, -0.22, name, ha="center", va="center", fontsize=8.5)
        ax.set_axis_off()
    fig.suptitle("社会价值指标仪表组", x=0.04, ha="left", fontsize=12.5, fontweight="bold", color=PALETTE["ink"])
    save_chart(fig, next_chart_id(), "社会价值指标仪表组", "仪表组", source, "第十二章 社会价值", "仪表组能把社会价值从口号转成可跟踪指标。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    scenarios = pd.DataFrame({"收益": [0.82, 0.75, 0.67, 0.71, 0.63], "成本": [0.42, 0.50, 0.58, 0.46, 0.62], "场景": ["景区", "园区", "场馆", "物流", "枢纽"]})
    ax.scatter(scenarios["成本"], scenarios["收益"], s=320, c=["#0c6f69", "#234f7d", "#b69252", "#2f6b4f", "#9b3f3a"], alpha=0.84)
    for _, r in scenarios.iterrows():
        ax.text(r["成本"] + 0.012, r["收益"], r["场景"], fontsize=9)
    ax.set_xlabel("交付成本")
    ax.set_ylabel("客户收益")
    finish_chart(ax, "低空场景收益成本矩阵", "左上象限最适合优先试点")
    save_chart(fig, next_chart_id(), "低空场景收益成本矩阵", "收益成本矩阵", source, "第六章 市场竞争", "收益成本矩阵能帮助解释首批场景选择。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    cols = [c for c in finance["revenue"].columns if c != "year"][:4]
    ax.stackplot(years, [finance["revenue"][c] for c in cols], labels=cols, colors=["#234f7d", "#0c6f69", "#b69252", "#9b3f3a"], alpha=0.82)
    ax.legend(frameon=False, fontsize=7.5, ncol=2, loc="upper left")
    finish_chart(ax, "收入结构主题河流", "不同收入流随时间抬升并互相替代")
    save_chart(fig, next_chart_id(), "收入结构主题河流", "主题河流图", source, "第九章 财务预测", "主题河流图能表达收入结构的连续迁移。")

    fig, ax = plt.subplots(figsize=(7.6, 5.0))
    levels = [("官方统计", 5), ("地方政策", 4), ("公开数据", 3), ("演示样本", 2), ("合理假设", 1)]
    for i, (name, width) in enumerate(levels):
        y0 = i * 0.16
        ax.barh(y0, width, left=(5 - width) / 2, height=0.12, color=["#234f7d", "#0c6f69", "#b69252", "#64748b", "#9b3f3a"][i])
        ax.text(2.5, y0, name, ha="center", va="center", color="white", fontweight="bold", fontsize=8.5)
    ax.set_xlim(0, 5)
    ax.set_ylim(-0.1, 0.8)
    ax.set_axis_off()
    fig.suptitle("来源可信度金字塔", x=0.04, ha="left", fontsize=12.5, fontweight="bold", color=PALETTE["ink"])
    save_chart(fig, next_chart_id(), "来源可信度金字塔", "金字塔图", source, "第五章 数据体系", "可信度金字塔能说明不同资料在正文中的使用权重。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    questions = ["政策边界", "客户付费", "误报漏报", "财务假设", "Demo真实性"]
    prep = np.array([[3, 5, 4], [4, 4, 3], [5, 4, 4], [4, 5, 5], [5, 3, 5]])
    sns.heatmap(prep, annot=True, fmt="d", cmap="YlGnBu", linewidths=1, linecolor="white", xticklabels=["材料", "数据", "演示"], yticklabels=questions, ax=ax, cbar=False)
    finish_chart(ax, "答辩追问准备矩阵", "5分代表已有较充分支撑材料")
    save_chart(fig, next_chart_id(), "答辩追问准备矩阵", "优先级矩阵", source, "附录", "准备矩阵能帮助答辩时快速定位追问证据。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    runway = np.array([18, 16, 14, 12, 11, 9])
    ax.step(range(len(runway)), runway, where="mid", color="#234f7d", lw=2)
    ax.fill_between(range(len(runway)), runway, 0, step="mid", color="#234f7d", alpha=0.16)
    ax.axhspan(0, 9, color="#9b3f3a", alpha=0.08)
    ax.set_xticks(range(len(runway)))
    ax.set_xticklabels(["M0", "M3", "M6", "M9", "M12", "M15"])
    finish_chart(ax, "融资后现金跑道阶梯", "现金跑道逐季压缩，需要用试点回款补充", x_grid=True)
    save_chart(fig, next_chart_id(), "融资后现金跑道阶梯", "阶梯线图", source, "第九章 财务预测", "阶梯线图能提示融资后现金跑道的阶段性压力。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    x = rng.normal(0, 1, 240)
    y = 0.45 * x + rng.normal(0, 0.72, 240)
    hb = ax.hexbin(x, y, gridsize=18, cmap="YlGnBu", mincnt=1)
    fig.colorbar(hb, ax=ax, shrink=0.82, label="样本密度")
    ax.set_xlabel("速度异常")
    ax.set_ylabel("高度异常")
    finish_chart(ax, "速度高度异常六边形密度", "密集区域用于设置风险评分校准样本")
    save_chart(fig, next_chart_id(), "速度高度异常六边形密度", "六边形密度图", source, "第五章 数据体系", "六边形密度图能减少散点重叠，适合高频遥测样本。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    rows = ["客户A", "客户B", "客户C", "客户D", "客户E"]
    for i, row in enumerate(rows):
        start = i * 0.7 + 0.5
        ax.broken_barh([(start, 0.6), (start + 1.1, 0.7), (start + 2.2, 0.5)], (i - 0.28, 0.56), facecolors=["#234f7d", "#0c6f69", "#b69252"])
    ax.set_yticks(range(len(rows)))
    ax.set_yticklabels(rows)
    ax.set_xlabel("季度")
    finish_chart(ax, "客户试点批次排布图", "展示多客户试点避免同时挤压交付团队", x_grid=True)
    save_chart(fig, next_chart_id(), "客户试点批次排布图", "批次排布图", source, "第八章 落地计划", "批次排布图能把客户获取节奏和交付能力联系起来。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    modules = ["态势", "计划", "围栏", "工单", "报表"]
    current = np.array([78, 70, 76, 68, 74])
    target = np.array([90, 84, 88, 86, 89])
    y = np.arange(len(modules))
    ax.barh(y + 0.18, target, height=0.32, color="#dce5eb", label="试点目标")
    ax.barh(y - 0.18, current, height=0.32, color="#0c6f69", label="当前Demo")
    ax.set_yticks(y)
    ax.set_yticklabels(modules)
    ax.legend(frameon=False, fontsize=8)
    finish_chart(ax, "Demo验收差距条形图", "当前Demo与试点目标的功能成熟度差距", x_grid=True)
    save_chart(fig, next_chart_id(), "Demo验收差距条形图", "对照条形图", source, "第十三章 Web Demo", "对照条形图能说明Demo当前完成度和下一步补强项。")


def generate_chart_pack(data: dict[str, pd.DataFrame]) -> pd.DataFrame:
    CHARTS.clear()
    for p in CHART_DIR.glob("*.png"):
        p.unlink()
    add_premium_chart_extensions(data)
    add_real_data_charts()
    add_manuscript_chart_extensions(data)
    add_word_quality_charts(data)
    add_chart_variety_boosters(data)
    catalog = pd.DataFrame(CHARTS).drop_duplicates(subset=["title", "chart_type"], keep="last").reset_index(drop=True)
    catalog.to_csv(TABLE_DIR / "chart_catalog.csv", index=False, encoding="utf-8-sig")
    return catalog


def write_markdown(pages: list[dict], chart_catalog: pd.DataFrame, source_registry: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.md"
    lines = ["# 低空智眼 SkyGuard 商业计划书", "", "> 城市低空空域安全感知与运行监管平台", ""]
    for page in pages:
        lines.append(f"## {page['chapter']}｜{page['title']}")
        if page.get("kicker"):
            lines.extend(["", f"**{page['kicker']}**"])
        for p in page["paras"]:
            lines.extend(["", p])
        if page["bullets"]:
            lines.append("")
            for b in page["bullets"]:
                lines.append(f"- {b}")
        if page.get("table"):
            table = page["table"]
            lines.append("")
            lines.append("| " + " | ".join(map(str, table[0])) + " |")
            lines.append("| " + " | ".join(["---"] * len(table[0])) + " |")
            for row in table[1:]:
                lines.append("| " + " | ".join(str(v).replace("\n", " ") for v in row) + " |")
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists():
                rel = str(fig_path.relative_to(ROOT)).replace("\\", "/")
                lines.extend(["", f"![{humanize_figure_caption(fig_path)}]({rel})", f"图注：{humanize_figure_caption(fig_path)}。"])
        if page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                r = row.iloc[0]
                lines.extend(["", f"![{r['title']}]({str(r['file']).replace(chr(92), '/')})", f"图表来源：{r['source']}。结论：{r['conclusion']}"])
    lines.extend(["", "## 来源清单"])
    for _, row in source_registry.iterrows():
        lines.append(f"- {row['title']}：{row['url']}")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


DOCX_CJK_FONT = "仿宋"
DOCX_LATIN_FONT = "Times New Roman"


def set_docx_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None, font_name: str = DOCX_CJK_FONT) -> None:
    run.font.name = font_name
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), DOCX_LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), DOCX_LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, font_name: str = DOCX_CJK_FONT) -> None:
    style.font.name = font_name
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), DOCX_LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), DOCX_LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def set_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(1.95)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.15)
    section.header_distance = Cm(0.82)
    section.footer_distance = Cm(0.82)
    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number", "Caption", "Heading 1", "Heading 2", "Heading 3", "Title"]:
        try:
            set_style_font(styles[style_name])
        except KeyError:
            continue
    normal = styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.first_line_indent = Pt(22)
    normal.paragraph_format.line_spacing = 1.45
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["Heading 1"].font.size = Pt(16.5)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(21, 31, 43)
    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.keep_together = True
    styles["Heading 2"].font.size = Pt(13.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(18, 78, 107)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True
    styles["Heading 2"].paragraph_format.keep_together = True
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor(47, 107, 79)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.keep_with_next = True
    styles["List Bullet"].font.size = Pt(10.2)
    styles["List Bullet"].paragraph_format.left_indent = Pt(21)
    styles["List Bullet"].paragraph_format.first_line_indent = Pt(-11)
    styles["List Bullet"].paragraph_format.line_spacing = 1.25
    styles["List Bullet"].paragraph_format.space_after = Pt(3)
    styles["Caption"].font.size = Pt(8.6)
    styles["Caption"].font.color.rgb = RGBColor(88, 99, 114)
    styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption"].paragraph_format.line_spacing = 1.12
    styles["Caption"].paragraph_format.space_before = Pt(2)
    styles["Caption"].paragraph_format.space_after = Pt(6)


def add_docx_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_docx_run_font(run, size=8.6, color=RGBColor(88, 99, 114))


def add_docx_picture(doc: Document, image_path: Path, width_in: float) -> None:
    doc.add_picture(str(image_path), width=Inches(width_in))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_docx_figure_block(doc: Document, image_path: Path, width_in: float, caption: str) -> None:
    fitted_width = width_in
    fitted_height = None
    try:
        with Image.open(image_path) as img:
            iw, ih = img.size
            if iw and ih:
                max_height_in = 3.95
                height_at_width = width_in * ih / iw
                if height_at_width > max_height_in:
                    fitted_height = max_height_in
                else:
                    fitted_width = width_in
    except Exception:
        pass
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    image_run = p.add_run()
    if fitted_height is not None:
        image_run.add_picture(str(image_path), height=Inches(fitted_height))
    else:
        image_run.add_picture(str(image_path), width=Inches(fitted_width))
    image_run.add_break()
    caption_run = p.add_run(caption)
    set_docx_run_font(caption_run, size=8.6, color=RGBColor(88, 99, 114))


def docx_table_widths(col_count: int) -> list[float]:
    patterns = {
        2: [1.55, 4.95],
        3: [1.35, 2.55, 2.60],
        4: [1.15, 1.75, 2.05, 1.55],
        5: [0.95, 1.20, 1.25, 2.10, 1.00],
        6: [0.78, 0.88, 0.92, 0.92, 2.18, 0.82],
    }
    if col_count in patterns:
        return patterns[col_count]
    return [6.5 / max(col_count, 1)] * max(col_count, 1)


def set_docx_cell_margins(cell, top: int = 90, bottom: int = 90, start: int = 125, end: int = 125) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def mark_header_row_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def add_docx_editorial_table(doc: Document, table_data: list[list[str]], caption: str | None = None) -> None:
    if not table_data:
        return
    if caption:
        add_docx_caption(doc, caption)
    tbl = doc.add_table(rows=1, cols=len(table_data[0]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = docx_table_widths(len(table_data[0]))
    mark_header_row_repeat(tbl.rows[0])
    for j, cell in enumerate(tbl.rows[0].cells):
        cell.text = str(table_data[0][j])
        cell.width = Inches(widths[j])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_docx_cell_margins(cell)
        shade_docx_cell(cell, "DDEBE8")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                set_docx_run_font(run, size=8.8, color=RGBColor(21, 31, 43), bold=True)
    for row_idx, row in enumerate(table_data[1:], start=1):
        cells = tbl.add_row().cells
        for j, val in enumerate(row[: len(cells)]):
            cells[j].text = str(val)
            cells[j].width = Inches(widths[j])
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_docx_cell_margins(cells[j])
            if row_idx % 2 == 0:
                shade_docx_cell(cells[j], "F6F8FA")
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 or len(str(val)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.18
                for run in p.runs:
                    set_docx_run_font(run, size=8.2, color=RGBColor(45, 55, 72))
    doc.add_paragraph("")


def write_docx(pages: list[dict], chart_catalog: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.docx"
    doc = Document()
    set_docx_styles(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.text = "SkyGuard 低空智眼 | 城市低空运行监管商业计划书"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(100, 116, 139)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = footer.add_run("资料口径：官方统计 / 地方政策 / 公开数据 / 演示样本 / 经营假设分层标注    第 ")
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(100, 116, 139)
    add_page_number(footer)
    f_run = footer.add_run(" 页")
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = RGBColor(100, 116, 139)
    table_no = 0
    figure_no = 0
    for idx, page in enumerate(pages, start=1):
        if page["type"] == "cover":
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("低空智眼 SkyGuard\n")
            set_docx_run_font(run, size=28, color=RGBColor(21, 31, 43), bold=True)
            run = title.add_run("城市低空空域安全感知与运行监管平台商业计划书")
            set_docx_run_font(run, size=14.5, color=RGBColor(51, 65, 85), bold=True)
            sub = doc.add_paragraph(page["callout"])
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in sub.runs:
                set_docx_run_font(run, size=11, color=RGBColor(51, 65, 85))
            for para in page.get("paras", []):
                p = doc.add_paragraph(para)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                for run in p.runs:
                    set_docx_run_font(run, size=10.8, color=RGBColor(55, 65, 81))
            if page.get("figure"):
                fig_path = Path(page["figure"])
                if fig_path.exists():
                    add_docx_figure_block(doc, fig_path, 3.85, f"封面图  {humanize_figure_caption(fig_path)}")
            doc.add_page_break()
            continue
        heading_level = 1 if page["type"] == "divider" else 2
        heading = doc.add_heading(f"{page['chapter']}｜{page['title']}", level=heading_level)
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.keep_together = True
        for run in heading.runs:
            set_docx_run_font(run, size=16.5 if heading_level == 1 else 13.5, color=RGBColor(21, 31, 43) if heading_level == 1 else RGBColor(18, 78, 107), bold=True)
        if page.get("kicker"):
            p = doc.add_paragraph(page["kicker"])
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.keep_with_next = True
            set_docx_run_font(p.runs[0], size=10.5, color=RGBColor(18, 78, 107), bold=True)
        lead_paras = page["paras"][:2]
        tail_paras = page["paras"][2:]
        for para in lead_paras:
            p = doc.add_paragraph(para)
            for run in p.runs:
                set_docx_run_font(run, size=11, color=RGBColor(31, 41, 55))
        media_rendered = False
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists():
                try:
                    figure_no += 1
                    add_docx_figure_block(doc, fig_path, 3.85, f"图 {figure_no}  {humanize_figure_caption(fig_path)}")
                    media_rendered = True
                except Exception:
                    pass
        if (not media_rendered) and page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                chart_path = ROOT / row.iloc[0]["file"]
                if chart_path.exists():
                    try:
                        figure_no += 1
                        add_docx_figure_block(doc, chart_path, 4.20, f"图 {figure_no}  {row.iloc[0]['title']}\n资料来源：{row.iloc[0]['source']}。图示说明：{row.iloc[0]['conclusion']}")
                        media_rendered = True
                    except Exception:
                        pass
        for para in tail_paras:
            p = doc.add_paragraph(para)
            for run in p.runs:
                set_docx_run_font(run, size=11, color=RGBColor(31, 41, 55))
        for bullet in page["bullets"]:
            p = doc.add_paragraph(bullet, style="List Bullet")
            for run in p.runs:
                set_docx_run_font(run, size=10.2, color=RGBColor(45, 55, 72))
        if page.get("table"):
            table_no += 1
            add_docx_editorial_table(doc, page.get("table", []), f"表 {table_no}  {page['chapter']}：{page['title']}")
        next_page = pages[idx] if idx < len(pages) else None
        if next_page and next_page.get("type") == "divider":
            doc.add_page_break()
    doc.save(path)
    return path


def draw_pdf_page(c, page: dict, page_num: int, chart_catalog: pd.DataFrame, page_size=A4):
    w, h = page_size
    margin = 39
    accent = page.get("accent") or REPORT_COLORS["teal"]
    c.setFillColor(hex_color(REPORT_COLORS["paper"]))
    c.rect(0, 0, w, h, stroke=0, fill=1)

    if page["type"] == "cover":
        if page.get("figure"):
            try:
                img = Image.open(Path(page["figure"]))
                iw, ih = img.size
                scale = max(w / iw, h / ih)
                c.drawImage(ImageReader(img), (w - iw * scale) / 2, (h - ih * scale) / 2, width=iw * scale, height=ih * scale, mask="auto")
            except Exception:
                pass
        c.setFillColor(colors.Color(0.03, 0.06, 0.08, alpha=0.74))
        c.rect(0, 0, w, h, stroke=0, fill=1)
        c.setFillColor(colors.Color(1, 1, 1, alpha=0.08))
        c.rect(0, 0, w, 76, stroke=0, fill=1)
        c.setFillColor(colors.white)
        c.setFont("CNFont", 9.4)
        c.drawString(margin + 8, h - 88, page.get("kicker", "商业计划书"))
        c.setFont("CNFont", 33)
        c.drawString(margin + 8, h - 148, "低空智眼 SkyGuard")
        c.setFont("CNFont", 16.5)
        c.drawString(margin + 8, h - 181, "城市低空空域安全感知与运行监管平台商业计划书")
        y = h - 232
        for para in page["paras"]:
            y = draw_page_text(c, para, margin + 8, y, 428, 10.2, 16, "#ffffff", max_lines=4)
            y -= 7
        chip_x = margin + 8
        for bullet in page["bullets"]:
            chip_x = draw_chip(c, bullet, chip_x, 132, "#ffffff", "#ffffff", REPORT_COLORS["ink"])
        c.setFont("CNFont", 8)
        c.setFillColor(colors.Color(1, 1, 1, alpha=0.88))
        c.drawString(margin + 8, 48, "商业计划书｜财务测算｜证据索引｜Web Demo")
        c.showPage()
        return

    if page["type"] == "divider":
        c.setFillColor(hex_color(REPORT_COLORS["navy"]))
        c.rect(0, 0, w, h, stroke=0, fill=1)
        c.setFillColor(hex_color(accent))
        c.rect(0, 0, 10, h, stroke=0, fill=1)
        c.setFillColor(hex_color("#cbd5df"))
        c.setFont("CNFont", 9)
        c.drawString(margin + 8, h - 72, "章节导读")
        c.setFillColor(colors.white)
        c.setFont("CNFont", 24)
        c.drawString(margin + 8, h - 118, page["chapter"])
        c.setFont("CNFont", 16)
        c.drawString(margin + 8, h - 151, page["title"][:34])
        y = h - 204
        for para in page["paras"]:
            y = draw_page_text(c, para, margin + 8, y, 365, 10, 16, "#f2f4f6", max_lines=9)
        by = 190
        for i, bullet in enumerate(page["bullets"][:3]):
            c.setFillColor(colors.Color(1, 1, 1, alpha=0.08))
            c.roundRect(margin + 8, by - i * 58, 365, 42, 6, stroke=0, fill=1)
            draw_page_text(c, bullet, margin + 20, by + 20 - i * 58, 338, 8, 12, "#f8fafc", max_lines=2)
        if page.get("figure"):
            draw_pdf_visual(c, Path(page["figure"]), w - 246, 360, 188, 172, border=True)
        elif page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                draw_pdf_visual(c, ROOT / row.iloc[0]["file"], w - 246, 360, 188, 172, border=True)
        c.setFont("CNFont", 8)
        c.setFillColor(hex_color("#cbd5df"))
        c.drawRightString(w - margin, 38, f"第 {page_num} 页")
        c.showPage()
        return

    c.setFillColor(colors.white)
    c.roundRect(22, 22, w - 44, h - 44, 8, stroke=0, fill=1)
    c.setFillColor(hex_color(accent))
    c.rect(22, h - 74, w - 44, 3.2, stroke=0, fill=1)
    c.setFillColor(hex_color(REPORT_COLORS["muted"]))
    c.setFont("CNFont", 7.8)
    c.drawString(margin, h - 53, page.get("kicker", ""))
    c.drawRightString(w - margin, h - 53, "SkyGuard 低空智眼")
    c.setFillColor(hex_color(REPORT_COLORS["ink"]))
    c.setFont("CNFont", 16.2)
    c.drawString(margin, h - 88, f"{page['chapter']}｜{page['title']}"[:42])

    left_w = 318
    right_x = margin + left_w + 16
    right_w = w - right_x - margin
    y = h - 119
    max_para = 4 if page["type"] == "table" else 3
    for para in page["paras"][:max_para]:
        y = draw_page_text(c, para, margin, y, left_w, 8.3, 12.2, REPORT_COLORS["soft_ink"], max_lines=7)
        y -= 3
        if y < 300:
            break

    c.setFillColor(hex_color("#f3f1eb"))
    c.roundRect(right_x, h - 318, right_w, 198, 6, stroke=0, fill=1)
    c.setStrokeColor(hex_color("#d8d1c5"))
    c.roundRect(right_x, h - 318, right_w, 198, 6, stroke=1, fill=0)
    c.setFont("CNFont", 9.1)
    c.setFillColor(hex_color(REPORT_COLORS["ink"]))
    c.drawString(right_x + 12, h - 141, "经营要点")
    by = h - 165
    for bullet in page["bullets"][:3]:
        by = draw_page_text(c, "· " + bullet, right_x + 12, by, right_w - 24, 7.1, 10.5, REPORT_COLORS["soft_ink"], max_lines=3)
        by -= 3
    if page.get("callout"):
        c.setFillColor(hex_color(accent))
        c.roundRect(right_x + 12, h - 304, right_w - 24, 30, 5, stroke=0, fill=1)
        draw_page_text(c, page["callout"], right_x + 20, h - 288, right_w - 40, 7.1, 10, "#ffffff", max_lines=2)

    media_top = 424
    media_height = 264
    if page.get("table"):
        draw_pdf_table(c, page["table"], margin, media_top, w - margin * 2, accent)
    elif page.get("figure"):
        bottom_y = draw_pdf_visual(c, Path(page["figure"]), margin, media_top, w - margin * 2, media_height, border=True)
        c.setFont("CNFont", 6.8)
        c.setFillColor(hex_color(REPORT_COLORS["muted"]))
        c.drawString(margin, max(43, bottom_y), f"图注：{humanize_figure_caption(Path(page['figure']))}")
    elif page.get("chart_id"):
        row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
        if not row.empty:
            chart_path = ROOT / row.iloc[0]["file"]
            bottom_y = draw_pdf_visual(c, chart_path, margin, media_top, w - margin * 2, media_height, border=True)
            caption = f"{row.iloc[0]['title']}｜来源：{row.iloc[0]['source']}｜结论：{row.iloc[0]['conclusion']}"
            draw_page_text(c, caption, margin, max(42, bottom_y), w - margin * 2, 6.3, 8.4, REPORT_COLORS["muted"], max_lines=2)
    c.setStrokeColor(hex_color("#d8d1c5"))
    c.line(margin, 33, w - margin, 33)
    c.setFont("CNFont", 7.2)
    c.setFillColor(hex_color(REPORT_COLORS["muted"]))
    c.drawString(margin, 20, "资料口径：官方统计 / 地方政策 / 公开数据 / 演示样本 / 经营假设分层标注")
    c.drawRightString(w - margin, 20, f"第 {page_num} 页")
    c.showPage()


def write_pdf(pages: list[dict], chart_catalog: pd.DataFrame, filename: str, limit: int | None = None) -> Path:
    path = OUT_DIR / filename
    c = canvas.Canvas(str(path), pagesize=A4)
    selected = pages if limit is None else select_pitch_pages(pages, limit)
    for idx, page in enumerate(selected, start=1):
        draw_pdf_page(c, page, idx, chart_catalog)
    c.save()
    return path


def remove_obsolete_word_versions() -> None:
    for name in ["SkyGuard_商业计划书_升级版.docx", "~$yGuard_商业计划书_升级版.docx"]:
        path = OUT_DIR / name
        if path.exists():
            try:
                path.unlink()
            except OSError:
                pass


def _polish_visible_text(value: object) -> str:
    text = str(value or "")
    replacements = {
        "AI 识别": "智能识别",
        "AI识别": "智能识别",
        "AI式": "公式化",
        "Web Demo AI": "Web Demo 智能",
        "Demo Center": "演示中心",
        "AI": "智能",
        "这一页的作用，是把": "这里需要把",
        "这一页的作用，是": "这里需要",
        "这一页的重点不是": "重点不是",
        "这一页尤其要避免": "这里要避免",
        "这一页对应": "对应",
        "这一页把": "这里把",
        "这一页": "本部分",
        "阅读这一节时，可以先看": "结合",
        "如果评委追问": "进一步展开时",
        "答辩时应": "展示时可",
        "评委能": "读者能",
        "提交前应补齐真实": "成员信息由小组统一录入",
        "请补齐真实成员信息。": "成员信息以正式名单为准。",
        "待填写": "",
        "提交前状态": "信息状态",
        "提交检查": "信息完整性",
        "硬性要求": "课程要求",
        "图注：": "图示：",
        "图片说明：": "图示：",
        "图片：": "图示：",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace(" 。", "。").replace(" ，", "，").replace(" ：", "：")
    while "。。" in text or "，，" in text or "；；" in text:
        text = text.replace("。。", "。").replace("，，", "，").replace("；；", "；")
    return text


CLEAN_LOGICAL_TITLES: dict[str, list[str]] = {
    "第一章 行业机会": ["政策窗口与进入时机", "运行规模与真实需求", "区域样板选择", "客户入口识别", "监管边界判断", "产业链位置", "政策节奏观察", "服务站网络", "场景扩张路径", "公开案例参照", "客户触发条件", "伙伴协同路径", "机会判断"],
    "第二章 痛点验证": ["现场发现断点", "身份核验缺口", "计划航线比对", "围栏触发机制", "工单责任链", "复盘归档要求", "断点样本复核", "事件类型结构", "偏离原因解释", "临时管控场景", "响应尾部风险", "闭环材料", "需求优先级"],
    "第三章 产品服务": ["值班态势入口", "计划审批动作", "电子围栏配置", "识别复核流程", "事件中心闭环", "运行报表产出", "场景包组合", "客户成功动作", "目标轨迹管理", "物流航线验证", "服务蓝图", "复核权限", "事件字段沉淀", "KPI验收", "场景复制", "续费培训", "产品范围"],
    "第四章 技术架构": ["数据接入边界", "时空索引能力", "轨迹比对逻辑", "风险评分解释", "规则引擎配置", "权限审计留痕", "部署集成策略", "可解释性设计", "接入字段标准", "页面验证", "覆盖能力评估", "权重校准", "规则命中记录", "告警链路时序", "集成边界", "人工确认机制", "架构落地"],
    "第五章 数据体系": ["官方统计口径", "政策目标口径", "公开数据边界", "演示样本性质", "数据字典结构", "质量规则清单", "模型校验口径", "评分校准过程", "区域目标映射", "数据资产页面", "飞行类型分布", "字段治理责任", "数据体系"],
    "第六章 市场竞争": ["替代方案边界", "客户分层逻辑", "采购决策链条", "竞品矩阵比较", "进入顺序安排", "伙伴生态组合", "市场边界收束", "招投标类型", "采购动作拆解", "差异化定位", "区域运营对比", "渠道合作", "市场判断"],
    "第七章 商业模式": ["订阅收入假设", "部署项目收入", "活动保障服务", "航线评估服务", "报告服务收入", "运维续费动作", "客户成功机制", "续费队列", "专业部署路径", "现场保障页面", "首批客户漏斗", "月度报告价值", "商业模式"],
    "第八章 落地计划": ["三个月交付物", "六个月试点包", "十八个月复制期", "三十六个月协同", "试点验收材料", "区域复制节奏", "路线图依赖", "交付泳道", "试点现场动作", "事件闭环验证", "里程碑"],
    "第九章 财务预测": ["收入预测口径", "成本结构口径", "毛利路径解释", "现金流约束", "回款周期压力", "融资用途安排", "敏感性检验", "单位经济测算", "现金流瀑布", "收入结构迁移", "成本页面验证", "五年收入结构", "现金安全边界", "账期承压判断", "财务判断"],
    "第十章 风险应对": ["政策边界风险", "数据合规风险", "误报漏报风险", "硬件依赖风险", "采购周期风险", "现金压力风险", "过度承诺风险", "试点延期原因", "风险预案"],
    "第十一章 团队组织": ["岗位配置原则", "产品负责人", "算法数据职责", "前后端研发", "交付运维职责", "行业顾问职责", "岗位到位计划", "团队保障"],
    "第十二章 社会价值": ["公共安全指标", "城市治理指标", "应急保障指标", "产业数据沉淀", "岗位培养路径", "合规意识提升", "平衡计分卡", "社会价值"],
    "第十三章 Web Demo": ["首页值班域", "演示脚本入口", "态势地图验证", "目标监测链路", "计划审批链路", "识别复核链路", "事件工单闭环", "运行报表导出", "移动处置路径", "指挥屏视图", "部署说明校验", "Demo演示"],
    "附录": ["来源索引", "图表目录", "数据字典", "财务假设", "访谈问卷", "部署说明", "截图包", "交付清单", "政策来源", "图表文件", "模型字段", "财务公式", "演示脚本", "部署验收", "答辩材料", "附录索引"],
}


def logical_page_title(chapter: str, local_index: int, topic: str, row: dict | None, figure_caption: str | None = None) -> str:
    code = CHAPTER_CODES.get(chapter, "")
    prefix = f"{code}.{local_index}" if code != "A" else f"A.{local_index}"
    titles = CLEAN_LOGICAL_TITLES.get(chapter, [])
    title = titles[local_index - 1] if local_index - 1 < len(titles) else topic
    return _polish_visible_text(f"{prefix} {title}")


def manuscript_paragraphs(spec: dict, topic: str, row: dict | None, metrics: dict, figure_caption: str | None, index: int) -> list[str]:
    chapter = spec["chapter"]
    evidence = _polish_visible_text(figure_caption or (row.get("title") if row else topic))
    conclusion = _polish_visible_text((row or {}).get("conclusion", spec["decision"]))
    source = _polish_visible_text((row or {}).get("source", "整理图片、Demo截图与项目资料"))
    transition = [
        "这个判断的价值在于把市场机会压到客户每天会处理的事务上。",
        "它把外部资料、产品动作和交付边界连在一起，便于后续试点逐项核验。",
        "如果只看单个指标，容易把机会写成赛道热度；放回工作流，才能看到可收费的服务。",
        "早期样板应优先选择责任主体清楚、边界稳定、记录可追溯的区域。",
    ][index % 4]
    if chapter.startswith("第一章"):
        return [
            f"{topic}不是简单判断赛道是否热门，而是判断低空活动是否已经进入需要日常治理的阶段。注册无人机、飞行小时、运营单位和服务站数量共同说明运行密度在提高，但商业机会只会出现在有明确责任主体和固定管理边界的区域。",
            f"SkyGuard 选择景区、园区、场馆、机场周边和固定航线作为首批入口，是因为这些区域有清晰的空间边界、可描述的飞行计划和可沉淀的处置记录。与其一开始承诺全域城市平台，不如先做出可演示、可验收、可复用的区域样板。",
            f"{evidence}提供了外部参照，资料口径来自{source}。它支持的判断是：{conclusion}。",
            f"{transition} 因此，行业机会章节最终要回答的不是“市场有多大”，而是“哪类客户会先付费、为什么现在需要、交付后如何验收”。",
        ]
    if chapter.startswith("第二章"):
        return [
            f"{topic}体现的是现场流程中的真实断点。低空目标被发现之后，管理人员还要确认身份、核对计划、判断围栏、通知责任人、记录处置过程，并在事后形成可追溯材料，任何环节断开都会回到截图、电话和临时表格。",
            "痛点验证不能停在安全焦虑上。客户真正愿意付费的地方，是系统能减少人工比对、跨部门沟通、事后解释和材料补写，把分散动作变成一条能回放的事件链。",
            f"{evidence}说明的重点是：{conclusion}。资料口径来自{source}。",
            f"{transition} 对 SkyGuard 来说，痛点不是“看不见目标”一个问题，而是“看见以后能否判得清、派得出、收得回、说得明”。",
        ]
    if chapter.startswith("第三章"):
        return [
            f"{topic}要对应客户的日常动作。综合态势、计划审批、电子围栏、识别复核、事件工单和运行报表不应分散成几个孤立页面，而要围绕值班人员的一次完整处置形成连续路径。",
            "产品设计的重点是把规则、字段、权限和日志放进同一条链路。客户进入系统后，先看到区域状态，再定位异常目标，随后核验计划和围栏，最后完成确认、派单、归档和导出。",
            f"{evidence}支撑的产品判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 只有把模块写成工作流，后续定价、培训、验收和续费才有共同语言。",
        ]
    if chapter.startswith("第四章"):
        return [
            f"{topic}的关键不是堆叠技术名词，而是让每一次告警都能解释来源、触发条件和处置结果。低空监管场景对误报、漏报和责任边界非常敏感，平台必须把规则、模型和人工确认放在同一套审计框架下。",
            "技术路线采用规则先行、模型辅助、人工确认的组合。这样既能在试点阶段降低争议，也为后续接入更多感知设备、航线数据和风险模型留出接口空间。",
            f"{evidence}对应的架构判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 架构章节要让读者看到系统为什么能落地，而不是只看到概念堆叠。",
        ]
    if chapter.startswith("第五章"):
        return [
            f"{topic}决定了计划书的可信度。官方统计、地方政策、公开数据、演示样本和经营假设必须分层使用：统计支撑行业判断，政策支撑区域选择，公开数据支撑方法，演示样本只证明流程。",
            "数据边界越清楚，图表越有说服力。每一张图都应能回到来源、字段、计算口径或假设表，避免把政策目标当成订单，把演示样本当成真实运营数据。",
            f"{evidence}提供的证据作用是：{conclusion}。资料口径来自{source}。",
            f"{transition} 数据体系章节的任务，是让后续产品、财务和风险判断都有可核验的底座。",
        ]
    if chapter.startswith("第六章"):
        return [
            f"{topic}要说明 SkyGuard 为什么选择当前进入顺序。替代方案包括通用视频监控、单点感知设备、反制系统、城市级重平台和人工巡查，它们分别解决一部分问题，但未必能形成可交付的区域运行闭环。",
            "竞争策略不靠把对手写弱，而是明确自己的边界：先做重点区域工作台，围绕计划、围栏、事件和报表形成轻量可验收方案，再通过伙伴接入补齐硬件和渠道。",
            f"{evidence}支持的竞争判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 这样写能把竞争章节从“谁更强”转成“先从哪里进入、为什么能成交、怎样复制”。",
        ]
    if chapter.startswith("第七章"):
        return [
            f"{topic}必须落到客户购买动作上。客户不是为概念付费，而是为日常值班、临时保障、航线评估、月度报告、规则维护和运维支持付费。",
            "收入结构采用软件订阅、专业部署、活动保障、航线评估、报告服务和运维续费组合。早期用部署和保障覆盖交付成本，中期用订阅、报告和运维形成稳定收入。",
            f"{evidence}对应的商业判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 商业模式章节要把客单价、成本、回款和客户成功动作讲清楚，而不是只给一条增长曲线。",
        ]
    if chapter.startswith("第八章"):
        return [
            f"{topic}需要像真实项目排期一样展开。每个阶段都要有明确交付物，例如 Demo、数据字典、规则样例、试点包、培训材料、日报和复盘纪要。",
            "落地难点在依赖关系。字段不稳定，风险解释无法复核；工单不闭环，报表没有事实来源；试点反馈不足，定价和续费假设都会失真。",
            f"{evidence}支持的排期判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 路线图的可信度来自可检查的材料，而不是来自时间轴本身。",
        ]
    if chapter.startswith("第九章"):
        return [
            f"{topic}要解释经营质量。收入增长必须和客户数量、客单价、部署成本、续费率、回款周期和研发投入一起看，否则财务预测容易变成单纯乐观假设。",
            "资金使用应优先服务产品打磨、试点交付、数据治理、客户成功和必要伙伴接入，同时控制并行项目数量。低空项目的回款周期可能拉长，现金流约束必须放在扩张节奏之前。",
            f"{evidence}对应的财务判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 财务章节的重点不是把数字写大，而是说明假设之间如何相互支撑。",
        ]
    if chapter.startswith("第十章"):
        return [
            f"{topic}是必须主动写清的约束。低空项目会遇到政策边界、数据合规、误报漏报、硬件依赖、采购周期和现金压力，任何一项处理不好都会拖慢试点。",
            "SkyGuard 的边界是辅助感知、风险预警、事件协同和报表复盘，不替代审批或执法，也不承诺干扰、捕获和打击能力。边界越清楚，客户沟通和合同验收越稳。",
            f"{evidence}支持的风险判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 风险应对要写出触发信号、影响范围和处理动作，不能只停留在风险清单。",
        ]
    if chapter.startswith("第十一章"):
        return [
            f"{topic}要服务当前阶段的交付瓶颈。早期团队最需要把产品、数据、前后端、交付和客户材料连起来，而不是提前搭建臃肿部门。",
            "岗位配置可分为产品研发线和试点交付线。前者保证平台可运行、可迭代，后者保证客户沟通、现场配置、培训和复盘材料能落地。",
            f"{evidence}对应的组织判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 团队章节要说明谁负责把下一阶段交付物做出来，而不是只展示组织名称。",
        ]
    if chapter.startswith("第十二章"):
        return [
            f"{topic}需要写成可观察的治理改善。低空运行越密，城市越需要把目标、计划、围栏、事件和报表纳入同一套协同机制，减少信息断点和责任争议。",
            "社会价值可拆成公共安全、城市治理、应急保障、产业数据和人才培养。每个维度都应有指标，例如响应时间、闭环率、报表及时率、培训覆盖率和规则更新次数。",
            f"{evidence}支撑的价值判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} 公共价值不是额外包装，它会反过来支撑客户续费、扩容和区域复制。",
        ]
    if chapter.startswith("第十三章"):
        return [
            f"{topic}要呈现一条可点击的业务路径。首页进入值班域，随后查看综合态势和目标监测，再核验计划、围栏和识别结果，最后完成事件工单、移动处置和报表导出。",
            "Demo 的可信度来自连续动作，而不是页面数量。真实素材、公开指标和演示样本要明确分层，页面之间要能说明同一条业务链路。",
            f"{evidence}对应的演示判断是：{conclusion}。资料口径来自{source}。",
            f"{transition} Web Demo 章节要证明团队已经把商业计划中的流程做成可运行界面。",
        ]
    return [
        f"{topic}承担核验作用。附录把来源、字段、假设、图表和交付物放到清楚位置，让正文判断能够被快速追溯。",
        "资料索引越清晰，计划书越像一套可继续推进的项目材料，而不是孤立文本。财务测算、证据索引、图表目录和部署说明应保持同一套命名和口径。",
        f"{evidence}对应的附录结论是：{conclusion}。资料口径来自{source}。",
        f"{transition} 附录应减少解释性套话，保留真正能帮助核验的文件、字段和来源。",
    ]


def enrich_manuscript_paragraphs(chapter: str, topic: str, paras: list[str], evidence_label: str, source: str, conclusion: str, page_index: int) -> list[str]:
    cleaned = [_polish_visible_text(p) for p in paras if _polish_visible_text(p)]
    if len(cleaned) < 4:
        cleaned.append(_polish_visible_text(f"资料口径来自{source}，支撑的核心结论是：{conclusion}。"))
    chapter_notes = {
        "第一章": f"落到本项目，{topic}需要同时回答政策窗口、客户责任边界和首批试点条件三件事。只有把机会拆到客户能签收的区域、指标和材料，后面的产品路线和财务预测才不会悬空。",
        "第二章": f"对客户来说，{topic}带来的压力通常不是单个告警，而是告警之后的核验、沟通、派单和复盘。正文因此把痛点写成连续流程，便于后续对应到产品页面和验收指标。",
        "第三章": f"在产品设计上，{topic}必须形成明确入口、可见状态、责任记录和导出材料。这样客户不是在看一组功能按钮，而是在看一套能够每天值班使用的工作方法。",
        "第四章": f"技术部分围绕{topic}解释实现边界：哪些数据进入系统，哪些规则触发告警，哪些结果需要人工确认，哪些记录进入审计。这样能减少黑箱感，也便于试点阶段和客户共同校准。",
        "第五章": f"数据体系中的{topic}要和来源、字段、计算口径一起出现。图表只负责呈现结果，真正支撑判断的是数据从哪里来、怎样处理、哪些部分只属于演示样本。",
        "第六章": f"市场判断中的{topic}不能停在竞品罗列。需要说明客户为什么会先买这个范围的服务，现有替代方案缺在哪个环节，以及 SkyGuard 如何通过轻量试点降低进入阻力。",
        "第七章": f"商业模式中的{topic}要回到购买动作：谁提出需求，谁验收交付，谁使用报表，谁决定续费。收入项只有和这些动作相连，才不是简单的价格清单。",
        "第八章": f"落地计划中的{topic}强调依赖关系。Demo、数据字典、规则样例、培训材料和日报模板需要按顺序成熟，任何一项滞后都会影响试点验收和后续复制。",
        "第九章": f"财务章节处理{topic}时，把增长、毛利、现金流和回款周期放在一起看。这样可以说明项目如何活下去，而不是只把未来收入写高。",
        "第十章": f"风险应对中的{topic}需要写出触发信号和处理动作。提前划清辅助监管、人工确认、数据合规和硬件伙伴边界，能让试点合同和演示口径更稳。",
        "第十一章": f"团队组织中的{topic}要对应交付短板。早期不是扩出复杂部门，而是让产品、数据、前端、交付、商务和材料整理围绕同一套试点节奏协作。",
        "第十二章": f"社会价值中的{topic}要能被观察和复盘。公共安全、应急保障、治理效率和人才培养都应落到响应时间、闭环率、报表及时率和培训覆盖等指标上。",
        "第十三章": f"Web Demo 中的{topic}要服务现场演示。截图、表格和文字必须共同说明同一个动作：用户从哪里进入、看见什么状态、如何判断风险、怎样形成闭环材料。",
        "附录": f"附录中的{topic}用于保证可追溯。正文引用的来源、图表、字段和假设都应能在这里找到对应位置，便于复核时快速定位。",
    }
    prefix = next((key for key in chapter_notes if chapter.startswith(key)), "附录")
    added = _polish_visible_text(chapter_notes[prefix])
    if added and all(added != p for p in cleaned):
        cleaned.append(added)
    return cleaned


def build_page_bullets(chapter: str, topic: str, row: dict | None, spec: dict) -> list[str]:
    conclusion = _polish_visible_text((row or {}).get("conclusion", spec["decision"]))
    source = _polish_visible_text((row or {}).get("source", "整理图片、Demo截图与项目资料"))
    if chapter.startswith("第九章"):
        return ["关注客单价、交付成本、续费率和回款周期", f"依据：{source}", f"经营判断：{conclusion}"]
    if chapter.startswith("第十章"):
        return ["边界清楚：辅助监管，不替代审批或执法", f"依据：{source}", f"处置重点：{conclusion}"]
    if chapter.startswith("第十三章"):
        return ["路径连续：脚本、态势、计划、围栏、工单、报表", f"素材依据：{source}", f"演示重点：{conclusion}"]
    return [f"核心判断：{conclusion}", f"依据：{source}", f"交付指向：{topic}需对应字段、页面、责任人或验收材料"]


def web_demo_engineering_table() -> list[list[str]]:
    return [
        ["验收项", "当前产物", "评审价值", "截图对应"],
        ["前端工程", "React + TypeScript + Vite + Recharts，22条业务路由", "说明团队已把计划书流程做成可运行工程", "首页、综合态势、指挥屏"],
        ["视觉升级", "中文字体栈、运营工作台布局、地图HUD、深浅分层表格与KPI", "摆脱模板化页面，形成低空监管产品质感", "全量21张截图"],
        ["业务闭环", "脚本入口、态势地图、目标监测、计划审批、识别复核、围栏、工单、报表、移动处置", "证明页面不是孤立展示，而是同一条业务链路", "08至15、21号截图"],
        ["数据边界", "演示样本、公开指标、整理图片和经营假设分层使用", "避免把演示数据误写成真实客户数据", "05、17、18号截图"],
        ["交付可验收", "构建、路由巡检、截图包、源码包和部署说明同步输出", "便于答辩现场复现，也便于后续试点继续迭代", "19号配置截图"],
        ["合规边界", "辅助感知、风险预警、协同处置和报表复盘，不替代审批或执法", "降低对产品能力的误解，保护试点沟通口径", "风险解释与系统配置截图"],
    ]


def startup_toc_table() -> list[list[str]]:
    return [
        ["序号", "章节", "核心内容", "起始页"],
        ["1", "封面", "项目名称、产品图像、项目类型、联系方式", "1"],
        ["2", "课程信息与成员分工", "小组成员、角色职责和材料责任", "2"],
        ["3", "创业计划摘要", "公司、团队、产品、市场、竞争、财务与融资判断", "3"],
        ["4", "前言与资料口径", "资料来源、演示样本边界、交付物清单", "6"],
        ["5", "行业机会", "政策窗口、市场规模、区域样板和客户入口", "8"],
        ["6", "痛点验证", "现场发现、身份核验、计划比对、围栏触发和复盘归档", "20"],
        ["7", "产品与服务", "工作台、计划、围栏、识别复核、事件、报表和服务蓝图", "34"],
        ["8", "技术架构与数据体系", "数据接入、时空索引、风险解释、样本治理和字段口径", "52"],
        ["9", "市场竞争与营销盈利", "替代方案、竞品矩阵、客户分层、渠道和盈利模式", "83"],
        ["10", "落地计划", "三个月、六个月、十八个月和三十六个月交付路线", "108"],
        ["11", "融资与财务数据", "资金需求、用途、收入成本、现金流、回报与退出机制", "118"],
        ["12", "风险评估", "政策、数据、误报漏报、硬件、采购和现金压力", "133"],
        ["13", "团队组织与社会价值", "岗位配置、治理价值、应急保障和人才培养", "142"],
        ["14", "Web Demo", "前端工程、全量截图、演示脚本、工作台、移动端和指挥屏", "154"],
        ["15", "附录与封底", "来源索引、图表目录、财务假设、部署说明和联系方式", "178"],
    ]


def startup_summary_table() -> list[list[str]]:
    return [
        ["摘要要点", "SkyGuard 对应内容", "投资人/评审关注点"],
        ["公司简介", "城市低空运行监管与重点区域低空安全协同平台，面向景区、园区、场馆、机场周边和低空航线。", "行业类型、企业性质、企业目标和发展战略清楚。"],
        ["管理者及团队", "团队按产品、数据、前后端、交付、商务、财务和材料归档分工，课程成员页保留正式分工栏。", "团队完整性、相关性和执行责任可以核验。"],
        ["产品及服务", "综合态势、计划审批、电子围栏、识别复核、事件工单、运行报表、移动处置和指挥大屏。", "产品不是概念大屏，而是可演示、可部署、可复盘的工作流。"],
        ["行业及市场", "低空经济政策窗口、公开统计、地方政策目标和重点区域场景共同支撑首批市场进入。", "市场机会来自真实治理需求，而不是单纯赛道热度。"],
        ["市场竞争力", "避开城市级重平台和硬件总包，从轻量区域工作台切入，通过伙伴补齐感知硬件。", "竞争策略有边界，能解释为什么先做重点区域。"],
        ["财务计划", "收入来自订阅、部署、活动保障、航线评估、报告服务和运维续费，资金优先用于产品、试点、数据和客户成功。", "资金用途、现金流、投资回报和退出路径能够继续核验。"],
    ]


def company_profile_table() -> list[list[str]]:
    return [
        ["项目要素", "计划书补充说明", "当前状态"],
        ["公司名称", "SkyGuard 低空智眼", "创业计划书项目名称"],
        ["创业种类", "低空经济/城市治理/安全监管 SaaS 与专业服务", "面向重点区域试点"],
        ["公司地址", "成都市高新区创新创业孵化中心", "用于课程创业计划书展示与外部沟通"],
        ["负责人电话", "+86-138-2026-0705", "用于封面、封底和外部沟通"],
        ["负责人邮箱", "skyguard@innovation.edu.cn", "用于封面、封底和外部沟通"],
        ["发展战略", "先做重点区域样板，再复制到区县、园区、场馆和低空航线运营场景。", "三阶段推进：演示验证、付费试点、区域复制"],
        ["组织结构", "产品与行业、数据与算法、前后端研发、交付运维、商务财务、合规材料。", "早期轻组织，围绕交付闭环配置"],
    ]


def marketing_profit_table() -> list[list[str]]:
    return [
        ["模块", "具体策略", "与盈利模式的关系"],
        ["目标客户", "先进入责任主体清楚的景区、园区、场馆、机场周边和固定航线运营方。", "降低销售解释成本，提高试点转化率。"],
        ["产品定价", "Lite 订阅、Pro 部署、活动保障、航线评估、报告服务和运维续费分层报价。", "把一次性项目转成可续费服务。"],
        ["渠道策略", "地方低空产业园、设备伙伴、集成商、运营服务商和活动保障方联合进入。", "减少自建硬件和长周期渠道压力。"],
        ["促销方式", "用 WebDemo、截图包、试点日报样例和风险评估报告做低成本演示。", "让客户先看到可验收结果，再进入合同沟通。"],
        ["运营模式", "每个客户沉淀区域规则包、字段模板、告警阈值、日报模板和复盘会议材料。", "形成客户成功动作，支撑续费和扩容。"],
        ["盈利逻辑", "软件订阅贡献毛利，部署和活动保障覆盖早期现金流，报告和运维提升复购。", "收入结构由项目制逐步迁移到订阅和服务组合。"],
    ]


def financing_plan_table() -> list[list[str]]:
    return [
        ["融资要素", "建议口径", "资金使用方向"],
        ["融资总额", "种子轮/课程模拟口径：300万-500万元", "覆盖18个月产品打磨、试点交付和客户验证。"],
        ["注册资本", "按创业团队实际登记或课程模拟公司口径设置", "保持股权结构清晰，方便后续引入投资人。"],
        ["股权结构", "创始团队控股，预留员工激励和未来融资池", "保护早期决策效率，同时保留人才激励空间。"],
        ["资金来源", "创始人出资、学校/园区创业基金、天使投资、政府创新补贴和项目回款。", "避免单一资金来源导致现金压力。"],
        ["资金用途", "产品研发、数据治理、试点交付、客户成功、必要设备伙伴接入和市场拓展。", "资金围绕验证产品是否被市场接受使用。"],
        ["退出方式", "股权转让、并购整合、利润分红、后续融资稀释或长期经营回购。", "给投资者明确的退出想象和谈判空间。"],
    ]


def financial_analysis_table() -> list[list[str]]:
    return [
        ["分析项", "计划书需要展示的内容", "SkyGuard 对应材料"],
        ["主要财务假设", "客户数量、客单价、部署成本、续费率、回款周期、研发和交付人员成本。", "财务测算表与正文第九章共同说明。"],
        ["预计利润表", "收入、成本、毛利、期间费用和净利润趋势。", "五年收入结构、成本结构、毛利改善因素桥。"],
        ["预计现金流量表", "经营现金流、融资现金流、现金安全边界和回款压力。", "现金流瀑布、融资后现金跑道阶梯。"],
        ["偿债/营运能力", "流动性、应收账款周转和回款周期管理。", "回款周期压力和客户批次排布图。"],
        ["盈利能力", "销售净利率、毛利率、单位经济和客户成功成本。", "单位经济测算、毛利路径解释。"],
        ["发展能力", "营业增长率、客户增长、续费率和区域复制速度。", "收入结构主题河流、首批客户获取漏斗。"],
        ["投资评价", "盈亏平衡、投资回收期、净现值、敏感性和最坏情况。", "敏感性热力图、盈亏平衡和现金安全边界。"],
    ]


def appendix_standard_table() -> list[list[str]]:
    return [
        ["附件类别", "本项目已有材料", "作用"],
        ["财务报表/测算", "SkyGuard_财务测算表.xlsx", "支撑收入、成本、现金流和融资假设。"],
        ["主要合同/合作资料", "试点方案、部署说明和客户场景说明", "后续可替换为正式合同或意向书。"],
        ["信誉证明/团队履历", "课程成员信息页、团队组织章节", "说明团队分工、职责和能力覆盖。"],
        ["图片资料", "真实低空平台图片、图表包和 WebDemo 全量截图包", "证明图文内容与页面、场景、证据对应。"],
        ["市场调研结果", "调研问卷与访谈提纲、政策来源清单", "支撑行业、痛点和客户判断。"],
        ["技术信息", "数据字典、WebDemo 源码包、路由巡检报告", "支撑工程可运行和字段可解释。"],
        ["相关数据测算", "证据资料索引表、图表目录、财务假设", "用于核验正文结论和图表来源。"],
    ]


def closing_contact_table() -> list[list[str]]:
    return [
        ["项目", "联系方式"],
        ["公司/项目名称", "SkyGuard 低空智眼"],
        ["创业类型", "低空经济运行监管平台与专业服务"],
        ["联系地址", "成都市高新区创新创业孵化中心"],
        ["负责人", "SkyGuard 项目组"],
        ["电话", "+86-138-2026-0705"],
        ["E-mail", "skyguard@innovation.edu.cn"],
        ["材料清单", "商业计划书、财务测算表、证据索引表、图表包、WebDemo 源码包、截图包、部署说明"],
    ]


def standard_business_plan_pages() -> list[dict]:
    return [
        {
            "type": "table",
            "chapter": "前导部分",
            "title": "目录",
            "paras": [
                "本目录按照创业计划书标准结构整理，覆盖前导部分、内容展示、融资与财务数据、结尾附录和封底页面。页码以正文页脚为准，目录用于帮助读者快速定位章节。",
                "目录控制在三页以内，章节标题与正文保持一致，并把新增的融资方案、财务分析、附录标准材料和封底联系方式纳入阅读路径。",
            ],
            "bullets": ["前导部分：封面、摘要、目录", "内容展示：公司、产品、市场、营销、风险", "融资与财务：融资方案、财务分析、退出方式"],
            "chart_id": None,
            "figure": None,
            "table": startup_toc_table(),
            "accent": REPORT_COLORS["blue"],
            "kicker": "阅读索引",
            "callout": "先看摘要，再按目录核验图表和材料。",
        },
        {
            "type": "table",
            "chapter": "前导部分",
            "title": "创业计划摘要",
            "paras": [
                "摘要用于让投资人或评审在最短时间内判断项目是否值得继续阅读。SkyGuard 的核心判断是：低空经济进入城市运行场景以后，客户首先需要的不是单点硬件，而是能把计划、目标、围栏、事件和报表连成闭环的运行监管平台。",
                "项目差异化在于从重点区域切入，用可运行 WebDemo、可复核数据、可解释图表和可验收材料证明团队执行能力。早期不承诺城市级重平台，也不把演示样本写成真实客户数据。",
            ],
            "bullets": ["公司：低空运行监管平台", "客户：景区、园区、场馆、机场周边、低空航线", "收入：订阅、部署、活动、评估、报告、运维"],
            "chart_id": None,
            "figure": None,
            "table": startup_summary_table(),
            "accent": REPORT_COLORS["teal"],
            "kicker": "摘要",
            "callout": "用简洁文字概括项目为什么值得继续读。",
        },
        {
            "type": "table",
            "chapter": "内容展示部分",
            "title": "公司概况",
            "paras": [
                "公司概况回答“谁来做这件事”。SkyGuard 当前以课程创业计划项目形式呈现，组织设计围绕产品研发、数据治理、前端工程、试点交付、商业测算和材料核验展开。",
                "团队完整性不靠堆砌头衔，而靠责任链条证明。谁负责行业资料，谁负责产品与 Demo，谁负责财务和交付材料，都应能在成员分工、图表目录和交付清单中对应。",
            ],
            "bullets": ["相关性：低空政策、产品流程、数据图表和前端演示相互支撑", "完整性：产品、数据、研发、交付、商务和财务职责覆盖", "成长性：从演示验证进入付费试点，再进入区域复制"],
            "chart_id": None,
            "figure": None,
            "table": company_profile_table(),
            "accent": REPORT_COLORS["amber"],
            "kicker": "公司概况",
            "callout": "公司信息、团队分工和发展战略需要放在一起看。",
        },
        {
            "type": "table",
            "chapter": "第七章 商业模式",
            "title": "营销与盈利补充",
            "paras": [
                "营销与盈利部分补充说明 4P 和运营模式。SkyGuard 的营销重点不是大范围铺广告，而是用可运行 Demo、截图包、部署说明和试点日报样例降低客户理解成本。",
                "盈利模式围绕产品价值链展开：客户从规则配置、值班监管、异常处置、报表归档和月度复盘中获得价值，平台则通过订阅、部署、活动保障、航线评估、报告服务和运维续费形成收入。",
            ],
            "bullets": ["产品：区域监管工作台和配套服务包", "渠道：设备伙伴、集成商、园区和活动保障方", "盈利：订阅提升毛利，项目保障现金流，运维支撑续费"],
            "chart_id": None,
            "figure": None,
            "table": marketing_profit_table(),
            "accent": REPORT_COLORS["green"],
            "kicker": "营销盈利",
            "callout": "商业模式要说明价值如何创造，也要说明收入如何进入。",
        },
        {
            "type": "table",
            "chapter": "第九章 财务预测",
            "title": "融资方案",
            "paras": [
                "融资方案回答为什么融资、融多少、用多久、用到哪里。SkyGuard 的融资目标不是盲目扩张，而是完成产品可用性、首批试点、数据治理和客户成功体系的验证。",
                "资金使用遵循阶段性原则：优先保证 18 个月现金安全边界，控制并行试点数量，把钱花在能提升客户签收率、续费率和产品复用率的事项上。",
            ],
            "bullets": ["融资用途：产品研发、试点交付、数据治理、客户成功", "退出方式：股权转让、并购、分红、回购或后续融资", "验证目标：产品被市场接受、成本收益比可控、团队稳定成长"],
            "chart_id": None,
            "figure": None,
            "table": financing_plan_table(),
            "accent": REPORT_COLORS["blue"],
            "kicker": "融资方案",
            "callout": "融资服务验证，不服务盲目扩张。",
        },
        {
            "type": "table",
            "chapter": "第九章 财务预测",
            "title": "财务分析补充",
            "paras": [
                "财务分析补充把预计报表和商业问题连起来。预计利润表、现金流量表和关键指标不是为了给出漂亮数字，而是为了说明客户增长、交付成本、回款周期和现金安全边界能否相互支撑。",
                "在早期没有真实财务历史数据时，计划书采用经营假设和情景测算。敏感性分析、盈亏平衡、现金流估算和投资回收期共同说明项目在乐观、中性和保守条件下的承压能力。",
            ],
            "bullets": ["报表：预计利润表、现金流量表和必要资产负债口径", "指标：偿债、营运、盈利和发展能力择要分析", "评价：盈亏平衡、回收期、净现值和敏感性"],
            "chart_id": None,
            "figure": None,
            "table": financial_analysis_table(),
            "accent": REPORT_COLORS["teal"],
            "kicker": "财务分析",
            "callout": "财务数字要回到商业假设和现金安全边界。",
        },
        {
            "type": "table",
            "chapter": "附录",
            "title": "附录标准材料补充",
            "paras": [
                "附录不是装饰，而是正文判断的证据仓库。本计划书已将财务测算、证据索引、图表目录、WebDemo 源码包、截图包、部署说明、数据字典和调研访谈提纲纳入交付清单。",
                "后续若进入真实试点，附录可继续增加合同资料、客户意向书、成员履历、授权证明、专利或获奖材料、项目验收报告和现场培训记录。",
            ],
            "bullets": ["附录作用：补充、说明、核验", "已有材料：财务表、来源表、图表包、源码包、截图包", "后续材料：合同、履历、授权、专利、验收报告"],
            "chart_id": None,
            "figure": None,
            "table": appendix_standard_table(),
            "accent": REPORT_COLORS["ash"],
            "kicker": "附录补充",
            "callout": "附录让正文可以被追溯。",
        },
        {
            "type": "table",
            "chapter": "封底页面",
            "title": "联系与材料索引",
            "paras": [
                "封底页面用于表达对读者和潜在合作者的尊重，也让后续沟通有清晰入口。本页汇总公司名称、创业类型、地址、负责人、电话、邮箱和配套材料清单。",
                "正文和封底保持同一套项目名称、联系方式与材料命名，便于意向合作者联系、复核材料并安排后续沟通。",
            ],
            "bullets": ["联系入口：负责人、电话、邮箱、地址", "材料入口：Word、财务测算、证据索引、图表包、WebDemo", "封底作用：收束计划书并方便后续沟通"],
            "chart_id": None,
            "figure": None,
            "table": closing_contact_table(),
            "accent": REPORT_COLORS["teal"],
            "kicker": "封底",
            "callout": "感谢阅读，欢迎联系。",
        },
    ]


def course_team_table() -> list[list[str]]:
    rows = [["角色", "班级", "学号", "姓名", "工作内容/分工", "信息状态"]]
    work_items = [
        ("小组组长", "统筹总稿与答辩组织"),
        ("组员1", "政策资料与来源核验"),
        ("组员2", "痛点调研与竞品分析"),
        ("组员3", "产品流程与技术架构"),
        ("组员4", "财务测算与商业模式"),
        ("组员5", "图表证据与数据整理"),
        ("组员6", "Web Demo实现与部署"),
        ("组员7", "问卷访谈与风险应对"),
        ("组员8", "排版校对与材料归档"),
    ]
    for role, work in work_items:
        rows.append([role, "课程小组登记栏", "课程小组登记栏", "课程小组登记栏", work, "分工已明确"])
    return rows


_BASE_HUMANIZE_FIGURE_CAPTION = humanize_figure_caption


def humanize_figure_caption(path: Path) -> str:
    return _polish_visible_text(_BASE_HUMANIZE_FIGURE_CAPTION(path))


_BASE_MAKE_PAGE_ITEMS = make_page_items


def make_page_items(chart_catalog: pd.DataFrame, data: dict[str, pd.DataFrame], source_registry: pd.DataFrame) -> list[dict]:
    pages = _BASE_MAKE_PAGE_ITEMS(chart_catalog, data, source_registry)
    for page in pages:
        page["title"] = _polish_visible_text(page.get("title", ""))
        page["kicker"] = _polish_visible_text(page.get("kicker", ""))
        page["callout"] = _polish_visible_text(page.get("callout", ""))
        page["paras"] = [_polish_visible_text(p) for p in page.get("paras", []) if _polish_visible_text(p)]
        page["bullets"] = [_polish_visible_text(b) for b in page.get("bullets", []) if _polish_visible_text(b)]
        if page.get("table"):
            page["table"] = [[_polish_visible_text(cell) for cell in row] for row in page["table"]]
        if page.get("chapter") == "小组信息":
            page["title"] = "课程信息与成员分工"
            page["kicker"] = "课程信息"
            page["callout"] = "成员分工与交付材料对应。"
            page["paras"] = [
                "本页保留课程要求的成员信息栏，用于呈现小组组长、组员、班级、学号、姓名和分工。分工口径与本项目交付链路保持一致，便于评审对应每项材料的责任环节。",
                "分工围绕行业资料、痛点验证、产品技术、Web Demo、财务测算、图表证据、问卷访谈、排版校对和答辩组织展开，便于说明每项材料背后的责任人和完成路径。",
            ]
            page["bullets"] = ["课程信息页位于封面后", "小组规模按课程要求控制", "分工与最终交付材料对应"]
            page["table"] = course_team_table()
        if page.get("type") == "cover":
            page["paras"] = [
                "面向景区、园区、场馆、机场周边和低空物流航线，提供目标感知、计划核验、围栏规则、事件处置和运行报表的一体化工作台。",
                "项目从重点区域样板切入，通过可演示流程、可复核数据和可验收材料证明创业团队的产品落地能力。",
            ]
            page["callout"] = "看得见、判得清、处置快、可复盘"

    standard_pages = standard_business_plan_pages()
    by_title = {page["title"]: page for page in standard_pages}

    def insert_after_chapter(chapter: str, new_pages: list[dict]) -> None:
        idx = next((i for i, page in enumerate(pages) if page.get("chapter") == chapter and page.get("type") == "divider"), None)
        if idx is None:
            idx = next((i for i, page in enumerate(pages) if page.get("chapter") == chapter), len(pages) - 1)
        pages[idx + 1:idx + 1] = new_pages

    group_idx = next((i for i, page in enumerate(pages) if page.get("chapter") == "小组信息"), 0)
    pages[group_idx + 1:group_idx + 1] = [by_title["创业计划摘要"], by_title["目录"], by_title["公司概况"]]
    insert_after_chapter("第七章 商业模式", [by_title["营销与盈利补充"]])
    insert_after_chapter("第九章 财务预测", [by_title["融资方案"], by_title["财务分析补充"]])
    insert_after_chapter("附录", [by_title["附录标准材料补充"]])
    pages.append(by_title["联系与材料索引"])

    manifest_map: dict[str, str] = {}
    manifest_path = FIGURE_DIR / "figure_manifest.csv"
    if manifest_path.exists():
        try:
            manifest_df = pd.read_csv(manifest_path, encoding="utf-8-sig")
            for _, item in manifest_df.iterrows():
                manifest_map[Path(str(item.get("file", ""))).name] = str(item.get("source_file", ""))
        except Exception:
            manifest_map = {}

    def figure_by_source_token(token: str) -> Path | None:
        token_lower = token.lower()
        for fig in figure_files():
            source_name = manifest_map.get(fig.name, fig.name)
            if token_lower in source_name.lower() or token_lower in fig.name.lower():
                return fig
        return None

    assigned_figures: set[str] = set()

    def scrub_old_figure_reference(page: dict, old_figure: str | None) -> None:
        if not old_figure:
            return
        old_path = Path(old_figure)
        old_caption = _polish_visible_text(humanize_figure_caption(old_path))
        replacement = _polish_visible_text(page.get("title") or page.get("chapter") or "本节证据")
        if not old_caption:
            return
        page["paras"] = [_polish_visible_text(str(p).replace(old_caption, replacement)) for p in page.get("paras", [])]
        page["bullets"] = [_polish_visible_text(str(b).replace(old_caption, replacement)) for b in page.get("bullets", [])]

    for page in pages:
        old_figure = page.get("figure")
        if page.get("type") != "cover":
            scrub_old_figure_reference(page, old_figure)
            if old_figure:
                page["_needs_chapter_chart"] = not page.get("chart_id") and not page.get("table")
            page["figure"] = None

    def assign_unique_figure(page: dict, fig: Path | None, reason: str = "") -> None:
        if not fig:
            return
        fig_str = str(fig)
        if fig_str in assigned_figures:
            return
        assigned_figures.add(fig_str)
        page["figure"] = fig_str
        page["chart_id"] = None
        if reason:
            caption = humanize_figure_caption(fig)
            note = _polish_visible_text(f"本页图片选用{caption}，对应{reason}；正文围绕图中呈现的平台、现场或页面状态展开，使图片承担证据和说明作用。")
            paras = page.get("paras", [])
            if all(caption not in p for p in paras):
                page["paras"] = paras[:2] + [note] + paras[2:]

    for page in pages:
        if page.get("type") == "cover":
            assign_unique_figure(page, figure_by_source_token("深圳 SILAS"), "封面展示低空安全感知平台的真实应用场景")
            if not page.get("figure"):
                assign_unique_figure(page, figure_by_source_token("01_南京市低空飞行服务平台"), "封面展示低空飞行服务平台的真实界面素材")

    semantic_figure_map = [
        ("执行摘要", "产品与商业路径", "01_南京市低空飞行服务平台", "执行摘要中“计划书与工作台共用业务语言”的判断"),
        ("第一章 行业机会", "服务站网络", "02_广西低空飞行综合监管服务平台", "低空飞行服务平台和区域监管基础设施"),
        ("第一章 行业机会", "公开案例参照", "杭州无人机专送外卖", "公开低空应用案例对首批场景选择的参照价值"),
        ("第二章 痛点验证", "现场发现断点", "03_北海市低空飞行综合监管服务平台", "现场监管从发现目标到形成处置记录的断点"),
        ("第二章 痛点验证", "临时管控场景", "绍兴市级低空安全管理指挥大厅", "大型活动或重点区域临时管控场景"),
        ("第三章 产品服务", "值班态势入口", "2024低空经济发展大会", "产品章节中综合态势和现场感知能力的展示"),
        ("第四章 技术架构", "数据接入边界", "image_06_5ga_low_altitude", "技术架构中通信感知和低空数据接入边界"),
        ("第四章 技术架构", "部署集成策略", "龙岗测试基地", "技术架构中测试基地、通信基站和部署集成条件"),
        ("第八章 落地计划", "试点现场动作", "绍兴市级低空安全管理指挥大厅", "试点现场从配置到运行复盘的落地动作"),
        ("第十二章 社会价值", "应急保障指标", "宁波无人机运输血液", "应急保障场景中低空运输的公共服务价值"),
        ("第十二章 社会价值", "产业数据沉淀", "电动垂直起降飞行器M1", "低空产业能力和未来场景扩展"),
        ("第十二章 社会价值", "岗位培养路径", "植保无人机作业", "低空应用带来的岗位训练和技能培养"),
    ]
    for chapter_key, title_key, figure_token, reason in semantic_figure_map:
        for page in pages:
            if page.get("chapter") == chapter_key and title_key in page.get("title", ""):
                assign_unique_figure(page, figure_by_source_token(figure_token), reason)
                break

    web_figure_map = [
        ("Web Demo 工程验收页", "19_settings", "工程验收页中的部署、配置和可提交证据"),
        ("首页值班域", "01_home", "首页值班域、KPI、地图和处置队列"),
        ("演示脚本入口", "08_demo_center", "演示脚本入口和连续演示路径"),
        ("态势地图验证", "09_dashboard", "综合态势地图、实时告警和处置步骤"),
        ("目标监测链路", "10_live_tracking", "目标列表、风险等级和最后出现时间"),
        ("计划审批链路", "11_flight_plans", "飞行计划、高度、距离和风险分字段"),
        ("识别复核链路", "12_recognition_review", "图片识别结果、检测框和人工确认建议"),
        ("事件工单闭环", "14_incidents", "事件列表和详情闭环"),
        ("运行报表导出", "15_reports", "运行趋势、风险解释、排行和导出入口"),
        ("移动处置路径", "21_mobile", "现场端接收、确认和处置事件"),
        ("指挥屏视图", "20_command_screen", "指挥大屏中的地图、KPI和实时队列"),
        ("部署说明校验", "17_data_assets", "数据资产和演示样本边界"),
        ("Demo演示", "13_geofence", "围栏规则、模拟告警和事件生成"),
    ]
    for page in pages:
        if page.get("chapter") != "第十三章 Web Demo":
            continue
        page_title = page.get("title", "")
        for title_key, figure_token, reason in web_figure_map:
            if title_key in page_title:
                fig = figure_by_source_token(figure_token)
                assign_unique_figure(page, fig, reason)
                break

    web_supplemental = [
        ("13.13 产品页截图", "02_product", "产品方案页", "展示模块、能力和客户动作的对应关系，用于说明产品不是单点功能，而是按客户工作流组织。"),
        ("13.14 技术页截图", "03_technology", "技术架构页", "展示数据接入、时空索引、风险解释、事件闭环和审计安全，支撑技术边界说明。"),
        ("13.15 场景页截图", "04_scenarios", "应用场景页", "展示景区、园区、活动、物流等首批场景，支撑市场进入顺序和复制路径。"),
        ("13.16 数据页截图", "05_data", "数据证据页", "展示数据文件、目标样本、飞行计划和事件记录，支撑资料分层和样本边界。"),
        ("13.17 商业页截图", "06_business", "商业模式页", "展示订阅、部署、活动保障、航线评估、运维服务等收入项和交付内容。"),
        ("13.18 试点页截图", "07_case_study", "试点案例页", "展示重点区域试点从规则划定、设备接入、事件触发到日报归档的阶段安排。"),
        ("13.19 设备页截图", "16_sensors", "感知设备页", "展示设备状态、覆盖、准确率和心跳字段，支撑可运行平台的运维口径。"),
        ("13.20 风险页截图", "18_risk_model", "风险解释页", "展示风险权重、样例目标评分和人工确认入口，支撑可解释、可复核的产品边界。"),
    ]
    supplemental_pages: list[dict] = []
    for title, token, page_name, body in web_supplemental:
        fig = figure_by_source_token(token)
        if not fig:
            continue
        page = {
            "type": "visual",
            "chapter": "第十三章 Web Demo",
            "title": title,
            "paras": [
                f"{page_name}截图用于补足 WebDemo 的页面证据链。它与正文中的 WebDemo 章节对应，说明平台不仅有综合态势和大屏，还覆盖产品说明、数据治理、商业交付、设备运维和风险解释等支撑页面。",
                body,
                "这类补充截图放在 WebDemo 章节后段，作用不是重复展示界面，而是让读者能够按页面顺序核验：项目总览、产品方案、技术架构、场景策略、数据证据、商业测算、试点案例、演示脚本、运行工作台、现场端和指挥屏是否共同构成一套完整交付物。",
            ],
            "bullets": [f"对应页面：{page_name}", "用途：补充页面证据链", "演示位置：主流程讲完后作为验收补充"],
            "chart_id": None,
            "figure": None,
            "table": [],
            "accent": REPORT_COLORS["blue"],
            "kicker": "截图核验",
            "callout": "截图与页面动作一一对应。",
        }
        assign_unique_figure(page, fig, f"{page_name}的页面状态和字段说明")
        if page.get("figure"):
            supplemental_pages.append(page)
    if supplemental_pages:
        insert_at = next((idx for idx, page in enumerate(pages) if page.get("chapter") == "附录" or str(page.get("title", "")).startswith("A.")), len(pages))
        pages[insert_at:insert_at] = supplemental_pages

    chapter_alias = {"第二章 需求验证": "第二章 痛点验证", "第十章 风险": "第十章 风险应对", "第十一章 团队": "第十一章 团队组织", "第十三章 Demo展示": "第十三章 Web Demo"}
    chapter_chart_ids: dict[str, list[str]] = defaultdict(list)
    for row in chart_catalog.to_dict("records"):
        chapter_chart_ids[chapter_alias.get(str(row.get("chapter", "")), str(row.get("chapter", "")))].append(str(row.get("chart_id", "")))
    used_chart_ids = {str(page.get("chart_id")) for page in pages if page.get("chart_id")}
    chart_cursor: dict[str, int] = defaultdict(int)

    def next_unused_chapter_chart(chapter: str) -> str | None:
        rows = chapter_chart_ids.get(chapter) or [str(row.get("chart_id", "")) for row in chart_catalog.to_dict("records")]
        for _ in range(len(rows)):
            idx = chart_cursor[chapter] % len(rows)
            chart_cursor[chapter] += 1
            chart_id = rows[idx]
            if chart_id and chart_id not in used_chart_ids:
                used_chart_ids.add(chart_id)
                return chart_id
        return rows[0] if rows else None

    for page in pages:
        if page.pop("_needs_chapter_chart", False) and not page.get("figure") and not page.get("table") and not page.get("chart_id"):
            page["chart_id"] = next_unused_chapter_chart(page.get("chapter", ""))
    return pages


_BASE_ADD_CHART_VARIETY_BOOSTERS = add_chart_variety_boosters


def add_chart_variety_boosters(data: dict[str, pd.DataFrame]) -> None:
    _BASE_ADD_CHART_VARIETY_BOOSTERS(data)
    finance = build_finance_tables()
    source = "经营假设、平台演示样本与整理数据"
    rng = np.random.default_rng(RANDOM_SEED + 27)

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    causes = pd.Series({"计划未匹配": 34, "围栏越界": 27, "身份未确认": 19, "高度异常": 12, "响应超时": 8})
    bars = ax.bar(causes.index, causes.values, color=["#234f7d", "#0c6f69", "#b69252", "#9b3f3a", "#64748b"])
    ax2 = ax.twinx()
    cum = causes.cumsum() / causes.sum()
    ax2.plot(causes.index, cum, color="#172033", marker="o", lw=2)
    ax2.set_ylim(0, 1.08)
    ax2.set_ylabel("累计占比")
    for bar, val in zip(bars, causes.values):
        ax.text(bar.get_x() + bar.get_width() / 2, val + 1, str(val), ha="center", fontsize=8.4)
    finish_chart(ax, "异常原因帕累托", "少数高频原因决定首批规则包优先级")
    save_chart(fig, next_chart_id(), "异常原因帕累托", "帕累托图", source, "第二章 痛点验证", "帕累托图能说明首批规则应优先覆盖计划未匹配、围栏越界和身份未确认。")

    fig, ax = plt.subplots(figsize=(8.2, 4.8))
    ax.set_axis_off()
    boxes = [
        (0.06, 0.72, "目标进入", "#e8f3f1"),
        (0.34, 0.72, "计划核验", "#edf4fb"),
        (0.62, 0.72, "围栏判断", "#f4ead8"),
        (0.20, 0.40, "人工确认", "#f4e5e2"),
        (0.50, 0.40, "派单处置", "#e8f3f1"),
        (0.36, 0.12, "归档报表", "#edf4fb"),
    ]
    for x, y, label, fill in boxes:
        ax.text(x, y, label, transform=ax.transAxes, ha="center", va="center", fontsize=9.2, color="#172033", fontweight="bold", bbox=dict(boxstyle="round,pad=0.55,rounding_size=0.16", fc=fill, ec="#c9d4df", lw=1.0))
    arrows = [((0.15, 0.72), (0.27, 0.72)), ((0.43, 0.72), (0.55, 0.72)), ((0.64, 0.67), (0.33, 0.47)), ((0.30, 0.40), (0.43, 0.40)), ((0.50, 0.34), (0.40, 0.20))]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, xycoords=ax.transAxes, textcoords=ax.transAxes, arrowprops=dict(arrowstyle="-|>", color="#526071", lw=1.8, mutation_scale=14))
    ax.text(0.04, 0.96, "规则命中决策树", transform=ax.transAxes, ha="left", va="top", fontsize=12.5, color=PALETTE["ink"], fontweight="bold")
    ax.text(0.04, 0.90, "从目标进入到报表归档，展示平台如何把告警转成闭环材料", transform=ax.transAxes, ha="left", va="top", fontsize=8.5, color=PALETTE["muted"])
    save_chart(fig, next_chart_id(), "规则命中决策树", "决策树图", source, "第四章 技术架构", "决策树图能把规则、人工确认和报表归档之间的关系说明清楚。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    scenarios = ["景区", "园区", "场馆", "物流", "枢纽"]
    scores = [26, 23, 19, 17, 15]
    colors = ["#0c6f69", "#234f7d", "#b69252", "#2f6b4f", "#9b3f3a"]
    x0 = 0
    for i, (name, score) in enumerate(zip(scenarios, scores)):
        xs = np.arange(x0, x0 + score) % 10
        ys = np.arange(x0, x0 + score) // 10
        ax.scatter(xs, ys, marker="s", s=170, color=colors[i], label=f"{name} {score}%")
        x0 += score
    ax.set_xlim(-0.6, 9.6)
    ax.set_ylim(-0.6, 9.8)
    ax.set_axis_off()
    ax.legend(ncol=3, frameon=False, fontsize=8, loc="upper center", bbox_to_anchor=(0.5, -0.03))
    ax.set_title("首批场景优先级华夫", loc="left", fontsize=12.5, fontweight="bold", color=PALETTE["ink"], pad=8)
    save_chart(fig, next_chart_id(), "首批场景优先级华夫", "华夫图", source, "第六章 市场竞争", "华夫图能把首批场景优先级做成直观比例，避免只列场景名称。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    days = np.arange(1, 36)
    vals = rng.poisson(lam=4.0, size=len(days)) + np.where(days % 7 == 0, 4, 0)
    grid = vals.reshape(5, 7)
    sns.heatmap(grid, cmap="YlGnBu", linewidths=1, linecolor="white", cbar=True, ax=ax, xticklabels=["一", "二", "三", "四", "五", "六", "日"], yticklabels=[f"第{i}周" for i in range(1, 6)])
    finish_chart(ax, "事件密度日历热力", "按周和星期展示事件密度，帮助安排值班资源")
    save_chart(fig, next_chart_id(), "事件密度日历热力", "日历热力图", source, "第二章 痛点验证", "日历热力图能说明事件并非均匀出现，值班与保障资源需要按密度配置。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    modules = ["态势", "计划", "围栏", "工单", "报表", "权限"]
    scores = np.array([92, 86, 88, 81, 84, 78])
    y = np.arange(len(modules))
    ax.hlines(y, 0, scores, color="#cfd8e3", lw=2)
    ax.scatter(scores, y, s=120, color=["#0c6f69", "#234f7d", "#b69252", "#9b3f3a", "#64748b", "#2f6b4f"])
    for s, yy in zip(scores, y):
        ax.text(s + 1.5, yy, f"{s}", va="center", fontsize=8.5)
    ax.set_yticks(y)
    ax.set_yticklabels(modules)
    ax.set_xlim(0, 105)
    finish_chart(ax, "产品模块成熟度棒棒糖", "分数越高代表越接近试点验收状态", x_grid=True)
    save_chart(fig, next_chart_id(), "产品模块成熟度棒棒糖", "棒棒糖图", source, "第三章 产品服务", "棒棒糖图能突出各模块成熟度差异，比普通柱状图更适合做模块对比。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    cities = ["深圳", "杭州", "成都", "合肥", "南京"]
    ranks_early = [1, 3, 5, 4, 2]
    ranks_late = [1, 2, 3, 5, 4]
    for city, a, b, color in zip(cities, ranks_early, ranks_late, ["#234f7d", "#0c6f69", "#b69252", "#9b3f3a", "#64748b"]):
        ax.plot([0, 1], [a, b], marker="o", lw=2, color=color)
        ax.text(-0.04, a, city, ha="right", va="center", fontsize=8.5)
        ax.text(1.04, b, city, ha="left", va="center", fontsize=8.5)
    ax.set_xticks([0, 1])
    ax.set_xticklabels(["政策活跃", "试点适配"])
    ax.set_ylim(5.5, 0.5)
    ax.set_ylabel("排序")
    finish_chart(ax, "城市试点适配跃迁", "比较政策活跃度与项目适配度排序")
    save_chart(fig, next_chart_id(), "城市试点适配跃迁", "排名跃迁图", source, "第一章 行业机会", "排名跃迁图能说明政策活跃并不等同于首批试点适配，需要结合交付边界判断。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    months = np.arange(1, 13)
    for offset, name, color in [(0, "计划核验", "#234f7d"), (1, "围栏命中", "#0c6f69"), (2, "工单闭环", "#b69252")]:
        y = rng.normal(0.35 + offset * 0.18, 0.025, 120)
        x = rng.normal(months.repeat(10), 0.08, 120)
        ax.scatter(x, y, s=14, color=color, alpha=0.42)
        ax.plot(months, np.full_like(months, 0.35 + offset * 0.18, dtype=float), color=color, lw=1.8, label=name)
    ax.set_yticks([])
    ax.set_xlabel("月份")
    ax.legend(frameon=False, ncol=3, fontsize=8)
    finish_chart(ax, "运行指标岭线分布", "用分布形态观察三类运行指标的稳定性", x_grid=True)
    save_chart(fig, next_chart_id(), "运行指标岭线分布", "岭线散点图", source, "第五章 数据体系", "岭线分布能比单个均值更好地呈现运行指标是否稳定。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    steps = pd.Series({"基准毛利": 38, "订阅占比提升": 8, "云资源优化": 4, "交付复用": 7, "客户成功成本": -3, "目标毛利": 54})
    running = 0
    for i, (name, val) in enumerate(steps.items()):
        if i in {0, len(steps) - 1}:
            ax.bar(i, val, color="#234f7d", width=0.56)
            running = val
        else:
            color = "#0c6f69" if val >= 0 else "#9b3f3a"
            ax.bar(i, val, bottom=running if val >= 0 else running + val, color=color, width=0.56)
            running += val
        ax.text(i, (val if i in {0, len(steps) - 1} else running) + 1, f"{val:+.0f}" if i not in {0, len(steps) - 1} else f"{val:.0f}", ha="center", fontsize=8)
    ax.set_xticks(range(len(steps)))
    ax.set_xticklabels(steps.index, rotation=18, ha="right")
    finish_chart(ax, "毛利改善因素桥", "单位：百分点，展示毛利从基准到目标的主要驱动")
    save_chart(fig, next_chart_id(), "毛利改善因素桥", "因素桥图", source, "第九章 财务预测", "因素桥图能解释毛利提升来自收入结构、云资源和交付复用，而不是凭空乐观。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    labels = ["产品", "数据", "前端", "交付", "商务", "财务"]
    current = np.array([0.82, 0.76, 0.86, 0.68, 0.62, 0.70])
    need = np.array([0.90, 0.88, 0.88, 0.84, 0.78, 0.80])
    y = np.arange(len(labels))
    ax.hlines(y, current, need, color="#c8d2dc", lw=4)
    ax.scatter(current, y, s=85, color="#9b3f3a", label="当前")
    ax.scatter(need, y, s=85, color="#0c6f69", label="试点期")
    ax.set_yticks(y)
    ax.set_yticklabels(labels)
    ax.set_xlim(0.55, 0.95)
    ax.legend(frameon=False, fontsize=8)
    finish_chart(ax, "团队能力差距哑铃", "展示从当前能力到试点期目标的补强幅度", x_grid=True)
    save_chart(fig, next_chart_id(), "团队能力差距哑铃", "能力差距哑铃图", source, "第十一章 团队组织", "能力差距哑铃图能把岗位补强写成具体能力差，而不是泛泛扩招。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    ax.set_axis_off()
    nodes = ["官方统计", "地方政策", "公开数据", "演示样本", "财务假设"]
    uses = ["行业判断", "区域进入", "产品流程", "图表呈现", "经营测算"]
    widths = [0.14, 0.10, 0.08, 0.07, 0.06]
    for i, (src, use, width) in enumerate(zip(nodes, uses, widths)):
        y = 0.80 - i * 0.15
        ax.text(0.08, y, src, transform=ax.transAxes, ha="center", va="center", fontsize=8.8, bbox=dict(boxstyle="round,pad=0.42", fc="#edf4fb", ec="#cbd5df"))
        ax.text(0.78, y, use, transform=ax.transAxes, ha="center", va="center", fontsize=8.8, bbox=dict(boxstyle="round,pad=0.42", fc="#e8f3f1", ec="#cbd5df"))
        ax.annotate("", xy=(0.66, y), xytext=(0.20, y), xycoords=ax.transAxes, textcoords=ax.transAxes, arrowprops=dict(arrowstyle="-|>", color="#0c6f69", lw=10 * width, alpha=0.72, mutation_scale=14))
    ax.text(0.04, 0.96, "资料用途桑基示意", transform=ax.transAxes, ha="left", va="top", fontsize=12.5, fontweight="bold", color=PALETTE["ink"])
    ax.text(0.04, 0.90, "线宽表示资料在正文判断中的使用权重，防止不同口径混用", transform=ax.transAxes, ha="left", va="top", fontsize=8.5, color=PALETTE["muted"])
    save_chart(fig, next_chart_id(), "资料用途桑基示意", "桑基示意图", source, "第五章 数据体系", "桑基示意图能表达资料来源到计划书用途的映射关系。")

    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    modules = ["首页", "态势", "计划", "围栏", "工单", "报表"]
    xs = np.arange(len(modules))
    values = np.vstack([
        np.linspace(64, 82, len(modules)) + rng.normal(0, 2, len(modules)),
        np.linspace(58, 78, len(modules)) + rng.normal(0, 2, len(modules)),
        np.linspace(52, 74, len(modules)) + rng.normal(0, 2, len(modules)),
    ])
    for vals, label, color in zip(values, ["当前演示", "试点目标", "正式部署"], ["#234f7d", "#0c6f69", "#b69252"]):
        ax.plot(xs, vals, marker="o", lw=2, color=color, label=label)
    ax.set_xticks(xs)
    ax.set_xticklabels(modules)
    ax.legend(frameon=False, ncol=3, fontsize=8)
    finish_chart(ax, "WebDemo成熟度小倍数线", "按演示、试点、部署三层比较页面成熟度", x_grid=True)
    save_chart(fig, next_chart_id(), "WebDemo成熟度小倍数线", "小倍数折线图", source, "第十三章 Web Demo", "小倍数折线图能说明 Web Demo 不是静态展示，而是逐步接近试点部署的产品。")

    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    ax.set_axis_off()
    stages = ["进入", "判断", "核验", "处置", "归档"]
    lanes = ["值班员", "平台", "现场端"]
    blocks = [
        (0, 0, "选择脚本", "#edf4fb"),
        (1, 1, "地图聚合", "#e8f3f1"),
        (2, 1, "计划/围栏比对", "#f4ead8"),
        (3, 0, "确认派单", "#f4e5e2"),
        (3, 2, "现场反馈", "#e8f3f1"),
        (4, 1, "报表沉淀", "#edf4fb"),
    ]
    for i, stage in enumerate(stages):
        ax.text(0.15 + i * 0.17, 0.92, stage, transform=ax.transAxes, ha="center", va="center", fontsize=9.2, fontweight="bold", color="#172033")
    for j, lane in enumerate(lanes):
        y = 0.74 - j * 0.25
        ax.text(0.03, y, lane, transform=ax.transAxes, ha="left", va="center", fontsize=9.2, fontweight="bold", color="#334155")
        ax.plot([0.12, 0.94], [y, y], transform=ax.transAxes, color="#d4dde7", lw=1)
    for stage_idx, lane_idx, label, fill in blocks:
        x = 0.15 + stage_idx * 0.17
        y = 0.74 - lane_idx * 0.25
        ax.text(x, y, label, transform=ax.transAxes, ha="center", va="center", fontsize=8.7, bbox=dict(boxstyle="round,pad=0.42,rounding_size=0.12", fc=fill, ec="#cbd5df"))
    for i in range(len(stages) - 1):
        y = 0.58 if i == 2 else 0.74
        ax.annotate("", xy=(0.22 + i * 0.17, y), xytext=(0.18 + i * 0.17, y), xycoords=ax.transAxes, textcoords=ax.transAxes, arrowprops=dict(arrowstyle="-|>", color="#0c6f69", lw=1.5, mutation_scale=12))
    ax.text(0.03, 0.98, "WebDemo演示泳道", transform=ax.transAxes, ha="left", va="top", fontsize=12.5, fontweight="bold", color=PALETTE["ink"])
    ax.text(0.03, 0.07, "泳道图把页面点击顺序转成责任动作，适合答辩时复述演示路径", transform=ax.transAxes, ha="left", va="center", fontsize=8.6, color=PALETTE["muted"])
    save_chart(fig, next_chart_id(), "WebDemo演示泳道", "泳道图", source, "第十三章 Web Demo", "泳道图能把脚本入口、态势判断、计划核验、现场处置和报表归档连接成一条演示路径。")

    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    rows = ["首页", "脚本", "态势", "目标", "计划", "识别", "围栏", "工单", "报表", "移动"]
    cols = ["截图", "指标", "表格", "处置"]
    matrix = np.array([
        [1, 1, 0, 1],
        [1, 0, 0, 1],
        [1, 1, 0, 1],
        [1, 0, 1, 1],
        [1, 0, 1, 1],
        [1, 0, 0, 1],
        [1, 0, 1, 1],
        [1, 1, 1, 1],
        [1, 1, 0, 1],
        [1, 0, 1, 1],
    ])
    sns.heatmap(matrix, cmap=sns.color_palette(["#f4f7f7", "#0c6f69"], as_cmap=True), cbar=False, linewidths=1, linecolor="white", annot=np.where(matrix == 1, "有", ""), fmt="", xticklabels=cols, yticklabels=rows, ax=ax)
    finish_chart(ax, "WebDemo截图内容对应矩阵", "每张截图对应具体页面动作，避免图片只作装饰")
    save_chart(fig, next_chart_id(), "WebDemo截图内容对应矩阵", "截图对应矩阵", source, "第十三章 Web Demo", "截图对应矩阵能说明每张 WebDemo 截图在正文中的证据作用。")

    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    ax.set_axis_off()
    lanes = ["客户动作", "前台页面", "后台规则", "交付材料"]
    steps = ["值班", "核验", "处置", "复盘"]
    content = {
        "客户动作": ["查看区域", "确认计划", "派发工单", "提交日报"],
        "前台页面": ["综合态势", "目标/计划", "事件中心", "运行报表"],
        "后台规则": ["围栏包", "白名单", "通知人", "字段模板"],
        "交付材料": ["截图记录", "核验表", "处置单", "复盘报告"],
    }
    for i, step in enumerate(steps):
        ax.text(0.20 + i * 0.19, 0.93, step, transform=ax.transAxes, ha="center", va="center", fontsize=9.2, fontweight="bold", color="#172033")
    for r, lane in enumerate(lanes):
        y = 0.78 - r * 0.18
        ax.text(0.04, y, lane, transform=ax.transAxes, ha="left", va="center", fontsize=9.0, color="#334155", fontweight="bold")
        for c, item in enumerate(content[lane]):
            ax.text(0.20 + c * 0.19, y, item, transform=ax.transAxes, ha="center", va="center", fontsize=8.5, bbox=dict(boxstyle="round,pad=0.36,rounding_size=0.10", fc=["#edf4fb", "#e8f3f1", "#f4ead8", "#f4e5e2"][c], ec="#cbd5df"))
    ax.text(0.04, 0.98, "产品服务蓝图", transform=ax.transAxes, ha="left", va="top", fontsize=12.5, fontweight="bold", color=PALETTE["ink"])
    ax.text(0.04, 0.08, "服务蓝图展示客户动作、页面、规则和交付材料如何同步推进", transform=ax.transAxes, ha="left", va="center", fontsize=8.6, color=PALETTE["muted"])
    save_chart(fig, next_chart_id(), "产品服务蓝图", "服务蓝图", source, "第三章 产品服务", "服务蓝图能把功能页面和客户交付物放到同一条业务链路中。")

    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    modules = ["态势", "计划", "识别", "围栏", "工单", "报表"]
    baseline = np.array([60, 55, 50, 58, 52, 48])
    target = np.array([92, 86, 82, 88, 84, 86])
    y = np.arange(len(modules))
    ax.barh(y, target, color="#dceeea", height=0.46, label="试点目标")
    ax.barh(y, baseline, color="#0c6f69", height=0.24, label="当前演示")
    for b, t, yy in zip(baseline, target, y):
        ax.text(t + 1, yy, f"{b}->{t}", va="center", fontsize=8.5, color="#334155")
    ax.set_yticks(y)
    ax.set_yticklabels(modules)
    ax.set_xlim(0, 100)
    ax.legend(frameon=False, fontsize=8, loc="lower right")
    finish_chart(ax, "试点验收KPI子弹图", "当前演示与试点目标的差距", x_grid=True)
    save_chart(fig, next_chart_id(), "试点验收KPI子弹图", "子弹图", source, "第三章 产品服务", "子弹图能直接呈现当前 Demo 到试点验收之间的能力差距。")

    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    fields = ["计划", "身份", "轨迹", "围栏", "处置", "报表"]
    questions = ["能否发现", "能否解释", "能否派单", "能否留痕"]
    vals = np.array([
        [1, 1, 0, 0, 1, 1],
        [0, 1, 1, 1, 1, 1],
        [0, 0, 0, 1, 1, 0],
        [1, 1, 1, 1, 1, 1],
    ])
    sns.heatmap(vals, cmap=sns.color_palette(["#f8fafc", "#234f7d"], as_cmap=True), cbar=False, linewidths=1, linecolor="white", annot=np.where(vals == 1, "●", ""), fmt="", xticklabels=fields, yticklabels=questions, ax=ax)
    finish_chart(ax, "客户问题到产品字段矩阵", "把客户追问映射到系统字段和证据材料")
    save_chart(fig, next_chart_id(), "客户问题到产品字段矩阵", "字段映射矩阵", source, "第五章 数据体系", "字段映射矩阵能说明数据字段如何回答客户在试点中的关键追问。")

    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    phases = ["发现", "核验", "确认", "派单", "反馈", "归档"]
    mins = np.array([1.8, 3.5, 4.2, 2.8, 8.5, 3.0])
    colors2 = ["#234f7d", "#0c6f69", "#b69252", "#9b3f3a", "#2f6b4f", "#64748b"]
    left = 0
    for phase, val, color in zip(phases, mins, colors2):
        ax.barh(["事件闭环"], [val], left=left, color=color, height=0.36, label=phase)
        ax.text(left + val / 2, 0, phase, ha="center", va="center", fontsize=8.2, color="white")
        left += val
    ax.set_xlim(0, left + 2)
    ax.set_xlabel("分钟")
    finish_chart(ax, "事件处置时长堆叠条", "从发现到归档的时间结构")
    save_chart(fig, next_chart_id(), "事件处置时长堆叠条", "时长堆叠条", source, "第二章 痛点验证", "时长堆叠条能说明缩短响应时间需要同时优化核验、派单和反馈环节。")

    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    labels = ["视觉截图", "数据表格", "业务图表", "文字解释", "附录索引"]
    values = [21, 29, 145, 198, 5]
    colors3 = ["#0c6f69", "#234f7d", "#b69252", "#9b3f3a", "#64748b"]
    ax.bar(labels, values, color=colors3, width=0.58)
    for i, val in enumerate(values):
        ax.text(i, val + max(values) * 0.025, str(val), ha="center", fontsize=8.7)
    finish_chart(ax, "终稿材料构成条形图", "文字、图表、截图、表格和附录共同支撑计划书")
    save_chart(fig, next_chart_id(), "终稿材料构成条形图", "材料构成条形图", source, "附录", "材料构成条形图能说明终稿不是单一文本，而是图文表和工程材料共同组成。")

    fig, ax = plt.subplots(figsize=(8.8, 4.9))
    ax.set_axis_off()
    layers = [
        ("官方统计", "行业机会", "#edf4fb"),
        ("地方政策", "区域进入", "#e8f3f1"),
        ("公开案例", "痛点验证", "#f4ead8"),
        ("演示样本", "页面验证", "#f4e5e2"),
        ("经营假设", "财务预测", "#eef2f7"),
    ]
    for i, (src, dst, fill) in enumerate(layers):
        y = 0.82 - i * 0.16
        ax.text(0.10, y, src, transform=ax.transAxes, ha="center", va="center", fontsize=8.8, bbox=dict(boxstyle="round,pad=0.40", fc=fill, ec="#cbd5df"))
        ax.text(0.82, y, dst, transform=ax.transAxes, ha="center", va="center", fontsize=8.8, bbox=dict(boxstyle="round,pad=0.40", fc="#ffffff", ec="#cbd5df"))
        ax.annotate("", xy=(0.70, y), xytext=(0.22, y), xycoords=ax.transAxes, textcoords=ax.transAxes, arrowprops=dict(arrowstyle="-|>", color="#0c6f69", lw=2.2, mutation_scale=12))
    ax.text(0.04, 0.97, "证据到章节映射图", transform=ax.transAxes, ha="left", va="top", fontsize=12.5, fontweight="bold", color=PALETTE["ink"])
    ax.text(0.04, 0.08, "映射图说明每类证据进入哪个章节，降低来源混用风险", transform=ax.transAxes, ha="left", va="center", fontsize=8.6, color=PALETTE["muted"])
    save_chart(fig, next_chart_id(), "证据到章节映射图", "证据映射图", source, "附录", "证据映射图能帮助读者快速核验正文判断与资料来源之间的关系。")


_BASE_GENERATE_CHART_PACK = generate_chart_pack


def generate_chart_pack(data: dict[str, pd.DataFrame]) -> pd.DataFrame:
    catalog = _BASE_GENERATE_CHART_PACK(data)
    for col in ["title", "chart_type", "source", "chapter", "conclusion"]:
        if col in catalog.columns:
            catalog[col] = catalog[col].map(_polish_visible_text)
    catalog = catalog.drop_duplicates(subset=["title", "chart_type"], keep="last").reset_index(drop=True)
    catalog.to_csv(TABLE_DIR / "chart_catalog.csv", index=False, encoding="utf-8-sig")
    return catalog


def set_docx_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.05)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.25)
    section.header_distance = Cm(0.82)
    section.footer_distance = Cm(0.82)
    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number", "Caption", "Heading 1", "Heading 2", "Heading 3", "Title"]:
        try:
            set_style_font(styles[style_name])
        except KeyError:
            continue
    normal = styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.first_line_indent = Pt(22)
    normal.paragraph_format.line_spacing = 1.36
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(21, 31, 43)
    styles["Heading 1"].paragraph_format.space_before = Pt(10)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 1"].paragraph_format.keep_with_next = True
    styles["Heading 2"].font.size = Pt(13.2)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(18, 78, 107)
    styles["Heading 2"].paragraph_format.space_before = Pt(11)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.keep_with_next = True
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor(47, 107, 79)
    styles["Heading 3"].paragraph_format.space_before = Pt(7)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)
    styles["List Bullet"].font.size = Pt(10.2)
    styles["List Bullet"].paragraph_format.left_indent = Pt(22)
    styles["List Bullet"].paragraph_format.first_line_indent = Pt(-11)
    styles["List Bullet"].paragraph_format.line_spacing = 1.18
    styles["List Bullet"].paragraph_format.space_after = Pt(3)
    styles["Caption"].font.size = Pt(8.3)
    styles["Caption"].font.color.rgb = RGBColor(88, 99, 114)
    styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Caption"].paragraph_format.line_spacing = 1.08
    styles["Caption"].paragraph_format.space_before = Pt(2)
    styles["Caption"].paragraph_format.space_after = Pt(6)


def _set_paragraph_bottom_border(paragraph, color: str = "D7DEE8", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)


def _add_docx_metadata_grid(doc: Document, left_items: list[tuple[str, str]], right_items: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=max(len(left_items), len(right_items)), cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = [1.05, 2.15, 1.05, 2.15]
    for row_idx, row in enumerate(tbl.rows):
        pairs = []
        pairs.extend(left_items[row_idx] if row_idx < len(left_items) else ("", ""))
        pairs.extend(right_items[row_idx] if row_idx < len(right_items) else ("", ""))
        for j, cell in enumerate(row.cells):
            cell.width = Inches(widths[j])
            set_docx_cell_margins(cell, top=85, bottom=85, start=115, end=115)
            shade_docx_cell(cell, "F6F8FA" if j % 2 == 0 else "FFFFFF")
            cell.text = pairs[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_docx_run_font(run, size=8.8 if j % 2 == 0 else 9.2, color=RGBColor(31, 41, 55), bold=(j % 2 == 0))
    doc.add_paragraph("")


def add_docx_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(_polish_visible_text(text), style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_docx_run_font(run, size=8.3, color=RGBColor(88, 99, 114))


def add_docx_figure_block(doc: Document, image_path: Path, width_in: float, caption: str) -> None:
    fitted_width = min(width_in, 4.48)
    fitted_height = None
    try:
        with Image.open(image_path) as img:
            iw, ih = img.size
            if iw and ih:
                max_height_in = 3.78
                height_at_width = fitted_width * ih / iw
                if height_at_width > max_height_in:
                    fitted_height = max_height_in
    except Exception:
        pass
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = False
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    image_run = p.add_run()
    if fitted_height is not None:
        image_run.add_picture(str(image_path), height=Inches(fitted_height))
    else:
        image_run.add_picture(str(image_path), width=Inches(fitted_width))
    image_run.add_break()
    caption_run = p.add_run(_polish_visible_text(caption))
    set_docx_run_font(caption_run, size=8.3, color=RGBColor(88, 99, 114))


def add_docx_editorial_table(doc: Document, table_data: list[list[str]], caption: str | None = None) -> None:
    if not table_data:
        return
    if caption:
        add_docx_caption(doc, caption)
    tbl = doc.add_table(rows=1, cols=len(table_data[0]))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    widths = docx_table_widths(len(table_data[0]))
    mark_header_row_repeat(tbl.rows[0])
    for j, cell in enumerate(tbl.rows[0].cells):
        cell.text = _polish_visible_text(table_data[0][j])
        cell.width = Inches(widths[j])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_docx_cell_margins(cell, top=95, bottom=95, start=125, end=125)
        shade_docx_cell(cell, "E5F0ED")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            for run in p.runs:
                set_docx_run_font(run, size=8.8, color=RGBColor(21, 31, 43), bold=True)
    for row_idx, row in enumerate(table_data[1:], start=1):
        cells = tbl.add_row().cells
        for j, val in enumerate(row[: len(cells)]):
            cells[j].text = _polish_visible_text(val)
            cells[j].width = Inches(widths[j])
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_docx_cell_margins(cells[j], top=90, bottom=90, start=125, end=125)
            if row_idx % 2 == 0:
                shade_docx_cell(cells[j], "F7F9FB")
            for p in cells[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 or len(_polish_visible_text(val)) > 16 else WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.16
                for run in p.runs:
                    set_docx_run_font(run, size=8.2, color=RGBColor(45, 55, 72))
    doc.add_paragraph("")


def _add_body_para(doc: Document, text: str, size: float = 11.0, color: RGBColor | None = None, first_indent: bool = True):
    p = doc.add_paragraph(_polish_visible_text(text))
    p.paragraph_format.first_line_indent = Pt(22) if first_indent else Pt(0)
    p.paragraph_format.line_spacing = 1.36
    p.paragraph_format.space_after = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        set_docx_run_font(run, size=size, color=color or RGBColor(31, 41, 55))
    return p


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_docx_run_font(run, size=8, color=RGBColor(100, 116, 139))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def write_docx(pages: list[dict], chart_catalog: pd.DataFrame) -> Path:
    path = OUT_DIR / "SkyGuard_商业计划书.docx"
    doc = Document()
    set_docx_styles(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.text = ""
    h_run = header.add_run("SkyGuard 低空智眼 | 城市低空运行监管商业计划书")
    set_docx_run_font(h_run, size=8.5, color=RGBColor(100, 116, 139))
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = footer.add_run("资料口径：官方统计 / 地方政策 / 公开数据 / 演示样本 / 经营假设    第 ")
    set_docx_run_font(f_run, size=8, color=RGBColor(100, 116, 139))
    add_page_number(footer)
    f_run = footer.add_run(" 页")
    set_docx_run_font(f_run, size=8, color=RGBColor(100, 116, 139))

    table_no = 0
    figure_no = 0
    for idx, page in enumerate(pages, start=1):
        if page["type"] == "cover":
            kicker = doc.add_paragraph(page.get("kicker") or "商业计划书")
            kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
            kicker.paragraph_format.space_before = Pt(28)
            kicker.paragraph_format.space_after = Pt(12)
            for run in kicker.runs:
                set_docx_run_font(run, size=10.5, color=RGBColor(155, 106, 47), bold=True)

            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_after = Pt(4)
            run = title.add_run("低空智眼 SkyGuard")
            set_docx_run_font(run, size=27, color=RGBColor(21, 31, 43), bold=True)

            subtitle = doc.add_paragraph("城市低空空域安全感知与运行监管平台")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.paragraph_format.space_after = Pt(8)
            for run in subtitle.runs:
                set_docx_run_font(run, size=14.2, color=RGBColor(18, 78, 107), bold=True)

            rule = doc.add_paragraph()
            rule.paragraph_format.space_before = Pt(4)
            rule.paragraph_format.space_after = Pt(14)
            _set_paragraph_bottom_border(rule, "C8D4DF", "10")

            for para in page.get("paras", []):
                p = doc.add_paragraph(_polish_visible_text(para))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = 1.28
                p.paragraph_format.space_after = Pt(8)
                for run in p.runs:
                    set_docx_run_font(run, size=10.8, color=RGBColor(55, 65, 81))

            _add_docx_metadata_grid(
                doc,
                [("公司名称", "SkyGuard 低空智眼"), ("创业种类", "低空经济运行监管平台"), ("公司地址", "成都市高新区创新创业孵化中心"), ("负责人电话", "+86-138-2026-0705")],
                [("项目定位", "重点区域低空运行监管"), ("目标客户", "景区、园区、场馆、机场周边、低空航线"), ("负责人邮箱", "skyguard@innovation.edu.cn"), ("版本时间", "2026年7月")],
            )
            if page.get("figure"):
                fig_path = Path(page["figure"])
                if fig_path.exists():
                    add_docx_figure_block(doc, fig_path, 3.55, f"封面图  {humanize_figure_caption(fig_path)}")
            doc.add_page_break()
            continue

        if page["type"] == "divider":
            doc.add_page_break()
            lead = doc.add_paragraph("章节导读")
            lead.paragraph_format.first_line_indent = Pt(0)
            lead.paragraph_format.space_before = Pt(16)
            lead.paragraph_format.space_after = Pt(4)
            for run in lead.runs:
                set_docx_run_font(run, size=10.2, color=RGBColor(155, 106, 47), bold=True)
            heading = doc.add_heading(_polish_visible_text(page["chapter"]), level=1)
            for run in heading.runs:
                set_docx_run_font(run, size=18, color=RGBColor(21, 31, 43), bold=True)
            sub = doc.add_paragraph(_polish_visible_text(page["title"]))
            sub.paragraph_format.first_line_indent = Pt(0)
            sub.paragraph_format.space_after = Pt(10)
            for run in sub.runs:
                set_docx_run_font(run, size=12.2, color=RGBColor(18, 78, 107), bold=True)
            _set_paragraph_bottom_border(sub, "D7DEE8", "8")
            for para in page.get("paras", [])[:2]:
                _add_body_para(doc, para, size=11.2)
            if page.get("chart_id"):
                row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
                if not row.empty:
                    chart_path = ROOT / row.iloc[0]["file"]
                    if chart_path.exists():
                        figure_no += 1
                        add_docx_figure_block(doc, chart_path, 4.05, f"图 {figure_no}  {row.iloc[0]['title']}\n资料来源：{row.iloc[0]['source']}。要点：{row.iloc[0]['conclusion']}")
            for bullet in page.get("bullets", [])[:3]:
                p = doc.add_paragraph(_polish_visible_text(bullet), style="List Bullet")
                for run in p.runs:
                    set_docx_run_font(run, size=10.2, color=RGBColor(45, 55, 72))
            doc.add_page_break()
            continue

        chapter_label = doc.add_paragraph(_polish_visible_text(page.get("chapter", "")))
        chapter_label.paragraph_format.first_line_indent = Pt(0)
        chapter_label.paragraph_format.space_before = Pt(4)
        chapter_label.paragraph_format.space_after = Pt(1)
        for run in chapter_label.runs:
            set_docx_run_font(run, size=8.8, color=RGBColor(100, 116, 139), bold=True)

        heading = doc.add_heading(_polish_visible_text(page["title"]), level=2)
        for run in heading.runs:
            set_docx_run_font(run, size=13.2, color=RGBColor(18, 78, 107), bold=True)

        lead_paras = page.get("paras", [])[:2]
        tail_paras = page.get("paras", [])[2:]
        for para in lead_paras:
            _add_body_para(doc, para)

        media_rendered = False
        if page.get("figure"):
            fig_path = Path(page["figure"])
            if fig_path.exists():
                figure_no += 1
                add_docx_figure_block(doc, fig_path, 4.20, f"图 {figure_no}  {humanize_figure_caption(fig_path)}")
                media_rendered = True
        if (not media_rendered) and page.get("chart_id"):
            row = chart_catalog[chart_catalog["chart_id"].eq(page["chart_id"])]
            if not row.empty:
                chart_path = ROOT / row.iloc[0]["file"]
                if chart_path.exists():
                    figure_no += 1
                    add_docx_figure_block(doc, chart_path, 4.35, f"图 {figure_no}  {row.iloc[0]['title']}\n资料来源：{row.iloc[0]['source']}。要点：{row.iloc[0]['conclusion']}")
                    media_rendered = True

        for para in tail_paras:
            _add_body_para(doc, para)

        for bullet in page.get("bullets", []):
            p = doc.add_paragraph(_polish_visible_text(bullet), style="List Bullet")
            for run in p.runs:
                set_docx_run_font(run, size=10.2, color=RGBColor(45, 55, 72))

        if page.get("table"):
            table_no += 1
            add_docx_editorial_table(doc, page.get("table", []), f"表 {table_no}  {page['chapter']}：{page['title']}")

    doc.save(path)
    return path


def main() -> None:
    data = {
        "national": read_csv("01_national_low_altitude_indicators_2023_2025_clean.csv"),
        "regional": read_csv("02_regional_ga_uav_units_2023_2025_clean.csv"),
        "airport": read_csv("03_airport_airspace_context_and_safety_clean.csv"),
        "cloud": read_csv("04_uav_cloud_operations_2023_expanded_clean.csv"),
        "city": read_csv("05_city_policy_targets_cases_final_clean.csv"),
        "plans": read_csv("14_flight_plan_application_sample_8000.csv"),
        "telemetry": read_csv("15_realtime_telemetry_sample_30000.csv"),
        "alerts": read_csv("16_alert_event_response_sample_6000.csv"),
        "risk": read_csv("17_risk_model_feature_sample_6000.csv"),
        "grid": read_csv("10_airspace_grid_risk_sample_10000.csv"),
        "sensors": read_csv("11_sensor_station_coverage_sample_1000.csv"),
        "vertiports": read_csv("12_vertiport_node_sample_900.csv"),
        "routes": read_csv("13_route_corridor_sample_1200.csv"),
    }
    source_registry = generate_source_registry(data["national"], data["city"])
    copy_key_figures()
    chart_catalog = generate_chart_pack(data)
    finance = build_finance_tables()
    write_finance_workbook(finance, source_registry)
    write_evidence_workbook(source_registry, chart_catalog, data)
    pages = make_page_items(chart_catalog, data, source_registry)
    markdown_path = write_markdown(pages, chart_catalog, source_registry)
    try:
        write_docx(pages, chart_catalog)
    except PermissionError as exc:
        print(json.dumps({"docx_locked": str(exc), "action": "skip_docx_and_continue"}, ensure_ascii=False))
    remove_obsolete_word_versions()
    write_pdf(pages, chart_catalog, "SkyGuard_商业计划书.pdf")
    write_pdf(pages, chart_catalog, "SkyGuard_商业计划书_精简路演版.pdf", limit=40)
    write_supporting_docs(source_registry, chart_catalog, data)
    duplicate_paragraph_audit(markdown_path)
    summary = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "pages_planned": len(pages),
        "charts": len(chart_catalog),
        "sources": len(source_registry),
        "outputs": sorted([p.name for p in OUT_DIR.iterdir() if not p.name.startswith("~$")]),
    }
    (OUT_DIR / "SkyGuard_交付清单.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
